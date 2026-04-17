# -*- coding: utf-8 -*-
# FBS warehouses/stocks
FBS_WAREHOUSES_URL = "https://marketplace-api.wildberries.ru/api/v3/warehouses"
FBS_STOCKS_BY_WAREHOUSE_URL = "https://marketplace-api.wildberries.ru/api/v3/stocks/{warehouseId}"
# Supplies API warehouses (for labels tool)
SUPPLIES_WAREHOUSES_URL = "https://supplies-api.wildberries.ru/api/v1/warehouses"
# Prices API
DISCOUNTS_PRICES_API_URL = "https://discounts-prices-api.wildberries.ru/api/v2/list/goods/filter"
# Alternative prices API
PRICES_API_URL = "https://marketplace-api.wildberries.ru/api/v2/list/goods/filter"
# Product history API (including price history)
PRODUCT_HISTORY_API_URL = "https://product-history.wildberries.ru/products/history"
# Alternative product history API
PRODUCT_HISTORY_API_URL_ALT = "https://product-history.wildberries.ru/products/history"
# Commission API
COMMISSION_API_URL = "https://common-api.wildberries.ru/api/v1/tariffs/commission"
DIMENSIONS_API_URL = "https://content-api.wildberries.ru/content/v1/cards/list"
WAREHOUSES_API_URL = "https://common-api.wildberries.ru/api/v1/tariffs/box"
import io
import os
import json
import uuid
import time
import random
import threading
import logging
from collections import defaultdict
from typing import Any

# -----------------------
# Long-running progress
# -----------------------
ORDERS_PROGRESS: dict[int, dict[str, object]] = {}
FINANCE_PROGRESS: dict[int, dict[str, object]] = {}
FINANCE_RESULTS: dict[int, dict[str, Any]] = {}  # ?????? ?????????? ??????????? ?????? ?? user_id
FINANCE_LOADING: dict[int, bool] = {}  # ???? ???????? ??????????? ??????

# -----------------------
# FBW planning in-memory cache (per user)
# -----------------------
# ???????? ??????? ???????? ???????, ???????????? ? /api/fbw/planning/data,
# ????? ??????? "???????? ???????" ?? ???????? ????????? ??????? ? WB.
FBW_PLANNING_DYNAMICS_CACHE: dict[int, dict[str, dict[str, object]]] = {}

def _set_orders_progress(user_id: int, total: int, done: int, key: str | None = None) -> None:
    try:
        current = ORDERS_PROGRESS.get(user_id) or {}
        if key is not None and current.get("key") not in (None, key):
            # New batch -> reset
            current = {"total": 0, "done": 0}
        prev_total = int(current.get("total", 0) or 0)
        prev_done = int(current.get("done", 0) or 0)
        new_total = max(prev_total, max(0, int(total)))
        new_done = max(prev_done, max(0, int(done)))
        ORDERS_PROGRESS[user_id] = {"key": key, "total": new_total, "done": new_done}
    except Exception:
        pass

def _clear_orders_progress(user_id: int, key: str | None = None) -> None:
    try:
        cur = ORDERS_PROGRESS.get(user_id)
        if cur is None:
            return
        if key is None or cur.get("key") == key:
            del ORDERS_PROGRESS[user_id]
    except Exception:
        pass

def _set_finance_progress(user_id: int, current: int, total: int, period: str = "") -> None:
    try:
        FINANCE_PROGRESS[user_id] = {"current": current, "total": total, "period": period}
    except Exception:
        pass

def _get_finance_progress(user_id: int) -> dict[str, object]:
    return FINANCE_PROGRESS.get(user_id) or {"current": 0, "total": 0, "period": ""}

def _clear_finance_progress(user_id: int) -> None:
    try:
        if user_id in FINANCE_PROGRESS:
            del FINANCE_PROGRESS[user_id]
    except Exception:
        pass
from datetime import datetime, timedelta, timezone
from typing import List, Dict, Any, Tuple

import requests
from flask import Flask, render_template, request, send_file, redirect, url_for, session, flash, jsonify, send_from_directory, has_request_context, Response
import xlwt
import jwt
from io import BytesIO
from openpyxl import Workbook, load_workbook
import xlrd
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from sqlalchemy import text
from flask_login import (
    LoginManager,
    UserMixin,
    login_required,
    login_user,
    logout_user,
    current_user,
)

APP_VERSION = "1.0.1"

from utils.constants import (
    LAST_RESULTS_CACHE_MAX_BYTES,
    WB_ORDERS_FETCH_MAX_PAGES,
    WB_ORDERS_FETCH_MAX_PAGES_INTRADAY,
    WB_ORDERS_PAGE_SLEEP_S,
    WB_ORDERS_PAGE_SLEEP_INTRADAY_S,
    ORDERS_TODAY_CACHE_TTL_SECONDS,
)
from utils.cache import period_cache_day_entry_is_fresh

# --- Throttling for WB supplies API ---
_last_supplies_api_call_ts: float = 0.0
_SUPPLIES_API_MIN_INTERVAL_S: float = float(os.getenv("SUPPLIES_API_MIN_INTERVAL_S", "2.0"))

def _supplies_api_throttle() -> None:
    """Ensure at most ~30 req/min (min interval ~2s) across supplies endpoints."""
    global _last_supplies_api_call_ts
    if _SUPPLIES_API_MIN_INTERVAL_S <= 0:
        return
    now = time.time()
    delta = now - _last_supplies_api_call_ts
    if delta < _SUPPLIES_API_MIN_INTERVAL_S:
        time.sleep(_SUPPLIES_API_MIN_INTERVAL_S - delta)
    _last_supplies_api_call_ts = time.time()

# -------------------- ??? ???????? ????? --------------------
DEFAULT_MARGIN_SETTINGS = {
    "tax": 6.0,         # ?????
    "storage": 0.5,     # ????????
    "receiving": 1.0,   # ???????
    "acquiring": 1.7,   # ?????????
    "scheme": "FBW",   # ????? ?????? ? WB
    "warehouse": "",    # ????? ????????
    "warehouse_coef": 0.0,  # ??????????? ????????? ??????
    "localization_index": 1.0,  # ?????? ???????????
}

def _get_cache_dir() -> str:
    cache_dir = os.path.join(os.path.dirname(__file__), "cache")
    os.makedirs(cache_dir, exist_ok=True)
    return cache_dir

def load_user_margin_settings(user_id: int) -> dict:
    try:
        cache_dir = _get_cache_dir()
        path = os.path.join(cache_dir, f"margin_settings_{user_id}.json")
        if os.path.exists(path):
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f) or {}
            result = DEFAULT_MARGIN_SETTINGS.copy()
            for key, default_val in DEFAULT_MARGIN_SETTINGS.items():
                val = data.get(key, default_val)
                if key in ["scheme", "warehouse"]:
                    # ????????? ????????
                    result[key] = str(val or default_val)
                else:
                    # ???????? ????????
                    try:
                        result[key] = float(val)
                    except Exception:
                        result[key] = default_val
            return result
    except Exception as e:
        print(f"?????? ?????? ???????? ?????: {e}")
    return DEFAULT_MARGIN_SETTINGS.copy()

def save_user_margin_settings(user_id: int, settings: dict) -> dict:
    normalized = DEFAULT_MARGIN_SETTINGS.copy()
    for key, default_val in DEFAULT_MARGIN_SETTINGS.items():
        val = settings.get(key, default_val)
        if key in ["scheme", "warehouse"]:
            # ????????? ????????
            normalized[key] = str(val or default_val)
        else:
            # ???????? ????????
            try:
                normalized[key] = float(val)
            except Exception:
                normalized[key] = default_val
    try:
        cache_dir = _get_cache_dir()
        path = os.path.join(cache_dir, f"margin_settings_{user_id}.json")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(normalized, f, ensure_ascii=False)
    except Exception as e:
        print(f"?????? ?????????? ???????? ?????: {e}")
    return normalized

def _read_version() -> str:
    try:
        path = os.path.join(os.path.dirname(__file__), "VERSION")
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                v = f.read().strip()
                return v or APP_VERSION
    except Exception:
        pass
    return APP_VERSION

def _read_changelog_md() -> str:
    """Read Markdown changelog (CHANGELOG.md). If missing, try convert from JSON fallback."""
    md_path = os.path.join(os.path.dirname(__file__), "CHANGELOG.md")
    if os.path.isfile(md_path):
        try:
            with open(md_path, "r", encoding="utf-8") as f:
                return f.read()
        except Exception:
            return ""
    # Fallback: convert from old changelog.json if exists
    try:
        json_path = os.path.join(os.path.dirname(__file__), "changelog.json")
        if os.path.isfile(json_path):
            with open(json_path, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data, list) and data:
                parts: List[str] = [f"# ?????????? ? ?????????\n\n??????? ??????: {_read_version()}\n"]
                for e in data:
                    ver = str(e.get("version") or "").strip()
                    date = str(e.get("date") or "").strip()
                    parts.append(f"\n## ?????? {ver} ? {date}\n")
                    if e.get("html"):
                        parts.append(e["html"])  # embed existing html as-is
                        parts.append("\n")
                    else:
                        notes = e.get("notes") or []
                        for n in notes:
                            parts.append(f"- {n}")
                        parts.append("\n")
                return "\n".join(parts)
    except Exception:
        pass
    # Default stub
    return f"# ?????????? ? ?????????\n\n??????? ??????: {_read_version()}\n\n## ?????? {_read_version()} ? {datetime.now(MOSCOW_TZ).strftime('%d.%m.%Y')}\n- ?????????????? ??????\n"

def _write_version(version: str) -> None:
    try:
        path = os.path.join(os.path.dirname(__file__), "VERSION")
        with open(path, "w", encoding="utf-8") as f:
            f.write((version or "").strip() or "0.0.0")
    except Exception:
        pass

def _write_changelog_md(content: str) -> None:
    try:
        path = os.path.join(os.path.dirname(__file__), "CHANGELOG.md")
        with open(path, "w", encoding="utf-8") as f:
            f.write(content or "")
    except Exception:
        pass

app = Flask(__name__)
app.secret_key = os.getenv("FLASK_SECRET", "dev-secret-change-me")
app.config["SQLALCHEMY_DATABASE_URI"] = os.getenv("DATABASE_URL", f"sqlite:///{os.path.join(os.path.dirname(__file__), 'app.db')}")
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False

# ????????? ???????????
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()  # ????? ? ???????
    ]
)
logger = logging.getLogger(__name__)
# ????????? ?????? ??? ?????????? ?????? ???????????
app.config["PERMANENT_SESSION_LIFETIME"] = 86400  # 24 ????
app.config["REMEMBER_COOKIE_DURATION"] = timedelta(days=30)
app.config["SESSION_COOKIE_SECURE"] = False  # ??? ??????????
app.config["SESSION_COOKIE_HTTPONLY"] = True
app.config["SESSION_COOKIE_SAMESITE"] = "Lax"
app.config["SESSION_COOKIE_NAME"] = "fuckberry_session"
app.config["SESSION_COOKIE_PATH"] = "/"

# ??????????? db ?? models.py ? ?????????????? ? ???????????
from models import db, User, Notification, PurchasePrice, delete_user_with_related
db.init_app(app)

from utils.wb_token import (
    user_may_call_wb_api,
    effective_wb_api_token,
    token_for_wb_request,
    is_user_valid_for_app as _is_user_valid_now,
    wb_api_key_expiry_summary,
    wb_api_key_expiry_banner,
)

login_manager = LoginManager(app)
login_manager.login_view = "login"
login_manager.login_message = "??????????, ??????? ? ??????? ??? ??????? ? ???? ????????."
login_manager.login_message_category = "info"
login_manager.session_protection = "strong"  # ?????? ??????
login_manager.refresh_view = "login"  # ???????? ??? ?????????? ??????
login_manager.needs_refresh_message = "??????????, ??????? ? ??????? ??? ??????? ? ???? ????????."
login_manager.needs_refresh_message_category = "info"

# --- Register Blueprints ---
from blueprints.auth import auth_bp
from blueprints.changelog import changelog_bp
from blueprints.profile import profile_bp
from blueprints.notifications import notifications_bp
from blueprints.orders import orders_bp
from blueprints.coefficients import coefficients_bp
from blueprints.fbs import fbs_bp
from blueprints.fbs_supplies import fbs_supplies_bp
from blueprints.dbs import dbs_bp
from blueprints.fbs_stock import fbs_stock_bp
from blueprints.fbw import fbw_bp
from blueprints.fbw_planning import fbw_planning_bp
from blueprints.products import products_bp
from blueprints.stocks import stocks_bp
from blueprints.reports import reports_bp
from blueprints.tools import tools_bp
from blueprints.admin import admin_bp

app.register_blueprint(auth_bp)
app.register_blueprint(changelog_bp)
app.register_blueprint(profile_bp)
app.register_blueprint(notifications_bp)
app.register_blueprint(orders_bp)
app.register_blueprint(coefficients_bp)
app.register_blueprint(fbs_bp)
app.register_blueprint(fbs_supplies_bp)
app.register_blueprint(dbs_bp)
app.register_blueprint(fbs_stock_bp)
app.register_blueprint(fbw_bp)
app.register_blueprint(fbw_planning_bp)
app.register_blueprint(products_bp)
app.register_blueprint(stocks_bp)
app.register_blueprint(reports_bp)
app.register_blueprint(tools_bp)
app.register_blueprint(admin_bp)

# --- DB init helpers (portable across common DBs) ---
def _ensure_schema_users_validity_columns() -> None:
    try:
        engine = db.engine
        dialect = getattr(engine, "name", getattr(engine.dialect, "name", "")) or getattr(engine.dialect, "name", "")
        # Use transactional context so DDL is committed on PG/MySQL
        with engine.begin() as conn:
            if dialect == "sqlite":
                try:
                    rows = conn.execute(text("PRAGMA table_info(users)")).fetchall()
                    cols = {r[1] for r in rows}
                    if "valid_from" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN valid_from DATE"))
                    if "valid_to" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN valid_to DATE"))
                    if "phone" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN phone VARCHAR(64)"))
                    if "email" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN email VARCHAR(120)"))
                    if "shipper_name" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN shipper_name VARCHAR(255)"))
                    if "shipper_address" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN shipper_address VARCHAR(255)"))
                    if "contact_person" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN contact_person VARCHAR(255)"))
                    if "display_name" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN display_name VARCHAR(255)"))
                    if "tax_rate" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN tax_rate REAL"))
                    if "org_display_name" not in cols:
                        conn.execute(text("ALTER TABLE users ADD COLUMN org_display_name VARCHAR(255)"))
                except Exception:
                    pass
            elif dialect in ("postgresql", "postgres"):
                # IF NOT EXISTS works on Postgres
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS valid_from DATE"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS valid_to DATE"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS phone VARCHAR(64)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS email VARCHAR(120)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS shipper_name VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS shipper_address VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS contact_person VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS display_name VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS tax_rate FLOAT"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS org_display_name VARCHAR(255)"))
                except Exception:
                    pass
            elif dialect in ("mysql", "mariadb"):
                # MySQL 8.0+ supports IF NOT EXISTS
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS valid_from DATE"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS valid_to DATE"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS phone VARCHAR(64)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS email VARCHAR(120)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS shipper_name VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS shipper_address VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS contact_person VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS display_name VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS tax_rate FLOAT"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN IF NOT EXISTS org_display_name VARCHAR(255)"))
                except Exception:
                    pass
            else:
                # Best-effort generic
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN valid_from DATE"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN valid_to DATE"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN phone VARCHAR(64)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN tax_rate FLOAT"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN email VARCHAR(120)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN shipper_name VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN shipper_address VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN contact_person VARCHAR(255)"))
                except Exception:
                    pass
                try:
                    conn.execute(text("ALTER TABLE users ADD COLUMN org_display_name VARCHAR(255)"))
                except Exception:
                    pass
    except Exception:
        pass


# Note: Flask 3.x removed before_first_request. We init DB in __main__ when run as a script.

API_URL = "https://statistics-api.wildberries.ru/api/v1/supplier/orders"
SALES_API_URL = "https://statistics-api.wildberries.ru/api/v1/supplier/sales"
CACHE_DIR = os.path.join(os.path.dirname(__file__), "cache")
if not os.path.isdir(CACHE_DIR):
    os.makedirs(CACHE_DIR, exist_ok=True)

# ????? ??? ?????????????? ?????????????? ?????????? ???? (?? ?????????????)
_supplies_cache_updating: dict[int, bool] = {}
_orders_cache_updating: dict[int, bool] = {}

# ?????????? ??????????????? ???? ????????: ?? ????????? ?????????
SUPPLIES_CACHE_AUTO = os.getenv("SUPPLIES_CACHE_AUTO", "0") == "1"

FBS_NEW_URL = "https://marketplace-api.wildberries.ru/api/v3/orders/new"
FBS_ORDERS_URL = "https://marketplace-api.wildberries.ru/api/v3/orders"
FBS_ORDERS_STATUS_URL = "https://marketplace-api.wildberries.ru/api/v3/orders/status"
FBS_SUPPLIES_LIST_URL = "https://marketplace-api.wildberries.ru/api/v3/supplies"
FBS_SUPPLY_INFO_URL = "https://marketplace-api.wildberries.ru/api/v3/supplies/{supplyId}"
FBS_SUPPLY_ORDERS_URL = "https://marketplace-api.wildberries.ru/api/v3/supplies/{supplyId}/orders"

# DBS (Delivery by Seller) API
DBS_NEW_URL = "https://marketplace-api.wildberries.ru/api/v3/dbs/orders/new"
DBS_STATUS_URL = "https://marketplace-api.wildberries.ru/api/v3/dbs/orders/status"
DBS_ORDERS_URL = "https://marketplace-api.wildberries.ru/api/v3/dbs/orders"

SELLER_INFO_URL = "https://common-api.wildberries.ru/api/v1/seller-info"

ACCEPT_COEFS_URL = "https://supplies-api.wildberries.ru/api/v1/acceptance/coefficients"
# FBW supplies API
FBW_SUPPLIES_LIST_URL = "https://supplies-api.wildberries.ru/api/v1/supplies"
FBW_SUPPLY_DETAILS_URL = "https://supplies-api.wildberries.ru/api/v1/supplies/{id}"
FBW_SUPPLY_GOODS_URL = "https://supplies-api.wildberries.ru/api/v1/supplies/{id}/goods"
FBW_SUPPLY_PACKAGE_URL = "https://supplies-api.wildberries.ru/api/v1/supplies/{id}/package"
# Wildberries Content API: cards list
WB_CARDS_LIST_URL = "https://content-api.wildberries.ru/content/v2/get/cards/list"
WB_CARDS_UPDATE_URL = "https://content-api.wildberries.ru/content/v2/cards/update"
STOCKS_API_URL = "https://statistics-api.wildberries.ru/api/v1/supplier/stocks"
FIN_REPORT_URL = "https://statistics-api.wildberries.ru/api/v5/supplier/reportDetailByPeriod"

# Wildberries Content API: cards list
WB_CARDS_LIST_URL = "https://content-api.wildberries.ru/content/v2/get/cards/list"


# Timezone helpers (Moscow)
try:
    from zoneinfo import ZoneInfo  # Python 3.9+
    MOSCOW_TZ = ZoneInfo("Europe/Moscow")
except Exception:  # Fallback to fixed UTC+3 if zoneinfo unavailable
    MOSCOW_TZ = timezone(timedelta(hours=3))

def to_moscow(dt: datetime | None) -> datetime | None:
    if dt is None:
        return None
    try:
        # If datetime is naive, consider it already in Moscow time (many WB fields come without TZ but are MSK)
        if dt.tzinfo is None:
            return dt.replace(tzinfo=MOSCOW_TZ)
        # If it has timezone (e.g., Z/UTC), convert to Moscow
        return dt.astimezone(MOSCOW_TZ)
    except Exception:
        return dt


def format_int_thousands(value: Any) -> str:
    try:
        return f"{int(value):,}".replace(",", " ")
    except Exception:
        return str(value)


def _parse_iso_datetime(value: str | None) -> datetime | None:
    if not value:
        return None
    try:
        # Support trailing 'Z'
        s = value.replace("Z", "+00:00")
        return datetime.fromisoformat(s)
    except Exception:
        try:
            # Fallback common formats
            return datetime.strptime(value, "%Y-%m-%dT%H:%M:%S%z")
        except Exception:
            try:
                return datetime.strptime(value, "%Y-%m-%dT%H:%M:%S")
            except Exception:
                return None


def _fmt_dt_moscow(value: str | None, with_time: bool = True) -> str:
    dt = _parse_iso_datetime(value)
    if not dt:
        return ""
    msk = to_moscow(dt) or dt
    return msk.strftime("%d.%m.%Y %H:%M") if with_time else msk.strftime("%d.%m.%Y")


def _fbw_status_from_id(status_id: Any) -> str | None:
    """
    ??????????? statusID ?? API Wildberries ? ????????? ?????? ???????? ????????????:
    1 ? ?? ?????????????
    2 ? ?????????????
    3 ? ???????? ?????????
    4 ? ???? ???????
    5 ? ???????
    6 ? ????????? ?? ???????
    """
    if status_id is None:
        return None
    try:
        sid = int(status_id)
        status_map = {
            1: "?? ?????????????",
            2: "?????????????",
            3: "???????? ?????????",
            4: "???? ???????",
            5: "???????",
            6: "????????? ?? ???????",
        }
        return status_map.get(sid)
    except (ValueError, TypeError):
        return None


def fetch_fbw_supplies_list(token: str, days_back: int = 90) -> list[dict[str, Any]]:
    if not token:
        return []
    date_to = datetime.now(MOSCOW_TZ).date()
    date_from = date_to - timedelta(days=days_back)
    # Some WB endpoints treat 'till' as exclusive. Add +1 day to include entire current day.
    date_till = date_to + timedelta(days=1)
    body = {
        "dates": [
            {
                "from": date_from.strftime("%Y-%m-%d"),
                "till": date_till.strftime("%Y-%m-%d"),
                "type": "createDate",
            }
        ]
        # statusIDs optional; omit to include all
    }
    # Try Bearer first
    headers1 = {"Authorization": f"Bearer {token}"}
    try:
        _supplies_api_throttle()
        resp = post_with_retry(FBW_SUPPLIES_LIST_URL, headers1, body)
        items = resp.json() or []
    except Exception:
        headers2 = {"Authorization": f"{token}"}
        _supplies_api_throttle()
        resp = post_with_retry(FBW_SUPPLIES_LIST_URL, headers2, body)
        items = resp.json() or []
    # Sort by createDate desc
    def _key(it: dict[str, Any]):
        return _parse_iso_datetime(str(it.get("createDate") or "")) or datetime.min.replace(tzinfo=MOSCOW_TZ)
    items.sort(key=_key, reverse=True)
    return items


def fetch_fbw_supply_details(token: str, supply_id: int | str) -> dict[str, Any] | None:
    if not token or not supply_id:
        return None
    url = FBW_SUPPLY_DETAILS_URL.format(id=supply_id)
    headers1 = {"Authorization": f"Bearer {token}"}
    try:
        resp = get_with_retry(url, headers1, params={})
        return resp.json()
    except Exception:
        headers2 = {"Authorization": f"{token}"}
        try:
            resp = get_with_retry(url, headers2, params={})
            return resp.json()
        except Exception:
            return None


def fetch_fbw_supply_goods(token: str, supply_id: int | str, limit: int = 200, offset: int = 0) -> list[dict[str, Any]]:
    if not token or not supply_id:
        return []
    url = FBW_SUPPLY_GOODS_URL.format(id=supply_id)
    headers1 = {"Authorization": f"Bearer {token}"}
    params = {"limit": limit, "offset": offset}
    try:
        resp = get_with_retry(url, headers1, params=params)
        return resp.json() or []
    except Exception:
        headers2 = {"Authorization": f"{token}"}
        try:
            resp = get_with_retry(url, headers2, params=params)
            return resp.json() or []
        except Exception:
            return []


def fetch_fbw_supply_packages(token: str, supply_id: int | str) -> list[dict[str, Any]]:
    if not token or not supply_id:
        return []
    url = FBW_SUPPLY_PACKAGE_URL.format(id=supply_id)
    headers1 = {"Authorization": f"Bearer {token}"}
    try:
        resp = get_with_retry(url, headers1, params={})
        return resp.json() or []
    except Exception:
        headers2 = {"Authorization": f"{token}"}
        try:
            resp = get_with_retry(url, headers2, params={})
            return resp.json() or []
        except Exception:
            return []


def fetch_fbw_last_supplies(token: str, limit: int = 15) -> list[dict[str, Any]]:
    base_list = fetch_fbw_supplies_list(token)
    supplies: list[dict[str, Any]] = []
    
    # ????????? ???????????? ?????? ??? ??????????? (?????????? ?????? ??????????????? ????, ?? ??????)
    cached = load_fbw_supplies_cache() or {}
    cached_items = cached.get("items") or []
    cached_map: dict[str, dict[str, Any]] = {}
    for item in cached_items:
        sid = str(item.get("supply_id") or item.get("supplyID") or item.get("supplyId") or item.get("id") or "")
        if sid:
            cached_map[sid] = item
    # ???? ?? ????? ??????, ?? ???????? ?? ?????? `limit` ???????? ??? ???????????
    for it in base_list:
        if len(supplies) >= max(0, int(limit)):
            break
        supply_id = it.get("supplyID") or it.get("supplyId") or it.get("id")
        supply_id_str = str(supply_id or "")
        
        # ????????? ??? ??? ??????????? (?????? ??????????????? ????, ?????? ?????? ????? ?? ???????)
        cached_item = cached_map.get(supply_id_str)
        
        # ?????? ???????? ?????????? ?????? (??????? ?????????? ????????)
        _supplies_api_throttle()
        details = fetch_fbw_supply_details(token, supply_id)
        details = details or {}  # ??????????, ??? details - ??? ???????
        # Normalize fields; prefer details when available, fallback to list fields
        create_date = details.get("createDate") or it.get("createDate")
        supply_date = details.get("supplyDate") or it.get("supplyDate")
        fact_date = details.get("factDate") or it.get("factDate")
        # ??????????? ?????????? statusID ?? API (????? ???????? ????????)
        status_id = details.get("statusID") or it.get("statusID")
        status_name = _fbw_status_from_id(status_id)
        
        # ???? statusID ?? ??? ??????????, ??????? ????????? ????
        if not status_name:
            status_name = (
                details.get("statusName") 
                or details.get("status") 
                or details.get("statusText")
                or it.get("statusName") 
                or it.get("status")
                or it.get("statusText")
                or ""
            )
        # ???? ?????? ??? ??? ??????, ???????? ?????????? ??? ?? ?????? ?????
        if not status_name:
            # ?????????, ???? ?? factDate - ???? ????, ?????? ???????? ???????
            if fact_date:
                status_name = "???????"
            # ????? ????????? supplyDate - ???? ????, ?????? ?????????????
            elif supply_date:
                # ?????????, ?????? ?? ???????? ????
                try:
                    planned_dt = _parse_iso_datetime(str(supply_date))
                    if planned_dt:
                        planned_dt_msk = to_moscow(planned_dt) if planned_dt else None
                        if planned_dt_msk:
                            today = datetime.now(MOSCOW_TZ).date()
                            planned_date = planned_dt_msk.date()
                            if planned_date < today:
                                status_name = "???????? ?????????"
                            else:
                                status_name = "?????????????"
                except Exception:
                    status_name = "?????????????"
            else:
                status_name = "?? ?????????????"
        # ???? ?????? ??????, ?? ???? ???????????? ??????, ????? ???????????? ??? ??? fallback
        if not status_name and cached_item:
            cached_status_val = str(cached_item.get("status", "")).strip()
            if cached_status_val:
                status_name = cached_status_val
        # ???????? ????????? ("?? ?????????????") ?? ??????????
        if "?? ?????????????" in str(status_name or ""):
            continue

        warehouse_name = details.get("warehouseName") or it.get("warehouseName") or ""
        # ????????? ???? ????????: ???? boxTypeName ???????????, ?????????? boxTypeID ? ??????????? ? ???????? ????????
        box_type = details.get("boxTypeName") or it.get("boxTypeName")
        if not box_type:
            box_type_id = details.get("boxTypeID") or it.get("boxTypeID")
            if box_type_id is not None:
                # ??????????? ID ? ????????
                box_type_map = {1: "??? ???????", 2: "??????"}
                box_type = box_type_map.get(int(box_type_id), str(box_type_id))
            else:
                box_type = ""
        total_qty = details.get("quantity")
        accepted_qty = details.get("acceptedQuantity")
        acceptance_cost = details.get("acceptanceCost")
        paid_coef = details.get("paidAcceptanceCoefficient")
        
        # ???? ???? ???????????? ?????????? ???????, ????????? ???
        package_count = None
        if cached_item and "package_count" in cached_item:
            package_count = cached_item["package_count"]
        
        supply_data = {
            "supply_id": supply_id_str,
            "type": str(box_type) if box_type is not None else "",
            "created_at": _fmt_dt_moscow(create_date, with_time=False),
            "total_goods": int(total_qty) if isinstance(total_qty, (int, float)) and total_qty is not None else None,
            "accepted_goods": int(accepted_qty) if isinstance(accepted_qty, (int, float)) and accepted_qty is not None else None,
            "warehouse": warehouse_name or "",
            "acceptance_coefficient": paid_coef,
            "acceptance_cost": acceptance_cost,
            "planned_date": _fmt_dt_moscow(supply_date, with_time=False),
            "fact_date": _fmt_dt_moscow(fact_date, with_time=True),
            "status": status_name or "",
        }
        
        if package_count is not None:
            supply_data["package_count"] = package_count
            
        supplies.append(supply_data)

    return supplies
def fetch_fbw_supplies_range(token: str, offset: int, limit: int) -> list[dict[str, Any]]:
    base_list = fetch_fbw_supplies_list(token)
    if offset < 0:
        offset = 0
    # offset/limit ????????? ?? *?????????* ?????? (??????? ?????????),
    # ?? ??? ??????????? ?? ????? ??????????? "?? ?????????????".
    end = offset + max(0, int(limit))
    slice_ids = base_list[offset:end]
    supplies: list[dict[str, Any]] = []
    
    # ????????? ???????????? ?????? ??? ??????????? (?????????? ?????? ??????????????? ????, ?? ??????)
    cached = load_fbw_supplies_cache() or {}
    cached_items = cached.get("items") or []
    cached_map: dict[str, dict[str, Any]] = {}
    for item in cached_items:
        sid = str(item.get("supply_id") or item.get("supplyID") or item.get("supplyId") or item.get("id") or "")
        if sid:
            cached_map[sid] = item
    
    for it in slice_ids:
        supply_id = it.get("supplyID") or it.get("supplyId") or it.get("id")
        supply_id_str = str(supply_id or "")
        
        # ????????? ??? ??? ??????????? (?????? ??????????????? ????, ?????? ?????? ????? ?? ???????)
        cached_item = cached_map.get(supply_id_str)
        
        # ?????? ???????? ?????????? ?????? (??????? ?????????? ????????)
        _supplies_api_throttle()
        details = fetch_fbw_supply_details(token, supply_id)
        details = details or {}  # ??????????, ??? details - ??? ???????
        create_date = details.get("createDate") or it.get("createDate")
        supply_date = details.get("supplyDate") or it.get("supplyDate")
        fact_date = details.get("factDate") or it.get("factDate")
        # ??????????? ?????????? statusID ?? API (????? ???????? ????????)
        status_id = details.get("statusID") or it.get("statusID")
        status_name = _fbw_status_from_id(status_id)
        
        # ???? statusID ?? ??? ??????????, ??????? ????????? ????
        if not status_name:
            status_name = (
                details.get("statusName") 
                or details.get("status") 
                or details.get("statusText")
                or it.get("statusName") 
                or it.get("status")
                or it.get("statusText")
                or ""
            )
        # ???? ?????? ??? ??? ??????, ???????? ?????????? ??? ?? ?????? ?????
        if not status_name:
            # ?????????, ???? ?? factDate - ???? ????, ?????? ???????? ???????
            if fact_date:
                status_name = "???????"
            # ????? ????????? supplyDate - ???? ????, ?????? ?????????????
            elif supply_date:
                # ?????????, ?????? ?? ???????? ????
                try:
                    planned_dt = _parse_iso_datetime(str(supply_date))
                    if planned_dt:
                        planned_dt_msk = to_moscow(planned_dt) if planned_dt else None
                        if planned_dt_msk:
                            today = datetime.now(MOSCOW_TZ).date()
                            planned_date = planned_dt_msk.date()
                            if planned_date < today:
                                status_name = "???????? ?????????"
                            else:
                                status_name = "?????????????"
                except Exception:
                    status_name = "?????????????"
            else:
                status_name = "?? ?????????????"
        # ???? ?????? ??????, ?? ???? ???????????? ??????, ????? ???????????? ??? ??? fallback
        if not status_name and cached_item:
            cached_status_val = str(cached_item.get("status", "")).strip()
            if cached_status_val:
                status_name = cached_status_val
        warehouse_name = details.get("warehouseName") or it.get("warehouseName") or ""
        # ????????? ???? ????????: ???? boxTypeName ???????????, ?????????? boxTypeID ? ??????????? ? ???????? ????????
        box_type = details.get("boxTypeName") or it.get("boxTypeName")
        if not box_type:
            box_type_id = details.get("boxTypeID") or it.get("boxTypeID")
            if box_type_id is not None:
                # ??????????? ID ? ????????
                box_type_map = {1: "??? ???????", 2: "??????"}
                box_type = box_type_map.get(int(box_type_id), str(box_type_id))
            else:
                box_type = ""
        total_qty = details.get("quantity")
        accepted_qty = details.get("acceptedQuantity")
        acceptance_cost = details.get("acceptanceCost")
        paid_coef = details.get("paidAcceptanceCoefficient")
        
        # ???? ???? ???????????? ?????????? ???????, ????????? ???
        package_count = None
        if cached_item and "package_count" in cached_item:
            package_count = cached_item["package_count"]
        
        supply_data = {
            "supply_id": supply_id_str,
            "type": str(box_type) if box_type is not None else "",
            "created_at": _fmt_dt_moscow(create_date, with_time=False),
            "total_goods": int(total_qty) if isinstance(total_qty, (int, float)) and total_qty is not None else None,
            "accepted_goods": int(accepted_qty) if isinstance(accepted_qty, (int, float)) and accepted_qty is not None else None,
            "warehouse": warehouse_name or "",
            "acceptance_coefficient": paid_coef,
            "acceptance_cost": acceptance_cost,
            "planned_date": _fmt_dt_moscow(supply_date, with_time=False),
            "fact_date": _fmt_dt_moscow(fact_date, with_time=True),
            "status": status_name or "",
        }
        
        if package_count is not None:
            supply_data["package_count"] = package_count
            
        supplies.append(supply_data)
    return supplies


def format_money_ru(value: Any) -> str:
    try:
        return f"{float(value):,.2f}".replace(",", " ").replace(".", ",")
    except Exception:
        return str(value)


app.jinja_env.filters["num_space"] = format_int_thousands
app.jinja_env.filters["money_ru"] = format_money_ru


def format_dmy(date_str: str) -> str:
    try:
        return datetime.strptime(date_str, "%Y-%m-%d").strftime("%d.%m.%Y")
    except Exception:
        return date_str or ""


def extract_nm(value: Any) -> str:
    try:
        import re
        s = str(value)
        m = re.search(r"(\d{7,12})", s)
        return m.group(1) if m else ""
    except Exception:
        return ""


app.jinja_env.filters["extract_nm"] = extract_nm


def days_left_from_str(date_str: str | None) -> int | None:
    try:
        if not date_str:
            return None
        # ???????? ???????? ??????? "??.??.????" (?? ??? ??????????? planned_date)
        dt = datetime.strptime(date_str.strip(), "%d.%m.%Y")
        today = datetime.now(MOSCOW_TZ).date()
        diff = (dt.date() - today).days
        return diff
    except Exception:
        return None


app.jinja_env.filters["days_left"] = days_left_from_str

def _merge_package_counts(items: list[dict[str, Any]], cached_items: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    try:
        cache_map: dict[str, int] = {}
        for it in (cached_items or []):
            sid = str(it.get("supply_id") or it.get("supplyID") or it.get("supplyId") or it.get("id") or "")
            pc = it.get("package_count")
            try:
                pc_int = int(pc)
            except Exception:
                pc_int = 0
            if sid and pc_int > 0:
                cache_map[sid] = pc_int
        merged: list[dict[str, Any]] = []
        for it in items:
            sid = str(it.get("supply_id") or it.get("supplyID") or it.get("supplyId") or it.get("id") or "")
            if sid in cache_map and (not it.get("package_count") or int(it.get("package_count") or 0) == 0):
                # Copy to avoid mutating original
                new_it = dict(it)
                new_it["package_count"] = cache_map[sid]
                merged.append(new_it)
            else:
                merged.append(it)
        return merged
    except Exception:
        return items
def _preload_package_counts(token: str, supplies: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """
    ?????????????? ????????? ?????????? ??????? ??? ????????, ??????? ??? ?? ????? ???? ??????????.
    ??? ????????? ?????????? ?????? ? ???????? ??? ?????????? ?????? ????????.
    """
    if not token or not supplies:
        return supplies
    
    # ??????? ???????? ??? ?????????? ? ?????????? ???????
    supplies_to_update = []
    for supply in supplies:
        if not supply.get("package_count") or int(supply.get("package_count") or 0) == 0:
            supplies_to_update.append(supply)
    
    if not supplies_to_update:
        return supplies
    
    # ????????? ?????????? ??????? ??? ????????? ???????? (???????????? ?????????? ????????)
    updated_supplies = []
    max_requests = 5  # ???????????? ?????????? ?????????????? API ????????
    
    for i, supply in enumerate(supplies):
        if supply in supplies_to_update and i < max_requests:
            try:
                supply_id = supply.get("supply_id") or supply.get("supplyID") or supply.get("supplyId") or supply.get("id")
                if supply_id:
                    packages = fetch_fbw_supply_packages(token, supply_id)
                    package_count = len(packages) if isinstance(packages, list) else 0
                    
                    # ??????? ????? ? ??????????? ??????????? ???????
                    updated_supply = dict(supply)
                    updated_supply["package_count"] = package_count
                    updated_supplies.append(updated_supply)
                else:
                    updated_supplies.append(supply)
            except Exception:
                # ? ?????? ?????? ????????? ???????? ??? ?????????
                updated_supplies.append(supply)
        else:
            updated_supplies.append(supply)
    
    return updated_supplies


def time_ago_ru(dt_val: Any) -> str:
    try:
        if dt_val is None:
            return ""
        if isinstance(dt_val, str):
            s = dt_val.strip()
            dt = parse_wb_datetime(s)
            if dt is None:
                # Try ISO first
                try:
                    dt = datetime.fromisoformat(s)
                except Exception:
                    dt = None
            if dt is None:
                # Try common RU formats
                for fmt in ("%d.%m.%Y %H:%M:%S", "%d.%m.%Y %H:%M", "%d.%m.%Y"):
                    try:
                        dt = datetime.strptime(s, fmt)
                        break
                    except Exception:
                        dt = None
            if dt is None:
                return ""
        elif isinstance(dt_val, datetime):
            dt = dt_val
        else:
            return ""
        # Convert both to Moscow time for consistent human delta
        dt = to_moscow(dt)
        now = datetime.now(MOSCOW_TZ)
        if dt > now:
            return "?????? ???"
        diff = now - dt
        days = diff.days
        seconds = diff.seconds
        hours = seconds // 3600
        minutes = (seconds % 3600) // 60
        if days > 0:
            return f"{days} ? {hours} ? ?????" if hours > 0 else f"{days} ? ?????"
        if hours > 0:
            return f"{hours} ? {minutes} ? ?????" if minutes > 0 else f"{hours} ? ?????"
        if minutes > 0:
            return f"{minutes} ? ?????"
        return "?????? ???"
    except Exception:
        return ""


app.jinja_env.filters["time_ago_ru"] = time_ago_ru

def _get_session_id() -> str:
    # For anonymous sessions only; with auth we key cache by user id
    sid = session.get("SID")
    if not sid:
        sid = uuid.uuid4().hex
        session["SID"] = sid
    return sid


def _cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"orders_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, f"orders_{_get_session_id()}.json")


def _cache_path_for_user_id(user_id: int) -> str:
    return os.path.join(CACHE_DIR, f"orders_user_{user_id}.json")


# Products cache helpers (per user)
def _products_cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"products_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "products_anon.json")


def load_products_cache() -> Dict[str, Any] | None:
    path = _products_cache_path_for_user()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None
def save_products_cache(payload: Dict[str, Any]) -> None:
    path = _products_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass

def _articles_cache_path_for_user() -> str:
    """?????????? ???? ? ????? ???? ????????? ??? ???????? ????????????"""
    if current_user.is_authenticated:
        return f"articles_cache_user_{current_user.id}.json"
    return "articles_cache.json"

def load_articles_cache() -> Dict[str, Any] | None:
    """????????? ??? ????????? ??? ???????? ????????????"""
    path = _articles_cache_path_for_user()
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except FileNotFoundError:
        return None
    except Exception:
        return None

def save_articles_cache(payload: Dict[str, Any]) -> None:
    """????????? ??? ????????? ??? ???????? ????????????"""
    path = _articles_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


# Stocks cache helpers (per user)
def _stocks_cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"stocks_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "stocks_anon.json")


def load_stocks_cache() -> Dict[str, Any] | None:
    path = _stocks_cache_path_for_user()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def save_stocks_cache(payload: Dict[str, Any]) -> None:
    path = _stocks_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


def load_stocks_cache_for_user(user_id: int) -> Dict[str, Any] | None:
    """????????? ??? ???????? ??? ??????????? ????????????"""
    path = os.path.join(CACHE_DIR, f"stocks_user_{user_id}.json")
    try:
        if os.path.isfile(path):
            with open(path, "r", encoding="utf-8") as f:
                return json.load(f)
    except Exception:
        pass
    return None


def save_stocks_cache_for_user(user_id: int, payload: Dict[str, Any]) -> None:
    """????????? ??? ???????? ??? ??????????? ????????????"""
    path = os.path.join(CACHE_DIR, f"stocks_user_{user_id}.json")
    try:
        enriched = dict(payload)
        enriched["_user_id"] = user_id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


def clear_stocks_cache_for_user(user_id: int) -> None:
    """??????? ??? ???????? ??? ??????????? ????????????"""
    path = os.path.join(CACHE_DIR, f"stocks_user_{user_id}.json")
    try:
        if os.path.exists(path):
            os.remove(path)
            print(f"Cleared stocks cache for user {user_id}")
    except Exception as e:
        print(f"Error clearing stocks cache for user {user_id}: {e}")


# FBS supplies cache helpers (per user)
def _fbs_supplies_cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"fbs_supplies_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "fbs_supplies_anon.json")


def load_fbs_supplies_cache() -> Dict[str, Any] | None:
    path = _fbs_supplies_cache_path_for_user()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def save_fbs_supplies_cache(payload: Dict[str, Any]) -> None:
    path = _fbs_supplies_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


# FBW supplies cache helpers (per user)
def _fbw_supplies_cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"fbw_supplies_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "fbw_supplies_anon.json")


def load_fbw_supplies_cache() -> Dict[str, Any] | None:
    path = _fbw_supplies_cache_path_for_user()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def save_fbw_supplies_cache(payload: Dict[str, Any]) -> None:
    path = _fbw_supplies_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


# ??????????? ??? ???????? ? ???????? (??? ??????? ?????????)
def _fbw_supplies_detailed_cache_path_for_user(user_id: int | None = None) -> str:
    if user_id is not None:
        return os.path.join(CACHE_DIR, f"fbw_supplies_detailed_user_{user_id}.json")
    # fallback ??? current_user ? ??????? ???????
    try:
        if current_user and getattr(current_user, "is_authenticated", False):
            return os.path.join(CACHE_DIR, f"fbw_supplies_detailed_user_{current_user.id}.json")
    except Exception:
        pass
    return os.path.join(CACHE_DIR, "fbw_supplies_detailed_anon.json")


def load_fbw_supplies_detailed_cache(user_id: int | None = None) -> Dict[str, Any] | None:
    """????????? ????????? ??? ????????. ?? ??????? ?? current_user ? ??????? ??????."""
    path = _fbw_supplies_detailed_cache_path_for_user(user_id)
    if not os.path.isfile(path):
        return None
    
    # ????????? ?????? ????? - ???? ?????? 50MB, ?? ?????????
    try:
        file_size = os.path.getsize(path)
        if file_size > 50 * 1024 * 1024:  # 50MB
            print(f"???? ???? ???????? ??????? ??????? ({file_size / 1024 / 1024:.1f}MB), ?????????? ????????")
            return None
    except Exception:
        pass
    
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"?????? ???????? ???? ????????: {e}")
        return None


def save_fbw_supplies_detailed_cache(payload: Dict[str, Any], user_id: int = None) -> None:
    path = _fbw_supplies_detailed_cache_path_for_user(user_id)
    try:
        enriched = dict(payload)
        if user_id:
            enriched["_user_id"] = user_id
        elif current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


def is_supplies_cache_fresh() -> bool:
    """?????????, ?????? ?? ??? ???????? (?????????? ?? ?? ????????? 24 ????)"""
    cached = load_fbw_supplies_detailed_cache()
    if not cached:
        return False
    
    last_update = cached.get("last_updated")
    if not last_update:
        return False
    
    try:
        last_update_dt = datetime.fromisoformat(last_update)
        now = datetime.now(MOSCOW_TZ)
        return (now - last_update_dt).total_seconds() < 24 * 3600  # 24 ????
    except Exception:
        return False


def build_supplies_detailed_cache(
    token: str,
    user_id: int | None = None,
    batch_size: int = 20,
    pause_seconds: float = 1.0,
    force_full: bool = False,
    days_back: int | None = None,
) -> Dict[str, Any]:
    """??????/????????? ????????? ??? ???????? ? ????????.

    ?????????:
    - ???? ???? ??? ??? force_full=True ? ????????? ?? 6 ??????? (180 ????)
    - ????? (??????????? ????????????) ? ????????? ?????? ????????? 10 ????
    - ????????? ???? ???????? ?? batch_size ? ?????? pause_seconds ????? ????????,
      ????? ?? ?????? ?????? API.
    """
    logger.info(f"?????? ????????? ??? ???????? ??? ???????????? {user_id}...")

    # ????????? ??????? ??? (???? ????) ??? ???????????????? ??????????
    existing_cache = None
    try:
        existing_cache = load_fbw_supplies_detailed_cache(user_id) or {}
    except Exception:
        existing_cache = None

    supplies_by_date: Dict[str, Dict[str, int]] = (
        (existing_cache.get("supplies_by_date") or {}) if existing_cache else {}
    )

    # ?????????? ??????? ???????
    if days_back is not None:
        period_days = int(days_back)
    else:
        if force_full or not supplies_by_date:
            period_days = 180
        else:
            period_days = 10

    # ??? ??????????????? ?????????? (10 ????) ??????? ?????? ?????? ?? ???? ??????
    # ????? ???????? ???????????? ??? ????????? ??????????
    if period_days == 10 and supplies_by_date:
        cutoff_date = (datetime.now(MOSCOW_TZ) - timedelta(days=10)).strftime("%Y-%m-%d")
        # ??????? ?????? ?? ????????? 10 ???? ?? ????????????? ????
        keys_to_remove = [date for date in supplies_by_date.keys() if date >= cutoff_date]
        for date in keys_to_remove:
            del supplies_by_date[date]
        print(f"??????? ?????? ?? {len(keys_to_remove)} ???? ??? ???????????????? ??????????")

    supplies_list = fetch_fbw_supplies_list(token, days_back=period_days)
    total_supplies = len(supplies_list)
    print(f"??????? {total_supplies} ???????? ?? {period_days} ????")

    processed_count = 0
    last_api_call_ts = 0.0

    for idx, supply in enumerate(supplies_list, start=1):
        supply_id = supply.get("supplyID") or supply.get("id")
        if not supply_id:
            continue

        try:
            # ????????? ?????? ?? ??????? ?????? ??????? ??????
            now = time.time()
            if now - last_api_call_ts < 0.1:
                time.sleep(0.1 - (now - last_api_call_ts))

            details = fetch_fbw_supply_details(token, supply_id)
            last_api_call_ts = time.time()
            if not details:
                continue

            # ????? ????????? ??????????? ????? ???????? ?????
            time.sleep(0.05)
            supply_goods = fetch_fbw_supply_goods(token, supply_id)
            last_api_call_ts = time.time()

            supply_date = details.get("supplyDate") or details.get("createDate")
            if supply_date:
                try:
                    if isinstance(supply_date, str):
                        if 'T' in supply_date:
                            supply_dt = datetime.fromisoformat(supply_date.replace('Z', '+00:00'))
                        else:
                            supply_dt = datetime.strptime(supply_date, "%Y-%m-%d")
                    else:
                        supply_dt = supply_date

                    supply_date_str = supply_dt.strftime("%Y-%m-%d")

                    day_bucket = supplies_by_date.get(supply_date_str) or {}
                    for good in supply_goods:
                        barcode = str(good.get("barcode", "")).strip()
                        qty = int(good.get("quantity", 0) or 0)
                        if not barcode or qty <= 0:
                            continue
                        day_bucket[barcode] = day_bucket.get(barcode, 0) + qty
                    supplies_by_date[supply_date_str] = day_bucket

                except Exception:
                    # ?????????? ??????????? ???????? ????????
                    pass

            processed_count += 1
            if processed_count % 10 == 0:
                print(f"?????????? {processed_count}/{total_supplies} ????????...")

            # ???????? ?????, ????? ?? ?????? ??????
            if batch_size > 0 and (idx % batch_size == 0) and pause_seconds > 0:
                time.sleep(pause_seconds)

        except Exception:
            # ????????? ? ????????? ???????? ??? ????? ??????? ??????
            continue

    # ????????? ?????
    if supplies_by_date:
        try:
            all_days = sorted(supplies_by_date.keys())
            print(f"??? ????????: {len(all_days)} ???? ? ?????????? (? {all_days[0]} ?? {all_days[-1]})")
        except Exception:
            print(f"??? ????????: {len(supplies_by_date)} ???? ? ??????????")
    else:
        print("??? ????????: 0 ???? ? ??????????")

    return {
        "supplies_by_date": supplies_by_date,
        "last_updated": datetime.now(MOSCOW_TZ).isoformat(),
        "total_supplies_processed": processed_count,
    }


# -------------------- Orders warm cache (6 months) --------------------
def _orders_cache_meta_path_for_user(user_id: int = None) -> str:
    if user_id:
        return os.path.join(CACHE_DIR, f"orders_warm_meta_user_{user_id}.json")
    elif current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"orders_warm_meta_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "orders_warm_meta_anon.json")


def load_orders_cache_meta() -> Dict[str, Any] | None:
    path = _orders_cache_meta_path_for_user()
    if not os.path.isfile(path):
        return None
    
    # ????????? ?????? ????? - ???? ?????? 10MB, ?? ?????????
    try:
        file_size = os.path.getsize(path)
        if file_size > 10 * 1024 * 1024:  # 10MB
            print(f"???? ?????????? ???? ??????? ??????? ??????? ({file_size / 1024 / 1024:.1f}MB), ?????????? ????????")
            return None
    except Exception:
        pass
    
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"?????? ???????? ?????????? ???? ???????: {e}")
        return None


def save_orders_cache_meta(payload: Dict[str, Any], user_id: int = None) -> None:
    path = _orders_cache_meta_path_for_user(user_id)
    try:
        enriched = dict(payload)
        if user_id:
            enriched["_user_id"] = user_id
        elif current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
def is_orders_cache_fresh() -> bool:
    meta = load_orders_cache_meta()
    if not meta:
        return False
    last_updated = meta.get("last_updated")
    if not last_updated:
        return False
    try:
        last_dt = datetime.fromisoformat(last_updated)
        now = datetime.now(MOSCOW_TZ)
        return (now - last_dt).total_seconds() < 24 * 3600
    except Exception:
        return False


def build_orders_warm_cache(token: str, user_id: int = None) -> Dict[str, Any]:
    """Warm up per-day orders cache for last 6 months in one go."""
    from_date = (datetime.now(MOSCOW_TZ).date() - timedelta(days=180)).strftime("%Y-%m-%d")
    to_date = datetime.now(MOSCOW_TZ).date().strftime("%Y-%m-%d")
    # Fetch all rows in range and persist into per-day cache
    raw = fetch_orders_range(token, from_date, to_date)
    rows = to_rows(raw, from_date, to_date)
    _update_period_cache_with_data(token, from_date, to_date, rows, user_id)
    meta = {
        "last_updated": datetime.now(MOSCOW_TZ).isoformat(),
        "date_from": from_date,
        "date_to": to_date,
        "total_orders_cached": len(rows),
        "cache_version": "1.0"
    }
    return meta


# Orders per-day cache helpers (per user)
def _orders_period_cache_path_for_user(user_id: int = None) -> str:
    if user_id:
        return os.path.join(CACHE_DIR, f"orders_period_user_{user_id}.json")
    elif current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"orders_period_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, f"orders_period_{_get_session_id()}.json")


def load_orders_period_cache(user_id: int = None) -> Dict[str, Any] | None:
    path = _orders_period_cache_path_for_user(user_id)
    print(f"Loading orders period cache from: {path}")
    if not os.path.isfile(path):
        print("Orders period cache file not found")
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            days_count = len(data.get('days', {}))
            print(f"Orders period cache loaded, days in cache: {days_count}")
            return data
    except Exception as e:
        print(f"Error loading orders period cache: {e}")
        return None


def save_orders_period_cache(payload: Dict[str, Any], user_id: int = None) -> None:
    path = _orders_period_cache_path_for_user(user_id)
    try:
        enriched = dict(payload)
        if user_id:
            enriched["_user_id"] = user_id
        elif current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        print(f"Saving orders period cache to: {path}")
        print(f"Days to save in cache: {len(enriched.get('days', {}))}")
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
        print("Orders period cache saved")
    except Exception as e:
        print(f"Error saving orders period cache: {e}")


def _normalize_date_str(date_str: str) -> str:
    try:
        dt = parse_date(date_str)
        return dt.strftime("%Y-%m-%d")
    except Exception:
        return date_str


def _order_row_day_iso(row: Dict[str, Any]) -> str:
    """YYYY-MM-DD prefix for charts/caches; safe if source-file key literals collide (mojibake)."""
    raw: Any = (
        row.get("_order_date")
        or row.get("\u0414\u0430\u0442\u0430")
        or row.get("date")
    )
    if raw is None:
        raw = row.get("????")
    if raw is None:
        return ""
    s = str(raw).strip()
    return s[:10] if s else ""


def _order_row_warehouse_label(row: Dict[str, Any]) -> str:
    for key in (
        "_warehouse_label",
        "_warehouse",
        "\u0421\u043a\u043b\u0430\u0434 \u043e\u0442\u0433\u0440\u0443\u0437\u043a\u0438",
        "warehouseName",
        "????? ????????",
    ):
        v = row.get(key)
        if v is not None and str(v).strip():
            return str(v).strip()
    return "\u041d\u0435 \u0443\u043a\u0430\u0437\u0430\u043d"


def _order_row_price_disc(row: Dict[str, Any]) -> float:
    raw: Any = row.get("_price")
    if raw is None:
        raw = row.get("\u0426\u0435\u043d\u0430 \u0441\u043e \u0441\u043a\u0438\u0434\u043a\u043e\u0439 \u043f\u0440\u043e\u0434\u0430\u0432\u0446\u0430")
    if raw is None:
        raw = row.get("priceWithDisc")
    if raw is None:
        raw = row.get("???? ?? ??????? ????????")
    try:
        return float(raw or 0)
    except (TypeError, ValueError):
        return 0.0


def _daterange_inclusive(start_date: str, end_date: str) -> list[str]:
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date)
    days: list[str] = []
    cur = start_dt
    while cur <= end_dt:
        days.append(cur.strftime("%Y-%m-%d"))
        cur += timedelta(days=1)
    return days


def get_orders_with_period_cache(
    token: str,
    date_from: str,
    date_to: str,
    *,
    bypass_today_ttl: bool = False,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    """Return (orders, cache_meta). Uses per-day cache and fetches only missing days.

    bypass_today_ttl: если True — день «сегодня» всегда тянется с WB (игнор TTL).
    cache_meta contains info like {"used_cache_days": int, "fetched_days": int}
    """
    # Load existing cache structure
    cache = load_orders_period_cache() or {}
    days_map: Dict[str, Any] = cache.get("days") or {}

    requested_days = _daterange_inclusive(date_from, date_to)

    today_iso = datetime.now(MOSCOW_TZ).date().strftime("%Y-%m-%d")
    days_to_fetch: list[str] = []
    for day in requested_days:
        entry = days_map.get(day)
        if day == today_iso:
            if bypass_today_ttl:
                days_to_fetch.append(day)
            elif entry and period_cache_day_entry_is_fresh(entry, ORDERS_TODAY_CACHE_TTL_SECONDS):
                pass
            else:
                days_to_fetch.append(day)
            continue
        if entry is None:
            days_to_fetch.append(day)

    _refetch_preview = (
        str(days_to_fetch)
        if len(days_to_fetch) <= 8
        else f"{len(days_to_fetch)} days from {min(days_to_fetch)} to {max(days_to_fetch)}"
    )
    print(
        f"Orders period cache: user range {date_from}..{date_to} ({len(requested_days)} d.), "
        f"days in cache file: {len(days_map)}, to refetch ({len(days_to_fetch)}): {_refetch_preview}"
    )

    collected_orders: list[dict[str, Any]] = []

    # Collect from cache first
    def _cached_orders(entry: Dict[str, Any]) -> list[dict[str, Any]]:
        """Backward-compatible extractor for orders list stored in a day cache entry.

        Older cache versions could store daily rows under different keys.
        Prefer the new key 'orders', but gracefully fall back to legacy ones.
        """
        if not isinstance(entry, dict):
            return []
        val = (
            entry.get("orders")
            or entry.get("orders_rows")
            or entry.get("rows")
            or entry.get("data")
        )
        return val if isinstance(val, list) else []
    for day in requested_days:
        entry = days_map.get(day)
        if entry and day not in days_to_fetch:
            collected_orders.extend(_cached_orders(entry))

    # Fetch missing days in one period request and split per day
    total_days = len(days_to_fetch)
    done_days = 0
    progress_key = f"{date_from}:{date_to}:{int(time.time())}"
    if current_user and current_user.is_authenticated:
        _set_orders_progress(current_user.id, total_days, done_days, key=progress_key)
    if days_to_fetch:
        try:
            # WB fetch + pagination cost scales with the requested calendar span. Only pull the
            # minimal range that covers missing/refetch days — not the full user-selected period.
            fetch_from = min(days_to_fetch)
            fetch_to = max(days_to_fetch)
            print(
                f"Fetching orders for minimal window {fetch_from}..{fetch_to} "
                f"({len(days_to_fetch)} day(s) to merge; user range was {date_from}..{date_to})"
            )
            raw = fetch_orders_range(token, fetch_from, fetch_to)
            all_rows = to_rows(raw, fetch_from, fetch_to)
            # Group by day
            by_day: Dict[str, list[dict[str, Any]]] = defaultdict(list)
            for r in all_rows:
                d = _order_row_day_iso(r)
                if d:
                    by_day[d].append(r)
            # For each missing day, update cache and progress
            for day in days_to_fetch:
                fetched_orders = by_day.get(day, [])
                days_map[day] = {
                    "orders": fetched_orders,
                    "updated_at": datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y %H:%M:%S"),
                }
                collected_orders.extend(fetched_orders)
                done_days += 1
                if current_user and current_user.is_authenticated:
                    _set_orders_progress(current_user.id, total_days, done_days, key=progress_key)
        except Exception as e:
            print(f"Error fetching orders with period cache: {e}")
            # Fallback: nothing added

    # Persist cache file if any changes were made
    if days_to_fetch:
        print(f"Saving cache for days: {days_to_fetch}")
        cache["days"] = days_map
        save_orders_period_cache(cache)
        print(f"Cache saved. Total cached days: {len(days_map)}")
    else:
        print("All requested days are already cached (or today within TTL), no fetch needed")

    if current_user and current_user.is_authenticated:
        _clear_orders_progress(current_user.id, key=progress_key)

    meta = {"used_cache_days": len(requested_days) - len(days_to_fetch), "fetched_days": len(days_to_fetch)}
    return collected_orders, meta


def _update_period_cache_with_data(
    token: str,
    date_from: str,
    date_to: str,
    orders: list[dict[str, Any]],
    user_id: int = None,
) -> None:
    """????????????? ????????? ??? ?? ???? ? ???????????????? ???????"""
    cache = load_orders_period_cache(user_id) or {}
    days_map: Dict[str, Any] = cache.get("days") or {}
    
    requested_days = _daterange_inclusive(date_from, date_to)
    
    # ?????????? ?????? ?? ????
    orders_by_day: Dict[str, list[dict[str, Any]]] = {}
    
    for order in orders:
        order_date = _order_row_day_iso(order) or (order.get("???? ??????") or "")
        if order_date:
            day_key = _normalize_date_str(order_date)
            if day_key not in orders_by_day:
                orders_by_day[day_key] = []
            orders_by_day[day_key].append(order)
    
    # ????????? ??? ??? ??????? ???
    for day in requested_days:
        days_map[day] = {
            "orders": orders_by_day.get(day, []),
            "updated_at": datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y %H:%M:%S"),
        }
    
    # ????????? ??????????? ???
    cache["days"] = days_map
    save_orders_period_cache(cache, user_id)


def _fbs_tasks_cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"fbs_tasks_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "fbs_tasks_anon.json")


def load_fbs_tasks_cache() -> Dict[str, Any] | None:
    path = _fbs_tasks_cache_path_for_user()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def load_fbs_tasks_cache_by_user_id(user_id: int) -> Dict[str, Any] | None:
    """Load FBS tasks cache by user ID (for background threads)"""
    path = os.path.join(CACHE_DIR, f"fbs_tasks_user_{user_id}.json")
    print(f"Loading FBS tasks cache from: {path}")
    if not os.path.isfile(path):
        print(f"FBS tasks cache file not found: {path}")
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            print(f"FBS tasks cache loaded successfully, {len(data.get('rows', []))} tasks found")
            return data
    except Exception as e:
        print(f"Error loading FBS tasks cache: {e}")
        return None


def save_fbs_tasks_cache(payload: Dict[str, Any]) -> None:
    path = _fbs_tasks_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass

def load_last_results() -> Dict[str, Any] | None:
    path = _cache_path_for_user()
    if not os.path.isfile(path):
        return None
    
    try:
        file_size = os.path.getsize(path)
        if file_size > LAST_RESULTS_CACHE_MAX_BYTES:
            lim_mb = LAST_RESULTS_CACHE_MAX_BYTES / (1024 * 1024)
            print(
                f"last_results cache too large ({file_size / 1024 / 1024:.1f} MB, limit {lim_mb:.0f} MB), skip load"
            )
            return None
    except Exception:
        pass
    
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception as e:
        print(f"?????? ???????? ????: {e}")
        return None
def save_last_results(payload: Dict[str, Any]) -> None:
    """
    ????????? ?????????? ? ???, ????????? ? ????????????? ???????.
    ??? ????????? ????????? ?????? ??? ?????? ??????? (??????, ?????????? ?????) ??? ?????????? ???? ?????.
    """
    path = _cache_path_for_user()
    try:
        # ????????? ???????????? ???, ???? ?? ????
        existing_cache = load_last_results() or {}
        
        # ?????????? ???????????? ??? ? ?????? ??????? (????? ?????? ????? ?????????)
        enriched = dict(existing_cache)
        enriched.update(payload)
        
        try:
            if current_user.is_authenticated:
                enriched["_user_id"] = current_user.id
                enriched["_username"] = getattr(current_user, "username", None)
        except Exception:
            # If current_user unavailable outside request context, ignore
            pass
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


# -------------------------
# Models & Auth
# -------------------------
# ?????? ????????????? ?? models.py ????

@login_manager.user_loader
def load_user(user_id: str):
    try:
        return User.query.get(int(user_id))
    except Exception:
        return None


@app.before_request
def _enforce_account_validity():
    # Allow unauthenticated pages and static assets
    endpoint = (request.endpoint or "")
    if endpoint in {"login", "logout", "favicon", "logo"} or endpoint.startswith("static"):
        return None
    
    # ????????????? ????????? ???????????? ??? ???? ???????, ????????? ???????????
    # ????????? ?????? ????????? ????????
    public_pages = ["/login", "/logout", "/favicon.ico", "/logo.png"]
    if not any(request.path.startswith(page) for page in public_pages):
        if not current_user.is_authenticated:
            # ???? ???????????? ?? ???????????, ?????????????? ?? ?????
            return redirect(url_for("login"))
    
    if current_user.is_authenticated:
        if not _is_user_valid_now(current_user):
            logout_user()
            # For API requests return JSON 401 to avoid HTML redirect in fetch()
            if request.path.startswith("/api/"):
                return jsonify({"error": "expired"}), 401
            flash("???? ???????? ??????? ?????? ?????")
            return redirect(url_for("login"))


def parse_date(date_str: str) -> datetime:
    """Parse date string in either YYYY-MM-DD or DD.MM.YYYY format"""
    try:
        # Try YYYY-MM-DD format first
        return datetime.strptime(date_str, "%Y-%m-%d")
    except ValueError:
        try:
            # Try DD.MM.YYYY format
            return datetime.strptime(date_str, "%d.%m.%Y")
        except ValueError:
            raise ValueError(f"Unable to parse date '{date_str}'. Expected formats: YYYY-MM-DD or DD.MM.YYYY")


def parse_wb_datetime(value: str) -> datetime | None:
    if not value:
        return None
    s = str(value)
    # Trim subseconds if present to 6 digits max for fromisoformat
    try:
        # Normalize Z to +00:00
        s_norm = s.replace("Z", "+00:00")
        dt = datetime.fromisoformat(s_norm[:26] + s_norm[26:])  # be forgiving on microseconds length
        return dt
    except Exception:
        try:
            return datetime.strptime(s[:19], "%Y-%m-%dT%H:%M:%S")
        except Exception:
            return None


def get_with_retry(url: str, headers: Dict[str, str], params: Dict[str, Any], max_retries: int = 3, timeout_s: int = 30) -> requests.Response:
    last_exc: Exception | None = None
    last_resp: requests.Response | None = None
    for attempt in range(max_retries):
        try:
            resp = requests.get(url, headers=headers, params=params, timeout=timeout_s)
            last_resp = resp
            if resp.status_code in (429, 500, 502, 503, 504):
                sleep_s = None
                if resp.status_code == 429:
                    # ??? ?????? 429 ????????? ????????? X-Ratelimit-Retry (WB API) ? Retry-After (???????????)
                    retry_header = resp.headers.get("X-Ratelimit-Retry") or resp.headers.get("Retry-After")
                    if retry_header is not None:
                        try:
                            sleep_s = float(retry_header)
                        except ValueError:
                            sleep_s = None
                
                if sleep_s is None:
                    # ???? ?????????? ???, ?????????? ???????????????? ????????
                    if resp.status_code == 429:
                        # ??? 429 ?????????? ????? ?????????? ????????
                        sleep_s = min(120, 30 * (attempt + 1))
                    else:
                        sleep_s = min(15, 0.8 * (2 ** attempt) + random.uniform(0, 0.7))
                
                time.sleep(sleep_s)
                continue
            resp.raise_for_status()
            return resp
        except requests.RequestException as exc:  # network or HTTP error
            last_exc = exc
            time.sleep(min(8, 0.5 * (2 ** attempt) + random.uniform(0, 0.5)))
            continue
    if last_exc:
        raise last_exc
    if last_resp is not None:
        raise requests.HTTPError(f"HTTP {last_resp.status_code} after {max_retries} retries", response=last_resp)
    raise RuntimeError("Request failed after retries")


def get_with_retry_json(url: str, headers: Dict[str, str], params: Dict[str, Any], max_retries: int = 3, timeout_s: int = 30) -> Any:
    resp = get_with_retry(url, headers, params, max_retries=max_retries, timeout_s=timeout_s)
    try:
        return resp.json()
    except Exception:
        raise RuntimeError("Invalid JSON from API")


def fetch_orders_page(token: str, date_from_iso: str, flag: int = 0) -> List[Dict[str, Any]]:
    headers = {"Authorization": f"Bearer {token}"}
    params = {"dateFrom": date_from_iso, "flag": flag}
    response = get_with_retry(API_URL, headers, params)
    return response.json()


def fetch_orders_range(token: str, start_date: str, end_date: str, days_back: int = 7) -> List[Dict[str, Any]]:
    """Fetch orders from WB by paginating with lastChangeDate, but filter by actual order date.
    
    Args:
        token: API ?????
        start_date: ????????? ???? (YYYY-MM-DD)
        end_date: ???????? ???? (YYYY-MM-DD)
        days_back: ?????????? ???? ????? ?? start_date ??? ?????? ???????? (?? ????????? 7 ????)
    """
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date)

    # ????????? ?????? ? ??????? ??? ??????? ???????, ??????? ????? ???? ????????? ?????
    # ??? ??????? ????? ???????????? ??????? ????? (1 ????) ??? ?????????
    extended_start = start_dt - timedelta(days=days_back)
    cursor_dt = datetime.combine(extended_start.date(), datetime.min.time())

    collected: List[Dict[str, Any]] = []
    seen_srid: set[str] = set()

    pages = 0
    now_date = datetime.now(MOSCOW_TZ).date()
    if end_dt.date() < now_date:
        lcd_stop_date = end_dt.date() + timedelta(days=1)
    else:
        lcd_stop_date = end_dt.date()

    span_inclusive = (end_dt.date() - start_dt.date()).days + 1
    intraday_short = end_dt.date() >= now_date and span_inclusive <= 2
    if intraday_short:
        max_pages = min(WB_ORDERS_FETCH_MAX_PAGES, WB_ORDERS_FETCH_MAX_PAGES_INTRADAY)
        page_sleep_s = WB_ORDERS_PAGE_SLEEP_INTRADAY_S
    else:
        max_pages = WB_ORDERS_FETCH_MAX_PAGES
        page_sleep_s = WB_ORDERS_PAGE_SLEEP_S

    while pages < max_pages:
        pages += 1
        before_cursor = cursor_dt
        page = fetch_orders_page(token, cursor_dt.strftime("%Y-%m-%dT%H:%M:%S"), flag=0)
        if not page:
            break
        try:
            page.sort(key=lambda x: parse_wb_datetime(x.get("lastChangeDate")) or datetime.min)
        except Exception:
            pass

        last_page_lcd: datetime | None = parse_wb_datetime(page[-1].get("lastChangeDate"))
        page_exceeds = bool(last_page_lcd and last_page_lcd.date() > lcd_stop_date)

        for item in page:
            srid = str(item.get("srid", ""))
            if srid and srid in seen_srid:
                continue
            # ??????? ?????????? ?? lastChangeDate ????? - ????? ??????????? ?? date ? to_rows
            if srid:
                seen_srid.add(srid)
            collected.append(item)

        if last_page_lcd is None:
            break
        if pages > 1 and last_page_lcd <= before_cursor:
            break
        cursor_dt = last_page_lcd
        if page_exceeds:
            break
        if page_sleep_s > 0:
            time.sleep(page_sleep_s)

    if pages >= max_pages:
        print(
            f"fetch_orders_range: max_pages={max_pages} start={start_date} end={end_date} "
            f"intraday_short_window={intraday_short}"
        )

    return collected


def fetch_sales_page(token: str, date_from_iso: str, flag: int = 0) -> List[Dict[str, Any]]:
    headers = {"Authorization": f"Bearer {token}"}
    params = {"dateFrom": date_from_iso, "flag": flag}
    response = get_with_retry(SALES_API_URL, headers, params)
    return response.json()


def fetch_sales_range(token: str, start_date: str, end_date: str) -> List[Dict[str, Any]]:
    start_dt = parse_date(start_date)
    end_dt = parse_date(end_date)
    cursor_dt = datetime.combine(start_dt.date(), datetime.min.time())

    collected: List[Dict[str, Any]] = []
    seen_id: set[str] = set()

    max_pages = 2000
    pages = 0
    while pages < max_pages:
        pages += 1
        page = fetch_sales_page(token, cursor_dt.strftime("%Y-%m-%dT%H:%M:%S"), flag=0)
        if not page:
            break
        try:
            page.sort(key=lambda x: parse_wb_datetime(x.get("lastChangeDate")) or datetime.min)
        except Exception:
            pass

        last_page_lcd: datetime | None = parse_wb_datetime(page[-1].get("lastChangeDate"))
        page_exceeds = last_page_lcd and last_page_lcd.date() > end_dt.date()

        for item in page:
            key = str(item.get("srid")) or f"{item.get('gNumber','')}_{item.get('barcode','')}_{item.get('date','')}"
            if key and key in seen_id:
                continue
            lcd = parse_wb_datetime(item.get("lastChangeDate"))
            if lcd and lcd.date() > end_dt.date():
                continue
            if key:
                seen_id.add(key)
            collected.append(item)

        if last_page_lcd is None:
            break
        cursor_dt = last_page_lcd
        if page_exceeds:
            break
        # Gentle delay between pages
        time.sleep(0.2)

    return collected


def _split_date_range(date_from: str, date_to: str, days_per_chunk: int = 7) -> List[tuple[str, str]]:
    """????????? ?????? ?? ????????? ?? ?????????? ?????????? ????.
    
    Returns:
        List of tuples (start_date, end_date) in YYYY-MM-DD format
    """
    try:
        start = datetime.strptime(date_from, "%Y-%m-%d").date()
        end = datetime.strptime(date_to, "%Y-%m-%d").date()
    except Exception:
        return [(date_from, date_to)]
    
    intervals = []
    current_start = start
    
    while current_start <= end:
        current_end = min(current_start + timedelta(days=days_per_chunk - 1), end)
        intervals.append((
            current_start.strftime("%Y-%m-%d"),
            current_end.strftime("%Y-%m-%d")
        ))
        current_start = current_end + timedelta(days=1)
    
    return intervals


def fetch_finance_report(token: str, date_from: str, date_to: str, limit: int = 100000, progress_callback=None) -> List[Dict[str, Any]]:
    """Fetch financial report details v5 with rrdid pagination.
    
    ????????? ?????? ?? ????????? ?? 7 ???? ??? ????????? ??????? API.

    According to docs, start with rrdid=0 and then pass last row's rrd_id until empty list is returned.
    date_from must be RFC3339 in MSK; we'll accept YYYY-MM-DD and convert to T00:00:00.
    date_to is YYYY-MM-DD (end date).
    
    Args:
        token: API token
        date_from: Start date in YYYY-MM-DD format
        date_to: End date in YYYY-MM-DD format
        limit: Maximum rows per request
        progress_callback: Optional callback function(current, total, current_period) for progress updates
    """
    headers = {"Authorization": f"Bearer {token}"}
    
    # ????????? ?????? ?? ????????? ?? 7 ????
    intervals = _split_date_range(date_from, date_to, days_per_chunk=7)
    total_intervals = len(intervals)
    all_rows: List[Dict[str, Any]] = []
    
    logging.info(f"???????? ???????? ??????????? ?????? ?? ?????? {date_from} - {date_to}, ??????????: {total_intervals}")
    for idx, (interval_from, interval_to) in enumerate(intervals, 1):
        logging.info(f"???????? ????????? {idx}/{total_intervals}: {interval_from} - {interval_to}")
        
        # ???????? callback ??? ?????????? ?????????
        if progress_callback:
            progress_callback(idx, total_intervals, f"{interval_from} - {interval_to}")
        
        # Compose RFC3339-like dateFrom in MSK start of day
        try:
            df_iso = datetime.strptime(interval_from, "%Y-%m-%d").strftime("%Y-%m-%dT00:00:00")
        except Exception:
            df_iso = f"{interval_from}T00:00:00"
        
        params_base: Dict[str, Any] = {"dateFrom": df_iso, "dateTo": interval_to, "limit": max(1, min(100000, int(limit)))}
        interval_rows: List[Dict[str, Any]] = []
        rrdid = 0
        interval_error = None
        page_count = 0
        
        while True:
            page_count += 1
            params = dict(params_base)
            params["rrdid"] = rrdid
            try:
                # ?????????? get_with_retry ???????? ??? ??????? ? ?????????? ??? ?????? 429
                resp = get_with_retry(FIN_REPORT_URL, headers, params, max_retries=3, timeout_s=30)
                data = resp.json()
            except requests.HTTPError as e:
                # ???????????? HTTP ??????, ??????? 429
                interval_error = str(e)
                error_str = str(e)
                is_429 = "429" in error_str or "Too Many Requests" in error_str or (hasattr(e, 'response') and e.response is not None and e.response.status_code == 429)
                
                if is_429:
                    # ??? ?????? 429 ?????????? ????????? ??????? ? ??????????????? ??????
                    # ???????? ???????? ????? ???????? ?? ????????? ??????
                    retry_after = 60  # ?? ????????? 60 ??????
                    if hasattr(e, 'response') and e.response is not None:
                        retry_header = e.response.headers.get('X-Ratelimit-Retry') or e.response.headers.get('Retry-After')
                        if retry_header:
                            try:
                                retry_after = int(float(retry_header))
                                logging.info(f"???????? ????? ???????? ?? ?????????: {retry_after} ??????")
                            except (ValueError, TypeError):
                                pass
                    
                    # ?????? ????????? ??????? ? ??????????????? ??????
                    max_429_retries = 3
                    retry_success = False
                    for retry_attempt in range(1, max_429_retries + 1):
                        wait_time = retry_after * retry_attempt  # ??????????? ????? ? ?????? ????????
                        logging.warning(f"?????? 429 (????? API) ??? ????????? {interval_from} - {interval_to}, ???????? {page_count} (rrdid={rrdid}). ??????? {retry_attempt}/{max_429_retries}, ????? {wait_time} ??????...")
                        time.sleep(wait_time)
                        
                        # ????????? ???????
                        try:
                            resp = get_with_retry(FIN_REPORT_URL, headers, params, max_retries=1, timeout_s=30)
                            data = resp.json()
                            logging.info(f"????????? ??????? {retry_attempt} ????? 429 ??????? ??? ????????? {interval_from} - {interval_to}")
                            interval_error = None  # ?????????? ??????, ??? ??? ????????? ??????? ???????
                            retry_success = True
                            break  # ??????? ?? ????? ????????? ???????
                        except Exception as e2:
                            error_str2 = str(e2)
                            is_429_2 = "429" in error_str2 or "Too Many Requests" in error_str2 or (hasattr(e2, 'response') and e2.response is not None and e2.response.status_code == 429)
                            if retry_attempt < max_429_retries:
                                logging.warning(f"????????? ??????? {retry_attempt} ????? 429 ?? ???????, ??????????...")
                                continue
                            else:
                                logging.error(f"??? {max_429_retries} ??????? ????? 429 ?? ??????? ??? ????????? {interval_from} - {interval_to}: {e2}")
                                interval_error = str(e2)
                    
                    if not retry_success:
                        # ???? ??? ??????? ?? ???????, ?????????? ????????
                        if rrdid == 0:
                            logging.error(f"?????????? ???????? {interval_from} - {interval_to} ??-?? ?????? 429 ??? ?????? ???????? ????? {max_429_retries} ???????")
                            break
                        # ???? ??? ?? ?????? ????????, ??????? ??????????
                        time.sleep(5)
                        continue
                else:
                    # ??? ?????? ??????
                    logging.warning(f"?????? ???????? ?????? ??? ????????? {interval_from} - {interval_to}, ???????? {page_count} (rrdid={rrdid}): {e}")
                    # ???? ??? ?????? ???????? (rrdid=0), ?????????? ???? ????????
                    if rrdid == 0:
                        logging.error(f"?????????? ???????? {interval_from} - {interval_to} ??-?? ?????? ??? ?????? ????????")
                        break
                    # ???? ??? ?? ?????? ????????, ??????? ?????????? (????? ???? ????????? ??????)
                    time.sleep(2)  # ????? ????? ????????? ????????
                    continue
            except Exception as e:
                # ???????????? ?????? ??????????
                interval_error = str(e)
                logging.warning(f"?????? ???????? ?????? ??? ????????? {interval_from} - {interval_to}, ???????? {page_count} (rrdid={rrdid}): {e}")
                # ???? ??? ?????? ???????? (rrdid=0), ?????????? ???? ????????
                if rrdid == 0:
                    logging.error(f"?????????? ???????? {interval_from} - {interval_to} ??-?? ?????? ??? ?????? ????????")
                    break
                # ???? ??? ?? ?????? ????????, ??????? ??????????
                time.sleep(2)
                continue
            
            if not isinstance(data, list) or not data:
                logging.info(f"???????? {interval_from} - {interval_to}: ??????? ?????? ????? ?? ???????? {page_count}")
                break
            interval_rows.extend(data)
            logging.debug(f"???????? {interval_from} - {interval_to}, ???????? {page_count}: ????????? {len(data)} ???????")
            
            try:
                last = data[-1]
                rrdid = int(last.get("rrd_id") or last.get("rrdid") or last.get("rrdId") or 0)
            except Exception:
                logging.warning(f"???????? {interval_from} - {interval_to}: ?? ??????? ???????? rrd_id ?? ????????? ??????")
                break
            # If received less than limit rows, it's the last page
            try:
                if len(data) < params_base.get("limit", 100000):
                    logging.info(f"???????? {interval_from} - {interval_to}: ???????? ?????? ??????? ??? ????? ({len(data)} < {params_base.get('limit', 100000)}), ??? ????????? ????????")
                    break
            except Exception:
                pass
            # ????????? ????? ????? ??????????
            time.sleep(0.5)
        
        if interval_rows:
            all_rows.extend(interval_rows)
            # ????????? ???????? ??? ? ??????????? ??????
            dates_in_interval = set()
            for row in interval_rows:
                try:
                    # ???????? ????? ???? ? ????????? ?????
                    date_str = row.get("doc_date") or row.get("date") or row.get("operation_date")
                    if date_str:
                        dates_in_interval.add(str(date_str)[:10])
                except Exception:
                    pass
            
            logging.info(f"???????? {interval_from} - {interval_to}: ????????? {len(interval_rows)} ??????? ?? {page_count} ???????(?)")
            if dates_in_interval:
                min_date = min(dates_in_interval) if dates_in_interval else "??????????"
                max_date = max(dates_in_interval) if dates_in_interval else "??????????"
                logging.info(f"???????? {interval_from} - {interval_to}: ???? ? ?????? ?? {min_date} ?? {max_date}")
        elif interval_error:
            logging.error(f"????????: ???????? {interval_from} - {interval_to} ?? ???????? ??-?? ??????: {interval_error}")
        else:
            logging.warning(f"???????? {interval_from} - {interval_to}: ?? ????????? ?? ????? ??????")
        
        # ????? ????? ??????????? ??? ????????? ???????
        # ??????????? ?????, ???????? ????? ?????????? ? ??????? ??????????? ??????
        if idx < total_intervals:
            # ??? ?????? ?????? ?????????, ??? ?????? ?????
            if interval_rows and len(interval_rows) > 15000:
                pause_time = 5  # ??????? ????? ????? ?????????? ? ??????? ??????? ??????
            elif interval_rows and len(interval_rows) > 10000:
                pause_time = 3
            elif interval_rows and len(interval_rows) > 5000:
                pause_time = 2.5
            else:
                pause_time = 2
            logging.debug(f"????? {pause_time} ??? ????? ????????? ?????????? (????????? {len(interval_rows) if interval_rows else 0} ???????)")
            time.sleep(pause_time)
    
    logging.info(f"???????? ??????????? ?????? ?????????. ????? ????????? {len(all_rows)} ??????? ?? {total_intervals} ??????????")
    
    return all_rows

def _process_finance_data(raw: List[Dict[str, Any]], req_from: str, req_to: str, user_id: int = None) -> Dict[str, Any]:
    """???????????? ????? ?????? ??????????? ?????? ? ?????????? ??????????? ???????.
    
    Args:
        raw: ?????? ????? ??????? ?? API
        req_from: ????????? ???? ???????
        req_to: ???????? ???? ???????
        user_id: ID ???????????? (??? ????????? tax_rate)
    
    Returns:
        ??????? ? ????????????? ??????? ? ?????????
    """
    from flask_login import current_user
    
    total_qty = 0
    total_sum = 0.0
    wbr_plus = 0.0
    wbr_minus = 0.0
    total_logistics = 0.0
    total_storage = 0.0
    total_acceptance = 0.0
    total_for_pay = 0.0
    total_buyouts = 0.0
    total_returns = 0.0
    total_acquiring = 0.0
    total_commission_wb = 0.0
    total_other_deductions = 0.0
    total_penalties = 0.0
    total_additional_payment = 0.0
    total_paid_delivery = 0.0
    x1 = x2 = x3 = x4 = x5 = x6 = x7 = x8 = 0.0
    k1 = k2 = k3 = k4 = k5 = k6 = k7 = k8 = k9 = 0.0
    u1 = u2 = u3 = u4 = u5 = u6 = u7 = u8 = u9 = u10 = u11 = u12 = u13 = u14 = 0.0
    buyout_oper_values_lower = {"???????", "?????? ?????????", "?????????? ???????", "????????? ??????"}
    
    for r in raw:
        try:
            total_qty += int(r.get("quantity") or 0)
        except Exception:
            pass
        try:
            total_sum += float(r.get("retail_amount") or 0.0)
        except Exception:
            pass
        try:
            total_logistics += float(r.get("delivery_rub") or 0.0)
        except Exception:
            pass
        try:
            total_storage += float(r.get("storage_fee") or 0.0)
        except Exception:
            pass
        try:
            total_acceptance += float(r.get("acceptance") or 0.0)
        except Exception:
            pass
        try:
            total_for_pay += float(r.get("ppvz_for_pay") or 0.0)
        except Exception:
            pass
        try:
            oper = (r.get("supplier_oper_name") or "").strip()
            if oper and oper.lower() in buyout_oper_values_lower:
                total_buyouts += float(r.get("retail_price") or 0.0)
        except Exception:
            pass
        try:
            oper_lc = (r.get("supplier_oper_name") or "").strip().lower()
            amt = float(r.get("retail_amount") or 0.0)
            if oper_lc in {"???????","?????? ?????????","?????????? ???????","????????? ??????"}:
                wbr_plus += amt
            elif oper_lc in {"???????","?????? ??????","?????????? ???????"}:
                wbr_minus += amt
        except Exception:
            pass
        try:
            oper = (r.get("supplier_oper_name") or "").strip()
            if oper == "???????":
                total_returns += float(r.get("retail_price") or 0.0)
        except Exception:
            pass
        try:
            dt_name = (r.get("doc_type_name") or "").strip()
            acq_pct = float(r.get("acquiring_percent") or 0.0)
            afee = float(r.get("acquiring_fee") or 0.0)
            if dt_name == "???????" and acq_pct > 0:
                total_acquiring += afee
            elif dt_name == "???????" and acq_pct > 0:
                total_acquiring -= afee
        except Exception:
            pass
        try:
            total_commission_wb += float(r.get("ppvz_vw") or 0.0)
        except Exception:
            pass
        try:
            total_other_deductions += float(r.get("deduction") or 0.0)
            total_other_deductions += float(r.get("additional_payment") or 0.0)
        except Exception:
            pass
        try:
            total_penalties += float(r.get("penalty") or 0.0)
        except Exception:
            pass
        try:
            total_additional_payment += float(r.get("additional_payment") or 0.0)
        except Exception:
            pass
        try:
            oper_l = (r.get("supplier_oper_name") or "").strip().lower()
            doc_l = (r.get("doc_type_name") or "").strip().lower()
            pay_val = float(r.get("ppvz_for_pay") or 0.0)
            if oper_l == "?????? ??????? ????????":
                total_paid_delivery += pay_val
            if oper_l == "??????????? ?????" and doc_l == "???????":
                x1 += pay_val
            if oper_l == "?????? ?????" and doc_l == "???????":
                x2 += pay_val
            if oper_l == "??????????? ?????" and doc_l == "???????":
                x3 += pay_val
            if oper_l == "?????? ?????" and doc_l == "???????":
                x4 += pay_val
            if oper_l == "????????? ??????????? ?????" and doc_l == "???????":
                x5 += pay_val
            if oper_l == "????????? ??????????? ?????" and doc_l == "???????":
                x6 += pay_val
            if oper_l == "???????????? ??????????? ??? ????????" and doc_l == "???????":
                x7 += pay_val
            if oper_l == "???????????? ??????????? ??? ????????" and doc_l == "???????":
                x8 += pay_val
            if oper_l == "???????":
                k1 += pay_val
            if oper_l == "?????? ?????????":
                k2 += pay_val
            if oper_l == "?????????? ???????":
                k3 += pay_val
            if oper_l == "????????? ??????" and doc_l == "???????":
                k4 += pay_val
            if oper_l == "???????":
                k5 += pay_val
            if oper_l == "?????? ??????":
                k6 += pay_val
            if oper_l == "????????? ??????" and doc_l == "???????":
                k7 += pay_val
            if oper_l == "?????????? ???????":
                k8 += pay_val
            if oper_l == "????????????? ??????????":
                k9 += pay_val
            if oper_l == "?????? ??????????? ??????" and doc_l == "???????":
                u1 += pay_val
            if oper_l == "??????????? ??????????? ??????" and doc_l == "???????":
                u2 += pay_val
            if oper_l == "?????? ??????????? ??????" and doc_l == "???????":
                u3 += pay_val
            if oper_l == "??????????? ??????????? ??????" and doc_l == "???????":
                u4 += pay_val
            if oper_l == "????????? ?????? ?? ????? ??? ????????" and doc_l == "???????":
                u5 += pay_val
            if oper_l == "????????? ?????? ?? ????? ??? ????????" and doc_l == "???????":
                u6 += pay_val
            if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                u7 += pay_val
            if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                u8 += pay_val
            if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                u9 += pay_val
            if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                u10 += pay_val
            if oper_l == "??????????? ??????" and doc_l == "???????":
                u11 += pay_val
            if oper_l == "??????????? ??????" and doc_l == "???????":
                u12 += pay_val
            if oper_l == "??????????? ??????" and doc_l == "???????":
                u13 += pay_val
            if oper_l == "??????????? ??????" and doc_l == "???????":
                u14 += pay_val
        except Exception:
            pass
    
    date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
    date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
    revenue_calc = total_buyouts - total_returns
    defect_comp = x1 + x2 - x3 - x4 + x5 - x6 + x7 - x8
    total_wb_realized = wbr_plus - wbr_minus
    commission_total = (
        revenue_calc
        - (k1 + k2 + k3 + k4)
        + (k5 + k6 + k7 + k8)
        - total_acquiring
    )
    damage_comp = u1 + u2 - u3 - u4 + u5 - u6 + u7 + u8 - u9 - u10 + u11 - u12 + u13 - u14
    
    total_deductions = (
        commission_total + 
        total_acquiring + 
        total_logistics + 
        total_storage + 
        total_other_deductions + 
        total_acceptance - 
        defect_comp - 
        damage_comp - 
        total_paid_delivery + 
        total_penalties + 
        total_additional_payment
    )
    
    e3_correction = 0
    for r in raw:
        try:
            oper_name = (r.get("supplier_oper_name") or "").strip()
            if oper_name == "????????????? ??????????":
                e3_correction += float(r.get("ppvz_for_pay") or 0.0)
        except Exception:
            pass
    
    total_for_transfer = revenue_calc - total_deductions + e3_correction
    
    # ?????? ??????
    tax_amount = 0.0
    tax_rate = None
    if user_id:
        try:
            user = User.query.get(user_id)
            if user and user.tax_rate is not None:
                tax_rate = float(user.tax_rate)
                tax_amount = (total_wb_realized * tax_rate) / 100.0
        except Exception:
            pass
    
    return {
        "rows": raw,
        "total_qty": int(total_qty),
        "total_sum": round(total_sum, 2),
        "total_logistics": round(total_logistics, 2),
        "total_storage": round(total_storage, 2),
        "total_acceptance": round(total_acceptance, 2),
        "total_for_pay": round(total_for_pay, 2),
        "total_buyouts": round(total_buyouts, 2),
        "total_returns": round(total_returns, 2),
        "revenue": round(revenue_calc, 2),
        "total_wb_realized": round(total_wb_realized, 2),
        "total_commission": round(commission_total, 2),
        "total_acquiring": round(total_acquiring, 2),
        "total_commission_wb": round(total_commission_wb, 2),
        "total_other_deductions": round(total_other_deductions, 2),
        "total_penalties": round(total_penalties, 2),
        "total_defect_compensation": round(defect_comp, 2),
        "total_damage_compensation": round(damage_comp, 2),
        "total_paid_delivery": round(total_paid_delivery, 2),
        "total_additional_payment": round(total_additional_payment, 2),
        "total_deductions": round(total_deductions, 2),
        "total_for_transfer": round(total_for_transfer, 2),
        "tax_amount": round(tax_amount, 2),
        "tax_rate": tax_rate,
        "date_from_fmt": date_from_fmt,
        "date_to_fmt": date_to_fmt,
    }

def aggregate_finance_rows(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Aggregate rows by product (nm_id + supplierArticle) to qty and revenue.
    Fields per docs: realize/prices/quantities. We'll use 'quantity' and 'retail_amount' if present.
    """
    by_key: Dict[tuple[Any, Any], Dict[str, Any]] = {}
    for r in rows:
        nm_id = r.get("nm_id") or r.get("nmId") or r.get("nm")
        prod = r.get("supplier_article") or r.get("supplierArticle") or r.get("supplierArticleName") or ""
        key = (nm_id, prod)
        item = by_key.get(key)
        qty = 0
        rev = 0.0
        # common fields in WB v5: quantity, retail_amount
        try:
            qty = int(r.get("quantity") or r.get("sale_qty") or r.get("qty") or 0)
        except Exception:
            qty = 0
        try:
            rev = float(r.get("retail_amount") or r.get("retailAmount") or r.get("sum_price") or 0)
        except Exception:
            rev = 0.0
        if item is None:
            by_key[key] = {"nm_id": nm_id, "product": prod, "qty": max(qty, 0), "revenue": max(rev, 0.0)}
        else:
            item["qty"] = int(item.get("qty", 0)) + max(qty, 0)
            item["revenue"] = float(item.get("revenue", 0.0)) + max(rev, 0.0)
    items = list(by_key.values())
    items.sort(key=lambda x: (x.get("revenue") or 0.0), reverse=True)
    return items


def _normalize_and_group_orders(orders: List[Dict[str, Any]]) -> tuple:
    """
    ??????????? ? ?????????? ?????? ?? ???????.
    ??????? ???????? ??? ??????????? ??????? ? ?????????? ????????/????????? ????????, ?? ?????? ????????? WB.
    
    Returns:
        tuple: (counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product)
    """
    counts_total: Dict[str, int] = defaultdict(int)
    by_wh: Dict[str, Dict[str, int]] = defaultdict(lambda: defaultdict(int))
    revenue_total: Dict[str, float] = defaultdict(float)
    by_wh_sum: Dict[str, Dict[str, float]] = defaultdict(lambda: defaultdict(float))
    nm_by_product: Dict[str, Any] = {}
    barcode_by_product: Dict[str, Any] = {}
    supplier_article_by_product: Dict[str, Any] = {}
    
    # ???????? ??? ????????????: ?????? -> ??????? WB, ??????? ???????? -> ??????? WB
    barcode_to_nm: Dict[str, Any] = {}
    supplier_article_to_nm: Dict[str, Any] = {}
    
    def _row_barcode(row: Dict[str, Any]) -> Any:
        return (
            row.get("_barcode")
            or row.get("\u0411\u0430\u0440\u043a\u043e\u0434")
            or row.get("??????")
            or row.get("barcode")
        )

    def _row_nm_id(row: Dict[str, Any]) -> Any:
        return (
            row.get("_nm_id")
            or row.get("\u0410\u0440\u0442\u0438\u043a\u0443\u043b WB")
            or row.get("??????? WB")
            or row.get("nmId")
            or row.get("nmID")
        )

    def _row_supplier_article(row: Dict[str, Any]) -> Any:
        return (
            row.get("_supplier_article")
            or row.get("\u0410\u0440\u0442\u0438\u043a\u0443\u043b \u043f\u0440\u043e\u0434\u0430\u0432\u0446\u0430")
            or row.get("??????? ????????")
            or row.get("supplierArticle")
        )

    def _row_warehouse(row: Dict[str, Any]) -> str:
        return str(
            row.get("_warehouse")
            or row.get("\u0421\u043a\u043b\u0430\u0434 \u043e\u0442\u0433\u0440\u0443\u0437\u043a\u0438")
            or row.get("????? ????????")
            or row.get("warehouseName")
            or "?? ??????"
        )

    def _row_price(row: Dict[str, Any]) -> float:
        raw = row.get("_price")
        if raw is None:
            raw = row.get("\u0426\u0435\u043d\u0430 \u0441\u043e \u0441\u043a\u0438\u0434\u043a\u043e\u0439 \u043f\u0440\u043e\u0434\u0430\u0432\u0446\u0430")
        if raw is None:
            raw = row.get("???? ?? ??????? ????????")
        if raw is None:
            raw = row.get("priceWithDisc")
        try:
            return float(raw or 0)
        except (TypeError, ValueError):
            return 0.0

    # ?????? ??????: ???????? ????? ??? ????????????
    for r in (orders or []):
        if r.get("is_cancelled", False):
            continue
        barcode = _row_barcode(r)
        nmv = _row_nm_id(r)
        supplier_article = _row_supplier_article(r)
        
        # ????????? ?????: ???? ???? ??????? WB, ????????? ??? ? ???????? ? ????????? ????????
        if nmv:
            if barcode:
                barcode_to_nm[barcode] = nmv
            if supplier_article:
                supplier_article_to_nm[supplier_article] = nmv
    
    # ?????? ??????: ?????????? ? ?????????????? ????????????
    for r in (orders or []):
        # ?????????? ?????????? ?????? ? ??????
        if r.get("is_cancelled", False):
            continue
        # ?????????? ?? ???????? WB (????????? 1 - ????? ??????????), ????? ?? ??????? (????????? 2), ????? ?? ???????? ???????? (????????? 3)
        barcode = _row_barcode(r)
        nmv = _row_nm_id(r)
        supplier_article = _row_supplier_article(r)
        
        # ????????????: ???? ???????? WB ???, ?? ???? ????????? ?????? ??? ??????? ????????, ?????????? ????????? ??????? WB
        if not nmv:
            if barcode and barcode in barcode_to_nm:
                nmv = barcode_to_nm[barcode]
            elif supplier_article and supplier_article in supplier_article_to_nm:
                nmv = supplier_article_to_nm[supplier_article]
        
        # ?????????? ???? ??? ???????????
        # ?????????? ??????? WB ??? ???????? ?????????????, ??? ??? ?? ????? ????????
        if nmv:
            prod_key = f"NM_{nmv}"
        elif barcode:
            prod_key = f"BARCODE_{barcode}"
        elif supplier_article:
            prod_key = str(supplier_article)
        else:
            prod_key = "?? ??????"
        
        wh = _row_warehouse(r)
        counts_total[prod_key] += 1
        by_wh[prod_key][wh] += 1
        price = _row_price(r)
        revenue_total[prod_key] += price
        by_wh_sum[prod_key][wh] += price
        
        # ????????? ?????? ? ??????, ?????? ???????? ??????? ???????? ?? ????????? (??????????)
        if nmv:
            nm_by_product[prod_key] = nmv
        if barcode:
            barcode_by_product[prod_key] = barcode
        if supplier_article:
            # ?????? ????????? ??????? ???????? ?? ????????? (??????????)
            supplier_article_by_product[prod_key] = supplier_article
    
    return (counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product)


def to_rows(data: List[Dict[str, Any]], start_date: str, end_date: str) -> List[Dict[str, Any]]:
    """??????????? ?????? ??????? ? ????????? ?? ???????? ???? ?????? (date), ??????? ??????????."""
    start = parse_date(start_date).date()
    end = parse_date(end_date).date()

    rows: List[Dict[str, Any]] = []
    for sale in data:
        # ????????? ?? ???????? ???? ?????? (date), ? ?? ?? lastChangeDate
        date_str = str(sale.get("date", ""))[:10]
        try:
            d = datetime.strptime(date_str, "%Y-%m-%d").date()
        except ValueError:
            continue
        if not (start <= d <= end):
            continue
        # ?????? ???????? ??? ??????, ??????? ??????????
        is_cancelled = sale.get("isCancel")
        is_cancelled_bool = is_cancelled is True or str(is_cancelled).lower() in ('true', '1', '??????')
        
        _wh_raw = sale.get("warehouseName")
        _warehouse_label = (
            str(_wh_raw).strip()
            if _wh_raw is not None and str(_wh_raw).strip()
            else "\u041d\u0435 \u0443\u043a\u0430\u0437\u0430\u043d"
        )
        rows.append({
            # Canonical fields (stable, encoding-independent) for server-side grouping
            "_warehouse": sale.get("warehouseName"),
            "_order_date": date_str,
            "_warehouse_label": _warehouse_label,
            "_supplier_article": sale.get("supplierArticle"),
            "_nm_id": sale.get("nmId"),
            "_barcode": sale.get("barcode"),
            "_price": sale.get("priceWithDisc"),
            "????": date_str,
            "???? ? ????? ?????????? ?????????? ? ???????": sale.get("lastChangeDate"),
            "????? ????????": sale.get("warehouseName"),
            "??? ?????? ???????? ???????": sale.get("warehouseType"),
            "??????": sale.get("countryName"),
            "?????": sale.get("oblastOkrugName"),
            "??????": sale.get("regionName"),
            "??????? ????????": sale.get("supplierArticle"),
            "??????? WB": sale.get("nmId"),
            "??????": sale.get("barcode"),
            "?????????": sale.get("category"),
            "???????": sale.get("subject"),
            "?????": sale.get("brand"),
            "?????? ??????": sale.get("techSize"),
            "????? ????????": sale.get("incomeID"),
            "??????? ????????": sale.get("isSupply"),
            "??????? ??????????": sale.get("isRealization"),
            "???? ??? ??????": sale.get("totalPrice"),
            "?????? ????????": sale.get("discountPercent"),
            "?????? WB": sale.get("spp"),
            "???? ? ?????? ???? ??????": sale.get("finishedPrice"),
            "???? ?? ??????? ????????": sale.get("priceWithDisc"),
            "?????? ??????": sale.get("isCancel"),
            "???? ? ????? ?????? ??????": sale.get("cancelDate"),
            "ID ???????": sale.get("sticker"),
            "????? ??????": sale.get("gNumber"),
            "?????????? ID ??????": sale.get("srid"),
            "is_cancelled": is_cancelled_bool,  # ????????? ???? ?????? ??? ????????
        })
    return rows


def to_sales_rows(data: List[Dict[str, Any]], start_date: str, end_date: str) -> List[Dict[str, Any]]:
    start = parse_date(start_date).date()
    end = parse_date(end_date).date()
    rows: List[Dict[str, Any]] = []
    for sale in data:
        date_str = str(sale.get("date", ""))[:10]
        try:
            d = datetime.strptime(date_str, "%Y-%m-%d").date()
        except ValueError:
            continue
        if not (start <= d <= end):
            continue
        rows.append({
            "????": date_str,
            "???? ? ????? ?????????? ?????????? ? ???????": sale.get("lastChangeDate"),
            "????? ????????": sale.get("warehouseName"),
            "??????? ????????": sale.get("supplierArticle"),
            "??????? WB": sale.get("nmId"),
            "??????": sale.get("barcode"),
            "???? ? ?????? ???? ??????": sale.get("finishedPrice"),
        })
    return rows


def aggregate_daily(rows: List[Dict[str, Any]]):
    count_by_day: Dict[str, int] = defaultdict(int)
    revenue_by_day: Dict[str, float] = defaultdict(float)

    for r in rows:
        day = _order_row_day_iso(r)
        if not day:
            continue
        price = _order_row_price_disc(r)
        count_by_day[day] += 1
        revenue_by_day[day] += price

    labels = sorted(count_by_day.keys())
    counts = [count_by_day[d] for d in labels]
    revenues = [round(revenue_by_day[d], 2) for d in labels]
    return labels, counts, revenues


def aggregate_daily_counts_and_revenue(rows: List[Dict[str, Any]]):
    count_by_day: Dict[str, int] = defaultdict(int)
    cancelled_count_by_day: Dict[str, int] = defaultdict(int)
    revenue_by_day: Dict[str, float] = defaultdict(float)
    for r in rows:
        day = _order_row_day_iso(r)
        if not day:
            continue
        is_cancelled = r.get("is_cancelled", False)
        
        # ???????????? ????? ?????????? ???????
        count_by_day[day] += 1
        
        # ???????????? ?????????? ?????? ????????
        if is_cancelled:
            cancelled_count_by_day[day] += 1
        
        # ??????? ??????? ?????? ? ???????? ???????
        if not is_cancelled:
            revenue_by_day[day] += _order_row_price_disc(r)
    return count_by_day, revenue_by_day, cancelled_count_by_day


def build_union_series(orders_counts: Dict[str, int], sales_counts: Dict[str, int],
                       orders_rev: Dict[str, float], sales_rev: Dict[str, float]):
    labels = sorted(set(orders_counts.keys()) | set(sales_counts.keys()))
    o_counts = [orders_counts.get(d, 0) for d in labels]
    s_counts = [sales_counts.get(d, 0) for d in labels]
    o_rev = [round(orders_rev.get(d, 0.0), 2) for d in labels]
    s_rev = [round(sales_rev.get(d, 0.0), 2) for d in labels]
    return labels, o_counts, s_counts, o_rev, s_rev


def aggregate_by_warehouse(rows: List[Dict[str, Any]]) -> List[Tuple[str, int]]:
    counts: Dict[str, int] = defaultdict(int)
    for r in rows:
        warehouse = _order_row_warehouse_label(r)
        counts[warehouse] += 1
    return sorted(counts.items(), key=lambda kv: kv[1], reverse=True)


def aggregate_by_warehouse_dual(orders_rows: List[Dict[str, Any]], sales_rows: List[Dict[str, Any]]):
    orders_map: Dict[str, int] = defaultdict(int)
    sales_map: Dict[str, int] = defaultdict(int)
    for r in orders_rows:
        warehouse = _order_row_warehouse_label(r)
        orders_map[warehouse] += 1
    for r in sales_rows:
        warehouse = _order_row_warehouse_label(r)
        sales_map[warehouse] += 1
    all_wh = sorted(set(orders_map.keys()) | set(sales_map.keys()))
    summary = []
    for w in all_wh:
        summary.append({"warehouse": w, "orders": orders_map.get(w, 0), "sales": sales_map.get(w, 0)})
    # ????????? ?? ???????
    summary.sort(key=lambda x: x["orders"], reverse=True)
    return summary

def aggregate_by_warehouse_orders_only(orders_rows: List[Dict[str, Any]]):
    orders_map: Dict[str, int] = defaultdict(int)
    for r in orders_rows:
        # ?????????? ?????????? ?????? ? ?????????? ?? ???????
        if r.get("is_cancelled", False):
            continue
        warehouse = _order_row_warehouse_label(r)
        orders_map[warehouse] += 1
    summary = []
    for w in sorted(orders_map.keys()):
        summary.append({"warehouse": w, "orders": orders_map.get(w, 0)})
    # ????????? ?? ???????
    summary.sort(key=lambda x: x["orders"], reverse=True)
    return summary


def aggregate_top_products(rows: List[Dict[str, Any]], limit: int = 15) -> List[Dict[str, Any]]:
    counts: Dict[str, int] = defaultdict(int)
    revenue_by_product: Dict[str, float] = defaultdict(float)
    nm_by_product: Dict[str, Any] = {}
    barcode_by_product: Dict[str, Any] = {}
    supplier_article_by_product: Dict[str, Any] = {}
    for r in rows:
        # ?????????? ?????????? ?????? ? ??? ???????
        if r.get("is_cancelled", False):
            continue
        product = r.get("??????? ????????") or r.get("??????? WB") or r.get("??????") or "?? ??????"
        product = str(product)
        counts[product] += 1
        revenue_by_product[product] += _order_row_price_disc(r)
        nm = r.get("??????? WB") or r.get("nmId") or r.get("nmID")
        if product not in nm_by_product and nm:
            nm_by_product[product] = nm
        barcode = r.get("??????")
        if product not in barcode_by_product and barcode:
            barcode_by_product[product] = barcode
        supplier_article = r.get("??????? ????????")
        if product not in supplier_article_by_product and supplier_article:
            supplier_article_by_product[product] = supplier_article
    # photo map
    nm_to_photo: Dict[Any, Any] = {}
    try:
        prod_cached = load_products_cache() or {}
        for it in (prod_cached.get("items") or []):
            nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
            photo = it.get("photo") or it.get("img")
            if nmv is not None and nmv not in nm_to_photo:
                nm_to_photo[nmv] = photo
    except Exception:
        nm_to_photo = {}

    items = [{
        "product": p,
        "qty": c,
        "nm_id": nm_by_product.get(p),
        "barcode": barcode_by_product.get(p),
        "supplier_article": supplier_article_by_product.get(p),
        "sum": round(revenue_by_product.get(p, 0.0), 2),
        "photo": nm_to_photo.get(nm_by_product.get(p))
    } for p, c in counts.items()]
    items.sort(key=lambda x: x["qty"], reverse=True)
    return items[:limit]


def aggregate_top_products_sales(rows: List[Dict[str, Any]], warehouse: str | None = None, limit: int = 50) -> List[Tuple[str, int]]:
    counts: Dict[str, int] = defaultdict(int)
    for r in rows:
        if warehouse and _order_row_warehouse_label(r) != warehouse:
            continue
        product = r.get("??????? ????????") or r.get("??????? WB") or r.get("??????") or "?? ??????"
        counts[str(product)] += 1
    return sorted(counts.items(), key=lambda kv: kv[1], reverse=True)[:limit]
def aggregate_top_products_orders(rows: List[Dict[str, Any]], warehouse: str | None = None, limit: int = 50) -> List[Dict[str, Any]]:
    counts: Dict[str, int] = defaultdict(int)
    revenue_by_product: Dict[str, float] = defaultdict(float)
    nm_by_product: Dict[str, Any] = {}
    barcode_by_product: Dict[str, Any] = {}
    supplier_article_by_product: Dict[str, Any] = {}
    for r in rows:
        # ?????????? ?????????? ?????? ? ??? ???????
        if r.get("is_cancelled", False):
            continue
        if warehouse and _order_row_warehouse_label(r) != warehouse:
            continue
        product = r.get("??????? ????????") or r.get("??????? WB") or r.get("??????") or "?? ??????"
        product = str(product)
        counts[product] += 1
        revenue_by_product[product] += _order_row_price_disc(r)
        nm = r.get("??????? WB") or r.get("nmId") or r.get("nmID")
        if product not in nm_by_product and nm:
            nm_by_product[product] = nm
        barcode = r.get("??????")
        if product not in barcode_by_product and barcode:
            barcode_by_product[product] = barcode
        supplier_article = r.get("??????? ????????")
        if product not in supplier_article_by_product and supplier_article:
            supplier_article_by_product[product] = supplier_article
    # Enrich with product photos from cache
    nm_to_photo: Dict[Any, Any] = {}
    try:
        prod_cached = load_products_cache() or {}
        for it in (prod_cached.get("items") or []):
            nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
            photo = it.get("photo") or it.get("img")
            if nmv is not None and nmv not in nm_to_photo:
                nm_to_photo[nmv] = photo
    except Exception:
        nm_to_photo = {}

    # Load stocks data for current user
    stocks_data = {}
    try:
        stocks_cached = load_stocks_cache()
        if stocks_cached and stocks_cached.get("_user_id"):
            for stock_item in stocks_cached.get("items", []):
                barcode = stock_item.get("barcode")
                stock_warehouse = stock_item.get("warehouse", "")
                qty = int(stock_item.get("qty", 0) or 0)
                
                if barcode:
                    if warehouse:
                        # ???? ?????? ?????????? ?????, ????????? ?????? ?? ????? ??????
                        if (stock_warehouse == warehouse or 
                            (warehouse in stock_warehouse) or 
                            (stock_warehouse in warehouse)):
                            stocks_data[barcode] = stocks_data.get(barcode, 0) + qty
                    else:
                        # ???? ?? ?????? ?????, ????????? ?? ???? ???????
                        stocks_data[barcode] = stocks_data.get(barcode, 0) + qty
    except Exception:
        stocks_data = {}

    items = [{
        "product": p,
        "qty": c,
        "nm_id": nm_by_product.get(p),
        "barcode": barcode_by_product.get(p),
        "supplier_article": supplier_article_by_product.get(p),
        "sum": round(revenue_by_product.get(p, 0.0), 2),
        "photo": nm_to_photo.get(nm_by_product.get(p)),
        "stock_qty": stocks_data.get(barcode_by_product.get(p), 0)
    } for p, c in counts.items()]
    items.sort(key=lambda x: x["qty"], reverse=True)
    return items[:limit]


def _extract_created_at(obj: Any) -> datetime:
    if not isinstance(obj, dict):
        return datetime.min
    val = (
        obj.get("createdAt")
        or obj.get("createAt")
        or obj.get("created")
        or obj.get("createDt")
        or obj.get("createdDt")
        or obj.get("createdDate")
        or obj.get("dateCreated")
        or obj.get("orderCreateDate")
        or obj.get("date")
        or obj.get("created_at")
        or obj.get("time")
        or ""
    )
    # Numeric timestamp support (ms or s)
    try:
        if isinstance(val, (int, float)):
            ts = float(val)
            if ts > 1e12:  # milliseconds
                return datetime.fromtimestamp(ts / 1000)
            if ts > 1e9:   # seconds
                return datetime.fromtimestamp(ts)
        s = str(val).strip()
        if s.isdigit():
            ts = float(s)
            if ts > 1e12:
                return datetime.fromtimestamp(ts / 1000)
            if ts > 1e9:
                return datetime.fromtimestamp(ts)
        dt = parse_wb_datetime(s)
        return dt or datetime.min
    except Exception:
        return datetime.min


def fetch_fbs_new_orders(token: str) -> List[Dict[str, Any]]:
    # Marketplace API expects the token without 'Bearer'
    headers = {"Authorization": f"{token}"}
    resp = get_with_retry(FBS_NEW_URL, headers, params={})
    data = resp.json()
    # Normalize to list of orders from various shapes
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        if isinstance(data.get("orders"), list):
            return data["orders"]
        inner = data.get("data")
        if isinstance(inner, list):
            return inner
        if isinstance(inner, dict) and isinstance(inner.get("orders"), list):
            return inner["orders"]
    return []


def fetch_fbs_orders(token: str, limit: int = 100, next_cursor: str | None = None) -> Dict[str, Any]:
    params: Dict[str, Any] = {"limit": limit, "next": 0 if next_cursor is None else next_cursor}
    # Try both auth styles. Some WB tenants expect bare token, ?????? ? Bearer
    headers_list = [
        {"Authorization": f"{token}"},
        {"Authorization": f"Bearer {token}"},
    ]
    last_err: Exception | None = None
    for hdrs in headers_list:
        try:
            resp = get_with_retry(FBS_ORDERS_URL, hdrs, params=params)
            data = resp.json()
            # consider non-empty result a success
            arr = (data.get("orders") if isinstance(data, dict) else None) or []
            if isinstance(arr, list) and arr:
                return data
            # even if empty, return once tried both
            last_data = data
        except Exception as e:
            last_err = e
            continue
    if last_err:
        raise last_err
    return last_data  # type: ignore[name-defined]


def fetch_fbs_statuses(token: str, order_ids: List[int]) -> Dict[str, Any]:
    headers_list = [
        {"Authorization": f"{token}"},
        {"Authorization": f"Bearer {token}"},
    ]
    bodies = [
        {"orders": order_ids},
        {"orders": [{"id": oid} for oid in order_ids]},
    ]
    last_err: Exception | None = None
    for hdrs in headers_list:
        for body in bodies:
            try:
                resp = post_with_retry(FBS_ORDERS_STATUS_URL, hdrs, json_body=body)
                return resp.json()
            except Exception as e:
                last_err = e
                continue
    if last_err:
        raise last_err
    return {}


def fetch_dbs_new_orders(token: str) -> List[Dict[str, Any]]:
    """Fetch new DBS orders."""
    headers = {"Authorization": f"{token}"}
    resp = get_with_retry(DBS_NEW_URL, headers, params={})
    data = resp.json()
    if isinstance(data, list):
        return data
    if isinstance(data, dict):
        arr = data.get("orders")
        if isinstance(arr, list):
            return arr
        inner = data.get("data")
        if isinstance(inner, list):
            return inner
        if isinstance(inner, dict) and isinstance(inner.get("orders"), list):
            return inner["orders"]
    return []


def fetch_dbs_statuses(token: str, order_ids: List[int]) -> Dict[str, Any]:
    headers_list = [
        {"Authorization": f"{token}"},
        {"Authorization": f"Bearer {token}"},
    ]
    bodies = [
        {"orders": order_ids},
    ]
    last_err: Exception | None = None
    for hdrs in headers_list:
        for body in bodies:
            try:
                resp = post_with_retry(DBS_STATUS_URL, hdrs, json_body=body)
                return resp.json()
            except Exception as e:
                last_err = e
                continue
    if last_err:
        raise last_err
    return {}


def fetch_dbs_orders(
    token: str,
    limit: int = 1000,
    next_cursor: Any | None = None,
    date_from_ts: int | None = None,
    date_to_ts: int | None = None,
) -> Dict[str, Any]:
    """Fetch completed DBS assembly orders after sale or cancellation."""
    params: Dict[str, Any] = {
        "limit": limit,
        "next": 0 if next_cursor is None else next_cursor,
    }
    if date_from_ts is not None:
        params["dateFrom"] = int(date_from_ts)
    if date_to_ts is not None:
        params["dateTo"] = int(date_to_ts)
    headers_list = [
        {"Authorization": f"{token}"},
        {"Authorization": f"Bearer {token}"},
    ]
    last_err: Exception | None = None
    last_data: Dict[str, Any] | None = None
    attempts: list[tuple[dict[str, Any], str]] = [(params, "sec")]
    if date_from_ts is not None and date_to_ts is not None:
        params_ms = dict(params)
        try:
            params_ms["dateFrom"] = int(date_from_ts) * 1000
            params_ms["dateTo"] = int(date_to_ts) * 1000
        except Exception:
            params_ms = params
        attempts.append((params_ms, "ms"))
    for hdrs in headers_list:
        for p, tag in attempts:
            try:
                resp = get_with_retry(DBS_ORDERS_URL, hdrs, params=p)
                data = resp.json() if (resp.text or "").strip() else {}
                last_data = data if isinstance(data, dict) else {"data": data}
                return last_data
            except Exception as e:
                last_err = e
                continue
    if last_err:
        raise last_err
    return last_data or {}


def to_dbs_rows(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    """Normalize DBS orders to a simple table row format for UI."""
    rows: List[Dict[str, Any]] = []
    for it in items:
        try:
            created_raw = (
                it.get("createdAt")
                or it.get("dateCreated")
                or it.get("date")
            )
            try:
                _dt = parse_wb_datetime(str(created_raw)) if created_raw else None
                _dt_msk = to_moscow(_dt) if _dt else None
                created_str = _dt_msk.strftime("%d.%m.%Y %H:%M") if _dt_msk else (str(created_raw) if created_raw else "")
            except Exception:
                created_str = str(created_raw) if created_raw else ""
            nm_id = it.get("nmId") or it.get("nmID")
            article = it.get("article") or it.get("vendorCode") or ""
            price = (
                it.get("finalPrice")
                or it.get("convertedFinalPrice")
                or it.get("salePrice")
                or it.get("price")
            )
            addr = None
            adr = it.get("address") or {}
            if isinstance(adr, dict):
                addr = adr.get("fullAddress") or None
            status_val = (
                it.get("status")
                or it.get("supplierStatus")
                or it.get("wbStatus")
            )
            status_name_val = (
                it.get("statusName")
                or it.get("supplierStatusName")
                or it.get("wbStatusName")
                or status_val
            )
            rows.append({
                "orderId": it.get("id") or it.get("orderId") or it.get("ID"),
                "????? ? ???? ??????": f"{it.get('id') or it.get('orderId') or ''} | {created_str}".strip(" |"),
                "???????????? ??????": article,
                "????": price,
                "?????": addr or "",
                "nm_id": nm_id,
                "status": status_val,
                "statusName": status_name_val,
            })
        except Exception:
            continue
    return rows


def fetch_fbs_latest_orders(token: str, want_count: int = 30, page_limit: int = 200, max_pages: int = 20) -> tuple[List[Dict[str, Any]], Any]:
    """Fetch multiple pages and return most recent `want_count` items by created time.

    WB API ?????? ???????? ?? ????????? next. ?????? ???????? (next=0) ????? ????????? ?????? ??????,
    ??????? ???? ?? ?????????, ???????? ? ????? ????? ????????? ?? ????.
    """
    collected: List[Dict[str, Any]] = []
    cursor: Any = 0
    pages = 0
    last_next: Any = None
    while pages < max_pages:
        page = fetch_fbs_orders(token, limit=page_limit, next_cursor=cursor)
        items, next_cursor = _normalize_fbs_orders_page(page)
        if not items:
            break
        collected.extend(items)
        pages += 1
        if not next_cursor or next_cursor == cursor:
            break
        cursor = next_cursor
        last_next = next_cursor
    try:
        collected.sort(key=_extract_created_at, reverse=True)
    except Exception:
        pass
    return collected[:want_count], last_next


def _merge_statuses_for_items(token: str, items: List[Dict[str, Any]]):
    ids: List[int] = []
    for it in items:
        oid = it.get("id") or it.get("orderId") or it.get("ID")
        try:
            if oid is not None:
                ids.append(int(oid))
        except Exception:
            continue
    if not ids:
        return
    st = fetch_fbs_statuses(token, ids[:1000])
    arr = st.get("orders") if isinstance(st, dict) else None
    if arr is None:
        arr = st.get("data") if isinstance(st, dict) else None
    if arr is None:
        arr = st if isinstance(st, list) else []
    map_st: Dict[int, Any] = {}
    if isinstance(arr, list):
        for x in arr:
            try:
                map_st[int(x.get("id") or x.get("orderId") or 0)] = x
            except Exception:
                continue
    for it in items:
        try:
            oid = int(it.get("id") or it.get("orderId") or it.get("ID") or 0)
            stx = map_st.get(oid) or {}
            # WB ????? ?????????? ?????? ???? ??? ????????
            status_val = (
                stx.get("status")
                or stx.get("supplierStatus")
                or stx.get("wbStatus")
                or stx.get("state")
            )
            status_name_val = (
                stx.get("statusName")
                or stx.get("supplierStatusName")
                or stx.get("wbStatusName")
                or stx.get("stateName")
                or status_val
            )
            if status_name_val:
                it["statusName"] = status_name_val
            if status_val:
                it["status"] = status_val
        except Exception:
            continue


def get_orders_with_status(token: str, need_count: int = 30, start_next: Any = None) -> tuple[List[Dict[str, Any]], Any]:
    collected: List[Dict[str, Any]] = []
    cursor: Any = 0 if (start_next is None or start_next == "" or start_next == "null") else start_next
    last_next: Any = None
    safety_pages = 0
    while len(collected) < need_count and safety_pages < 5:
        page = fetch_fbs_orders(token, limit=200, next_cursor=cursor)
        items, next_cursor = _normalize_fbs_orders_page(page)
        if not items:
            break
        collected.extend(items)
        last_next = next_cursor
        safety_pages += 1
        if not next_cursor or next_cursor == cursor:
            break
        cursor = next_cursor
    try:
        collected.sort(key=_extract_created_at, reverse=True)
    except Exception:
        pass
    result = collected[:need_count]
    _merge_statuses_for_items(token, result)
    return result, last_next


def _normalize_fbs_orders_page(page: Any) -> tuple[list[dict], str | None]:
    try:
        if isinstance(page, list):
            return page, None
        if isinstance(page, dict):
            items = page.get("orders") or page.get("data") or []
            if not isinstance(items, list):
                items = []
            next_cursor = page.get("next") or page.get("cursor") or None
            return items, next_cursor
    except Exception:
        pass
    return [], None

def to_fbs_rows(orders: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    rows: List[Dict[str, Any]] = []
    for o in orders:
        if not isinstance(o, dict):
            order_num = str(o)
            rows.append({
                "????? ? ???? ???????": f"{order_num}",
                "???????????? ??????": "",
                "????????? ??????": 0,
                "?????": "",
            })
            continue
        # ????? ??????? ? ID
        order_id = o.get("ID") or o.get("id") or o.get("orderId") or o.get("gNumber") or ""
        # ???? ? createdAt (??????????? ? ???????? ???? ????????)
        ca_raw = o.get("createdAt") or o.get("dateCreated") or o.get("date")
        ca_dt = parse_wb_datetime(str(ca_raw))
        ca_dt_msk = to_moscow(ca_dt) if ca_dt else None
        created_at = ca_dt_msk.strftime("%d.%m.%Y %H:%M") if ca_dt_msk else str(ca_raw or "")[:10]
        # ???????????? ? article
        article = o.get("article") or ""
        # ???? ? price ??? ???? ????????? ?????
        raw_price = o.get("convertedPrice")
        try:
            price_value = int(raw_price) // 100
        except Exception:
            try:
                price_value = int(float(raw_price)) // 100
            except Exception:
                price_value = 0
        # ????? ? offices
        warehouse = ""
        offices = o.get("offices")
        if isinstance(offices, list) and offices:
            first = offices[0]
            if isinstance(first, dict):
                warehouse = first.get("name") or first.get("officeName") or first.get("address") or ""
            else:
                warehouse = str(first)
        elif isinstance(offices, dict):
            warehouse = offices.get("name") or offices.get("officeName") or offices.get("address") or ""
        elif isinstance(offices, str):
            warehouse = offices
        if not warehouse:
            warehouse = o.get("warehouseName") or o.get("warehouse") or o.get("warehouseId") or ""

        rows.append({
            "????? ? ???? ???????": f"{order_id} | {created_at}".strip(" |"),
            "???????????? ??????": article,
            "????????? ??????": price_value,
            "????": price_value,
            "?????": warehouse,
            # placeholders for enrichment from products cache
            "nm_id": o.get("nmID") or o.get("nmId") or None,
            "barcode": None,
            "photo": None,
            "orderId": order_id,  # ????????? orderId ??? JavaScript
        })
    return rows


def fetch_seller_info(token: str) -> Dict[str, Any] | None:
    if not token:
        return None
    # WB /api/v1/seller-info: use one request format only (Bearer),
    # because endpoint has strict per-seller rate limits.
    headers1 = {"Authorization": f"Bearer {token}"}
    try:
        resp = get_with_retry(SELLER_INFO_URL, headers1, params={})
        return resp.json()
    except Exception:
        return None
def decode_token_info(token: str) -> Dict[str, Any] | None:
    """Decode JWT token to extract creation and expiration information"""
    if not token:
        return None
    
    try:
        # Decode JWT token without verification (we only need the payload)
        decoded = jwt.decode(token, options={"verify_signature": False})
        
        token_info = {}
        
        # Extract creation date (iat - issued at)
        if 'iat' in decoded:
            iat_timestamp = decoded['iat']
            token_info['created_at'] = datetime.fromtimestamp(iat_timestamp, tz=MOSCOW_TZ)
        
        # Extract expiration date (exp - expiration)
        if 'exp' in decoded:
            exp_timestamp = decoded['exp']
            token_info['expires_at'] = datetime.fromtimestamp(exp_timestamp, tz=MOSCOW_TZ)
            
            # Calculate days until expiration
            now = datetime.now(MOSCOW_TZ)
            if token_info['expires_at'] > now:
                days_left = (token_info['expires_at'] - now).days
                token_info['days_until_expiry'] = days_left
                token_info['is_expired'] = False
            else:
                token_info['days_until_expiry'] = 0
                token_info['is_expired'] = True
        
        # Extract other useful information if available
        if 'sub' in decoded:
            token_info['subject'] = decoded['sub']
        if 'iss' in decoded:
            token_info['issuer'] = decoded['iss']
        
        return token_info
        
    except Exception as e:
        print(f"Error decoding token: {e}")
        return None


def fetch_acceptance_coefficients(token: str) -> List[Dict[str, Any]] | None:
    if not token:
        return None
    # Try Bearer first
    headers1 = {"Authorization": f"Bearer {token}"}
    try:
        resp = get_with_retry(ACCEPT_COEFS_URL, headers1, params={})
        return resp.json()
    except Exception:
        # Fallback raw token
        headers2 = {"Authorization": f"{token}"}
        try:
            resp = get_with_retry(ACCEPT_COEFS_URL, headers2, params={})
            return resp.json()
        except Exception:
            return None


def build_acceptance_grid(items: List[Dict[str, Any]], days: int = 14):
    # Prepare date list: today + next N days
    today = datetime.now().date()
    date_objs = [today + timedelta(days=i) for i in range(days + 1)]
    date_keys = [d.strftime("%Y-%m-%d") for d in date_objs]
    date_labels = [d.strftime("%d-%m") for d in date_objs]

    # Filter only box type '??????' (boxTypeID == 2) for robustness also match by name
    filtered: List[Dict[str, Any]] = []
    for it in items or []:
        try:
            bt_id = it.get("boxTypeID")
            bt_name = str(it.get("boxTypeName") or "").lower()
            if (bt_id == 2) or ("?????" in bt_name):
                filtered.append(it)
        except Exception:
            continue

    # Unique warehouses from filtered
    warehouses: List[str] = sorted({str(it.get("warehouseName") or "") for it in filtered if it})

    # Map: (warehouse, date_key) -> record
    grid: Dict[str, Dict[str, Dict[str, Any]]] = {w: {} for w in warehouses}

    for it in filtered:
        try:
            wname = str(it.get("warehouseName") or "")
            dkey = str(it.get("date") or "")[:10]
            if wname not in grid or dkey not in date_keys:
                continue
            raw_coef = it.get("coefficient")
            try:
                coef_val = float(raw_coef)
            except Exception:
                coef_val = None
            grid[wname][dkey] = {
                "coef": coef_val,
                "allow": bool(it.get("allowUnload")),
            }
        except Exception:
            continue

    # Fill empty cells
    for w in warehouses:
        for dkey in date_keys:
            if dkey not in grid[w]:
                grid[w][dkey] = {"coef": None, "allow": None}

    # Sort warehouses by number of non-negative coefficients (>=0) across the horizon
    def count_non_negative(w: str) -> int:
        count = 0
        for dkey in date_keys:
            coef = grid[w][dkey].get("coef")
            try:
                if coef is not None and float(coef) >= 0:
                    count += 1
            except Exception:
                continue
        return count

    warehouses.sort(key=lambda w: count_non_negative(w), reverse=True)

    return warehouses, date_keys, date_labels, grid


def fetch_fbs_warehouses(token: str) -> list[dict[str, Any]]:
    if not token:
        return []
    headers1 = {"Authorization": f"Bearer {token}"}
    try:
        resp = get_with_retry(FBS_WAREHOUSES_URL, headers1, params={})
        return resp.json() or []
    except Exception:
        headers2 = {"Authorization": f"{token}"}
        try:
            resp = get_with_retry(FBS_WAREHOUSES_URL, headers2, params={})
            return resp.json() or []
        except Exception:
            return []


def fetch_fbs_stocks_by_warehouse(token: str, warehouse_id: int, skus: list[str]) -> list[dict[str, Any]]:
    if not token or not warehouse_id or not skus:
        return []
    url = FBS_STOCKS_BY_WAREHOUSE_URL.format(warehouseId=warehouse_id)
    headers1 = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    body = {"skus": skus[:1000]}
    try:
        resp = post_with_retry(url, headers1, body)
        data = resp.json() or {}
        return data.get("stocks") or []
    except Exception:
        headers2 = {"Authorization": f"{token}", "Content-Type": "application/json"}
        resp = post_with_retry(url, headers2, body)
        data = resp.json() or {}
        return data.get("stocks") or []


def update_fbs_stocks_by_warehouse(token: str, warehouse_id: int, items: list[dict[str, Any]]) -> None:
    # items: {"sku": str, "amount": int}
    if not token or not warehouse_id or not items:
        raise ValueError("bad_args")
    url = FBS_STOCKS_BY_WAREHOUSE_URL.format(warehouseId=warehouse_id)
    
    headers1 = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    headers2 = {"Authorization": f"{token}", "Content-Type": "application/json"}
    
    print(f"Updating stocks for warehouse {warehouse_id}, items: {len(items)}")
    
    # Validate items before sending
    valid_items = []
    for item in items:
        if isinstance(item, dict) and 'sku' in item and 'amount' in item:
            sku = str(item['sku']).strip()
            amount = int(item['amount'])
            if sku and amount >= 0:  # Only include valid items with non-negative amounts
                valid_items.append({"sku": sku, "amount": amount})
            else:
                print(f"Skipping invalid item: {item}")
        else:
            print(f"Skipping malformed item: {item}")
    
    if not valid_items:
        print("No valid items to update")
        return
    
    print(f"Valid items to update: {len(valid_items)}")
    
    # Split into batches to avoid rate limiting (max 200 items per batch)
    batch_size = 200
    batches = [valid_items[i:i + batch_size] for i in range(0, len(valid_items), batch_size)]
    print(f"Split into {len(batches)} batches of up to {batch_size} items each")
    
    # Try different body formats for WB API
    body_formats = [
        lambda batch: batch,  # Direct array without wrapper
        lambda batch: {"stocks": batch},  # With "stocks" wrapper
        lambda batch: {"data": batch},  # With "data" wrapper
    ]
    
    total_updated = 0
    
    for batch_idx, batch in enumerate(batches):
        print(f"Processing batch {batch_idx + 1}/{len(batches)} ({len(batch)} items)")
        
        success = False
        for format_idx, format_func in enumerate(body_formats):
            body = format_func(batch)
            print(f"Trying body format {format_idx + 1}/3 for batch {batch_idx + 1}")
            
            # Try with Bearer token first
            for attempt in range(3):  # Max 3 retries per format
                try:
                    print(f"Making PUT request to: {url} (attempt {attempt + 1}/3)")
                    resp = requests.put(url, headers=headers1, json=body, timeout=30)
                    print(f"Response status: {resp.status_code}")
                    
                    if resp.status_code == 204:
                        print(f"Successfully updated batch {batch_idx + 1} with Bearer token")
                        total_updated += len(batch)
                        success = True
                        break
                    elif resp.status_code == 429:
                        # Rate limit exceeded - wait and retry
                        retry_after = 60  # Default wait time
                        retry_header = resp.headers.get('X-Ratelimit-Retry') or resp.headers.get('Retry-After')
                        if retry_header:
                            try:
                                retry_after = int(float(retry_header))
                            except (ValueError, TypeError):
                                pass
                        
                        wait_time = retry_after * (attempt + 1)  # Exponential backoff
                        print(f"Rate limit exceeded (429), waiting {wait_time} seconds before retry...")
                        time.sleep(wait_time)
                        continue
                    else:
                        print(f"Response text: {resp.text[:200]}")
                        print(f"Response headers: {dict(resp.headers)}")

                    # ???? WB ???????? 400 IncorrectRequestBody, ????????? PUT ????????? ???
                    # ????????????: ???????? ? payload ??? ???????? format_func.
                    if resp.status_code == 400:
                        break
                        
                        # Try without Bearer token if Bearer failed
                        if attempt == 0:  # Only try without Bearer on first attempt
                            print(f"Trying without Bearer token...")
                            resp2 = requests.put(url, headers=headers2, json=body, timeout=30)
                            print(f"Response status (no Bearer): {resp2.status_code}")
                            if resp2.status_code == 204:
                                print(f"Successfully updated batch {batch_idx + 1} without Bearer token")
                                total_updated += len(batch)
                                success = True
                                break
                            elif resp2.status_code == 429:
                                # Rate limit exceeded - wait and retry
                                retry_after = 60
                                retry_header = resp2.headers.get('X-Ratelimit-Retry') or resp2.headers.get('Retry-After')
                                if retry_header:
                                    try:
                                        retry_after = int(float(retry_header))
                                    except (ValueError, TypeError):
                                        pass
                                wait_time = retry_after * (attempt + 1)
                                print(f"Rate limit exceeded (429), waiting {wait_time} seconds before retry...")
                                time.sleep(wait_time)
                                continue
                            else:
                                print(f"Response text (no Bearer): {resp2.text[:200]}")
                                break
                    # ??? ???? ????????? HTTP-?????? ?????????? ??????? ??? ???????? format_func.
                    break
                        
                except requests.RequestException as e:
                    print(f"Request error on attempt {attempt + 1}: {e}")
                    if attempt < 2:  # Retry on network errors
                        time.sleep(2 * (attempt + 1))
                        continue
                    else:
                        raise
                except Exception as e:
                    print(f"Error with body format {format_idx + 1}, attempt {attempt + 1}: {e}")
                    if attempt < 2:
                        time.sleep(2 * (attempt + 1))
                        continue
                    else:
                        break
            
            if success:
                break
        
        if not success:
            raise requests.HTTPError(f"Failed to update batch {batch_idx + 1} after trying all formats")
        
        # Add delay between batches to avoid rate limiting
        if batch_idx < len(batches) - 1:
            delay = 2  # 2 seconds between batches
            print(f"Waiting {delay} seconds before next batch...")
            time.sleep(delay)
    
    print(f"Successfully updated {total_updated} stocks for warehouse {warehouse_id}")


def _fbs_stock_cache_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"fbs_stock_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "fbs_stock_anon.json")


def load_fbs_stock_cache() -> Dict[str, Any] | None:
    path = _fbs_stock_cache_path_for_user()
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def save_fbs_stock_cache(payload: Dict[str, Any]) -> None:
    path = _fbs_stock_cache_path_for_user()
    try:
        enriched = dict(payload)
        if current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


def _auto_update_settings_path_for_user() -> str:
    if current_user.is_authenticated:
        return os.path.join(CACHE_DIR, f"auto_update_settings_user_{current_user.id}.json")
    return os.path.join(CACHE_DIR, "auto_update_settings_anon.json")


# Notification system functions
def create_notification(user_id: int, title: str, message: str, notification_type: str, data: dict = None, created_at: datetime = None) -> Notification:
    """Create a new notification for a user"""
    # ?????????? ?????????? ????? ??? ??????? ?????????? ?????
    notification_time = created_at if created_at else datetime.now(MOSCOW_TZ)
    
    notification = Notification(
        user_id=user_id,
        title=title,
        message=message,
        notification_type=notification_type,
        data=json.dumps(data) if data else None,
        created_at=notification_time
    )
    db.session.add(notification)
    db.session.commit()
    return notification


def get_unread_notifications_count(user_id: int) -> int:
    """Get count of unread notifications for a user"""
    return Notification.query.filter_by(user_id=user_id, is_read=False).count()


def get_user_notifications(user_id: int, limit: int = 20) -> List[Notification]:
    """Get recent notifications for a user"""
    return Notification.query.filter_by(user_id=user_id)\
        .order_by(Notification.created_at.desc())\
        .limit(limit).all()


def mark_notification_as_read(notification_id: int, user_id: int) -> bool:
    """Mark a notification as read"""
    notification = Notification.query.filter_by(id=notification_id, user_id=user_id).first()
    if notification:
        notification.is_read = True
        db.session.commit()
        return True
    return False


def mark_all_notifications_as_read(user_id: int) -> int:
    """Mark all notifications as read for a user"""
    count = Notification.query.filter_by(user_id=user_id, is_read=False).update({'is_read': True})
    db.session.commit()
    return count


def cleanup_old_notifications(days: int = 30) -> int:
    """Clean up notifications older than specified days"""
    with app.app_context():
        cutoff_date = datetime.now(MOSCOW_TZ) - timedelta(days=days)
        count = Notification.query.filter(Notification.created_at < cutoff_date).delete()
        db.session.commit()
        return count


def _log_wb_notification_error(context: str, user_id: int, exc: BaseException) -> None:
    """
    ?????? WB ??? ??????? ????????????: 401/403/429 ? ??????? ????????,
    ?? ???????? ??? ?????? traceback.
    """
    if isinstance(exc, requests.HTTPError) and exc.response is not None:
        code = exc.response.status_code
        if code in (401, 403):
            print(
                f"{context}: user_id={user_id} ? WB ?????? {code} "
                f"(????? ??????????????, ??????? ??? ??? ??????? ? API ??????? FBS), ???????? ?????????"
            )
            return
        if code == 429:
            print(f"{context}: user_id={user_id} ? WB 429 (????? ????????), ???????? ?????????")
            return
    print(f"{context}: user_id={user_id}: {exc}")
    import traceback

    traceback.print_exc()


def check_fbs_new_orders_for_notifications():
    """Check for new FBS orders and create notifications for all active users"""
    with app.app_context():
        try:
            # Get all active users with WB tokens
            users = User.query.filter_by(is_active=True).filter(User.wb_token.isnot(None)).all()
            
            for user in users:
                try:
                    if not user_may_call_wb_api(user):
                        continue
                    # Get last check time and processed orders from cache
                    cache_path = os.path.join(CACHE_DIR, f"fbs_notifications_user_{user.id}.json")
                    last_check = None
                    processed_order_ids = set()
                    
                    if os.path.exists(cache_path):
                        with open(cache_path, 'r', encoding='utf-8') as f:
                            cache_data = json.load(f)
                            last_check_str = cache_data.get('last_check')
                            if last_check_str:
                                last_check = datetime.fromisoformat(last_check_str.replace('Z', '+00:00'))
                            # ????????? ?????? ??? ???????????? ???????
                            processed_order_ids = set(cache_data.get('processed_order_ids', []))
                    
                    # If no previous check, check last 5 minutes
                    if not last_check:
                        last_check = datetime.now(MOSCOW_TZ) - timedelta(minutes=5)
                    
                    # Fetch new orders (??????????? ?????????? ??? ??????????)
                    new_orders, _ = fetch_fbs_latest_orders(user.wb_token, want_count=100)
                    
                    # Filter orders created after last check and not yet processed
                    new_orders_since_check = []
                    for order in new_orders:
                        order_id = str(order.get('id', ''))
                        if not order_id or order_id == 'Unknown':
                            continue
                        
                        # ?????????? ??? ???????????? ??????
                        if order_id in processed_order_ids:
                            continue
                        
                        order_time = _parse_iso_datetime(str(order.get('createdAt', '')))
                        # ???? ????? ???????? ???? ? ??? ????? last_check, ??? ???? ??????? ??? (????? ?? ?????? ??????)
                        if order_time and order_time > last_check:
                            new_orders_since_check.append(order)
                        elif not order_time:
                            # ???? ??????? ???, ?? ?????? ??? ? ???????????? - ???? ????????? (?? ?????? ??????? ? API)
                            new_orders_since_check.append(order)
                    
                    # Create notifications for new orders (with duplicate check)
                    if new_orders_since_check:
                        for order in new_orders_since_check:
                            order_id = str(order.get('id', ''))
                            if not order_id or order_id == 'Unknown':
                                continue
                            
                            # ?????????, ?? ?????????? ?? ??? ??????????? ??? ????? ??????
                            # ???? ?? ???? ? ????????? data ????? JSON
                            existing_notifications = Notification.query.filter_by(
                                user_id=user.id,
                                notification_type="fbs_new_order"
                            ).all()
                            
                            order_exists = False
                            for notif in existing_notifications:
                                try:
                                    notif_data = json.loads(notif.data) if notif.data else {}
                                    if str(notif_data.get('order_id', '')) == order_id:
                                        order_exists = True
                                        break
                                except (json.JSONDecodeError, TypeError):
                                    continue
                            
                            if order_exists:
                                # ??????????? ??? ??????????, ??????????
                                processed_order_ids.add(order_id)
                                continue
                            
                            order_time = _parse_iso_datetime(str(order.get('createdAt', '')))
                            # ???????????? ????? ? ?????????? ????? ??? ??????????? ???????????
                            moscow_time = to_moscow(order_time) if order_time else None
                            time_str = moscow_time.strftime('%H:%M') if moscow_time else 'Unknown'
                            
                            create_notification(
                                user_id=user.id,
                                title="????? ????? FBS",
                                message=f"???????? ????? ????? #{order_id}",
                                notification_type="fbs_new_order",
                                data={
                                    'order_id': order_id,
                                    'order_data': order,
                                    'created_at': order.get('createdAt')
                                },
                                created_at=datetime.now(MOSCOW_TZ)
                            )
                            
                            # ????????? ????? ? ?????? ????????????
                            processed_order_ids.add(order_id)
                    
                    # Update last check time and processed orders list
                    current_time = datetime.now(MOSCOW_TZ)
                    # ???????????? ?????? ?????? ???????????? ??????? (?????? ????????? 1000)
                    processed_list = list(processed_order_ids)[-1000:]
                    
                    with open(cache_path, 'w', encoding='utf-8') as f:
                        json.dump({
                            'last_check': current_time.isoformat(),
                            'checked_orders_count': len(new_orders_since_check),
                            'processed_order_ids': processed_list
                        }, f, ensure_ascii=False)
                        
                except Exception as e:
                    print(f"Error checking FBS orders for user {user.id}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
                    
        except Exception as e:
            print(f"Error in FBS notifications check: {e}")
            import traceback
            traceback.print_exc()


def check_dbs_new_orders_for_notifications():
    """Check for new DBS orders and create notifications for all active users"""
    with app.app_context():
        try:
            # Get all active users with WB tokens
            users = User.query.filter_by(is_active=True).filter(User.wb_token.isnot(None)).all()
            
            for user in users:
                try:
                    if not user_may_call_wb_api(user):
                        continue
                    # Get last check time from cache
                    cache_path = os.path.join(CACHE_DIR, f"dbs_notifications_user_{user.id}.json")
                    last_check = None
                    if os.path.exists(cache_path):
                        with open(cache_path, 'r', encoding='utf-8') as f:
                            cache_data = json.load(f)
                            last_check_str = cache_data.get('last_check')
                            if last_check_str:
                                last_check = datetime.fromisoformat(last_check_str.replace('Z', '+00:00'))
                    
                    # If no previous check, check last 5 minutes
                    if not last_check:
                        last_check = datetime.now(MOSCOW_TZ) - timedelta(minutes=5)
                    
                    # Fetch new orders
                    new_orders = fetch_dbs_new_orders(user.wb_token)
                    
                    # Filter orders created after last check
                    new_orders_since_check = []
                    for order in new_orders:
                        order_time = _extract_created_at(order)
                        if order_time and order_time > last_check:
                            new_orders_since_check.append(order)
                    
                    # Create notifications for new orders
                    if new_orders_since_check:
                        for order in new_orders_since_check:
                            order_id = order.get('id') or order.get('orderId') or order.get('ID') or 'Unknown'
                            order_time = _extract_created_at(order)
                            # ???????????? ????? ? ?????????? ????? ??? ??????????? ???????????
                            moscow_time = to_moscow(order_time) if order_time else None
                            time_str = moscow_time.strftime('%H:%M') if moscow_time else 'Unknown'
                            
                            create_notification(
                                user_id=user.id,
                                title="????? ????? DBS",
                                message=f"???????? ????? ????? DBS #{order_id}",
                                notification_type="dbs_new_order",
                                data={
                                    'order_id': order_id,
                                    'order_data': order,
                                    'created_at': order.get('createdAt') or order.get('dateCreated') or order.get('date')
                                },
                                created_at=datetime.now(MOSCOW_TZ)
                            )
                    
                    # Update last check time
                    current_time = datetime.now(MOSCOW_TZ)
                    with open(cache_path, 'w', encoding='utf-8') as f:
                        json.dump({
                            'last_check': current_time.isoformat(),
                            'checked_orders_count': len(new_orders_since_check)
                        }, f, ensure_ascii=False)
                        
                except Exception as e:
                    _log_wb_notification_error("DBS ??????????? (??????)", user.id, e)
                    continue
                    
        except Exception as e:
            print(f"Error in DBS notifications check: {e}")


def check_version_updates():
    """Check for version updates and create notifications"""
    with app.app_context():
        try:
            version_file = "VERSION"
            version_cache_file = os.path.join(CACHE_DIR, "version_cache.json")
            
            # Read current version
            current_version = None
            if os.path.exists(version_file):
                try:
                    with open(version_file, 'r', encoding='utf-8') as f:
                        current_version = f.read().strip()
                except Exception as e:
                    return
            
            if not current_version:
                return
            
            # Read cached version
            cached_version = None
            if os.path.exists(version_cache_file):
                try:
                    with open(version_cache_file, 'r', encoding='utf-8') as f:
                        cache_data = json.load(f)
                        cached_version = cache_data.get('version')
                except Exception as e:
                    pass
            
            print(f"Version check: current={current_version}, cached={cached_version}")
        
            # If version changed, create notifications for all active users
            if cached_version and current_version != cached_version:
                print(f"Version changed from {cached_version} to {current_version}")
                active_users = User.query.filter_by(is_active=True).all()
                print(f"Found {len(active_users)} active users")
                
                for user in active_users:
                    # Check if user already has notification for this version
                    existing_notification = Notification.query.filter_by(
                        user_id=user.id,
                        notification_type="version_update",
                        data=json.dumps({"version": current_version})
                    ).first()
                    
                    if not existing_notification:
                        print(f"Creating version notification for user {user.id}")
                        # ??????? ??????????? ? ?????? ???????? ??????????? ????????? ??????
                        create_notification(
                            user_id=user.id,
                            title="?????????? ???????",
                            message=f"????? ?????????? {current_version}",
                            notification_type="version_update",
                            data={"version": current_version, "previous_version": cached_version},
                            created_at=datetime.now(MOSCOW_TZ)
                        )
                    else:
                        print(f"User {user.id} already has notification for version {current_version}")
            
            # Update version cache
            try:
                cache_data = {"version": current_version, "last_check": datetime.now(MOSCOW_TZ).isoformat()}
                with open(version_cache_file, 'w', encoding='utf-8') as f:
                    json.dump(cache_data, f, ensure_ascii=False, indent=2)
            except Exception as e:
                print(f"Error updating version cache: {e}")
                
        except Exception as e:
            print(f"Error in check_version_updates: {e}")


# Global variable to track monitoring state
_monitoring_started = False
_last_cache_refresh_hour = None
_last_fbs_check_at = None
_last_dbs_check_at = None
_last_version_check_at = None
_last_stocks_refresh_at = None

def start_notification_monitoring():
    """Start background monitoring for notifications"""
    global _monitoring_started, _last_cache_refresh_hour, _last_fbs_check_at, _last_dbs_check_at, _last_version_check_at, _last_stocks_refresh_at
    
    if _monitoring_started:
        return
    
    def monitor_loop():
        global _last_cache_refresh_hour, _last_fbs_check_at, _last_dbs_check_at, _last_version_check_at, _last_stocks_refresh_at
        while True:
            try:
                current_time = datetime.now()
                print(f"Running notification checks at {current_time.strftime('%H:%M:%S')}")

                # ???????? ??????? ????????, ????? ??? ?? ????????????? ? ??????? ????????? ????? ??????.
                # ?? ????? ??? ??????????? ?????? 30 ??????.
                if _last_fbs_check_at is None or (current_time - _last_fbs_check_at).total_seconds() >= 300:  # 5 ???
                    check_fbs_new_orders_for_notifications()
                    _last_fbs_check_at = current_time
                if _last_dbs_check_at is None or (current_time - _last_dbs_check_at).total_seconds() >= 300:  # 5 ???
                    check_dbs_new_orders_for_notifications()
                    _last_dbs_check_at = current_time
                if _last_version_check_at is None or (current_time - _last_version_check_at).total_seconds() >= 1800:  # 30 ???
                    check_version_updates()
                    _last_version_check_at = current_time
                
                # Clean up old notifications every hour
                if current_time.minute == 0:
                    cleanup_old_notifications()
                
                # Auto-refresh stocks: ?????? ??? ? 30 ????? (?? ????? ?????????? ???????),
                # ????? ?? ??????????? ????????? ??? ??? ?????? ??? ? 30 ??????.
                if _last_stocks_refresh_at is None or (current_time - _last_stocks_refresh_at).total_seconds() >= 1800:
                    print(f"Triggering auto stocks refresh at {current_time.strftime('%H:%M:%S')}")
                    try:
                        auto_refresh_stocks_for_all_users()
                    except Exception as e:
                        print(f"Error in auto stocks refresh: {e}")
                    _last_stocks_refresh_at = current_time
                
                # Auto-refresh supplies and orders cache every 2 hours (at 0:00, 2:00, 4:00, etc.)
                if current_time.hour % 2 == 0 and current_time.minute == 0 and _last_cache_refresh_hour != current_time.hour:
                    _last_cache_refresh_hour = current_time.hour
                    print(f"Triggering auto cache refresh (supplies & orders) at {current_time.strftime('%H:%M:%S')}")
                    try:
                        auto_refresh_supplies_cache_for_all_users()
                    except Exception as e:
                        print(f"Error in auto supplies cache refresh: {e}")
                    try:
                        auto_refresh_orders_cache_for_all_users()
                    except Exception as e:
                        print(f"Error in auto orders cache refresh: {e}")
                    
            except Exception as e:
                print(f"Error in monitoring loop: {e}")
            time.sleep(30)  # Check every 30 seconds for faster testing
    
    # Start monitoring in a separate thread
    # ????? ?? ????????? ??????? WB-???????? ????? ????? ?????? ???????? (????????, ? ?????? ??????? ??????),
    # ??????????? last-run "??? ??????". ????????? ??????? ????? ?????? ????? ?????????.
    init_now = datetime.now()
    _last_fbs_check_at = init_now
    _last_dbs_check_at = init_now
    _last_version_check_at = init_now
    _last_stocks_refresh_at = init_now

    monitor_thread = threading.Thread(target=monitor_loop, daemon=True)
    monitor_thread.start()
    _monitoring_started = True
    print("Notification monitoring started")


def auto_refresh_stocks_for_all_users():
    """????????????? ????????? ??????? ??? ???? ????????????? ? ????????"""
    try:
        # ??????? ???????? ?????????? ??? ?????? ? ????? ??????
        with app.app_context():
            # User ??? ????????? ? ???? ?????
            
            # ???????? ???? ????????????? ? ????????
            try:
                users_with_tokens = User.query.filter(User.wb_token.isnot(None), User.wb_token != '').all()
            except Exception as e:
                print(f"Error querying users: {e}")
                return
            
            if not users_with_tokens:
                print("No users with tokens found for auto stocks refresh")
                return
            
            print(f"Auto-refreshing stocks for {len(users_with_tokens)} users")
            
            for i, user in enumerate(users_with_tokens):
                try:
                    if not user_may_call_wb_api(user):
                        continue
                    # ?????????, ????? ?? ????????? ??? (???? ?? ???????)
                    cached = load_stocks_cache_for_user(user.id)
                    should_refresh = True
                    
                    if cached and cached.get("_user_id") == user.id:
                        updated_at = cached.get("updated_at")
                        if updated_at:
                            try:
                                cache_time = datetime.strptime(updated_at, "%d.%m.%Y %H:%M:%S")
                                # ???? ??????? ??????????? ????? 25 ????? ?????, ??????????
                                if (datetime.now() - cache_time).total_seconds() < 1500:  # 25 ?????
                                    should_refresh = False
                                    print(f"Skipping auto-refresh for user {user.id} - cache is fresh")
                            except Exception:
                                pass
                    
                    if should_refresh:
                        print(f"Auto-refreshing stocks for user {user.id}")
                        
                        # ????????? ???????? ????? ????????? ??? ????????? 429 ??????
                        if i > 0:
                            time.sleep(2)  # 2 ??????? ????? ?????????
                        
                        raw = fetch_stocks_resilient(user.wb_token)
                        items = normalize_stocks(raw)
                        now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
                        save_stocks_cache_for_user(user.id, {"items": items, "updated_at": now_str})
                        print(f"Auto-refresh completed for user {user.id}: {len(items)} items at {now_str}")
                    
                except requests.HTTPError as e:
                    if e.response and e.response.status_code == 401:
                        print(f"User {user.id}: Invalid token (401) - skipping")
                    elif e.response and e.response.status_code == 429:
                        print(f"User {user.id}: Rate limit exceeded (429) - will retry later")
                    else:
                        print(f"Error auto-refreshing stocks for user {user.id}: {e}")
                    continue
                except Exception as e:
                    print(f"Error auto-refreshing stocks for user {user.id}: {e}")
                    continue
            
            print("Auto stocks refresh cycle completed")
        
    except Exception as e:
        print(f"Error in auto_refresh_stocks_for_all_users: {e}")


def auto_refresh_supplies_cache_for_all_users():
    """????????????? ????????? ??? ???????? ??? ???? ????????????? ? ????????"""
    try:
        with app.app_context():
            users_with_tokens = User.query.filter(User.wb_token.isnot(None), User.wb_token != '').all()
            
            if not users_with_tokens:
                print("No users with tokens found for auto supplies cache refresh")
                return
            
            print(f"Auto-refreshing supplies cache for {len(users_with_tokens)} users")
            
            for i, user in enumerate(users_with_tokens):
                try:
                    if not user_may_call_wb_api(user):
                        continue
                    # ?????????, ????? ?? ????????? ??? (???? ?? ???????)
                    cached = load_fbw_supplies_detailed_cache(user.id)
                    should_refresh = True
                    
                    if cached:
                        last_updated = cached.get("last_updated")
                        if last_updated:
                            try:
                                last_update_dt = datetime.fromisoformat(last_updated)
                                # ???? ??? ?????????? ????? 1.5 ????? ?????, ??????????
                                if (datetime.now(MOSCOW_TZ) - last_update_dt).total_seconds() < 5400:  # 1.5 ????
                                    should_refresh = False
                                    print(f"Skipping auto-refresh supplies cache for user {user.id} - cache is fresh")
                            except Exception:
                                pass
                    
                    if should_refresh:
                        print(f"Auto-refreshing supplies cache for user {user.id}")
                        
                        # ????????? ???????? ????? ?????????
                        if i > 0:
                            time.sleep(2)
                        
                        has_cache = bool(cached)
                        cache_data = build_supplies_detailed_cache(
                            user.wb_token,
                            user.id,
                            batch_size=10,
                            pause_seconds=2.0,
                            force_full=not has_cache,
                            days_back=(180 if not has_cache else 10),
                        )
                        save_fbw_supplies_detailed_cache(cache_data, user.id)
                        print(f"Auto-refresh supplies cache completed for user {user.id}")
                    
                except Exception as e:
                    print(f"Error auto-refreshing supplies cache for user {user.id}: {e}")
                    continue
            
            print("Auto supplies cache refresh cycle completed")
        
    except Exception as e:
        print(f"Error in auto_refresh_supplies_cache_for_all_users: {e}")


def auto_refresh_orders_cache_for_all_users():
    """????????????? ????????? ??? ??????? ??? ???? ????????????? ? ????????"""
    try:
        with app.app_context():
            users_with_tokens = User.query.filter(User.wb_token.isnot(None), User.wb_token != '').all()
            
            if not users_with_tokens:
                print("No users with tokens found for auto orders cache refresh")
                return
            
            print(f"Auto-refreshing orders cache for {len(users_with_tokens)} users")
            
            for i, user in enumerate(users_with_tokens):
                try:
                    if not user_may_call_wb_api(user):
                        continue
                    # ?????????, ????? ?? ????????? ??? (???? ?? ???????)
                    cached_path = _orders_cache_meta_path_for_user(user.id)
                    should_refresh = True
                    
                    if os.path.isfile(cached_path):
                        try:
                            with open(cached_path, "r", encoding="utf-8") as f:
                                cached = json.load(f)
                                last_updated = cached.get("last_updated")
                                if last_updated:
                                    try:
                                        last_update_dt = datetime.fromisoformat(last_updated)
                                        # ???? ??? ?????????? ????? 1.5 ????? ?????, ??????????
                                        if (datetime.now(MOSCOW_TZ) - last_update_dt).total_seconds() < 5400:  # 1.5 ????
                                            should_refresh = False
                                            print(f"Skipping auto-refresh orders cache for user {user.id} - cache is fresh")
                                    except Exception:
                                        pass
                        except Exception:
                            pass
                    
                    if should_refresh:
                        print(f"Auto-refreshing orders cache for user {user.id}")
                        
                        # ????????? ???????? ????? ?????????
                        if i > 0:
                            time.sleep(2)
                        
                        meta = build_orders_warm_cache(user.wb_token, user.id)
                        save_orders_cache_meta(meta, user.id)
                        print(f"Auto-refresh orders cache completed for user {user.id}")
                    
                except Exception as e:
                    print(f"Error auto-refreshing orders cache for user {user.id}: {e}")
                    continue
            
            print("Auto orders cache refresh cycle completed")
        
    except Exception as e:
        print(f"Error in auto_refresh_orders_cache_for_all_users: {e}")


def load_auto_update_settings() -> Dict[str, Any]:
    path = _auto_update_settings_path_for_user()
    if not os.path.isfile(path):
        return {
            "global": {
                "url": "",
                "interval": 60,
                "enabled": False,
                "lastCheck": None,
                "history": []
            },
            "warehouses": {}
        }
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            # Migrate old format to new format
            if "url" in data and "warehouses" not in data:
                return {
                    "global": {
                        "url": data.get("url", ""),
                        "interval": data.get("interval", 60),
                        "enabled": data.get("enabled", False),
                        "lastCheck": data.get("lastCheck"),
                        "history": data.get("history", [])
                    },
                    "warehouses": {}
                }
            return data
    except Exception:
        return {
            "global": {
                "url": "",
                "interval": 60,
                "enabled": False,
                "lastCheck": None,
                "history": []
            },
            "warehouses": {}
        }


def save_auto_update_settings(settings: Dict[str, Any], user_id: int = None) -> None:
    if user_id is None:
        path = _auto_update_settings_path_for_user()
    else:
        path = os.path.join(CACHE_DIR, f"auto_update_settings_user_{user_id}.json")
    
    try:
        enriched = dict(settings)
        if user_id is not None:
            enriched["_user_id"] = user_id
        elif hasattr(current_user, 'is_authenticated') and current_user.is_authenticated:
            enriched["_user_id"] = current_user.id
        with open(path, "w", encoding="utf-8") as f:
            json.dump(enriched, f, ensure_ascii=False)
    except Exception:
        pass


def test_remote_file(url: str) -> Dict[str, Any]:
    """Test connection to remote file and return file info"""
    try:
        response = requests.head(url, timeout=30)
        response.raise_for_status()
        
        size = int(response.headers.get('content-length', 0))
        last_modified = response.headers.get('last-modified', '')
        
        return {
            "size": size,
            "lastModified": last_modified,
            "accessible": True
        }
    except Exception as e:
        return {
            "error": str(e),
            "accessible": False
        }
# Remote stock XLS: header aliases (Unicode escapes — safe in any source encoding)
_REMOTE_STOCK_BARCODE_HEADERS = frozenset({
    "barcode",
    "\u0431\u0430\u0440\u043a\u043e\u0434",
})
_REMOTE_STOCK_QUANTITY_HEADERS = frozenset({
    "quantity",
    "qty",
    "\u043a\u043e\u043b-\u0432\u043e",
    "\u043a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e",
})


def download_and_process_remote_file(url: str, user_id: int) -> Dict[str, Any]:
    """Download remote file and process it for stock updates"""
    try:
        # Download file
        response = requests.get(url, timeout=30)
        response.raise_for_status()
        
        # Parse Excel file
        import xlrd
        workbook = xlrd.open_workbook(file_contents=response.content)
        sheet = workbook.sheet_by_index(0)
        
        # Find column indices
        header_row = sheet.row_values(0)
        barcode_col = -1
        quantity_col = -1
        
        for i, cell in enumerate(header_row):
            cell_value = str(cell).strip().lower()
            if cell_value in _REMOTE_STOCK_BARCODE_HEADERS:
                barcode_col = i
            elif cell_value in _REMOTE_STOCK_QUANTITY_HEADERS:
                quantity_col = i
        
        if barcode_col == -1 or quantity_col == -1:
            return {
                "success": False,
                "error": (
                    "Required columns not found in Excel: need headers "
                    "'barcode' / '\u0431\u0430\u0440\u043a\u043e\u0434' and "
                    "'quantity' / '\u043a\u043e\u043b-\u0432\u043e' / '\u043a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e'"
                ),
            }
        
        # Parse data
        file_data = {}
        for row_idx in range(1, sheet.nrows):
            row = sheet.row_values(row_idx)
            if len(row) > max(barcode_col, quantity_col):
                barcode = str(row[barcode_col]).strip()
                # Remove .0 from barcode if it exists (Excel sometimes adds .0 to numbers)
                if barcode.endswith('.0'):
                    barcode = barcode[:-2]
                
                try:
                    quantity = int(float(row[quantity_col]))
                    if barcode:
                        file_data[barcode] = quantity
                except (ValueError, TypeError):
                    continue
        
        return {"success": True, "data": file_data}
        
    except Exception as e:
        return {"success": False, "error": str(e)}
def get_reserved_quantities_from_fbs_tasks(user_id: int) -> Dict[str, int]:
    """Get reserved quantities from FBS tasks for each barcode"""
    print(f"=== GET_RESERVED_QUANTITIES_FROM_FBS_TASKS CALLED ===")
    print(f"User ID: {user_id}")
    
    try:
        # Load FBS tasks cache using user_id instead of current_user
        cached_tasks = load_fbs_tasks_cache_by_user_id(user_id) or {}
        tasks_rows = cached_tasks.get("rows") or []
        
        print(f"Found {len(tasks_rows)} FBS tasks")
        
        # Debug: print first few tasks to see structure
        if tasks_rows:
            print(f"First task structure: {tasks_rows[0]}")
            print(f"All task keys: {list(tasks_rows[0].keys()) if tasks_rows else 'No tasks'}")
        
        # Aggregate reserved quantities by barcode
        reserved_quantities = {}
        
        for i, task in enumerate(tasks_rows):
            print(f"Task {i}: {task}")
            
            barcode = task.get("barcode")
            if not barcode:
                print(f"Task {i}: No barcode found, skipping")
                continue
                
            print(f"Task {i}: Found barcode {barcode}")
                
            # Get quantity from task (usually 1 per task, but could be more)
            quantity = 1  # Default quantity per task
            _qty_raw = task.get("quantity") or task.get("\u043a\u043e\u043b\u0438\u0447\u0435\u0441\u0442\u0432\u043e")
            if _qty_raw is not None:
                try:
                    quantity = int(_qty_raw)
                except (ValueError, TypeError):
                    quantity = 1
            
            # Add to reserved quantities
            if barcode in reserved_quantities:
                reserved_quantities[barcode] += quantity
            else:
                reserved_quantities[barcode] = quantity
            
            print(f"Task {i}: Added {quantity} to reserved for barcode {barcode}")
        
        print(f"Reserved quantities: {reserved_quantities}")
        return reserved_quantities
        
    except Exception as e:
        print(f"Error getting reserved quantities from FBS tasks: {e}")
        import traceback
        traceback.print_exc()
        return {}


def adjust_stock_quantities_for_reserved(file_data: Dict[str, int], user_id: int) -> Dict[str, int]:
    """Adjust stock quantities by subtracting reserved quantities from FBS tasks"""
    print(f"=== ADJUST_STOCK_QUANTITIES_FOR_RESERVED CALLED ===")
    print(f"Original file data: {file_data}")
    print(f"User ID: {user_id}")
    
    # Get reserved quantities from FBS tasks
    reserved_quantities = get_reserved_quantities_from_fbs_tasks(user_id)
    print(f"Reserved quantities returned: {reserved_quantities}")
    
    # Check if we have any reserved quantities
    if not reserved_quantities:
        print("WARNING: No reserved quantities found from FBS tasks!")
        return file_data
    
    # Adjust quantities
    adjusted_data = {}
    for barcode, quantity in file_data.items():
        reserved = reserved_quantities.get(barcode, 0)
        adjusted_quantity = max(0, quantity - reserved)  # Don't go below 0
        
        if reserved > 0:
            print(f"Barcode {barcode}: original={quantity}, reserved={reserved}, adjusted={adjusted_quantity}")
        else:
            print(f"Barcode {barcode}: original={quantity}, no reservations")
        
        adjusted_data[barcode] = adjusted_quantity
    
    print(f"Adjusted data: {adjusted_data}")
    return adjusted_data


def update_stocks_from_remote_data(file_data: Dict[str, int], user_id: int, enabled_warehouse_ids: list = None) -> int:
    """Update stocks in the cache based on remote file data"""
    print(f"=== UPDATE_STOCKS_FROM_REMOTE_DATA CALLED ===")
    print(f"File data: {file_data}")
    print(f"User ID: {user_id}")
    print(f"Enabled warehouse IDs: {enabled_warehouse_ids}")
    
    # Adjust stock quantities by subtracting reserved quantities from FBS tasks
    adjusted_file_data = adjust_stock_quantities_for_reserved(file_data, user_id)
    print(f"Using adjusted file data: {adjusted_file_data}")
    
    try:
        # Get user token from database
        from app import app, User
        
        with app.app_context():
            print("Getting user from database...")
            user = User.query.get(user_id)
            print(f"User found: {user}")
            
            if not user or not user.wb_token:
                print(f"No token found for user {user_id}")
                return 0
            if not user_may_call_wb_api(user):
                print(f"WB API skipped for user {user_id}: ???? ???????? ?????? ?? ? ?????????? ??????? ??? ?????? ?????????")
                return 0
            
            token = user.wb_token
            print(f"Token found: {token[:10]}...")
            
            # Get warehouses
            print("Fetching warehouses...")
            warehouses = fetch_fbs_warehouses(token)
            print(f"Warehouses found: {warehouses}")
            
            if not warehouses:
                print("No warehouses found")
                return 0
            
            total_updated = 0
            
            # Update stocks for each warehouse
            for i, warehouse in enumerate(warehouses):
                warehouse_id = warehouse.get("id") or warehouse.get("warehouseId") or warehouse.get("warehouseID")
                
                if not warehouse_id:
                    print(f"No warehouse ID found for warehouse: {warehouse}")
                    continue
                
                # Skip warehouse if not in enabled list
                if enabled_warehouse_ids and warehouse_id not in enabled_warehouse_ids:
                    print(f"Warehouse {warehouse_id} not enabled for auto-update, skipping")
                    continue
                
                print(f"Processing warehouse {i+1}/{len(warehouses)}: {warehouse}")
                print(f"Warehouse ID: {warehouse_id}")
                
                # Get current stocks for this warehouse to check which SKUs are valid
                try:
                    print(f"Getting current stocks for warehouse {warehouse_id}...")
                    current_stocks = fetch_fbs_stocks_by_warehouse(token, int(warehouse_id), list(adjusted_file_data.keys()))
                    print(f"Current stocks: {current_stocks}")
                    
                    # Filter file data to only include SKUs that exist in this warehouse
                    valid_skus = {stock.get('sku') for stock in current_stocks if stock.get('sku')}
                    print(f"Valid SKUs for warehouse {warehouse_id}: {valid_skus}")
                    
                    # Prepare stock updates only for valid SKUs
                    stock_updates = []
                    for barcode, quantity in adjusted_file_data.items():
                        if barcode in valid_skus:
                            # Filter out negative quantities as WB API doesn't accept them
                            if quantity >= 0:
                                stock_updates.append({"sku": barcode, "amount": quantity})
                            else:
                                print(f"SKU {barcode} has negative quantity {quantity}, setting to 0")
                                stock_updates.append({"sku": barcode, "amount": 0})
                        else:
                            print(f"SKU {barcode} not found in warehouse {warehouse_id}, skipping")
                    
                    print(f"Prepared {len(stock_updates)} stock updates for warehouse {warehouse_id}")
                    
                    if not stock_updates:
                        print(f"No valid stock updates for warehouse {warehouse_id}")
                        continue
                    
                    # Update stocks via WB API for this warehouse
                    print(f"Updating {len(stock_updates)} stocks for warehouse {warehouse_id}")
                    update_fbs_stocks_by_warehouse(token, int(warehouse_id), stock_updates)
                    total_updated += len(stock_updates)
                    print(f"Successfully updated {len(stock_updates)} stocks for warehouse {warehouse_id}")
                    
                    # Add delay between requests to avoid rate limiting
                    if i < len(warehouses) - 1:  # Don't delay after the last warehouse
                        print("Waiting 2 seconds before next warehouse...")
                        time.sleep(2)
                        
                except Exception as e:
                    print(f"Error updating stocks for warehouse {warehouse_id}: {e}")
                    import traceback
                    traceback.print_exc()
                    continue
            
            print(f"Total updated {total_updated} stocks for user {user_id}")
            return total_updated
        
    except Exception as e:
        print(f"Error updating stocks from remote data: {e}")
        import traceback
        traceback.print_exc()
        return 0


def auto_update_worker():
    """Background worker for automatic stock updates"""
    while True:
        try:
            # Get all user settings
            settings_files = [f for f in os.listdir(CACHE_DIR) if f.startswith('auto_update_settings_user_')]
            
            for settings_file in settings_files:
                try:
                    user_id = int(settings_file.replace('auto_update_settings_user_', '').replace('.json', ''))
                    settings_path = os.path.join(CACHE_DIR, settings_file)
                    
                    with open(settings_path, 'r', encoding='utf-8') as f:
                        settings = json.load(f)
                    
                    global_settings = settings.get('global', settings)  # Support old format
                    if not global_settings.get('enabled'):
                        continue
                    
                    # Check if it's time to check
                    last_check = global_settings.get('lastCheck')
                    interval_minutes = global_settings.get('interval', 60)
                    
                    if last_check:
                        last_check_time = datetime.fromisoformat(last_check)
                        if datetime.now() - last_check_time < timedelta(minutes=interval_minutes):
                            continue
                    
                    # Get enabled warehouses with their URLs
                    enabled_warehouses = []
                    for warehouse_id, warehouse_settings in settings.get('warehouses', {}).items():
                        if warehouse_settings.get('enabled') and warehouse_settings.get('url'):
                            enabled_warehouses.append({
                                'warehouseId': warehouse_id,
                                'url': warehouse_settings['url']
                            })
                    
                    if not enabled_warehouses:
                        print(f"No enabled warehouses with URLs for user {user_id}")
                        continue
                    
                    print(f"Auto update for user {user_id}, processing {len(enabled_warehouses)} warehouses")
                    
                    total_processed = 0
                    total_updated = 0
                    all_success = True
                    
                    # Process each warehouse individually
                    for warehouse in enabled_warehouses:
                        warehouse_id = warehouse['warehouseId']
                        url = warehouse['url']
                        
                        print(f"Processing warehouse {warehouse_id} with URL: {url}")
                        
                        # Test file
                        file_info = test_remote_file(url)
                        if not file_info.get('accessible'):
                            print(f"File not accessible for warehouse {warehouse_id}: {file_info.get('error', 'Unknown error')}")
                            all_success = False
                            continue
                        
                        current_size = file_info.get('size', 0)
                        current_modified = file_info.get('lastModified', '')
                        last_size = settings.get('warehouses', {}).get(warehouse_id, {}).get('lastFileSize', 0)
                        last_modified = settings.get('warehouses', {}).get(warehouse_id, {}).get('lastFileModified', '')
                        
                        print(f"File check for warehouse {warehouse_id}:")
                        print(f"  Current: size={current_size}, modified={current_modified}")
                        print(f"  Last: size={last_size}, modified={last_modified}")
                        
                        # Check if file has changed
                        if current_size == last_size and current_modified == last_modified:
                            print(f"File unchanged for warehouse {warehouse_id}, skipping update")
                            continue
                        
                        print(f"File changed for warehouse {warehouse_id}, processing...")
                        
                        # Download and process file
                        result = download_and_process_remote_file(url, user_id)
                        
                        if result['success']:
                            print(f"File data for warehouse {warehouse_id}: {len(result['data'])} items")
                            
                            # Update stocks for this specific warehouse
                            updated_count = update_stocks_from_remote_data(result['data'], user_id, [int(warehouse_id)])
                            print(f"Updated count for warehouse {warehouse_id}: {updated_count}")
                            
                            total_processed += len(result['data'])
                            total_updated += updated_count
                            
                            # Update warehouse-specific file info
                            if 'warehouses' not in settings:
                                settings['warehouses'] = {}
                            if warehouse_id not in settings['warehouses']:
                                settings['warehouses'][warehouse_id] = {}
                            settings['warehouses'][warehouse_id]['lastFileSize'] = current_size
                            settings['warehouses'][warehouse_id]['lastFileModified'] = current_modified
                        else:
                            print(f"File processing failed for warehouse {warehouse_id}: {result['error']}")
                            all_success = False
                    
                    # Update settings
                    global_settings['lastCheck'] = datetime.now().isoformat()
                    global_settings['history'].insert(0, {
                        "timestamp": datetime.now().isoformat(),
                        "success": all_success,
                        "message": (
                            f"Auto-update: processed {total_processed} file rows, "
                            f"updated {total_updated} stock rows (remote FBS file)"
                        ),
                    })
                    global_settings['history'] = global_settings['history'][:50]  # Keep last 50 entries
                    save_auto_update_settings(settings, user_id)
                    
                except Exception as e:
                    print(f"Error processing auto-update for user {user_id}: {e}")
                    continue
            
            # Sleep for 1 minute before next check
            time.sleep(60)
            
        except Exception as e:
            print(f"Error in auto-update worker: {e}")
            time.sleep(60)


# Start auto-update worker in background
auto_update_thread = threading.Thread(target=auto_update_worker, daemon=True)
auto_update_thread.start()

# Context processor to add organization info to all templates
# Organization (seller) info is fetched from WB API.
# Doing it synchronously on every page render can block the UI if WB is slow.
_seller_info_cache: dict[int, dict[str, Any]] = {}
_seller_info_refresh_users: set[int] = set()
_seller_info_cache_lock = threading.Lock()
_SELLER_INFO_TTL_SECONDS = 6 * 60 * 60  # 6 hours
_seller_info_rate_limited_until: dict[int, float] = {}
_SELLER_INFO_RATE_LIMIT_COOLDOWN_SECONDS = 15 * 60  # 15 minutes
_seller_info_last_attempt_at: dict[int, float] = {}
_SELLER_INFO_MIN_INTERVAL_SECONDS = 60  # WB doc: 1 request per minute per seller


def _seller_info_cache_path(user_id: int) -> str:
    return os.path.join(CACHE_DIR, f"seller_info_user_{user_id}.json")


def _load_seller_info_cache(user_id: int) -> dict[str, Any] | None:
    cached = _seller_info_cache.get(user_id)
    if cached:
        return cached
    path = _seller_info_cache_path(user_id)
    if not os.path.isfile(path):
        return None
    try:
        with open(path, "r", encoding="utf-8") as f:
            payload = json.load(f)
        if isinstance(payload, dict):
            _seller_info_cache[user_id] = payload
            return payload
    except Exception:
        return None
    return None


def _save_seller_info_cache(user_id: int, payload: dict[str, Any]) -> None:
    _seller_info_cache[user_id] = payload
    path = _seller_info_cache_path(user_id)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(payload, f, ensure_ascii=False)
    except Exception:
        return


def _is_seller_info_cache_fresh(payload: dict[str, Any] | None) -> bool:
    if not payload:
        return False
    try:
        return (time.time() - float(payload.get("cached_at", 0) or 0)) < _SELLER_INFO_TTL_SECONDS
    except Exception:
        return False


def _is_placeholder_org_name(name: str | None) -> bool:
    val = (name or "").strip()
    if not val:
        return True
    lowered = val.lower()
    if lowered in ("\u043f\u0440\u043e\u0444\u0438\u043b\u044c", "profile", "???????"):
        return True
    # Broken-encoding fallback from older deployments.
    if "?" in val:
        return True
    return False


def _get_light_supplies_cache_info(user_id: int) -> dict[str, Any] | None:
    """
    Lightweight info for profile page without loading the entire large supplies cache JSON.
    """
    path = _fbw_supplies_detailed_cache_path_for_user(user_id)
    if not os.path.isfile(path):
        return None
    try:
        modified_ts = os.path.getmtime(path)
        modified_dt = datetime.fromtimestamp(modified_ts, tz=MOSCOW_TZ)
        return {
            "last_updated": modified_dt.isoformat(),
            "total_supplies": 0,
            "is_fresh": (datetime.now(MOSCOW_TZ) - modified_dt).total_seconds() < 24 * 3600,
            "cache_period_from": None,
            "cache_period_to": None,
        }
    except Exception:
        return None


def _start_seller_info_refresh_bg(user_id: int, token: str) -> None:
    if not token:
        return

    with _seller_info_cache_lock:
        last_attempt = float(_seller_info_last_attempt_at.get(user_id, 0) or 0)
        if time.time() - last_attempt < _SELLER_INFO_MIN_INTERVAL_SECONDS:
            return
        limited_until = float(_seller_info_rate_limited_until.get(user_id, 0) or 0)
        if time.time() < limited_until:
            return
        if user_id in _seller_info_refresh_users:
            return
        _seller_info_last_attempt_at[user_id] = time.time()
        _seller_info_refresh_users.add(user_id)

    def _worker() -> None:
        try:
            seller_info = fetch_seller_info(token)
            if not seller_info:
                return
            organization_name = (
                seller_info.get("name")
                or seller_info.get("companyName")
                or seller_info.get("supplierName")
                or "???????????"
            )
            with _seller_info_cache_lock:
                payload = {
                    "organization_name": organization_name,
                    "seller_info": seller_info,
                    "cached_at": time.time(),
                }
                _save_seller_info_cache(user_id, payload)
                _seller_info_rate_limited_until.pop(user_id, None)
        except Exception:
            # Avoid frequent retries when WB limits seller requests.
            _seller_info_rate_limited_until[user_id] = time.time() + _SELLER_INFO_RATE_LIMIT_COOLDOWN_SECONDS
            # Never block page rendering because of seller info issues
            return
        finally:
            with _seller_info_cache_lock:
                _seller_info_refresh_users.discard(user_id)

    threading.Thread(target=_worker, daemon=True).start()


@app.context_processor
def inject_organization_info():
    """Add organization information to all templates for navbar"""
    user_display_name = None
    if current_user.is_authenticated:
        user_display_name = current_user.display_name
        manual_org = (getattr(current_user, "org_display_name", None) or "").strip()
        if manual_org and not _is_placeholder_org_name(manual_org):
            return {"organization_name": manual_org, "user_display_name": user_display_name}

    wb_tok = effective_wb_api_token(current_user)
    if current_user.is_authenticated:
        with _seller_info_cache_lock:
            cached = _load_seller_info_cache(current_user.id)
            limited_until = float(_seller_info_rate_limited_until.get(current_user.id, 0) or 0)
        if cached and cached.get("organization_name") and not _is_placeholder_org_name(cached.get("organization_name")):
            return {
                "organization_name": cached.get("organization_name") or "\u041f\u0440\u043e\u0444\u0438\u043b\u044c",
                "user_display_name": user_display_name,
            }
        if time.time() < limited_until:
            return {
                "organization_name": "\u041e\u0440\u0433\u0430\u043d\u0438\u0437\u0430\u0446\u0438\u044f (WB API limit)",
                "user_display_name": user_display_name,
            }
        # Organization title is missing/placeholder: run one background refresh attempt.
        if wb_tok:
            _start_seller_info_refresh_bg(current_user.id, wb_tok)

    return {"organization_name": "\u041f\u0440\u043e\u0444\u0438\u043b\u044c", "user_display_name": user_display_name}

@app.route("/", methods=["GET", "POST"]) 
def root():
    if request.method == "POST":
        return redirect(url_for("index"), code=307)
    return redirect(url_for("index"))


@app.route("/orders", methods=["GET", "POST"]) 
@login_required
def index():
    error = None
    # Orders
    orders = []
    total_orders = 0
    total_active_orders = 0
    total_cancelled_orders = 0
    total_revenue = 0.0

    # Chart series
    daily_labels: List[str] = []
    daily_orders_counts: List[int] = []
    daily_orders_cancelled_counts: List[int] = []
    daily_orders_revenue: List[float] = []

    # Warehouses combined
    warehouse_summary_dual: List[Dict[str, Any]] = []

    # TOPs
    top_products: List[Tuple[str, int]] = []  # by orders (existing)
    top_products_orders_filtered: List[Tuple[str, int]] = []  # by orders and warehouse filter
    warehouses: List[str] = []
    selected_warehouse: str = request.args.get("warehouse", "")

    updated_at: str = ""
    date_from = request.form.get("date_from", "")
    date_to = request.form.get("date_to", "")
    include_orders = True
    if request.method == "POST":
        force_refresh = request.form.get("force_refresh") is not None
    date_from_fmt = format_dmy(date_from)
    date_to_fmt = format_dmy(date_to)

    # ?????: ????? ?? ?????, ????? ?? ??????? ????????????
    token = token_for_wb_request(current_user, request.form.get("token"))

    # ???? GET ? ??????? ???????? ????????? ?????????? ?? ????
    top_mode = "orders"
    cache_info = None
    if request.method == "GET":
        cached = load_last_results()
        # Use cache only if it belongs to this user (by user_id) and user has token
        if cached and current_user.is_authenticated and cached.get("_user_id") == current_user.id and (current_user.wb_token or ""):
            date_from = cached.get("date_from", date_from)
            date_to = cached.get("date_to", date_to)
            date_from_fmt = format_dmy(date_from)
            date_to_fmt = format_dmy(date_to)
            orders = cached.get("orders", [])
            total_orders = cached.get("total_orders", 0)
            total_active_orders = cached.get("total_active_orders", 0)
            total_cancelled_orders = cached.get("total_cancelled_orders", 0)
            total_revenue = cached.get("total_revenue", 0.0)
            top_products = cached.get("top_products", [])
            # charts
            daily_labels = cached.get("daily_labels", [])
            daily_orders_counts = cached.get("daily_orders_counts", [])
            daily_orders_cancelled_counts = cached.get("daily_orders_cancelled_counts", [])
            daily_orders_revenue = cached.get("daily_orders_revenue", [])
            warehouse_summary_dual = cached.get("warehouse_summary_dual", [])
            updated_at = cached.get("updated_at", "")
            # default mode when loading cache
            top_mode = cached.get("top_mode", "orders")
            # Add cache info for display
            cache_info = {"used_cache_days": 0, "fetched_days": 0}  # Will be calculated if needed
            
            # Fallback: ???? ? ???? ??? ????? ??????, ???????????? ?? ?? orders
            if total_active_orders == 0 and total_cancelled_orders == 0 and orders:
                total_active_orders = len([o for o in orders if not o.get("is_cancelled", False)])
                total_cancelled_orders = len([o for o in orders if o.get("is_cancelled", False)])
            
            # Fallback: ???? ? ???? ??? ?????? ?? ?????????? ??????? ??? ????????, ???????????? ??
            if not daily_orders_cancelled_counts and orders:
                _, _, o_cancelled_counts_map = aggregate_daily_counts_and_revenue(orders)
                daily_orders_cancelled_counts = [o_cancelled_counts_map.get(d, 0) for d in daily_labels]

    if request.method == "POST":
        if not token:
            error = "??????? ????? API"
        elif not date_from or not date_to:
            error = "???????? ????"
        else:
            try:
                df = parse_date(date_from)
                dt = parse_date(date_to)
                # normalize inverted range
                if df > dt:
                    date_from, date_to = date_to, date_from
            except ValueError:
                error = "???????? ?????? ???"

        if not error:
            try:
                if force_refresh:
                    # ?????????????? ?????????? - ????????? ??? ?????? ????? API, ????????? ???
                    raw_orders = fetch_orders_range(token, date_from, date_to)
                    orders = to_rows(raw_orders, date_from, date_to)
                    total_orders = len(orders)
                    total_active_orders = len([o for o in orders if not o.get("is_cancelled", False)])
                    total_cancelled_orders = len([o for o in orders if o.get("is_cancelled", False)])
                    total_revenue = round(sum(float(o.get("???? ?? ??????? ????????") or 0) for o in orders if not o.get("is_cancelled", False)), 2)
                    # ????????? ??? ?????????????
                    _update_period_cache_with_data(token, date_from, date_to, orders)
                else:
                    # ??????? ?????????? - ?????????? ??? ?? ????
                    orders, _meta = get_orders_with_period_cache(
                        token, date_from, date_to
                    )
                    total_orders = len(orders)
                    total_active_orders = len([o for o in orders if not o.get("is_cancelled", False)])
                    total_cancelled_orders = len([o for o in orders if o.get("is_cancelled", False)])
                    total_revenue = round(sum(float(o.get("???? ?? ??????? ????????") or 0) for o in orders if not o.get("is_cancelled", False)), 2)
                    cache_info = _meta

                # Aggregates for charts
                o_counts_map, o_rev_map, o_cancelled_counts_map = aggregate_daily_counts_and_revenue(orders)
                daily_labels = sorted(k for k in o_counts_map if k)
                daily_orders_counts = [o_counts_map.get(d, 0) for d in daily_labels]
                daily_orders_cancelled_counts = [o_cancelled_counts_map.get(d, 0) for d in daily_labels]
                daily_orders_revenue = [round(o_rev_map.get(d, 0.0), 2) for d in daily_labels]

                # Warehouses combined summary
                warehouse_summary_dual = aggregate_by_warehouse_orders_only(orders)

                # Top products (by orders)
                top_mode = "orders"
                top_products = aggregate_top_products(orders, limit=15)

                # ????????? ????? ? ??????? ???????????? ??? ???????
                if current_user.is_authenticated and token:
                    try:
                        # ?????????, ????????? ?? ?????
                        token_changed = current_user.wb_token != token
                        current_user.wb_token = token
                        db.session.commit()
                        
                        # ???? ????? ????????? ??? ??? ???????? ???????, ?????? ????? ???
                        if SUPPLIES_CACHE_AUTO and (token_changed or not is_supplies_cache_fresh()):
                            print(f"????? ????????? ??? ??? ???????, ????????? ?????????? ???? ???????? (auto={SUPPLIES_CACHE_AUTO})...")
                            # ????????? ? ???? (?? ????????? ???????? ??????)
                            import threading
                            def build_cache_background():
                                try:
                                    # ???? ???? ??? ? ?????? ????????????? ?? 6 ???, ????? ????????? 10 ????
                                    has_cache = bool(load_fbw_supplies_detailed_cache(current_user.id))
                                    cache_data = build_supplies_detailed_cache(
                                        token,
                                        current_user.id,
                                        batch_size=10,           # ?????? ?????
                                        pause_seconds=2.0,       # ??????? ?????
                                        force_full=not has_cache,
                                        days_back=(180 if not has_cache else 10),
                                    )
                                    save_fbw_supplies_detailed_cache(cache_data, current_user.id)
                                    print(f"??? ???????? ??????? ???????? ??? ???????????? {current_user.id}")
                                except Exception as e:
                                    print(f"?????? ?????????? ???? ????????: {e}")

                            thread = threading.Thread(target=build_cache_background)
                            thread.daemon = True
                            thread.start()

                        # ???? ????? ????????? ??? ??? ??????? ???????, ????????? ??? ???????
                        if token_changed or not is_orders_cache_fresh():
                            print("????????? ???????? ???? ??????? (6 ???????)...")
                            import threading
                            def warm_orders_cache_bg():
                                try:
                                    meta = build_orders_warm_cache(token)
                                    save_orders_cache_meta(meta)
                                    print("??? ??????? ????????")
                                except Exception as e:
                                    print(f"?????? ????????? ???? ???????: {e}")
                            t2 = threading.Thread(target=warm_orders_cache_bg)
                            t2.daemon = True
                            t2.start()
                            
                    except Exception:
                        db.session.rollback()
                updated_at = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
                date_from_fmt = format_dmy(date_from)
                date_to_fmt = format_dmy(date_to)
                save_last_results({
                    "date_from": date_from,
                    "date_to": date_to,
                    "orders": orders,
                    "total_orders": total_orders,
                    "total_active_orders": total_active_orders,
                    "total_cancelled_orders": total_cancelled_orders,
                    "total_revenue": total_revenue,
                    "daily_labels": daily_labels,
                    "daily_orders_counts": daily_orders_counts,
                    "daily_orders_cancelled_counts": daily_orders_cancelled_counts,
                    "daily_orders_revenue": daily_orders_revenue,
                    "warehouse_summary_dual": warehouse_summary_dual,
                    "top_products": top_products,
                    "top_mode": top_mode,
                    "updated_at": updated_at,
                })
            except requests.HTTPError as http_err:
                error = f"?????? API: {http_err.response.status_code}"
            except Exception as exc:  # noqa: BLE001
                error = f"??????: {exc}"

    # Build warehouses list and filtered ORDERS TOP from current orders
    warehouses = sorted({_order_row_warehouse_label(r) for r in orders})
    top_products_orders_filtered = aggregate_top_products_orders(
        orders, selected_warehouse or None, limit=50
    )

    return render_template(
        "index.html",
        error=error,
        token=token,
        date_from=date_from,
        date_to=date_to,
        date_from_fmt=date_from_fmt,
        date_to_fmt=date_to_fmt,
        # Orders table remains orders-only
        orders=orders,
        # KPIs
        total_orders=total_orders,
        total_active_orders=total_active_orders,
        total_cancelled_orders=total_cancelled_orders,
        total_revenue=total_revenue,
        updated_at=updated_at,
        # Charts
        daily_labels=daily_labels,
        daily_orders_counts=daily_orders_counts,
        daily_orders_cancelled_counts=daily_orders_cancelled_counts,
        daily_orders_revenue=daily_orders_revenue,
        # Warehouses dual
        warehouse_summary_dual=warehouse_summary_dual,
        # TOPs
        top_products=top_products,
        warehouses=warehouses,
        selected_warehouse=selected_warehouse,
        top_products_orders_filtered=top_products_orders_filtered,
        include_orders=include_orders,
        top_mode=top_mode,
        # Cache info
        cache_info=cache_info,
    )
def _filter_fbw_display_items(items: list[dict[str, Any]] | None) -> list[dict[str, Any]]:
    """
    ????????? ?????? ???????? FBW ??? ??????????? ? UI:
    - ???????? ???????? ?? ???????? "?? ?????????????" (?????????).
    """
    if not items:
        return []
    result: list[dict[str, Any]] = []
    for it in items:
        try:
            status = str((it or {}).get("status") or "")
            if "?? ?????????????" in status:
                continue
            result.append(it)
        except Exception:
            # ?? ?????? ?????? ?? ???????? ??????, ???? ???-?? ????? ?? ???
            result.append(it)
    return result


@app.route("/fbw", methods=["GET"]) 
@login_required
def fbw_supplies_page():
    """
    ???????? ???????? ???????? FBW ?????, ??? ???????? API ????????.
    ?????? ??????????? ? ???? ????? JavaScript ????? ???????? ????????.
    """
    token = effective_wb_api_token(current_user)
    error = None
    supplies: list[dict[str, Any]] = []
    generated_at = ""
    
    # ????????? ?????? ??? ??? ???????? ?????????? ??????????? (???? ????)
    # ?? ?????? ?????????? ???????? ? API - ??? ????? ???????? ? ???? ????? JavaScript
    cached = load_fbw_supplies_cache() or {}
    if cached and cached.get("_user_id") == (current_user.id if current_user.is_authenticated else None):
                supplies = cached.get("items", [])
                generated_at = cached.get("updated_at", "")
    
    if not token:
        error = "??????? API ????? ? ???????"

    # ???????? ????????? ("?? ?????????????") ?? ?????? ??? ???????????
    supplies = _filter_fbw_display_items(supplies)

    return render_template(
        "fbw_supplies.html",
        error=error,
        supplies=supplies,
        generated_at=generated_at,
    )


# ??????? /fbw/planning ????????? ? blueprints/fbw_planning.py
@app.route("/api/fbw/planning/products", methods=["GET"])
@login_required
def api_fbw_planning_products():
    """API ??? ????????? ?????? ??????? ? ????????? ??? ???????????? ????????"""
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        # ?????? ????????? ?????? ?????? ??? ????????????
        # ????????? ??? ??????
        raw_cards = fetch_all_cards(token, page_limit=100)
        products = normalize_cards_response({"cards": raw_cards})
        save_products_cache({"items": products, "_user_id": current_user.id})
        
        # ????????? ?????? ??????? ? ????????? - ?????????? ?? ?? ?????? ??? ? ?? ???????? /products
        products_with_barcodes = []
        for product in products:
            # ???????? ?????? ????? ??? ?? ??? ?? ???????? /products
            barcode = product.get("barcode")
            
            # ???? ???? ??????, ????????? ?????
            if barcode:
                products_with_barcodes.append({
                    "barcode": str(barcode),
                    "name": product.get("supplier_article") or "??? ????????",  # ?????????? supplier_article ??? ????????
                    "nm_id": product.get("nm_id"),
                    "supplier_article": product.get("supplier_article") or "??? ????????",
                    "photo": product.get("photo")
                })
        
        return jsonify({
            "success": True,
            "products": products_with_barcodes,
            "count": len(products_with_barcodes)
        })
        
    except requests.HTTPError as http_err:
        return jsonify({"error": "api_error", "message": f"?????? API: {http_err.response.status_code}"}), 502
    except Exception as exc:
        return jsonify({"error": "server_error", "message": f"??????: {str(exc)}"}), 500


@app.route("/api/fbw/warehouses", methods=["GET"])
@login_required
def api_fbw_warehouses():
    """API ??? ????????? ?????? ??????? FBW"""
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        # ?????? ? API Wildberries ??? ????????? ?????? ???????
        headers = {
            "Authorization": token,
            "Content-Type": "application/json"
        }
        
        response = requests.get(
            "https://supplies-api.wildberries.ru/api/v1/warehouses",
            headers=headers,
            timeout=30
        )
        response.raise_for_status()
        
        warehouses_data = response.json()
        
        # ???????????? ?????? ???????
        warehouses = []
        print(f"DEBUG: ???????? ?????? ???????: {len(warehouses_data)} ?????????")
        
        for warehouse in warehouses_data:
            warehouse_id = warehouse.get("ID")
            warehouse_name = warehouse.get("name")
            
            # ?????????, ??? ? ??? ???? ???????? ??????
            if warehouse_id and warehouse_name:
                warehouses.append({
                    "id": warehouse_id,
                    "name": warehouse_name,
                    "city": warehouse.get("city", ""),
                    "address": warehouse.get("address", ""),
                    "is_sorting_center": warehouse.get("isSortingCenter", False)
                })
            else:
                print(f"DEBUG: ???????? ????? - ID: {warehouse_id}, Name: {warehouse_name}")
        
        print(f"DEBUG: ?????????? ???????: {len(warehouses)}")
        
        return jsonify({
            "success": True,
            "warehouses": warehouses,
            "count": len(warehouses)
        })
        
    except requests.HTTPError as http_err:
        return jsonify({"error": "api_error", "message": f"?????? API: {http_err.response.status_code}"}), 502
    except Exception as exc:
        return jsonify({"error": "server_error", "message": f"??????: {str(exc)}"}), 500
@app.route("/api/fbw/planning/stocks", methods=["GET"])
@login_required
def api_fbw_planning_stocks():
    """API ??? ????????? ???????? ??????? ?? ?????????? ?????? ??? ???????????? ????????"""
    token = effective_wb_api_token(current_user)
    warehouse_id = request.args.get("warehouse_id")
    
    if not token:
        return jsonify({"error": "no_token"}), 401
    if not warehouse_id:
        return jsonify({"error": "no_warehouse", "message": "?? ?????? ID ??????"}), 400
    
    try:
        from datetime import datetime
        # ????????? ??? ???????? ????? ?????????????? ???????????
        cached = load_stocks_cache()
        should_refresh = True
        
        if cached and cached.get("_user_id") == current_user.id:
            # ?????????, ????? ????????? ??? ??????????? ???????
            updated_at = cached.get("updated_at")
            if updated_at:
                try:
                    # ?????? ????? ?????????? ?? ????
                    cache_time = datetime.strptime(updated_at, "%d.%m.%Y %H:%M:%S")
                    # ???? ??????? ??????????? ????? 10 ????? ?????, ?????????? ???
                    if (datetime.now() - cache_time).total_seconds() < 600:  # 10 ?????
                        should_refresh = False
                        print(f"=== ???????????? ????????: ?????????? ???????????? ??????? ===")
                        print(f"??? ????????: {updated_at}")
                except Exception as e:
                    print(f"?????? ???????? ??????? ????: {e}")
        
        if should_refresh:
            print("=== ???????????? ????????: ?????????????? ?????????? ???????? ===")
            print(f"????????????: {current_user.id}, ?????: {warehouse_id}")
            try:
                raw_stocks = fetch_stocks_resilient(token)
                stocks = normalize_stocks(raw_stocks)
                now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
                save_stocks_cache({"items": stocks, "_user_id": current_user.id, "updated_at": now_str})
                print(f"??????? ????????? ??? ???????????? ????????: {len(stocks)} ??????? ? {now_str}")
            except requests.HTTPError as e:
                if e.response and e.response.status_code == 429:
                    print("=== ???????????? ????????: ?????? 429, ?????????? ??? ===")
                    if cached and cached.get("_user_id") == current_user.id:
                        stocks = cached.get("items", [])
                        print(f"?????????? ???????????? ???????: {len(stocks)} ???????")
                    else:
                        return jsonify({"error": "rate_limit", "message": "???????? ????? ???????? ? API. ?????????? ?????."}), 429
                else:
                    raise
        else:
            stocks = cached.get("items", []) if cached else []
            print(f"?????????? ???????????? ???????: {len(stocks)} ???????")
            
        # ?????????, ??? ? ??? ???? ???????
        if not stocks:
            return jsonify({"error": "no_stocks", "message": "??? ?????? ?? ????????. ?????????? ?????."}), 500
        
        # ???????? ???????? ?????? ?? API ???????
        warehouse_name = None
        try:
            # ????????? ?????? ??????? ??? ????????? ????????
            warehouses_response = requests.get(
                "https://supplies-api.wildberries.ru/api/v1/warehouses",
                headers={"Authorization": token, "Content-Type": "application/json"},
                timeout=30
            )
            if warehouses_response.status_code == 200:
                warehouses_data = warehouses_response.json()
                for warehouse in warehouses_data:
                    if str(warehouse.get("ID")) == str(warehouse_id):
                        warehouse_name = warehouse.get("name")
                        break
        except Exception as e:
            print(f"?????? ????????? ???????? ??????: {e}")
        
        # Fallback ?? ID ???? ???????? ?? ???????
        if not warehouse_name:
            warehouse_name = f"????? {warehouse_id}"
        
        # ?????????? ?????????? - ?????????, ????? ?????? ???? ? ??????
        unique_warehouses = set()
        for stock in stocks:
            warehouse = stock.get("warehouse", "")
            if warehouse:
                unique_warehouses.add(warehouse)
        
        print(f"=== DEBUG: ????????? ?????? ? ?????? ???????? ===")
        for wh in sorted(unique_warehouses):
            print(f"  - '{wh}'")
        print(f"???? ?????: '{warehouse_name}' (ID: {warehouse_id})")
        
        # ????????? ??????? ?? ?????????? ?????? ??? ?? ???? ???????
        warehouse_stocks = {}
        for stock in stocks:
            stock_warehouse = stock.get("warehouse", "")
            barcode = stock.get("barcode")
            
            if barcode:
                # ???? ??????????? ??????? ?? ???? ???????
                if warehouse_id == "all":
                    # ????????? ??????? ?? ??????? ?? ???? ???????
                    if barcode in warehouse_stocks:
                        warehouse_stocks[barcode] += int(stock.get("qty", 0) or 0)
                    else:
                        warehouse_stocks[barcode] = int(stock.get("qty", 0) or 0)
                else:
                    # ?????????? ?? ???????? ?????? (? ?????? ??????????????)
                    if _warehouse_names_match(stock_warehouse, warehouse_name):
                        # ????????? ??????? ?? ??????? ?? ???? ??????
                        if barcode in warehouse_stocks:
                            warehouse_stocks[barcode] += int(stock.get("qty", 0) or 0)
                        else:
                            warehouse_stocks[barcode] = int(stock.get("qty", 0) or 0)
        
        # ???????? ????? ?????????? ?? ???? ??? ?????????? ??????? ?????
        now_str = cached.get("updated_at") if cached else datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        
        if warehouse_id == "all":
            print(f"??????? ???????? ?? ???? ???????: {len(warehouse_stocks)}")
            return jsonify({
                "success": True,
                "stocks": warehouse_stocks,
                "warehouse_id": "all",
                "warehouse_name": "??? ??????",
                "updated_at": now_str
            })
        else:
            print(f"??????? ???????? ??? ?????? '{warehouse_name}': {len(warehouse_stocks)}")
            return jsonify({
                "success": True,
                "stocks": warehouse_stocks,
                "warehouse_id": warehouse_id,
                "warehouse_name": warehouse_name,
                "updated_at": now_str
            })
        
    except requests.HTTPError as http_err:
        print(f"=== ?????? API ? api_fbw_planning_stocks ===")
        print(f"HTTP Error: {http_err.response.status_code}")
        print(f"Response: {http_err.response.text}")
        return jsonify({"error": "api_error", "message": f"?????? API: {http_err.response.status_code}"}), 502
    except Exception as exc:
        print(f"=== ?????? ? api_fbw_planning_stocks ===")
        print(f"Exception: {str(exc)}")
        print(f"Exception type: {type(exc)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "server_error", "message": f"??????: {str(exc)}"}), 500

def _normalize_warehouse_name(name: str) -> str:
    """
    ??????????? ???????? ??????, ?????? ????? ? ?????.
    ??? ???????? ???????????? ??????, ??????? ???? ?????????????
    (????????, "???????????? - ????????????? 12" -> "???????????? - ????????????? 14").
    
    ???????:
    - "???????????? - ????????????? 14" -> "???????????? - ?????????????"
    - "???????????? - ????????????? 12" -> "???????????? - ?????????????"
    - "???????????? - ????????????? 14" -> "???????????? - ?????????????"
    - "?????? - ???????? 1" -> "?????? - ????????"
    """
    if not name:
        return ""
    import re
    # ??????? ????? ? ????? (?????? + ?????)
    normalized = re.sub(r'\s+\d+$', '', name.strip())
    return normalized

def _warehouse_names_match(name1: str, name2: str) -> bool:
    """
    ?????????, ????????????? ?? ??? ???????? ?????? ???? ?????.
    ????????? ?????? ??????????, ????????? ?????????? ? ?????????????? ???????.
    ????? ????????? ?????? ???????? ????????? (?????????????/?????????????).
    """
    if not name1 or not name2:
        return False
    
    # ?????? ??????????
    if name1 == name2:
        return True
    
    # ????????? ?????????? (???? ???????? ??????)
    if name1 in name2 or name2 in name1:
        return True
    
    # ??????????????? ?????????? (??? ??????????????? ???????)
    norm1 = _normalize_warehouse_name(name1)
    norm2 = _normalize_warehouse_name(name2)
    if norm1 and norm2:
        # ?????? ?????????? ??????????????? ????????
        if norm1 == norm2:
            return True
        # ?????????, ??? ??????????????? ???????? ?????????? ?????? ??????????
        # (????????, "?????????????" ? "?????????????")
        # ????? ??????? ????? ??? ????????? 2-3 ????????
        if len(norm1) > 5 and len(norm2) > 5:
            base1 = norm1[:-3] if len(norm1) > 3 else norm1
            base2 = norm2[:-3] if len(norm2) > 3 else norm2
            if base1 == base2:
                return True
        # ????? ?????????, ??? ???? ??????????????? ???????? ???????? ??????
        if norm1 in norm2 or norm2 in norm1:
            return True
    
    return False


def _normalize_barcode_key(value: Any) -> str:
    """??????????? ?????? ? ?????????? ?????, ???????????? ? /tools/prices."""
    s = str(value or "").strip()
    if s.endswith(".0"):
        try:
            int(float(s))
            s = s[:-2]
        except Exception:
            pass
    return s


def _calc_tools_prices_margin_map(
    token: str,
    products: list[dict[str, Any]],
    user_id: int,
) -> dict[str, dict[str, float]]:
    """??????? ?????????????? ?? ??? ?? ????????, ??? ? /tools/prices."""
    if not token or not products or not user_id:
        return {}

    try:
        settings = load_user_margin_settings(user_id) or {}
        scheme = str(settings.get("scheme") or "FBW")
        tax_pct = float(settings.get("tax") or 0)
        storage_pct = float(settings.get("storage") or 0)
        receiving_pct = float(settings.get("receiving") or 0)
        acquiring_pct = float(settings.get("acquiring") or 0)
        warehouse_coef = float(settings.get("warehouse_coef") or 0)
        localization_index = float(settings.get("localization_index") or 1.0) or 1.0

        # ??????? ?? ???????? ?????? ??????? (?? ????? ??? ??????? ???????????? ? ??? ?????? ??? ??????? ??)
        barcode_keys_set: set[str] = set()
        for p in products:
            bc = _normalize_barcode_key(p.get("barcode"))
            if bc:
                barcode_keys_set.add(bc)

        # ?????????? ???? ???????????? ?????? ?? ?????? ???????? (??????? ? ????? IN ? SQLite)
        purchase_price_by_barcode: dict[str, float] = {}
        try:
            barcode_list = list(barcode_keys_set)
            _in_chunk = 450
            for i in range(0, len(barcode_list), _in_chunk):
                chunk = barcode_list[i : i + _in_chunk]
                if not chunk:
                    continue
                rows = (
                    PurchasePrice.query.filter_by(user_id=user_id)
                    .filter(PurchasePrice.barcode.in_(chunk))
                    .all()
                )
                for rec in rows:
                    bc = _normalize_barcode_key(rec.barcode)
                    if bc:
                        purchase_price_by_barcode[bc] = float(rec.price)
        except Exception:
            purchase_price_by_barcode = {}

        # ???? ? ???????? WB
        nm_ids: list[int] = []
        for p in products:
            try:
                nmv = p.get("nm_id") or p.get("nmID")
                if nmv is not None:
                    nm_ids.append(int(nmv))
            except Exception:
                continue
        nm_ids = list({x for x in nm_ids})

        # ??????? ??????? ????? ??????????? ???????? ?????????????? ?? ??.
        # ??? ??????????? ?????????? ???? ? /tools/prices.
        try:
            from models import MarginCalculation

            if nm_ids:
                nm_to_pct: dict[int, float] = {}
                nm_to_net: dict[int, float] = {}
                _nm_chunk = 450
                for i in range(0, len(nm_ids), _nm_chunk):
                    chunk = nm_ids[i : i + _nm_chunk]
                    part = (
                        MarginCalculation.query.filter_by(user_id=user_id)
                        .filter(MarginCalculation.nm_id.in_(chunk))
                        .all()
                    )
                    for rec in part:
                        try:
                            nm_key = int(rec.nm_id)
                        except Exception:
                            continue
                        if rec.profit_pct is not None:
                            try:
                                nm_to_pct[nm_key] = float(rec.profit_pct)
                            except Exception:
                                pass
                        if rec.profit_net is not None:
                            try:
                                nm_to_net[nm_key] = float(rec.profit_net)
                            except Exception:
                                pass

                if nm_to_pct:
                    result_from_db: dict[str, dict[str, float]] = {}
                    for p in products:
                        try:
                            nm_id = int(p.get("nm_id") or p.get("nmID"))
                        except Exception:
                            continue
                        if nm_id not in nm_to_pct:
                            continue
                        bc = _normalize_barcode_key(p.get("barcode"))
                        if not bc:
                            continue
                        result_from_db[bc] = {
                            "profit_pct": float(nm_to_pct[nm_id]),
                            "profit_net": float(nm_to_net.get(nm_id, 0.0)),
                        }
                    if result_from_db:
                        return result_from_db
        except Exception:
            pass

        prices_data = fetch_prices_data(token, nm_ids) if nm_ids else {}
        commission_data = fetch_commission_data(token) if products else {}

        def get_commission_percent(subject_id: Any) -> float | None:
            try:
                sid = int(subject_id)
            except Exception:
                return None
            item = commission_data.get(sid)
            if not item:
                return None
            if scheme == "FBS":
                return float(item.get("fbs_commission")) if item.get("fbs_commission") is not None else None
            if scheme == "C&C":
                return float(item.get("cc_commission")) if item.get("cc_commission") is not None else None
            if scheme == "DBS/DBW":
                return float(item.get("dbs_dbw_commission")) if item.get("dbs_dbw_commission") is not None else None
            if scheme == "EDBS":
                return float(item.get("edbs_commission")) if item.get("edbs_commission") is not None else None
            return float(item.get("fbw_commission")) if item.get("fbw_commission") is not None else None

        def calc_logistics_rub(volume: float) -> float:
            logistics_value = 0.0
            if volume > 0:
                if volume < 1:
                    if 0.001 <= volume <= 0.200:
                        logistics_value = 23
                    elif 0.201 <= volume <= 0.400:
                        logistics_value = 26
                    elif 0.401 <= volume <= 0.600:
                        logistics_value = 29
                    elif 0.601 <= volume <= 0.800:
                        logistics_value = 30
                    elif 0.801 <= volume <= 0.999:
                        logistics_value = 32
                else:
                    logistics_value = 46 + (max(0.0, volume - 1.0) * 14)
            if warehouse_coef > 0 and logistics_value > 0:
                logistics_value = logistics_value * warehouse_coef / 100.0
            if localization_index > 0 and logistics_value > 0:
                logistics_value = logistics_value * localization_index
            return float(logistics_value)

        result: dict[str, dict[str, float]] = {}
        for p in products:
            barcode = _normalize_barcode_key(p.get("barcode"))
            if not barcode:
                continue
            purchase_val = purchase_price_by_barcode.get(barcode)
            if purchase_val is None or purchase_val <= 0:
                continue

            try:
                nm_id = int(p.get("nm_id") or p.get("nmID"))
            except Exception:
                continue
            price_item = prices_data.get(nm_id) or {}
            discount_price = float(price_item.get("discount_price") or 0)
            if discount_price <= 0:
                continue

            commission_pct = get_commission_percent(p.get("subject_id"))
            commission_rub = discount_price * ((commission_pct or 0) / 100.0)
            tax_rub = discount_price * (tax_pct / 100.0)
            storage_rub = discount_price * (storage_pct / 100.0)
            receiving_rub = discount_price * (receiving_pct / 100.0)
            acquiring_rub = discount_price * (acquiring_pct / 100.0)

            dimensions = p.get("dimensions") or {}
            try:
                volume = float(dimensions.get("volume") or 0)
            except Exception:
                volume = 0.0
            logistics_rub = calc_logistics_rub(volume)

            total_expenses = commission_rub + tax_rub + logistics_rub + storage_rub + receiving_rub + acquiring_rub
            profit_net = discount_price - total_expenses - purchase_val
            profit_pct = (profit_net / purchase_val) * 100.0 if purchase_val > 0 else 0.0
            result[barcode] = {
                "profit_pct": round(float(profit_pct), 2),
                "profit_net": round(float(profit_net), 2),
            }

        return result
    except Exception as e:
        print(f"?????? ??????? ?????????????? ??? FBW planning: {e}")
        return {}

@app.route("/api/fbw/planning/data", methods=["GET"])
@login_required
def api_fbw_planning_data():
    """API ??? ????????? ???? ?????? ???????????? ???????? ????? ???????? (??????????? ?????? 429)"""
    token = effective_wb_api_token(current_user)
    warehouse_id = request.args.get("warehouse_id")
    date_from = request.args.get('date_from')
    date_to = request.args.get('date_to')
    
    if not token:
        return jsonify({"error": "no_token"}), 401
    if not warehouse_id:
        return jsonify({"error": "no_warehouse", "message": "?? ?????? ID ??????"}), 400
    if not date_from or not date_to:
        return jsonify({"error": "missing_dates", "message": "?? ??????? ????"}), 400
    
    try:
        from datetime import datetime
        
        # 1. ???????? ?????? ????????????
        print("=== ???????????? ????????: ????????? ?????? ???????????? ===")
        try:
            # ?????????? ?? ?? ??????? ??? ? ? ???????????? endpoint
            raw_cards = fetch_all_cards(token, page_limit=100)
            products = normalize_cards_response({"cards": raw_cards})
            print(f"????????? ???????: {len(products)}")
        except requests.HTTPError as e:
            if e.response and e.response.status_code == 429:
                print("=== ???????????? ????????: ?????? 429 ??? ???????? ???????, ?????????? ??? ===")
                # ???????? ????????? ?????? ?? ????
                try:
                    cached_products = load_products_cache()
                    if cached_products and cached_products.get("_user_id") == current_user.id:
                        products = cached_products.get("items", [])
                        print(f"?????????? ???????????? ??????: {len(products)}")
                    else:
                        products = []
                        print("??? ??????? ??????????, ?????????? ??? ???????")
                except Exception:
                    products = []
            else:
                print(f"?????? ???????? ???????: {e}")
                products = []
        except Exception as e:
            print(f"?????? ???????? ???????: {e}")
            products = []
        
        # 2. ???????? ??????? (???? ?????? ??? ???? ??????)
        print("=== ???????????? ????????: ????????? ??????? ===")
        cached = load_stocks_cache()
        should_refresh = True
        
        if cached and cached.get("_user_id") == current_user.id:
            updated_at = cached.get("updated_at")
            if updated_at:
                try:
                    cache_time = datetime.strptime(updated_at, "%d.%m.%Y %H:%M:%S")
                    if (datetime.now() - cache_time).total_seconds() < 600:  # 10 ?????
                        should_refresh = False
                        print(f"=== ???????????? ????????: ?????????? ???????????? ??????? ===")
                except Exception as e:
                    print(f"?????? ???????? ??????? ????: {e}")
        
        if should_refresh:
            print("=== ???????????? ????????: ?????????????? ?????????? ???????? ===")
            print(f"????????????: {current_user.id}, ?????: {warehouse_id}")
            try:
                raw_stocks = fetch_stocks_resilient(token)
                stocks = normalize_stocks(raw_stocks)
                now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
                save_stocks_cache({"items": stocks, "_user_id": current_user.id, "updated_at": now_str})
                print(f"??????? ????????? ??? ???????????? ????????: {len(stocks)} ??????? ? {now_str}")
            except requests.HTTPError as e:
                if e.response and e.response.status_code == 429:
                    print("=== ???????????? ????????: ?????? 429, ?????????? ??? ===")
                    if cached and cached.get("_user_id") == current_user.id:
                        stocks = cached.get("items", [])
                        print(f"?????????? ???????????? ???????: {len(stocks)} ???????")
                    else:
                        return jsonify({"error": "rate_limit", "message": "???????? ????? ???????? ? API. ?????????? ?????."}), 429
                else:
                    raise
        else:
            stocks = cached.get("items", []) if cached else []
            print(f"?????????? ???????????? ???????: {len(stocks)} ???????")
        
        if not stocks:
            return jsonify({"error": "no_stocks", "message": "??? ?????? ?? ????????. ?????????? ?????."}), 500
        
        # 3. ???????? ???????? ??????
        warehouse_name = None
        try:
            warehouses_response = requests.get(
                "https://supplies-api.wildberries.ru/api/v1/warehouses",
                headers={"Authorization": token, "Content-Type": "application/json"},
                timeout=30
            )
            if warehouses_response.status_code == 200:
                warehouses_data = warehouses_response.json()
                print(f"=== DEBUG: ???????? ?????? ???????: {len(warehouses_data)} ?????????")
                for warehouse in warehouses_data:
                    if str(warehouse.get("ID")) == str(warehouse_id):
                        warehouse_name = warehouse.get("name")
                        print(f"?????? ?????: ID={warehouse_id}, Name='{warehouse_name}'")
                        break
                if not warehouse_name:
                    print(f"????? ? ID {warehouse_id} ?? ?????? ? ?????? ???????")
                    print("????????? ??????:")
                    for wh in warehouses_data[:10]:  # ?????????? ?????? 10
                        print(f"  ID={wh.get('ID')}, Name='{wh.get('name')}'")
            elif warehouses_response.status_code == 429:
                print("=== ???????????? ????????: ?????? 429 ??? ???????? ??????? ===")
                # ?????????? fallback ????????
                warehouse_name = f"????? {warehouse_id}"
                print(f"?????????? fallback ????????: '{warehouse_name}'")
        except requests.HTTPError as e:
            if e.response and e.response.status_code == 429:
                print("=== ???????????? ????????: ?????? 429 ??? ???????? ??????? ===")
                warehouse_name = f"????? {warehouse_id}"
                print(f"?????????? fallback ????????: '{warehouse_name}'")
            else:
                print(f"?????? ????????? ???????? ??????: {e}")
                warehouse_name = f"????? {warehouse_id}"
        except Exception as e:
            print(f"?????? ????????? ???????? ??????: {e}")
            warehouse_name = f"????? {warehouse_id}"
        
        if not warehouse_name:
            warehouse_name = f"????? {warehouse_id}"
            print(f"?????????? fallback ????????: '{warehouse_name}'")
        
        # 4. ???????????? ??????? ??? ??????????? ?????? ? ???? ???????
        warehouse_stocks = {}
        all_warehouse_stocks = {}
        
        # ???????? ?????????? ???????? ??????? ?? ???????? ??? ???????
        unique_stock_warehouses = set()
        for stock in stocks:
            stock_warehouse = stock.get("warehouse", "")
            if stock_warehouse:
                unique_stock_warehouses.add(stock_warehouse)
        
        print(f"=== DEBUG: ?????????? ?????? ? ????????: {sorted(unique_stock_warehouses)}")
        print(f"=== DEBUG: ???? ??????? ??? ??????: '{warehouse_name}'")
        
        matched_stocks_count = 0
        for stock in stocks:
            stock_warehouse = stock.get("warehouse", "")
            barcode = stock.get("barcode")
            
            if barcode:
                qty = int(stock.get("qty", 0) or 0)
                
                # ??????? ?? ???? ???????
                if barcode in all_warehouse_stocks:
                    all_warehouse_stocks[barcode] += qty
                else:
                    all_warehouse_stocks[barcode] = qty
                
                # ??????? ?? ??????????? ?????? (? ?????? ??????????????)
                if _warehouse_names_match(stock_warehouse, warehouse_name):
                    matched_stocks_count += 1
                    if barcode in warehouse_stocks:
                        warehouse_stocks[barcode] += qty
                    else:
                        warehouse_stocks[barcode] = qty
        
        print(f"=== DEBUG: ??????? ???????? ??? ?????? '{warehouse_name}': {matched_stocks_count} ???????, ?????????? ???????: {len(warehouse_stocks)}")
        
        # 5. ???????? ??????
        print(f"????????? ?????? ??? ??????: '{warehouse_name}' ?? ?????? {date_from} - {date_to}")
        
        headers = {
            "Authorization": token,
            "Content-Type": "application/json"
        }
        
        # ???????????? ???? ? ?????? RFC3339 ??? API
        try:
            date_from_dt = datetime.strptime(date_from, "%d.%m.%Y")
            date_to_dt = datetime.strptime(date_to, "%d.%m.%Y")
            date_from_iso = date_from_dt.strftime("%Y-%m-%dT00:00:00")
            date_to_iso = date_to_dt.strftime("%Y-%m-%dT23:59:59")
        except ValueError:
            return jsonify({"error": "invalid_date_format"}), 400
        
        orders_url = "https://statistics-api.wildberries.ru/api/v1/supplier/orders"
        params = {
            "dateFrom": date_from_iso,
            "dateTo": date_to_iso
        }
        
        try:
            orders_response = requests.get(orders_url, headers=headers, params=params, timeout=30)
            orders_response.raise_for_status()
            orders_data = orders_response.json()
            print(f"???????? ???????: {len(orders_data)}")
        except requests.HTTPError as e:
            if e.response and e.response.status_code == 429:
                print("=== ???????????? ????????: ?????? 429 ??? ???????? ??????? ===")
                print("???????? ????? ???????? ? API ???????. ?????????? ??? ?????? ? ???????.")
                orders_data = []
            else:
                print(f"?????? ???????? ???????: {e}")
                orders_data = []
        except Exception as e:
            print(f"?????? ???????? ???????: {e}")
            orders_data = []
        
        # ?????????? ?????????? - ??????? ??? ?????? ? ???????
        unique_warehouses = set()
        for order in orders_data:
            wh_name = order.get("warehouseName")
            if wh_name:
                unique_warehouses.add(wh_name)
        
        print(f"?????????? ?????? ? ???????: {sorted(unique_warehouses)}")
        print(f"???? ?????? ??? ??????: '{warehouse_name}'")
        
        # ????????? ?????? ?? ?????? (?????????? ?????? ?????? ? ?????? ??????????????)
        filtered_orders = []
        matched_warehouses = set()
        for order in orders_data:
            order_warehouse = order.get("warehouseName")
            if order_warehouse and not order.get("isCancel", False):
                if _warehouse_names_match(order_warehouse, warehouse_name):
                    matched_warehouses.add(order_warehouse)
                    if order_warehouse != warehouse_name:
                        print(f"??????? ?????????? (???????? ??????????????): '{order_warehouse}' <-> '{warehouse_name}'")
                    filtered_orders.append(order)
        
        print(f"=== DEBUG: ??????? ??????? ??? ?????? '{warehouse_name}': {len(filtered_orders)}")
        if matched_warehouses:
            print(f"=== DEBUG: ??????????? ?????? ? ???????: {sorted(matched_warehouses)}")
        else:
            print(f"=== DEBUG: ?? ??????? ??????????. ????????? ????????????:")
            norm_warehouse = _normalize_warehouse_name(warehouse_name)
            print(f"  ??????????????? ???????? ???????? ??????: '{norm_warehouse}'")
            # ?????????? ??????????????? ???????? ?? ???????
            sample_norms = []
            for wh in list(unique_warehouses)[:5]:
                norm_wh = _normalize_warehouse_name(wh)
                sample_norms.append(f"'{wh}' -> '{norm_wh}'")
            print(f"  ??????? ??????????????? ???????? ?? ???????: {', '.join(sample_norms)}")
        
        # ?????????? ?????????? - ??????? ??????? ???????
        if filtered_orders:
            print("=== DEBUG: ??????? ????????? ??????? ===")
            for i, order in enumerate(filtered_orders[:3]):
                print(f"  ????? {i+1}: barcode={order.get('barcode')}, warehouseName={order.get('warehouseName')}")
        else:
            print("=== DEBUG: ?????? ?? ??????? ===")
            print("????????? ?????? 5 ??????? ?? ?????? ??????:")
            for i, order in enumerate(orders_data[:5]):
                print(f"  ????? {i+1}: barcode={order.get('barcode')}, warehouseName={order.get('warehouseName')}")
        
        # ?????????? ?????? ?? ???????? + ??????? ?? ???? ??? ?????? ????????
        from datetime import timedelta as _timedelta
        import statistics as _statistics
        orders_by_barcode: dict[str, int] = {}
        daily_by_barcode: dict[str, dict[str, int]] = {}  # barcode -> {YYYY-MM-DD: count}
        cancelled_orders = 0

        # ?????? ??? ??????? ????? (????????????)
        try:
            period_start = date_from_dt.date()
            period_end = date_to_dt.date()
        except Exception:
            period_start = None
            period_end = None

        for order in filtered_orders:
            barcode = str(order.get("barcode") or "").strip()
            if not barcode:
                print(f"????? ??? ???????: {order}")
                continue

            qty_raw = order.get("quantity")
            try:
                qty = int(qty_raw) if qty_raw is not None else 1
            except Exception:
                qty = 1
            if qty <= 0:
                qty = 1

            orders_by_barcode[barcode] = int(orders_by_barcode.get(barcode, 0)) + qty

            # date ? WB ?????? ????; ???? ????? ??? ? ????? lastChangeDate ??? fallback (????, ?? ????? ??? ??????)
            dt = parse_wb_datetime(order.get("date")) or parse_wb_datetime(order.get("lastChangeDate"))
            if dt is None:
                continue
            d = dt.date()
            if period_start and period_end:
                if d < period_start or d > period_end:
                    continue
            day_key = d.strftime("%Y-%m-%d")
            if barcode not in daily_by_barcode:
                daily_by_barcode[barcode] = {}
            daily_by_barcode[barcode][day_key] = int(daily_by_barcode[barcode].get(day_key, 0)) + qty
        
        # ???????????? ?????????? ?????? ??? ???????
        for order in orders_data:
            if order.get("isCancel", False):
                cancelled_orders += 1
        
        print(f"?????????? ??????? ?????????: {cancelled_orders}")
        
        print(f"?????? ????????????? ?? {len(orders_by_barcode)} ????????")
        
        # ???????????? ????? ?????????? ???????
        total_orders = sum(orders_by_barcode.values())
        print(f"????? ?????????? ???????: {total_orders}")

        # ????????? ??????? ???? ? ??? (????? ???????? ? ??????? ?? ??????? WB ????????)
        try:
            uid = int(current_user.id) if current_user.is_authenticated else None
            if uid is not None and period_start and period_end:
                cache_key = f"{warehouse_id}|{period_start.strftime('%Y-%m-%d')}|{period_end.strftime('%Y-%m-%d')}"
                if uid not in FBW_PLANNING_DYNAMICS_CACHE:
                    FBW_PLANNING_DYNAMICS_CACHE[uid] = {}
                FBW_PLANNING_DYNAMICS_CACHE[uid][cache_key] = {
                    "created_at": datetime.now().isoformat(),
                    "warehouse_id": str(warehouse_id),
                    "warehouse_name": str(warehouse_name),
                    "date_from": period_start.strftime("%Y-%m-%d"),
                    "date_to": period_end.strftime("%Y-%m-%d"),
                    "daily_by_barcode": daily_by_barcode,
                }
        except Exception:
            pass

        # ????????: ???? ????????? ???? ???????????? "???????" ??????? ??????
        # ?????????? ?????? ??????? ? ??????, ????? ?? ????????? JSON.
        anomalies: dict[str, dict[str, object]] = {}
        if period_start and period_end:
            total_days = (period_end - period_start).days + 1
        else:
            total_days = 0

        if total_days >= 7:
            for bc, day_map in (daily_by_barcode or {}).items():
                if not day_map:
                    continue
                # ???????? ??? ?? ???? ???????, ????? ????????? ????????? "????"
                values: list[int] = []
                if period_start and period_end:
                    cur = period_start
                    while cur <= period_end:
                        values.append(int(day_map.get(cur.strftime("%Y-%m-%d"), 0)))
                        cur = cur + _timedelta(days=1)
                else:
                    values = [int(v) for v in day_map.values()]

                nonzero = [v for v in values if v > 0]
                if len(nonzero) < 3:
                    continue

                try:
                    median_nz = float(_statistics.median(nonzero))
                except Exception:
                    median_nz = float(nonzero[len(nonzero) // 2])

                max_val = max(values) if values else 0
                if max_val <= 0:
                    continue

                # ??????? (??????? ? ??????????):
                # - ?????? ???? ???????? ????? ?? ???????
                # - ? ?????????? ????? (????? ?? ???????? 2->4)
                # ??????: "?????? 1..5, ?? ?????? 10" -> ????????.
                abs_min_spike = 10
                mult_spike = 2.0
                add_spike = 5.0
                is_anomaly = (max_val >= abs_min_spike) and (max_val >= (median_nz * mult_spike)) and ((max_val - median_nz) >= add_spike)
                if not is_anomaly:
                    continue

                # ??????? ???? ????????? (??????, ???? ?????????)
                max_day = None
                for day_str, v in day_map.items():
                    if int(v) == int(max_val):
                        max_day = day_str
                        break

                anomalies[bc] = {
                    "max": int(max_val),
                    "median_nonzero": float(median_nz),
                    "day": max_day,
                }
        
        # ???????? ????? ??????????
        now_str = cached.get("updated_at") if cached else datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        
        # ????????? ?????? ??????? ? ????????? (??? ? ???????????? endpoint)
        products_with_barcodes = []
        for product in products:
            barcode = product.get("barcode")
            if barcode:
                products_with_barcodes.append({
                    "barcode": str(barcode),
                    "name": product.get("supplier_article") or "??? ????????",
                    "nm_id": product.get("nm_id"),
                    "supplier_article": product.get("supplier_article") or "??? ????????",
                    "photo": product.get("photo")
                })

        # ?????????????? ?? /tools/prices (?????? ?? ??? ?? ??????/?????????? ????????????)
        margins_by_barcode = _calc_tools_prices_margin_map(token, products, current_user.id)
        
        # ?????????? ??? ?????? ????? ???????
        return jsonify({
            "success": True,
            "products": {
                "success": True,
                "products": products_with_barcodes,
                "count": len(products_with_barcodes)
            },
            "stocks": {
                "success": True,
                "stocks": warehouse_stocks,
                "total_stocks": sum(warehouse_stocks.values()),
                "unique_products": len(warehouse_stocks),
                "warehouse_id": warehouse_id,
                "warehouse_name": warehouse_name,
                "updated_at": now_str
            },
            "all_stocks": {
                "success": True,
                "stocks": all_warehouse_stocks,
                "total_stocks": sum(all_warehouse_stocks.values()),
                "unique_products": len(all_warehouse_stocks),
                "warehouse_id": "all",
                "warehouse_name": "??? ??????",
                "updated_at": now_str
            },
            "orders": {
                "success": True,
                "orders": orders_by_barcode,
                "anomalies": anomalies,
                "total_orders": total_orders,
                "unique_products": len(orders_by_barcode),
                "warehouse_id": warehouse_id,
                "warehouse_name": warehouse_name,
                "date_from": date_from,
                "date_to": date_to
            },
            "margins": margins_by_barcode,
        })
        
    except requests.HTTPError as http_err:
        print(f"=== ?????? API ? api_fbw_planning_data ===")
        print(f"HTTP Error: {http_err.response.status_code}")
        print(f"Response: {http_err.response.text}")
        
        # ??????????? ????????? ?????? 429
        if http_err.response and http_err.response.status_code == 429:
            # ???????? ??????? retry_after ?? ??????????
            retry_after = http_err.response.headers.get("Retry-After", "60")
            try:
                retry_seconds = int(retry_after)
            except (ValueError, TypeError):
                retry_seconds = 60
            
            return jsonify({
                "error": "rate_limit",
                "message": f"???????? ????? ???????? ? API Wildberries. ?????????? ????? {retry_seconds} ??????.",
                "retry_after": retry_seconds
            }), 429
        
        return jsonify({"error": "api_error", "message": f"?????? API: {http_err.response.status_code}"}), 502
    except Exception as exc:
        print(f"=== ?????? ? api_fbw_planning_data ===")
        print(f"Exception: {str(exc)}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": "server_error", "message": f"??????: {str(exc)}"}), 500

@app.route("/api/fbw/planning/orders", methods=["GET"])
@login_required
def api_fbw_planning_orders():
    """API ??? ????????? ??????? ?? ?????? ?? ?????? ??? ????????????"""
    warehouse_id = request.args.get('warehouse_id')
    date_from = request.args.get('date_from')
    date_to = request.args.get('date_to')
    
    if not warehouse_id or not date_from or not date_to:
        return jsonify({"error": "missing_parameters"}), 400
    
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        # ???????? ???????? ?????? ?? API ???????
        warehouse_name = None
        try:
            warehouses_response = requests.get(
                "https://supplies-api.wildberries.ru/api/v1/warehouses",
                headers={"Authorization": token, "Content-Type": "application/json"},
                timeout=30
            )
            if warehouses_response.status_code == 200:
                warehouses_data = warehouses_response.json()
                for warehouse in warehouses_data:
                    if str(warehouse.get("ID")) == str(warehouse_id):
                        warehouse_name = warehouse.get("name")
                        break
        except Exception as e:
            print(f"?????? ????????? ???????? ??????: {e}")
        
        if not warehouse_name:
            warehouse_name = f"????? {warehouse_id}"
        
        print(f"????????? ?????? ??? ??????: '{warehouse_name}' ?? ?????? {date_from} - {date_to}")
        
        # ????????? ?????? ????? API Wildberries
        headers = {
            "Authorization": token,
            "Content-Type": "application/json"
        }
        
        # ???????????? ???? ? ?????? RFC3339 ??? API
        from datetime import datetime
        try:
            # ?????? ???? ?? ??????? DD.MM.YYYY
            date_from_obj = datetime.strptime(date_from, "%d.%m.%Y")
            date_to_obj = datetime.strptime(date_to, "%d.%m.%Y")
            
            # ???????????? ? RFC3339 ?????? (? ???????? 00:00:00)
            date_from_rfc = date_from_obj.strftime("%Y-%m-%dT00:00:00")
            date_to_rfc = date_to_obj.strftime("%Y-%m-%dT23:59:59")
            
        except ValueError as e:
            return jsonify({"error": "invalid_date_format", "message": "???????? ?????? ????. ??????????? DD.MM.YYYY"}), 400
        
        # ????????? ??????
        orders_url = "https://statistics-api.wildberries.ru/api/v1/supplier/orders"
        orders_params = {
            "dateFrom": date_from_rfc,
            "dateTo": date_to_rfc
        }
        
        print(f"?????? ? API ???????: {orders_url} ? ???????????: {orders_params}")
        
        orders_response = requests.get(
            orders_url,
            headers=headers,
            params=orders_params,
            timeout=30
        )
        
        if orders_response.status_code != 200:
            print(f"?????? API ???????: {orders_response.status_code} - {orders_response.text}")
            return jsonify({
                "error": "orders_api_error", 
                "message": f"?????? API ???????: {orders_response.status_code}"
            }), 502
        
        orders_data = orders_response.json()
        print(f"???????? ???????: {len(orders_data)}")
        
        # ????????? ?????? ?? ??????
        filtered_orders = []
        
        # ???????? ?????????? ???????? ??????? ?? ??????? ??? ???????
        unique_warehouses = set()
        for order in orders_data:
            wh_name = order.get("warehouseName")
            if wh_name:
                unique_warehouses.add(wh_name)
        
        print(f"?????????? ?????? ? ???????: {sorted(unique_warehouses)}")
        print(f"???? ?????? ??? ??????: '{warehouse_name}'")
        
        for order in orders_data:
            order_warehouse = order.get("warehouseName")
            if order_warehouse:
                if _warehouse_names_match(order_warehouse, warehouse_name):
                    if order_warehouse != warehouse_name:
                        print(f"??????? ?????????? (???????? ??????????????): '{order_warehouse}' <-> '{warehouse_name}'")
                    filtered_orders.append(order)
        
        print(f"??????? ??????? ??? ?????? '{warehouse_name}': {len(filtered_orders)}")
        
        # ?????????? ?????? ?? ???????? ??? ???????? ??????????
        # ? API ??????? ?????? ????? = 1 ?????, ??????? ??????? ?????????? ???????
        # ????????? ?????????? ?????? (isCancel = true)
        orders_by_barcode = {}
        cancelled_orders = 0
        
        for order in filtered_orders:
            # ?????????, ?? ??????? ?? ?????
            is_cancelled = order.get("isCancel", False)
            if is_cancelled:
                cancelled_orders += 1
                continue  # ?????????? ?????????? ??????
            
            barcode = order.get("barcode")
            if barcode:
                if barcode not in orders_by_barcode:
                    orders_by_barcode[barcode] = 0
                # ?????? ????? = 1 ?????
                orders_by_barcode[barcode] += 1
        
        print(f"?????????? ??????? ?????????: {cancelled_orders}")
        
        # ?????????? ?????????? ? ??????????? ???????
        print(f"=== DEBUG: ??????????? ??????? ===")
        print(f"????? ??????????????? ???????: {len(filtered_orders)}")
        print(f"?????????? ????????: {len(orders_by_barcode)}")
        
        # ???????? ?????? ????????? ???????
        if filtered_orders:
            print("??????? ???????:")
            for i, order in enumerate(filtered_orders[:3]):
                print(f"  ????? {i+1}: barcode={order.get('barcode')}, quantity={order.get('quantity')}, warehouseName={order.get('warehouseName')}")
                print(f"    ??? ???? ??????: {list(order.keys())}")
                print(f"    ?????? ?????: {order}")
        
        # ???????? ?????? ????????? ??????????????? ???????
        if orders_by_barcode:
            print("??????? ??????????????? ???????:")
            for i, (barcode, qty) in enumerate(list(orders_by_barcode.items())[:5]):
                print(f"  {barcode}: {qty}")
        
        print(f"?????? ????????????? ?? {len(orders_by_barcode)} ????????")
        
        return jsonify({
            "success": True,
            "warehouse_id": warehouse_id,
            "warehouse_name": warehouse_name,
            "date_from": date_from,
            "date_to": date_to,
            "orders": orders_by_barcode,
            "total_orders": len(filtered_orders),
            "cancelled_orders": cancelled_orders,
            "unique_products": len(orders_by_barcode)
        })
        
    except Exception as e:
        print(f"?????? ????????? ???????: {e}")
        return jsonify({"error": "server_error", "message": str(e)}), 500


@app.route("/api/fbw/planning/orders/dynamics", methods=["GET"])
@login_required
def api_fbw_planning_orders_dynamics():
    """API: ???????? ??????? ?? ???? ??? ?????? ?????? (?? ???????) ? ?????? ?????? ? ???????."""
    warehouse_id = request.args.get("warehouse_id")
    barcode = (request.args.get("barcode") or "").strip()
    date_from = request.args.get("date_from")
    date_to = request.args.get("date_to")

    if not warehouse_id or not barcode or not date_from or not date_to:
        return jsonify({"error": "missing_parameters"}), 400

    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401

    try:
        from datetime import date as _date, timedelta as _timedelta

        start_dt = parse_date(str(date_from))
        end_dt = parse_date(str(date_to))
        if start_dt.date() > end_dt.date():
            return jsonify({"error": "invalid_period", "message": "date_from must be <= date_to"}), 400

        start_ymd = start_dt.strftime("%Y-%m-%d")
        end_ymd = end_dt.strftime("%Y-%m-%d")

        # 0) ??????? ????: ???? ???????????? ??? ?????????, ????? ??????? ???????? ?? in-memory ????
        try:
            uid = int(current_user.id) if current_user.is_authenticated else None
            if uid is not None:
                cache_key = f"{warehouse_id}|{start_ymd}|{end_ymd}"
                cached = (FBW_PLANNING_DYNAMICS_CACHE.get(uid) or {}).get(cache_key) or {}
                daily_by_barcode = cached.get("daily_by_barcode") if isinstance(cached, dict) else None
                if isinstance(daily_by_barcode, dict):
                    day_map = daily_by_barcode.get(barcode) or {}
                    series: list[dict[str, object]] = []
                    cur = start_dt.date()
                    total = 0
                    while cur <= end_dt.date():
                        day_key = cur.strftime("%Y-%m-%d")
                        c = int(day_map.get(day_key, 0) or 0)
                        series.append({"date": day_key, "count": c})
                        total += c
                        cur = cur + _timedelta(days=1)

                    return jsonify(
                        {
                            "success": True,
                            "source": "cache",
                            "warehouse_id": cached.get("warehouse_id", warehouse_id),
                            "warehouse_name": cached.get("warehouse_name", f"????? {warehouse_id}"),
                            "barcode": barcode,
                            "date_from": cached.get("date_from", start_ymd),
                            "date_to": cached.get("date_to", end_ymd),
                            "total": int(total),
                            "cancelled_skipped": 0,
                            "series": series,
                        }
                    )
        except Exception:
            pass

        # ???????? ???????? ?????? ?? API ??????? (??? ? ? ?????? planning endpoints)
        warehouse_name = None
        try:
            warehouses_response = requests.get(
                "https://supplies-api.wildberries.ru/api/v1/warehouses",
                headers={"Authorization": token, "Content-Type": "application/json"},
                timeout=30,
            )
            if warehouses_response.status_code == 200:
                warehouses_data = warehouses_response.json()
                for wh in warehouses_data:
                    if str(wh.get("ID")) == str(warehouse_id):
                        warehouse_name = wh.get("name")
                        break
        except Exception as e:
            print(f"?????? ????????? ???????? ?????? (dynamics): {e}")

        if not warehouse_name:
            warehouse_name = f"????? {warehouse_id}"

        # ????????? ?????? ? ?????????? (?? lastChangeDate), ????????? ?? ???????? ???? (date)
        # days_back ?????????, ????? ????????? "??????????" ??????????, ?? ?? ????????? ???????.
        raw_orders = fetch_orders_range(token, start_ymd, end_ymd, days_back=1)

        counts_by_day: dict[_date, int] = {}
        total = 0
        cancelled = 0

        for order in (raw_orders or []):
            try:
                if order.get("isCancel", False):
                    cancelled += 1
                    continue

                o_barcode = str(order.get("barcode") or "").strip()
                if o_barcode != barcode:
                    continue

                o_wh = (order.get("warehouseName") or "").strip()
                if not o_wh:
                    continue
                if not _warehouse_names_match(o_wh, warehouse_name):
                    continue

                dt = parse_wb_datetime(order.get("date")) or parse_wb_datetime(order.get("lastChangeDate"))  # fallback
                if dt is None:
                    continue
                d = dt.date()
                if d < start_dt.date() or d > end_dt.date():
                    continue

                qty_raw = order.get("quantity")
                try:
                    qty = int(qty_raw) if qty_raw is not None else 1
                except Exception:
                    qty = 1
                if qty <= 0:
                    qty = 1

                counts_by_day[d] = counts_by_day.get(d, 0) + qty
                total += qty
            except Exception:
                continue

        # ????????? ??? ?? ???? ???? ??????? (????????????)
        series: list[dict[str, object]] = []
        cur = start_dt.date()
        while cur <= end_dt.date():
            series.append({"date": cur.strftime("%Y-%m-%d"), "count": int(counts_by_day.get(cur, 0))})
            cur = cur + _timedelta(days=1)

        return jsonify(
            {
                "success": True,
                "source": "wb",
                "warehouse_id": warehouse_id,
                "warehouse_name": warehouse_name,
                "barcode": barcode,
                "date_from": start_ymd,
                "date_to": end_ymd,
                "total": int(total),
                "cancelled_skipped": int(cancelled),
                "series": series,
            }
        )
    except requests.HTTPError as e:
        if e.response and e.response.status_code == 429:
            retry_after = e.response.headers.get("X-Ratelimit-Retry") or e.response.headers.get("Retry-After") or "60"
            try:
                retry_s = int(float(retry_after))
            except Exception:
                retry_s = 60
            return jsonify({"error": "rate_limit", "message": "???????? ????? ???????? ? API Wildberries.", "retry_after": retry_s}), 429
        return jsonify({"error": "api_error", "message": str(e)}), 502
    except Exception as e:
        print(f"?????? ????????? ???????? ???????: {e}")
        return jsonify({"error": "server_error", "message": str(e)}), 500


@app.route("/api/fbw/planning/export-excel", methods=["POST"])
@login_required
def api_fbw_planning_export_excel():
    """??????? ??????????? ???????????? ? Excel ?????? XLS"""
    try:
        data = request.get_json()
        if not data or 'products' not in data:
            return jsonify({"error": "??? ?????? ??? ????????"}), 400
        
        products = data['products']
        warehouse_name = data.get('warehouse_name', '???????????_?????')
        
        # ????????? ?????? - ???????????? ?????? ??, ? ??????? ?????????? ??? ???????? ?????? 0
        products_to_export = [p for p in products if p.get('toSupply', 0) > 0]
        
        if not products_to_export:
            return jsonify({"error": "??? ??????? ??? ????????"}), 400
        
        # ??????? Excel ???? ? ??????? XLS (Excel 97-2003)
        workbook = xlwt.Workbook(encoding='utf-8')
        worksheet = workbook.add_sheet('???????????? ????????')
        
        # ?????
        header_style = xlwt.easyxf('font: bold on; align: horiz center;')
        number_style = xlwt.easyxf('align: horiz right;')
        
        # ?????????
        headers = [
            '?', '????????', '????????????', '??????? ???????', 
            '??????? ?? ???? ???????', '? ???? ?? ?????', '???????? ?? ??????', '??????????????',
            '?????? ? ????', '??????????? ???????', '???????????????', '????????? ?? ?????'
        ]
        
        for col, header in enumerate(headers):
            worksheet.write(0, col, header, header_style)
        
        # ??????
        for row, product in enumerate(products_to_export, 1):
            worksheet.write(row, 0, row)  # ?
            worksheet.write(row, 1, str(product.get('barcode', '')))  # ????????
            worksheet.write(row, 2, str(product.get('name', '')))  # ????????????
            worksheet.write(row, 3, product.get('currentStock', 0), number_style)  # ??????? ???????
            worksheet.write(row, 4, product.get('allStocks', 0), number_style)  # ??????? ?? ???? ???????
            worksheet.write(row, 5, product.get('inTransit', 0), number_style)  # ? ???? ?? ?????
            worksheet.write(row, 6, product.get('orderedInPeriod', 0), number_style)  # ???????? ?? ??????
            worksheet.write(row, 7, product.get('marginPct', ''), number_style)  # ??????????????
            worksheet.write(row, 8, round(product.get('salesPerDay', 0), 2), number_style)  # ?????? ? ????
            worksheet.write(row, 9, round(product.get('requiredStock', 0)), number_style)  # ??????????? ???????
            worksheet.write(row, 10, round(product.get('turnover', 0), 1), number_style)  # ???????????????
            worksheet.write(row, 11, round(product.get('toSupply', 0)), number_style)  # ????????? ?? ?????
        
        # ?????????? ?????? ???????
        for col in range(len(headers)):
            worksheet.col(col).width = 3000  # ????????? ??????
        
        # ??????? ???? ? ??????
        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        
        # ?????????? ??? ????? ? ????? ??????????? ???????
        now = datetime.now()
        day = now.strftime("%d.%m.%Y")
        time = now.strftime("%H_%M")
        total_quantity = sum(round(p.get('toSupply', 0)) for p in products_to_export)
        filename = f"{warehouse_name}_{day}_{time}_({total_quantity}).xls"
        
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.ms-excel'
        )
        
    except Exception as e:
        print(f"?????? ???????? ? Excel: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"?????? ????????: {str(e)}"}), 500
@app.route("/api/fbw/planning/supplies", methods=["GET"])
@login_required
def api_fbw_planning_supplies():
    """API ??? ????????? ???????? ?? ???????? '???????? ?????????' ??? ???????????? (???????????????? ??????)"""
    token = effective_wb_api_token(current_user)
    warehouse_name = request.args.get("warehouse_name")
    force_refresh = request.args.get("force_refresh", "false").lower() == "true"
    
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        from datetime import datetime
        print(f"?????? ???????? ??? ????????????: warehouse='{warehouse_name}', force_refresh={force_refresh}")
        
        # ????????? ??? ????????
        cached = load_fbw_supplies_cache() or {}
        
        # ?????????? ???????????? ?????? ?? ???????? ???????? /fbw, ???? ??? ????
        # ??? ??? ???? ??????? ? ???? ????????, ??? ??????? ????????
        cached_items = cached.get("items", [])
        cached_supplies_map = {}
        for item in cached_items:
            sid = str(item.get("supply_id") or "")
            if sid:
                cached_supplies_map[sid] = item
        
        # ??? ???????????? ?????????? ????????? ???, ????? ?? ?????? ?? ???????? ??? ???????? /fbw
        planning_cache_key = f"planning_supplies_{current_user.id}"
        cached_planning = cached.get(planning_cache_key, {})
        
        # ????????? ??????? ?????? ???????? ??? ????????? ID ???? ????????
        print("???????????? ????????: ????????? ?????? ID ???????? ? API")
        supplies_list = fetch_fbw_supplies_list(token, days_back=30)
        print(f"???????? ?????? ID ???????? ??? ????????????: {len(supplies_list)} ????????")
        
        # ??????? ????????? ???????? ?? ?????? ?? ???? - ??? ???????, ??? ????????? ?????? ??? ????
        # ????????? ???????? ?? ???? ???????? - ?????? ?? ????????? 60 ????
        from datetime import timedelta
        cutoff_date = datetime.now(MOSCOW_TZ) - timedelta(days=60)
        
        # ??????????????? ??????????: ?? ???? ???????? /fbw (???? ????) + ??????????? ?????? ???????? ??? ????.
        # ?????? ???????? ??? ?????? ? load_fbw_supplies_cache() ????????????? ? ????? ???????????? ??????
        # ????????? ? ????? ?????? ????? ???????? /fbw. ????????? ??? ??????? (supplies_by_date) ???????????
        # ?? ???? ??? ???????? supply_id, ??? ???? ???????? ?? ?? ????????.
        def _planning_supply_create_ts(sup: dict[str, Any]) -> float:
            cd = sup.get("createDate")
            if not cd:
                return 0.0
            try:
                dt = _parse_iso_datetime(str(cd))
                return float(dt.timestamp()) if dt else 0.0
            except Exception:
                return 0.0

        supplies_from_cache: list[dict[str, Any]] = []
        uncached_candidates: list[dict[str, Any]] = []
        for supply in supplies_list:
            supply_id = supply.get("supplyID") or supply.get("supplyId") or supply.get("id")
            supply_id_str = str(supply_id or "")
            if not supply_id_str:
                continue
            
            # ????????? ???? ????????
            create_date = supply.get("createDate")
            if create_date:
                try:
                    create_dt = _parse_iso_datetime(str(create_date))
                    if create_dt and create_dt < cutoff_date:
                        continue  # ?????????? ?????? ????????
                except Exception:
                    pass  # ???? ?? ??????? ??????????, ????????? ??????
            
            # ??????????????? ?????????? ?? ???? ??? ?????????
            # ?????: ??? ?????? ??????????????? ??????????, ????????????? ???????? ????? ? API
            cached_item = cached_supplies_map.get(supply_id_str)
            if cached_item:
                cached_warehouse = cached_item.get("warehouse", "").strip()
                cached_status = cached_item.get("status", "").strip()
                cached_type = cached_item.get("type", "").strip()
                
                # ???? ????? ?????? ? ?? ????????? - ????? ?????????? (? ?????? ??????????????)
                if cached_warehouse and warehouse_name:
                    if not _warehouse_names_match(cached_warehouse, warehouse_name):
                        continue  # ?????????? ???????? ?? ?????? ??????
                
                # ???? ?????? "???????" ? ???? - ?????????? (?? ? ???????? ???????? ??? ????? ???????? ????? API)
                # ??? ???????? ??????????????? ??????????
                if cached_status and "???????" in cached_status:
                    print(f"???????? {supply_id_str}: ? ???? ?????? '???????', ?????????? ?? ????? ??????????????? ??????????")
                    continue
                
                # ???? ??? "?????????" ??? "??? ???????" - ????? ??????????
                if cached_type and ("?????????" in cached_type or "??? ???????" in cached_type):
                    continue
                
                # ???? ??? ?? "??????" - ??????????
                if cached_type and "?????" not in cached_type.lower():
                    continue
                
                # ???? ??? ???????? ???????? - ????????? ??? ?????????? ?????????
                # ? ???????? ???????? ????? ????????? ?????????? ?????????? ? API
                supplies_from_cache.append(supply)
            else:
                # ??? ?????? ? ?????? ???? /fbw ? ??? ????? ????????? ????? API (?????/?????? ?????????? ?? ???????)
                uncached_candidates.append(supply)

        uncached_sorted = sorted(uncached_candidates, key=_planning_supply_create_ts, reverse=True)
        # ??????? ????? ????? ???????? ??? ???? (????? ??? ??? ?? ?????), ????? ??????????????? ?? ????
        PLANNING_UNCACHED_PRIORITY_MAX = 24
        MAX_SUPPLIES_TO_CHECK = 28
        seen_ids: set[str] = set()
        supplies_to_check: list[dict[str, Any]] = []
        for s in uncached_sorted[:PLANNING_UNCACHED_PRIORITY_MAX]:
            sid = str(s.get("supplyID") or s.get("supplyId") or s.get("id") or "")
            if sid and sid not in seen_ids:
                supplies_to_check.append(s)
                seen_ids.add(sid)
        for s in supplies_from_cache:
            sid = str(s.get("supplyID") or s.get("supplyId") or s.get("id") or "")
            if sid and sid not in seen_ids:
                supplies_to_check.append(s)
                seen_ids.add(sid)
        if len(supplies_to_check) > MAX_SUPPLIES_TO_CHECK:
            supplies_to_check = supplies_to_check[:MAX_SUPPLIES_TO_CHECK]

        print(
            f"???????????? ????????: ? ???????? {len(supplies_to_check)} ??. "
            f"(?? ???? /fbw: {len(supplies_from_cache)}, ??? ????: {len(uncached_candidates)}, ????? ? ?????? API: {len(supplies_list)})"
        )
        
        pending_supplies = []
        status_counts = {}
        
        for supply in supplies_to_check:
            supply_id = supply.get("supplyID") or supply.get("supplyId") or supply.get("id")
            supply_id_str = str(supply_id or "")
            if not supply_id_str:
                continue
            
            try:
                # ??????? ????????? ??? - ??? ??? ????? ???? ?????? ? ???
                cached_item = cached_supplies_map.get(supply_id_str)
                details = None
                details_status = None
                box_type_name = None
                box_type_id = None
                warehouse_from_details = None
                
                # ?????????? ??? ??? ?????????, ?? ????????? ????????? ?????? ????? API
                # ???? ? ???? ???? ?????? ? ?? ?? "???????", ?????????? ??? ??? ?????????
                details = None
                use_cache_for_status = False
                
                if cached_item:
                    cached_status = cached_item.get("status", "").strip()
                    cached_warehouse = cached_item.get("warehouse", "").strip()
                    
                    # ???? ? ???? ?????? ????? ?? "???????" ? ????? ?????????, ?????????? ???
                    # ??? ??????????? ??????? ?????????
                    if cached_status and "???????" not in cached_status and cached_warehouse:
                        if not warehouse_name or _warehouse_names_match(cached_warehouse, warehouse_name):
                            use_cache_for_status = True
                            details_status = cached_status
                            warehouse_from_details = cached_warehouse
                            box_type_name = cached_item.get("type", "").strip()
                            # ?????????? boxTypeID ?? ????
                            if box_type_name == "?????????" or box_type_name == "??? ???????":
                                box_type_id = 1
                            elif "?????" in box_type_name.lower():
                                box_type_id = 2
                            else:
                                box_type_id = None
                            print(f"???????? {supply_id}: ?????????? ?????? ?? ???? (??????: {details_status})")
                
                # ???? ?? ?????????? ???, ????????? ?????? ? API
                if not use_cache_for_status:
                    try:
                        _supplies_api_throttle()
                        details = fetch_fbw_supply_details(token, supply_id)
                        if not details:
                            print(f"???????? {supply_id}: ?? ??????? ????????? ?????? ? API, ??????????")
                            continue
                        
                        # ???????? ????? ?? ??????? ? ????????? ???
                        warehouse_from_details = details.get("warehouseName", "").strip()
                        if not warehouse_from_details:
                            print(f"???????? {supply_id}: ??? ???????? ?????? ? ???????, ??????????")
                            continue
                        
                        # ????????? ?????????? ?????? ????? ????????? ??????? ? ???? (? ?????? ??????????????)
                        if warehouse_name:
                            if not _warehouse_names_match(warehouse_from_details, warehouse_name):
                                print(f"???????? {supply_id} ?? ?????? ????? '{warehouse_from_details}' != '{warehouse_name}', ??????????")
                                continue
                        
                        # ???????? ?????? ?? ??????? API (????? ??????????)
                        details_status = details.get("statusName", "").strip()
                        if not details_status:
                            # ???? ?????? ??????, ???????? ?????????? ??? ?? ?????
                            supply_date = details.get("supplyDate")
                            fact_date = details.get("factDate")
                            if fact_date:
                                # ???? ???? ??????????? ???? - ???????? ???????
                                details_status = "???????"
                                print(f"???????? {supply_id}: ????????? ?????? '???????' ?? factDate={fact_date}")
                            elif supply_date:
                                try:
                                    planned_dt = _parse_iso_datetime(str(supply_date))
                                    if planned_dt:
                                        planned_dt_msk = to_moscow(planned_dt) if planned_dt else None
                                        if planned_dt_msk:
                                            today = datetime.now(MOSCOW_TZ).date()
                                            planned_date = planned_dt_msk.date()
                                            if planned_date < today:
                                                details_status = "???????? ?????????"
                                            else:
                                                details_status = "?????????????"
                                except Exception:
                                    details_status = "?????????????"
                            else:
                                details_status = "?? ?????????????"
                        
                        # ???????? ??? ?? ???????
                        box_type_name = details.get("boxTypeName", "").strip()
                        box_type_id = details.get("boxTypeID")
                    except requests.Timeout:
                        print(f"???????? {supply_id}: ??????? ??? ???????? ???????, ?????????? ??? ???? ????????")
                        if cached_item:
                            details_status = cached_item.get("status", "").strip() or "??????????"
                            warehouse_from_details = cached_item.get("warehouse", "").strip()
                            box_type_name = cached_item.get("type", "").strip()
                            if box_type_name == "?????????" or box_type_name == "??? ???????":
                                box_type_id = 1
                            elif "?????" in box_type_name.lower():
                                box_type_id = 2
                            else:
                                box_type_id = None
                        else:
                            print(f"???????? {supply_id}: ??? ????, ?????????? ??-?? ????????")
                            continue
                    except Exception as e:
                        print(f"???????? {supply_id}: ?????? ???????? ???????: {e}, ??????????")
                        continue
                
                # ????????: ????????? ?????? "???????" ?????
                # ??? ????? ?????? ???????? - ???????? ???????? ?? ?????? ???????????
                if "???????" in details_status:
                    print(f"???????? {supply_id}: ?????? '{details_status}' - ???????, ?????????? (?? ????????? ? ????????????)")
                    continue
                
                # ?????????? ?????? ?? ???? ?????? ??? ?????????????? ?????????? (????, ??????????)
                cached_item = cached_supplies_map.get(supply_id_str)
                
                status_counts[details_status] = status_counts.get(details_status, 0) + 1
                
                # ????????? ??? ???????? - ????????? "?????????" ?????
                # ???? boxTypeID = 1 ??? boxTypeName ???????? "?????????"/"??? ???????" - ??????????
                if box_type_id == 1 or (box_type_name and ("?????????" in box_type_name.lower() or "??? ???????" in box_type_name.lower())):
                    print(f"???????? {supply_id} ????? ??? '?????????' (boxTypeID={box_type_id}, boxTypeName='{box_type_name}'), ??????????")
                    continue
                
                # ???? boxTypeID = 2 ??? boxTypeName ???????? "?????" - ?????????, ????? ??????????
                if box_type_id is not None and box_type_id != 2:
                    if not box_type_name or "?????" not in box_type_name.lower():
                        print(f"???????? {supply_id} ?? ???????? ????? '??????' (boxTypeID={box_type_id}, boxTypeName='{box_type_name}'), ??????????")
                        continue
                elif box_type_name and "?????" not in box_type_name.lower():
                    print(f"???????? {supply_id} ?? ???????? ????? '??????' (boxTypeName='{box_type_name}'), ??????????")
                    continue
                
                # ?????????, ??? ?????? ???????? ??? ???????????? (?????? "?????????????" ? "???????? ?????????")
                # ?????? "???????" ??? ???????? ???? ? ????????
                if details_status not in ["?????????????", "???????? ?????????"]:
                    print(f"???????? {supply_id} ????? ?????? '{details_status}', ?? ???????? ??? ????????????, ??????????")
                    continue
                
                # ?????????, ??? ????? ???? (???????? ??? ???? ???? ??? ???????? ??? ????)
                if not warehouse_from_details:
                    continue
                
                # ??? ???????? ?? ???? ????????? ?????????? ?????? ??? ??? (?? ?????? ??????, ? ?????? ??????????????)
                if cached_item and warehouse_name:
                    if not _warehouse_names_match(warehouse_from_details, warehouse_name):
                        print(f"???????? {supply_id} ?? ???? ?? ?????? ????? '{warehouse_from_details}' != '{warehouse_name}', ??????????")
                        continue
                
                # ???????? ???? ?? ???? ??? ???????
                planned_date_str = ""
                if cached_item:
                    planned_date_str = cached_item.get("planned_date", "") or ""
                if not planned_date_str and details:
                    planned_date_str = _fmt_dt_moscow(details.get("supplyDate"), with_time=False) if details.get("supplyDate") else ""
                
                created_at_str = ""
                if cached_item:
                    created_at_str = cached_item.get("created_at", "") or ""
                if not created_at_str:
                    created_at_str = _fmt_dt_moscow(supply.get("createDate"), with_time=False) if supply.get("createDate") else ""
                
                # ????????? ?????? ?? ???????? ?????? ???? ???????? ?????? ??? ????????
                # ?????????? ??? ??? ?????????, ???? ????????
                supply_goods = []
                try:
                    # ???????? ???????? ?????? ?? ????
                    if cached_item and cached_item.get("goods"):
                        supply_goods = cached_item.get("goods", [])
                        print(f"???????? {supply_id}: ?????????? ?????? ?? ???? ({len(supply_goods)} ???????)")
                    else:
                        # ???? ? ???? ???, ????????? ? API
                        goods = fetch_fbw_supply_goods(token, supply_id)
                        for good in goods:
                            barcode = good.get("barcode", "").strip()
                            qty = int(good.get("quantity", 0) or 0)
                            if barcode and qty > 0:
                                supply_goods.append({
                                    "barcode": barcode,
                                    "quantity": qty,
                                    "name": good.get("name", ""),
                                    "article": good.get("article", "")
                                })
                except requests.Timeout:
                    print(f"???????? {supply_id}: ??????? ??? ???????? ???????, ?????????? ??? ???? ????????")
                    if cached_item and cached_item.get("goods"):
                        supply_goods = cached_item.get("goods", [])
                    # ???? ???? ???, ?????????? ??? ??????? (??? ?? ????????)
                except Exception as e:
                    print(f"?????? ???????? ??????? ???????? {supply_id}: {e}")
                    # ?????????? ??? ??????? - ??? ?? ???????? ??? ????????????
                
                # ???????? total_goods ?? ???? ??? ???????
                total_goods = 0
                if cached_item and cached_item.get("total_goods") is not None:
                    total_goods = int(cached_item.get("total_goods", 0) or 0)
                elif details:
                    total_goods = int(details.get("quantity", 0) or 0)
                
                # ????????? ???????? ? ??????
                pending_supplies.append({
                    "supply_id": supply_id_str,
                    "warehouse": warehouse_from_details,
                    "total_goods": total_goods,
                    "goods": supply_goods,
                    "planned_date": planned_date_str,
                    "created_at": created_at_str
                })
                
            except Exception as e:
                print(f"?????? ????????? ???????? {supply_id}: {e}")
                continue
        
        print(f"??????? ?????????? ???????? ??? ????????????: {len(pending_supplies)}")
        print(f"?????????? ?? ????????: {status_counts}")
        
        return jsonify({
            "success": True,
            "supplies": pending_supplies,
            "warehouse_name": warehouse_name,
            "count": len(pending_supplies)
        })
        
    except Exception as e:
        return jsonify({"error": "server_error", "message": str(e)}), 500


@app.route("/api/fbw/supplies", methods=["GET"])
@login_required
def api_fbw_supplies():
    # If cached=1, return cached items only (no API calls)
    if request.args.get("cached"):
        cached = load_fbw_supplies_cache() or {}
        items = cached.get("items") or []
        # ???????? ????????? ("?? ?????????????") ?? ?????? API
        items = _filter_fbw_display_items(items)
        updated_at = cached.get("updated_at", "")
        next_offset = cached.get("next_offset", len(items))
        return jsonify({"items": items, "updated_at": updated_at, "next_offset": next_offset})

    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        # Refresh from WB: fetch first 15 or a subsequent page for load-more
        offset = int(request.args.get("offset", "0"))
        limit = int(request.args.get("limit", "15"))
        if offset <= 0:
            items = fetch_fbw_last_supplies(token, limit=limit)
            next_offset = limit
        else:
            # Always derive items from the same globally sorted list to avoid gaps
            items = fetch_fbw_supplies_range(token, offset=offset, limit=limit)
            next_offset = offset + limit
        # Merge cached package_count so ??? ?? ???????? ????? ????????????
        cached_for_user = load_fbw_supplies_cache() or {}
        cached_items = cached_for_user.get("items") or []
        items = _merge_package_counts(items, cached_items)
        
        # ?????????????? ????????? ?????????? ??????? ??? ???????? ??? ???? ??????????
        # ??? ?????? ?????? ??? ?????? ????????, ????? ?? ????????? ????????
        if offset <= 0:
            items = _preload_package_counts(token, items)

        # ???????? ????????? ("?? ?????????????") ?? ??????
        items = _filter_fbw_display_items(items)
        
        updated_at = datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y %H:%M")
        if offset <= 0:
            save_fbw_supplies_cache({"items": items, "updated_at": updated_at, "next_offset": next_offset})
        return jsonify({"items": items, "updated_at": updated_at, "next_offset": next_offset})
    except requests.HTTPError as http_err:
        return jsonify({"error": f"api_{http_err.response.status_code}"}), http_err.response.status_code
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


@app.route("/fbs-stock", methods=["GET"]) 
@login_required
def fbs_stock_page():
    token = effective_wb_api_token(current_user)
    error = None
    warehouses: list[dict[str, Any]] = []
    updated_at = ""
    # Show cache only; refresh by button
    cached = load_fbs_stock_cache() or {}
    if cached and cached.get("_user_id") == (current_user.id if current_user.is_authenticated else None):
        warehouses = cached.get("warehouses", []) or []
        updated_at = cached.get("updated_at", "")
    return render_template("fbs_stock.html", error=error, warehouses=warehouses, updated_at=updated_at)
@app.route("/api/fbs-stock/refresh", methods=["POST"]) 
@login_required
def api_fbs_stock_refresh():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        # 1) collect SKUs from products cache
        prod_cached = load_products_cache() or {}
        products = prod_cached.get("products") or prod_cached.get("items") or []
        # products expected structure should have barcode(s) and nm_id, vendor code, photo
        skus: list[str] = []
        barcode_to_info: dict[str, dict[str, Any]] = {}
        for p in products:
            bcs: list[str] = []
            if isinstance(p.get("barcodes"), list):
                bcs = [str(x) for x in p.get("barcodes") if x]
            elif p.get("barcode"):
                bcs = [str(p.get("barcode"))]
            for bc in bcs:
                skus.append(bc)
                barcode_to_info[bc] = {
                    "nm_id": p.get("nm_id") or p.get("nmId") or p.get("nmID"),
                    "vendor_code": p.get("??????? ????????") or p.get("vendor_code") or p.get("vendorCode") or p.get("supplierArticle") or p.get("supplier_article"),
                    "photo": p.get("photo") or p.get("img") or None,
                }
        if not skus:
            return jsonify({"error": "no_skus_in_products_cache"}), 400
        # 2) fetch warehouses and build summary totals
        wlist = fetch_fbs_warehouses(token)
        warehouses: list[dict[str, Any]] = []
        # Maps for human-readable labels
        cargo_labels = {
            1: "??? (??????????????)",
            2: "??? (???????????????)",
            3: "???+ (????????????????)",
        }
        delivery_labels = {
            1: "FBS",
            2: "DBS",
            3: "DBW",
            5: "C&C",
            6: "EDBS",
        }
        for w in wlist:
            wid = w.get("id") or w.get("warehouseId") or w.get("warehouseID")
            wname = w.get("name") or w.get("warehouseName") or ""
            try:
                stocks = fetch_fbs_stocks_by_warehouse(token, int(wid), skus)
                total_amount = sum(int(s.get("amount") or 0) for s in stocks)
            except Exception:
                total_amount = 0
            warehouses.append({
                "id": wid,
                "name": wname,
                "cargoType": cargo_labels.get(int(w.get("cargoType") or 0), "-"),
                "deliveryType": delivery_labels.get(int(w.get("deliveryType") or 0), "-"),
                "total_amount": total_amount,
            })
        updated_at = datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y %H:%M")
        save_fbs_stock_cache({"warehouses": warehouses, "updated_at": updated_at})
        return jsonify({"warehouses": warehouses, "updated_at": updated_at})
    except requests.HTTPError as http_err:
        return jsonify({"error": f"api_{http_err.response.status_code}"}), http_err.response.status_code
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


@app.route("/api/fbs-stock/update", methods=["POST"]) 
@login_required
def api_fbs_stock_update():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        payload = request.get_json(silent=True) or {}
        items = payload.get("stocks") or []
        warehouse_id_payload = payload.get("warehouseId")
        if not isinstance(items, list) or not items:
            return jsonify({"error": "no_items"}), 400
        # determine warehouse id
        if warehouse_id_payload:
            warehouse_id = warehouse_id_payload
        else:
            warehouses = fetch_fbs_warehouses(token)
            if not warehouses:
                return jsonify({"error": "warehouse_not_found"}), 404
            warehouse_id = warehouses[0].get("id") or warehouses[0].get("warehouseId") or warehouses[0].get("warehouseID")
        # update
        update_fbs_stocks_by_warehouse(token, int(warehouse_id), items)
        return jsonify({"ok": True})
    except requests.HTTPError as http_err:
        try:
            err_text = http_err.response.text
        except Exception:
            err_text = ""
        return jsonify({"error": f"api_{http_err.response.status_code}", "detail": err_text}), http_err.response.status_code
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


@app.route("/api/fbs-stock/warehouse/<int:warehouse_id>", methods=["GET"]) 
@login_required
def api_fbs_stock_by_warehouse(warehouse_id: int):
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        # collect SKUs from products cache
        prod_cached = load_products_cache() or {}
        products = prod_cached.get("products") or prod_cached.get("items") or []
        skus: list[str] = []
        barcode_to_info: dict[str, dict[str, Any]] = {}
        for p in products:
            bcs: list[str] = []
            if isinstance(p.get("barcodes"), list):
                bcs = [str(x) for x in p.get("barcodes") if x]
            elif p.get("barcode"):
                bcs = [str(p.get("barcode"))]
            for bc in bcs:
                skus.append(bc)
                barcode_to_info[bc] = {
                    "nm_id": p.get("nm_id") or p.get("nmId") or p.get("nmID"),
                    "vendor_code": p.get("??????? ????????") or p.get("vendor_code") or p.get("vendorCode") or p.get("supplierArticle") or p.get("supplier_article"),
                    "photo": p.get("photo") or p.get("img") or None,
                }
        if not skus:
            return jsonify({"error": "no_skus_in_products_cache"}), 400
        stocks = fetch_fbs_stocks_by_warehouse(token, int(warehouse_id), skus)
        rows: list[dict[str, Any]] = []
        for st in stocks:
            sku = str(st.get("sku") or "")
            amount = int(st.get("amount") or 0)
            info = barcode_to_info.get(sku, {})
            rows.append({
                "photo": info.get("photo"),
                "vendor_code": info.get("vendor_code") or "",
                "nm_id": info.get("nm_id") or "",
                "barcode": sku,
                "amount": amount,
            })
        total_amount = sum(r.get("amount", 0) for r in rows)
        return jsonify({"rows": rows, "total_amount": total_amount})
    except requests.HTTPError as http_err:
        return jsonify({"error": f"api_{http_err.response.status_code}"}), http_err.response.status_code
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


@app.route("/api/fbs-stock/auto-update/settings", methods=["GET", "POST"])
@login_required
def api_fbs_stock_auto_update_settings():
    if request.method == "GET":
        try:
            settings = load_auto_update_settings()
            return jsonify(settings)
        except Exception as exc:
            return jsonify({"error": str(exc)}), 500
    
    # POST - save settings
    try:
        payload = request.get_json(silent=True) or {}
        interval = int(payload.get("interval", 60))
        enabled = bool(payload.get("enabled", False))
        warehouse_settings = payload.get("warehouses", {})
        
        # Load current settings
        settings = load_auto_update_settings()
        
        # Update global settings
        settings["global"] = {
            "interval": interval,
            "enabled": enabled,
            "lastCheck": settings["global"].get("lastCheck"),
            "history": settings["global"].get("history", [])
        }
        
        # Update warehouse-specific settings
        if warehouse_settings:
            settings["warehouses"] = warehouse_settings
        
        save_auto_update_settings(settings)
        return jsonify({"ok": True})
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500
@app.route("/api/fbs-stock/auto-update/test", methods=["POST"])
@login_required
def api_fbs_stock_auto_update_test():
    try:
        payload = request.get_json(silent=True) or {}
        url = payload.get("url", "").strip()
        
        if not url:
            return jsonify({"error": "no_url"}), 400
        
        # Test connection and get file info
        file_info = test_remote_file(url)
        return jsonify(file_info)
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/fbs-stock/auto-update/manual-update", methods=["POST"])
@login_required
def api_fbs_stock_auto_update_manual():
    print("=== MANUAL UPDATE ENDPOINT CALLED ===")
    try:
        payload = request.get_json(silent=True) or {}
        warehouses = payload.get("warehouses", [])
        
        print(f"Payload: {payload}")
        print(f"Warehouses: {warehouses}")
        
        if not warehouses:
            print("ERROR: No warehouses provided")
            return jsonify({"error": "no_warehouses"}), 400
        
        user_id = current_user.id if current_user.is_authenticated else None
        print(f"User ID: {user_id}")
        
        if not user_id:
            print("ERROR: No user ID")
            return jsonify({"error": "no_user"}), 401
        
        print(f"Manual update requested for user {user_id}, warehouses: {len(warehouses)}")
        
        total_processed = 0
        total_updated = 0
        results = []
        
        # Process each warehouse individually
        for warehouse in warehouses:
            warehouse_id = warehouse.get('warehouseId')
            url = warehouse.get('url', '').strip()
            
            if not url:
                print(f"ERROR: No URL provided for warehouse {warehouse_id}")
                results.append(f"warehouse {warehouse_id}: no URL")
                continue
            
            print(f"Processing warehouse {warehouse_id} with URL: {url}")
            
            # Download and process file
            print(f"Starting file download and processing for warehouse {warehouse_id}...")
            result = download_and_process_remote_file(url, user_id)
            print(f"File processing result for warehouse {warehouse_id}: {result}")
            
            if result['success']:
                print(f"File data for warehouse {warehouse_id}: {len(result['data'])} items")
                
                # Update stocks for this specific warehouse
                print(f"Calling update_stocks_from_remote_data for warehouse {warehouse_id}...")
                updated_count = update_stocks_from_remote_data(result['data'], user_id, [int(warehouse_id)])
                print(f"Updated count for warehouse {warehouse_id}: {updated_count}")
                
                total_processed += len(result['data'])
                total_updated += updated_count
                results.append(
                    f"warehouse {warehouse_id}: ok, rows={len(result['data'])}, stock_updates={updated_count}"
                )
            else:
                print(f"File processing failed for warehouse {warehouse_id}: {result['error']}")
                results.append(f"warehouse {warehouse_id}: error - {result['error']}")
        
        # Add to history
        settings = load_auto_update_settings()
        history_entry = {
            "timestamp": datetime.now().isoformat(),
            "success": True,
            "message": (
                f"Manual update: processed {total_processed} file rows, "
                f"updated {total_updated} stock rows (remote FBS file)"
            ),
        }
        settings['global']['history'].insert(0, history_entry)
        settings['global']['history'] = settings['global']['history'][:50]  # Keep last 50 entries
        settings['global']['lastCheck'] = datetime.now().isoformat()
        save_auto_update_settings(settings)
        
        print(f"Returning success: processed={total_processed}, updated={total_updated}")
        return jsonify({
            "processed": total_processed,
            "updated": total_updated,
            "details": results
        })

            
    except Exception as exc:
        print(f"Manual update error: {exc}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": str(exc)}), 500


@app.route("/api/fbw/supplies/<supply_id>/details", methods=["GET"]) 
@login_required
def api_fbw_supply_details(supply_id: str):
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        details = fetch_fbw_supply_details(token, supply_id) or {}
        goods = fetch_fbw_supply_goods(token, supply_id, limit=200, offset=0)
        packages = fetch_fbw_supply_packages(token, supply_id)
        return jsonify({"details": details, "goods": goods, "packages": packages})
    except requests.HTTPError as http_err:
        return jsonify({"error": f"api_{http_err.response_status}"}), 500
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


@app.route("/api/supplies/refresh-cache", methods=["POST"]) 
@login_required
def api_refresh_supplies_cache():
    """?????? ?????????? ???? ????????"""
    global _supplies_cache_updating
    
    user_id = current_user.id
    
    # ?????????, ?? ???? ?? ??? ?????????? ??? ????? ????????????
    if _supplies_cache_updating.get(user_id, False):
        return jsonify({
            "error": "??? ???????? ??? ???????????. ??????????, ????????? ?????????? ???????? ????????."
        }), 409
    
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        # ???? ??? ???? ?????????? ? ?? ???????? ??????, ????? ????????
        if _supplies_cache_updating.get(user_id):
            return jsonify({
                "success": True,
                "message": "?????????? ??? ???????????",
                "in_progress": True,
            })

        # ????????????? ???? ?????????? ??? ????? ????????????
        _supplies_cache_updating[user_id] = True

        # ????????? ?????????? ???? ? ??????? ??????
        import threading

        def build_cache_background():
            global _supplies_cache_updating
            try:
                has_cache = bool(load_fbw_supplies_detailed_cache(user_id))
                cache_data = build_supplies_detailed_cache(
                    token,
                    user_id,
                    batch_size=10,           # ?????? ?????
                    pause_seconds=2.0,       # ??????? ?????
                    force_full=not has_cache,
                    days_back=(180 if not has_cache else 10),
                )
                save_fbw_supplies_detailed_cache(cache_data, user_id)
                print(f"??? ???????? ??????? ???????? ??? ???????????? {user_id}")
            except Exception as e:
                print(f"?????? ??? ?????????? ???? ????????: {e}")
            finally:
                # ?????????? ???? ?????????? ??? ????? ????????????
                _supplies_cache_updating[user_id] = False

        thread = threading.Thread(target=build_cache_background)
        thread.daemon = True
        thread.start()

        # ????? ?????????? ?????????? ?????????? ????, ????? ????? ??? ????????
        cached_meta = load_fbw_supplies_detailed_cache(user_id) or {}
        return jsonify({
            "success": True,
            "message": "?????????? ???? ???????? ???????? ? ????",
            "in_progress": True,
            "total_supplies": cached_meta.get("total_supplies_processed", 0),
            "last_updated": cached_meta.get("last_updated"),
        })
    except Exception as exc:
        _supplies_cache_updating[user_id] = False  # ?????????? ???? ? ?????? ??????
        return jsonify({"error": str(exc)}), 500


@app.route("/api/orders/refresh-cache", methods=["POST"]) 
@login_required
def api_refresh_orders_cache():
    global _orders_cache_updating
    
    user_id = current_user.id
    
    # ?????????, ?? ???? ?? ??? ?????????? ??? ????? ????????????
    if _orders_cache_updating.get(user_id, False):
        return jsonify({
            "error": "??? ??????? ??? ???????????. ??????????, ????????? ?????????? ???????? ????????."
        }), 409
    
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        # ????????????? ???? ?????????? ??? ????? ????????????
        _orders_cache_updating[user_id] = True
        
        # ????????? ?????????? ???? ? ??????? ??????
        import threading
        
        def build_orders_cache_background():
            global _orders_cache_updating
            try:
                meta = build_orders_warm_cache(token, user_id)
                save_orders_cache_meta(meta, user_id)
                print(f"??? ??????? ??????? ???????? ??? ???????????? {user_id}")
            except Exception as e:
                print(f"?????? ??? ?????????? ???? ???????: {e}")
            finally:
                # ?????????? ???? ?????????? ??? ????? ????????????
                _orders_cache_updating[user_id] = False
        
        thread = threading.Thread(target=build_orders_cache_background)
        thread.daemon = True
        thread.start()
        
        return jsonify({
            "success": True,
            "message": "?????????? ???? ??????? ???????? ? ??????? ??????. ??? ????? ?????? ????????? ?????.",
            "total_orders": 0,  # ?????????? 0, ??? ??? ??????? ??? ????
            "last_updated": datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y %H:%M")
        })
    except Exception as exc:
        _orders_cache_updating[user_id] = False  # ?????????? ???? ? ?????? ??????
        return jsonify({"error": str(exc)}), 500


@app.route("/api/cache/status", methods=["GET"])
@login_required
def api_cache_status():
    """???????? ??????? ?????????? ????"""
    user_id = current_user.id
    return jsonify({
        "supplies_cache_updating": _supplies_cache_updating.get(user_id, False),
        "orders_cache_updating": _orders_cache_updating.get(user_id, False)
    })
@app.route("/api/fbw/supplies/<supply_id>/package-count", methods=["GET"]) 
@login_required
def api_fbw_supply_package_count(supply_id: str):
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        packages = fetch_fbw_supply_packages(token, supply_id)
        count = len(packages) if isinstance(packages, list) else 0
        # Try update cache to persist the count
        try:
            cached = load_fbw_supplies_cache() or {}
            items = cached.get("items") or []
            changed = False
            for it in items:
                sid = str(it.get("supply_id") or it.get("supplyID") or it.get("supplyId") or it.get("id") or "")
                if sid == str(supply_id):
                    it["package_count"] = int(count)
                    changed = True
                    break
            if changed:
                # Keep updated_at and next_offset intact
                save_fbw_supplies_cache({
                    "items": items,
                    "updated_at": cached.get("updated_at", ""),
                    "next_offset": cached.get("next_offset", 0),
                })
        except Exception:
            pass
        return jsonify({"package_count": int(count)})
    except requests.HTTPError as http_err:
        return jsonify({"error": f"api_{http_err.response_status}"}), 500
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500
@app.route("/api/orders-refresh", methods=["POST"]) 
@login_required
def api_orders_refresh():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    date_from = request.form.get("date_from", "")
    date_to = request.form.get("date_to", "")
    print(f"API orders-refresh: received date_from='{date_from}', date_to='{date_to}'")
    try:
        df = parse_date(date_from)
        dt = parse_date(date_to)
        print(f"API orders-refresh: parsed dates df={df}, dt={dt}")
        if df > dt:
            date_from, date_to = date_to, date_from
            print(f"API orders-refresh: swapped inverted range to date_from='{date_from}', date_to='{date_to}'")
    except ValueError as e:
        print(f"API orders-refresh: date parse error: {e}")
        return jsonify({"error": "bad_dates"}), 400
    try:
        # Orders
        force_refresh = request.form.get("force_refresh") is not None
        
        if force_refresh:
            # ?????????????? ?????????? - ????????? ??? ?????? ????? API, ????????? ???
            raw_orders = fetch_orders_range(token, date_from, date_to)
            orders = to_rows(raw_orders, date_from, date_to)
            total_orders = len(orders)
            total_active_orders = len([o for o in orders if not o.get("is_cancelled", False)])
            total_cancelled_orders = len([o for o in orders if o.get("is_cancelled", False)])
            total_revenue = round(sum(float(o.get("???? ?? ??????? ????????") or 0) for o in orders if not o.get("is_cancelled", False)), 2)
            # ????????? ??? ?????????????
            _update_period_cache_with_data(token, date_from, date_to, orders)
            meta = {"used_cache_days": 0, "fetched_days": len(_daterange_inclusive(date_from, date_to))}
        else:
            # ??????? ?????????? - ?????????? ??? ?? ????
            orders, meta = get_orders_with_period_cache(
                token, date_from, date_to
            )
            total_orders = len(orders)
            total_active_orders = len([o for o in orders if not o.get("is_cancelled", False)])
            total_cancelled_orders = len([o for o in orders if o.get("is_cancelled", False)])
            total_revenue = round(sum(float(o.get("???? ?? ??????? ????????") or 0) for o in orders if not o.get("is_cancelled", False)), 2)
        # Aggregates
        o_counts_map, o_rev_map, o_cancelled_counts_map = aggregate_daily_counts_and_revenue(orders)
        daily_labels = sorted(k for k in o_counts_map if k)
        daily_orders_counts = [o_counts_map.get(d, 0) for d in daily_labels]
        daily_orders_cancelled_counts = [o_cancelled_counts_map.get(d, 0) for d in daily_labels]
        daily_orders_revenue = [round(o_rev_map.get(d, 0.0), 2) for d in daily_labels]
        # Warehouses and TOPs
        warehouse_summary_dual = aggregate_by_warehouse_orders_only(orders)
        top_products = aggregate_top_products(orders, limit=15)
        top_mode = "orders"
        warehouses = sorted({_order_row_warehouse_label(r) for r in orders})
        top_products_orders_filtered = aggregate_top_products_orders(orders, None, limit=50)
        updated_at = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        # Save last results snapshot
        save_last_results({
            "date_from": date_from,
            "date_to": date_to,
            "orders": orders,
            "total_orders": total_orders,
            "total_active_orders": total_active_orders,
            "total_cancelled_orders": total_cancelled_orders,
            "total_revenue": total_revenue,
            "daily_labels": daily_labels,
            "daily_orders_counts": daily_orders_counts,
            "daily_orders_cancelled_counts": daily_orders_cancelled_counts,
            "daily_orders_revenue": daily_orders_revenue,
            "warehouse_summary_dual": warehouse_summary_dual,
            "top_products": top_products,
            "top_mode": top_mode,
            "updated_at": updated_at,
        })
        resp = {
            "total_orders": total_orders,
            "total_active_orders": total_active_orders,
            "total_cancelled_orders": total_cancelled_orders,
            "total_revenue": total_revenue,
            "daily_labels": daily_labels,
            "daily_orders_counts": daily_orders_counts,
            "daily_orders_cancelled_counts": daily_orders_cancelled_counts,
            "daily_orders_revenue": daily_orders_revenue,
            "warehouse_summary_dual": warehouse_summary_dual,
            "top_products": top_products,
            "warehouses": list(warehouses),
            "top_products_orders_filtered": top_products_orders_filtered,
            "updated_at": updated_at,
            "date_from_fmt": format_dmy(date_from),
            "date_to_fmt": format_dmy(date_to),
            "top_mode": top_mode,
        }
        try:
            resp["cache"] = {"used_cache_days": meta.get("used_cache_days", 0), "fetched_days": meta.get("fetched_days", 0)}
        except Exception:
            pass
        return jsonify(resp)
    except requests.HTTPError as http_err:
        status = 502
        try:
            if http_err.response is not None and http_err.response.status_code:
                status = http_err.response.status_code
        except Exception:
            status = 502
        # Graceful fallback for 429: return cached data if dates match
        if status == 429:
            cached = load_last_results() or {}
            if (
                cached.get("date_from") == date_from
                and cached.get("date_to") == date_to
                and cached.get("_user_id") == (current_user.id if current_user.is_authenticated else None)
            ):
                return jsonify({
                    "total_orders": cached.get("total_orders", 0),
                    "total_revenue": cached.get("total_revenue", 0),
                    "daily_labels": cached.get("daily_labels", []),
                    "daily_orders_counts": cached.get("daily_orders_counts", []),
                    "daily_sales_counts": cached.get("daily_sales_counts", []),
                    "daily_orders_revenue": cached.get("daily_orders_revenue", []),
                    "daily_sales_revenue": cached.get("daily_sales_revenue", []),
                    "warehouse_summary_dual": cached.get("warehouse_summary_dual", []),
                    "top_products": cached.get("top_products", []),
                    "warehouses": cached.get("warehouses", []),
                    "top_products_orders_filtered": cached.get("top_products_orders_filtered", []),
                    "updated_at": cached.get("updated_at", ""),
                    "date_from_fmt": format_dmy(date_from),
                    "date_to_fmt": format_dmy(date_to),
                    "top_mode": cached.get("top_mode", "orders"),
                    "rate_limited": True,
                }), 200
        return jsonify({"error": "http", "status": status}), status
    except Exception as exc:  # noqa: BLE001
        return jsonify({"error": str(exc)}), 500


@app.route("/api/orders-progress", methods=["GET"]) 
@login_required
def api_orders_progress():
    try:
        uid = current_user.id
        prog = ORDERS_PROGRESS.get(uid) or {"total": 0, "done": 0}
        return jsonify({"total": int(prog.get("total", 0)), "done": int(prog.get("done", 0))}), 200
    except Exception as exc:
        return jsonify({"total": 0, "done": 0, "error": str(exc)}), 200

@app.route("/profile", methods=["GET"]) 
@login_required
def profile():
    seller_info: Dict[str, Any] | None = None
    token_info: Dict[str, Any] | None = None
    supplies_cache_info: Dict[str, Any] | None = None
    orders_cache_info: Dict[str, Any] | None = None
    token = (current_user.wb_token or "").strip()
    token_info = decode_token_info(token) if token else None
    with _seller_info_cache_lock:
        cached = _load_seller_info_cache(current_user.id)
    if cached and cached.get("organization_name"):
        seller_info = cached.get("seller_info") or {
            "name": cached.get("organization_name"),
            "companyName": cached.get("organization_name"),
            "supplierName": cached.get("organization_name"),
        }
    if token:
        try:
            # ?????????? ? ???? ???????? (?????????? ????????)
            supplies_cache_info = _get_light_supplies_cache_info(current_user.id)
            
            # ?????????? ? ???? ??????? (?????????? ????????)
            orders_cache_info = None
            try:
                orders_meta = load_orders_cache_meta()
                if orders_meta:
                    orders_cache_info = {
                        "last_updated": orders_meta.get("last_updated"),
                        "date_from": orders_meta.get("date_from"),
                        "date_to": orders_meta.get("date_to"),
                        "total_orders_cached": orders_meta.get("total_orders_cached", 0),
                        "is_fresh": is_orders_cache_fresh(),
                        "cache_version": orders_meta.get("cache_version", "1.0")
                    }
            except Exception as e:
                print(f"?????? ???????? ???? ???????: {e}")
                orders_cache_info = None
        except Exception:
            supplies_cache_info = None
            orders_cache_info = None
    validity_status = None
    if current_user.valid_from or current_user.valid_to:
        today = datetime.now(MOSCOW_TZ).date()
        active = True
        if current_user.valid_from and today < current_user.valid_from:
            active = False
        if current_user.valid_to and today > current_user.valid_to:
            active = False
        validity_status = "active" if active and current_user.is_active else "inactive"
    return render_template(
        "profile.html",
        message=None,
        token=token,
        seller_info=seller_info,
        token_info=token_info,
        orders_cache_info=orders_cache_info,
        supplies_cache_info=supplies_cache_info,
        valid_from=current_user.valid_from.strftime("%d.%m.%Y") if current_user.valid_from else None,
        valid_to=current_user.valid_to.strftime("%d.%m.%Y") if current_user.valid_to else None,
        validity_status=validity_status,
    )


@app.route("/api/profile/organization/refresh", methods=["POST"])
@login_required
def api_profile_organization_refresh():
    """Force-refresh organization name from WB API for current user."""
    token_api = effective_wb_api_token(current_user)
    if not token_api:
        return jsonify({"success": False, "error": "????? WB API ??????????? ??? ??????????????"}), 400

    now_ts = time.time()
    last_attempt = float(_seller_info_last_attempt_at.get(current_user.id, 0) or 0)
    if now_ts - last_attempt < _SELLER_INFO_MIN_INTERVAL_SECONDS:
        wait_sec = int(_SELLER_INFO_MIN_INTERVAL_SECONDS - (now_ts - last_attempt))
        return jsonify({"success": False, "error": f"??????? ?????. ????????? ????? {max(1, wait_sec)} ???."}), 429
    _seller_info_last_attempt_at[current_user.id] = now_ts

    try:
        resp = requests.get(
            SELLER_INFO_URL,
            headers={"Authorization": f"Bearer {token_api}"},
            timeout=15,
        )
    except Exception as exc:
        return jsonify({"success": False, "error": f"?????? ??????? ? WB API: {exc}"}), 502

    if resp.status_code >= 400:
        if resp.status_code == 429:
            _seller_info_rate_limited_until[current_user.id] = time.time() + _SELLER_INFO_RATE_LIMIT_COOLDOWN_SECONDS
        return jsonify({"success": False, "error": f"{resp.status_code}: {resp.text[:300]}"}), resp.status_code

    try:
        seller_info = resp.json()
    except Exception:
        return jsonify({"success": False, "error": f"???????????? JSON: {resp.text[:300]}"}), 502

    if not isinstance(seller_info, dict):
        return jsonify({"success": False, "error": "WB API ?????? ??????/??????????? ?????", "raw": seller_info}), 502

    organization_name = (
        seller_info.get("name")
        or seller_info.get("companyName")
        or seller_info.get("supplierName")
    )
    if not organization_name:
        return jsonify({
            "success": False,
            "error": "? ?????? WB API ??? ????? name/companyName/supplierName",
            "keys": sorted(list(seller_info.keys())),
        }), 502

    with _seller_info_cache_lock:
        _save_seller_info_cache(
            current_user.id,
            {
                "organization_name": organization_name,
                "seller_info": seller_info,
                "cached_at": time.time(),
            },
        )

    try:
        if not (getattr(current_user, "org_display_name", None) or "").strip():
            current_user.org_display_name = (organization_name or "")[:255] or None
            db.session.commit()
    except Exception:
        db.session.rollback()

    return jsonify({
        "success": True,
        "organization_name": organization_name,
        "keys": sorted(list(seller_info.keys())),
    }), 200


@app.route("/profile/token", methods=["POST"]) 
@login_required
def profile_token():
    new_token = request.form.get("token", "").strip()
    try:
        current_user.wb_token = new_token or None
        db.session.commit()
        # Fetch seller info once after clicking "Save" with a token.
        # Internal rate-limit guards prevent excessive requests.
        if new_token:
            _start_seller_info_refresh_bg(current_user.id, new_token)
        if new_token:
            hint = []
            if not (current_user.phone and current_user.email and current_user.shipper_address):
                hint.append(" ????????? ???????, email ? ????? ?????? ??? ???????? ? ???????.")
            flash("????? ??????? ????????." + (hint[0] if hint else ""))
        else:
            flash("????? ??????")
    except Exception:
        db.session.rollback()
        flash("?????? ?????????? ??????")
    return redirect(url_for("profile"))


@app.route("/profile/org-display-name", methods=["POST"])
@login_required
def profile_org_display_name():
    """????????? ???????? ??????????? ??? ????? (??? ??????? ? WB API)."""
    val = (request.form.get("org_display_name") or "").strip()
    try:
        current_user.org_display_name = val or None
        db.session.commit()
        flash("???????? ? ????? ?????????" if val else "???????? ? ????? ????????")
    except Exception:
        db.session.rollback()
        flash("?????? ?????????? ????????")
    return redirect(url_for("profile"))


@app.route("/profile/shipping", methods=["POST"]) 
@login_required
def profile_shipping():
    current_user.shipper_name = (request.form.get("shipper_name") or "").strip() or None
    current_user.contact_person = (request.form.get("contact_person") or "").strip() or None
    current_user.phone = (request.form.get("phone") or "").strip() or None
    current_user.email = (request.form.get("email") or "").strip() or None
    current_user.shipper_address = (request.form.get("shipper_address") or "").strip() or None
    try:
        db.session.commit()
        flash("????????? ?????????")
    except Exception:
        db.session.rollback()
        flash("?????? ?????????? ??????????")
    return redirect(url_for("profile"))


@app.route("/profile/tax-rate", methods=["POST"]) 
@login_required
def profile_tax_rate():
    tax_rate_str = request.form.get("tax_rate", "").strip()
    try:
        if tax_rate_str:
            tax_rate = float(tax_rate_str)
            if tax_rate < 0 or tax_rate > 100:
                flash("????????? ?????? ?????? ???? ?? 0 ?? 100%")
                return redirect(url_for("profile"))
            current_user.tax_rate = tax_rate
        else:
            current_user.tax_rate = None
        db.session.commit()
        flash("????????? ?????? ?????????")
    except ValueError:
        flash("??????: ???????? ???????? ????????? ??????")
    except Exception:
        db.session.rollback()
        flash("?????? ?????????? ????????? ??????")
    return redirect(url_for("profile"))


@app.route("/profile/password", methods=["POST"]) 
@login_required
def profile_password():
    old_password = (request.form.get("old_password", "") or "").strip()
    new_password = (request.form.get("new_password", "") or "").strip()
    if not old_password or not new_password:
        flash("????????? ??? ????")
        return redirect(url_for("profile"))
    if current_user.password != old_password:
        flash("??????? ?????? ???????")
        return redirect(url_for("profile"))
    if len(new_password) < 4:
        flash("????? ?????? ??????? ???????? (???. 4 ???????)")
        return redirect(url_for("profile"))
    if new_password == old_password:
        flash("????? ?????? ????????? ? ???????")
        return redirect(url_for("profile"))
    try:
        current_user.password = new_password
        db.session.commit()
        flash("?????? ????????")
    except Exception:
        db.session.rollback()
        flash("?????? ?????????? ??????")
    return redirect(url_for("profile"))


@app.route("/export", methods=["POST"]) 
@login_required
def export_excel():
    token = token_for_wb_request(current_user, request.form.get("token"))
    date_from = request.form.get("date_from", "")
    date_to = request.form.get("date_to", "")

    if not token or not date_from or not date_to:
        return render_template(
            "index.html",
            error="??? ???????? ??????? ????? ? ????",
            token=token,
            date_from=date_from,
            date_to=date_to,
            orders=[],
            total_orders=0,
            total_revenue=0,
            daily_labels=[],
            daily_counts=[],
            daily_revenue=[],
            warehouse_summary=[],
            top_products=[],
        )

    # Normalize and validate dates
    try:
        df = parse_date(date_from)
        dt = parse_date(date_to)
        if df > dt:
            date_from, date_to = date_to, date_from
    except ValueError:
        return render_template(
            "index.html",
            error="???????? ?????? ???",
            token=token,
            date_from=date_from,
            date_to=date_to,
            orders=[],
            total_orders=0,
            total_revenue=0,
            daily_labels=[],
            daily_counts=[],
            daily_revenue=[],
            warehouse_summary=[],
            top_products=[],
        )

    # ????????? ??? ????? ???????? ? API, ????? ???????? ?????? 429
    cached = load_last_results()
    if (
        cached
        and current_user.is_authenticated
        and cached.get("_user_id") == current_user.id
        and cached.get("date_from") == date_from
        and cached.get("date_to") == date_to
    ):
        # ?????????? ?????? ?? ???? (??? ??? ?????????? ????? to_rows)
        rows = cached.get("orders", [])
    else:
        # ???? ??? ??????????? ??? ?? ?????????, ?????? ?????? ? API
        try:
            raw_data = fetch_orders_range(token, date_from, date_to)
            rows = to_rows(raw_data, date_from, date_to)
        except requests.HTTPError as http_err:
            # ???? ???????? ?????? 429 ??? ??????, ???????? ???????????? ??? ??? ????????? ???????
            if cached and current_user.is_authenticated and cached.get("_user_id") == current_user.id:
                rows = cached.get("orders", [])
                if not rows:
                    return render_template(
                        "index.html",
                        error=f"?????? API ??? ???????? (HTTP {http_err.response.status_code}). ??? ??????????. ?????????? ?????.",
                        token=token,
                        date_from=date_from,
                        date_to=date_to,
                        orders=[],
                        total_orders=0,
                        total_revenue=0,
                        daily_labels=[],
                        daily_counts=[],
                        daily_revenue=[],
                        warehouse_summary=[],
                        top_products=[],
                    )
            else:
                return render_template(
                    "index.html",
                    error=f"?????? API ??? ???????? (HTTP {http_err.response.status_code}). ????????? ?????? ?? ???????? ????????? ???????.",
                    token=token,
                    date_from=date_from,
                    date_to=date_to,
                    orders=[],
                    total_orders=0,
                    total_revenue=0,
                    daily_labels=[],
                    daily_counts=[],
                    daily_revenue=[],
                    warehouse_summary=[],
                    top_products=[],
                )

    wb = Workbook()
    ws = wb.active
    ws.title = "orders"

    if rows:
        headers = list(rows[0].keys())
        ws.append(headers)
        for r in rows:
            ws.append([r.get(h, "") for h in headers])
    else:
        ws.append(["??? ??????"])

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)

    filename = f"wb_orders_{date_from}_{date_to}.xlsx"
    return send_file(output, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


@app.route("/api/top-products-sales", methods=["GET"]) 
@login_required
def api_top_products_sales():
    warehouse = request.args.get("warehouse", "") or None
    cached = load_last_results()
    if not cached or not (current_user.is_authenticated and cached.get("_user_id") == current_user.id):
        return jsonify({"items": []})
    sales_rows = cached.get("sales_rows", [])
    items = aggregate_top_products_sales(sales_rows, warehouse, limit=50)
    return jsonify({"items": items})


@app.route("/api/top-products-orders", methods=["GET"]) 
@login_required
def api_top_products_orders():
    warehouse = request.args.get("warehouse", "") or None
    cached = load_last_results()
    if not cached or not (current_user.is_authenticated and cached.get("_user_id") == current_user.id):
        return jsonify({"items": [], "total_qty": 0, "total_sum": 0})
    orders = cached.get("orders", [])
    items = aggregate_top_products_orders(orders, warehouse, limit=50)
    
    # ???????????? ????? ????? ??? ?????????? ??????
    total_qty = 0
    total_sum = 0.0
    for item in items:
        total_qty += item.get("qty", 0)
        total_sum += item.get("sum", 0.0)
    
    return jsonify({
        "items": items,
        "total_qty": total_qty,
        "total_sum": round(total_sum, 2)
    })
@app.route("/report/sales", methods=["GET"]) 
@login_required
def report_sales_page():
    cached = load_last_results()
    token = effective_wb_api_token(current_user)
    
    # ????????? ??????? ???? ????? (???? ??? ???????)
    if token and current_user.is_authenticated:
        _start_stocks_update_bg(current_user.id, token)
    
    # ???????? ?? ????????? ??????????? ??????: ??? ??????, ???? ???????????? ?? ?????? ?????? ? ?? ?????? ?????????
    if not request.args.get("date_from") and not request.args.get("date_to"):
        return render_template(
            "report_sales.html",
            error=None,
            items=[],
            date_from_fmt="",
            date_to_fmt="",
            warehouse=None,
            warehouses=[],
            date_from_val="",
            date_to_val="",
        ), 200

    # Accept date range from query params; fallback to cached
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    warehouse = request.args.get("warehouse") or None

    if req_from and req_to and token:
        # Prefer cache if it matches the requested period and belongs to the user
        if (
            cached
            and current_user.is_authenticated
            and cached.get("_user_id") == current_user.id
            and cached.get("date_from") == req_from
            and cached.get("date_to") == req_to
        ):
            orders = cached.get("orders", [])
            try:
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
            except Exception:
                date_from_fmt = cached.get("date_from_fmt") or req_from
                date_to_fmt = cached.get("date_to_fmt") or req_to
        else:
            # Fetch fresh orders only if cache doesn't match
            try:
                raw_orders = fetch_orders_range(token, req_from, req_to)
                orders = to_rows(raw_orders, req_from, req_to)
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
            except Exception as exc:
                # Fallback to cache on error
                orders = (cached or {}).get("orders", [])
                date_from_fmt = (cached or {}).get("date_from_fmt") or (cached or {}).get("date_from")
                date_to_fmt = (cached or {}).get("date_to_fmt") or (cached or {}).get("date_to")
    else:
        # ??? ?????? ??????? ?? ??????????? ??? ? ???????? ???????? ??????
        orders = []
        date_from_fmt = ""
        date_to_fmt = ""

    warehouses = sorted({_order_row_warehouse_label(r) for r in orders}) if orders else []
    # Build matrix for client-side filtering (same as API)
    counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product = _normalize_and_group_orders(orders)
    # build photo map from products cache
    nm_to_photo: Dict[Any, Any] = {}
    try:
        prod_cached = load_products_cache() or {}
        for it in (prod_cached.get("items") or []):
            nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
            photo = it.get("photo") or it.get("img")
            if nmv is not None and nmv not in nm_to_photo:
                nm_to_photo[nmv] = photo
    except Exception:
        nm_to_photo = {}

    # Load stocks data for current user - ????????? ??????? ?? ???????
    stocks_by_warehouse = {}
    stocks_metadata = {}  # ?????????????? ?????????? ? ??????? ?? ????????
    try:
        stocks_cached = load_stocks_cache()
        # print(f"DEBUG: stocks_cached loaded: {stocks_cached is not None}")
        if stocks_cached and stocks_cached.get("_user_id"):
            # print(f"DEBUG: stocks_cached user_id: {stocks_cached.get('_user_id')}, current_user.id: {current_user.id if current_user.is_authenticated else 'not authenticated'}")
            items = stocks_cached.get("items", [])
            # print(f"DEBUG: stocks items count: {len(items)}")
            for stock_item in items:
                barcode = stock_item.get("barcode")
                stock_warehouse = stock_item.get("warehouse", "")
                qty = int(stock_item.get("qty", 0) or 0)
                vendor_code = stock_item.get("vendor_code", "")
                nm_id = stock_item.get("nm_id")
                
                if barcode:
                    if barcode not in stocks_by_warehouse:
                        stocks_by_warehouse[barcode] = {}
                    if barcode not in stocks_metadata:
                        stocks_metadata[barcode] = {
                            "vendor_code": vendor_code,
                            "nm_id": nm_id,
                            "barcode": barcode
                        }
                    
                    if stock_warehouse:
                        stocks_by_warehouse[barcode][stock_warehouse] = stocks_by_warehouse[barcode].get(stock_warehouse, 0) + qty
            # print(f"DEBUG: stocks_by_warehouse loaded: {len(stocks_by_warehouse)} barcodes")
        # else:
            # print("DEBUG: No stocks cache or wrong user")
    except Exception as e:
        # print(f"DEBUG: Error loading stocks: {e}")
        stocks_by_warehouse = {}
        stocks_metadata = {}

    def _build_items(target_wh: str | None, show_all: bool = False) -> List[Dict[str, Any]]:
        items_local: List[Dict[str, Any]] = []
        
        # Get all products that have sales or stocks
        all_products = set(counts_total.keys())
        if show_all:
            # Add ALL products from stocks metadata (including those with zero stock)
            for barcode, metadata in stocks_metadata.items():
                # Find product by barcode
                found_in_sales = False
                for prod, prod_barcode in barcode_by_product.items():
                    if prod_barcode == barcode:
                        all_products.add(prod)
                        found_in_sales = True
                        break
                
                # If barcode not found in sales, create a virtual product entry
                if not found_in_sales:
                    # Use vendor_code from stocks metadata
                    virtual_prod = metadata["vendor_code"] or f"????? ? ???????? {barcode}"
                    # Add to mappings
                    barcode_by_product[virtual_prod] = barcode
                    if metadata["nm_id"]:
                        nm_by_product[virtual_prod] = metadata["nm_id"]
                    if metadata["vendor_code"]:
                        supplier_article_by_product[virtual_prod] = metadata["vendor_code"]
                    all_products.add(virtual_prod)
            # print(f"DEBUG: show_all=True, total products with sales: {len(counts_total)}, total products with stocks: {len(stocks_by_warehouse)}, all_products: {len(all_products)}")
        
        for prod in all_products:
            qty = (by_wh.get(prod, {}).get(target_wh, 0) if target_wh else counts_total.get(prod, 0))
            
            # Include items with sales OR (if show_all) items with stocks
            if qty > 0 or (show_all and prod in barcode_by_product):
                s = (by_wh_sum.get(prod, {}).get(target_wh, 0.0) if target_wh else revenue_total.get(prod, 0.0))
                # Calculate stock quantity for the target warehouse
                barcode = barcode_by_product.get(prod)
                stock_qty = 0
                if barcode and barcode in stocks_by_warehouse:
                    if target_wh:
                        # If specific warehouse selected, sum only for that warehouse
                        for wh_name, wh_qty in stocks_by_warehouse[barcode].items():
                            if (wh_name == target_wh or 
                                (target_wh in wh_name) or 
                                (wh_name in target_wh)):
                                stock_qty += wh_qty
                    else:
                        # If no warehouse selected, sum all warehouses
                        stock_qty = sum(stocks_by_warehouse[barcode].values())
                
                # Only include if has sales or (if show_all) has any stock data
                if qty > 0 or (show_all and prod in barcode_by_product):
                    items_local.append({
                        "product": prod,
                        "qty": qty,
                        "nm_id": nm_by_product.get(prod),
                        "barcode": barcode,
                        "supplier_article": supplier_article_by_product.get(prod),
                        "sum": round(float(s or 0.0), 2),
                        "photo": nm_to_photo.get(nm_by_product.get(prod)),
                        "stock_qty": stock_qty
                    })
        
        # Sort by quantity (descending), then by stock quantity (descending)
        items_local.sort(key=lambda x: (x["qty"], x["stock_qty"]), reverse=True)
        # print(f"DEBUG: _build_items returning {len(items_local)} items, show_all={show_all}")
        return items_local
    show_all = request.args.get("show_all_products") == "on"
    items = _build_items(warehouse, show_all) if orders else []
    matrix = [{
        "product": p,
        "nm_id": nm_by_product.get(p),
        "barcode": barcode_by_product.get(p),
        "supplier_article": supplier_article_by_product.get(p),
        "total": counts_total[p],
        "by_wh": by_wh[p],
        "total_sum": round(float(revenue_total.get(p, 0.0)), 2),
        "by_wh_sum": by_wh_sum[p],
        "photo": nm_to_photo.get(nm_by_product.get(p)),
        "by_wh_stock": stocks_by_warehouse.get(barcode_by_product.get(p), {})
    } for p in counts_total.keys()] if orders else []
    return render_template(
        "report_sales.html",
        error=None,
        items=items,
        date_from_fmt=date_from_fmt,
        date_to_fmt=date_to_fmt,
        warehouse=warehouse,
        warehouses=warehouses,
        date_from_val=(request.args.get("date_from") or ""),
        date_to_val=(request.args.get("date_to") or ""),
        matrix=matrix,
    )
@app.route("/report/orders", methods=["GET"]) 
@login_required
def report_orders_page():
    cached = load_last_results()
    token = effective_wb_api_token(current_user)
    
    # ????????? ??????? ???? ????? (???? ??? ???????)
    if token and current_user.is_authenticated:
        _start_stocks_update_bg(current_user.id, token)
    
    # ???????? ?? ????????? ??????????? ??????: ??? ??????, ???? ???????????? ?? ?????? ?????? ? ?? ?????? ?????????
    if not request.args.get("date_from") and not request.args.get("date_to"):
        return render_template(
            "report_orders.html",
            error=None,
            items=[],
            date_from_fmt="",
            date_to_fmt="",
            warehouse=None,
            warehouses=[],
            date_from_val="",
            date_to_val="",
        ), 200

    # Accept date range from query params; fallback to cached
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    warehouse = request.args.get("warehouse") or None

    if req_from and req_to and token:
        # Prefer cache if it matches the requested period and belongs to the user
        if (
            cached
            and current_user.is_authenticated
            and cached.get("_user_id") == current_user.id
            and cached.get("date_from") == req_from
            and cached.get("date_to") == req_to
        ):
            orders = cached.get("orders", [])
            try:
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
            except Exception:
                date_from_fmt = cached.get("date_from_fmt") or req_from
                date_to_fmt = cached.get("date_to_fmt") or req_to
        else:
            # Fetch fresh orders only if cache doesn't match
            try:
                raw_orders = fetch_orders_range(token, req_from, req_to)
                orders = to_rows(raw_orders, req_from, req_to)
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
            except Exception as exc:
                # Fallback to cache on error
                orders = (cached or {}).get("orders", [])
                date_from_fmt = (cached or {}).get("date_from_fmt") or (cached or {}).get("date_from")
                date_to_fmt = (cached or {}).get("date_to_fmt") or (cached or {}).get("date_to")
    else:
        # ??? ?????? ??????? ?? ??????????? ??? ? ???????? ???????? ??????
        orders = []
        date_from_fmt = ""
        date_to_fmt = ""

    warehouses = sorted({_order_row_warehouse_label(r) for r in orders}) if orders else []
    # Build matrix for client-side filtering (same as API)
    counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product = _normalize_and_group_orders(orders)
    # build photo map from products cache
    nm_to_photo: Dict[Any, Any] = {}
    try:
        prod_cached = load_products_cache() or {}
        for it in (prod_cached.get("items") or []):
            nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
            photo = it.get("photo") or it.get("img")
            if nmv is not None and nmv not in nm_to_photo:
                nm_to_photo[nmv] = photo
    except Exception:
        nm_to_photo = {}

    # Load stocks data for current user - ????????? ??????? ?? ???????
    stocks_by_warehouse = {}
    stocks_metadata = {}  # ?????????????? ?????????? ? ??????? ?? ????????
    try:
        stocks_cached = load_stocks_cache()
        if stocks_cached and stocks_cached.get("_user_id"):
            items = stocks_cached.get("items", [])
            for stock_item in items:
                barcode = stock_item.get("barcode")
                stock_warehouse = stock_item.get("warehouse", "")
                qty = int(stock_item.get("qty", 0) or 0)
                vendor_code = stock_item.get("vendor_code", "")
                nm_id = stock_item.get("nm_id")
                
                if barcode:
                    if barcode not in stocks_by_warehouse:
                        stocks_by_warehouse[barcode] = {}
                    if barcode not in stocks_metadata:
                        stocks_metadata[barcode] = {
                            "vendor_code": vendor_code,
                            "nm_id": nm_id,
                            "barcode": barcode
                        }
                    
                    if stock_warehouse:
                        stocks_by_warehouse[barcode][stock_warehouse] = stocks_by_warehouse[barcode].get(stock_warehouse, 0) + qty
    except Exception as e:
        stocks_by_warehouse = {}
        stocks_metadata = {}

    def _build_items(target_wh: str | None, show_all: bool = False) -> List[Dict[str, Any]]:
        items_local: List[Dict[str, Any]] = []
        
        # Get all products that have orders
        all_products = set(counts_total.keys())
        if show_all:
            # Add ALL products from stocks metadata (including those with zero stock)
            for barcode, metadata in stocks_metadata.items():
                # Find product by barcode
                found_in_orders = False
                for prod, prod_barcode in barcode_by_product.items():
                    if prod_barcode == barcode:
                        all_products.add(prod)
                        found_in_orders = True
                        break
                
                # If barcode not found in orders, create a virtual product entry
                if not found_in_orders:
                    # ?????????? ???? ??????????? ?? ???????
                    virtual_prod = f"BARCODE_{barcode}"
                    # Add to mappings
                    barcode_by_product[virtual_prod] = barcode
                    if metadata["nm_id"]:
                        nm_by_product[virtual_prod] = metadata["nm_id"]
                    if metadata["vendor_code"]:
                        supplier_article_by_product[virtual_prod] = metadata["vendor_code"]
                    all_products.add(virtual_prod)
        
        for prod in all_products:
            qty = (by_wh.get(prod, {}).get(target_wh, 0) if target_wh else counts_total.get(prod, 0))
            
            # Include items with orders OR (if show_all) items with stocks
            if qty > 0 or (show_all and prod in barcode_by_product):
                s = (by_wh_sum.get(prod, {}).get(target_wh, 0.0) if target_wh else revenue_total.get(prod, 0.0))
                # Calculate stock quantity for the target warehouse
                barcode = barcode_by_product.get(prod)
                stock_qty = 0
                if barcode and barcode in stocks_by_warehouse:
                    if target_wh:
                        # If specific warehouse selected, sum only for that warehouse
                        stock_qty = stocks_by_warehouse[barcode].get(target_wh, 0)
                    else:
                        # If no warehouse selected, sum across all warehouses
                        stock_qty = sum(stocks_by_warehouse[barcode].values())
                
                # Get photo from cache
                nm_id = nm_by_product.get(prod)
                photo = nm_to_photo.get(nm_id) if nm_id else None
                
                # ?????????? ???????? ?????? ??? ???????????: ?????????? ?????????? ??????? ????????, ???? ????
                display_name = supplier_article_by_product.get(prod)
                if not display_name:
                    # ???? ???????? ???????? ???, ??????? ??????? ?? ????? ???????????
                    if prod.startswith("BARCODE_"):
                        display_name = f"????? ? ???????? {prod.replace('BARCODE_', '')}"
                    elif prod.startswith("NM_"):
                        display_name = f"????? ? ????????? WB {prod.replace('NM_', '')}"
                    else:
                        display_name = prod
                
                items_local.append({
                    "product": display_name,
                    "qty": qty,
                    "sum": round(s, 2),
                    "warehouse": target_wh or "??? ??????",
                    "stock_qty": stock_qty,
                    "nm_id": nm_id,
                    "barcode": barcode,
                    "supplier_article": supplier_article_by_product.get(prod),
                    "photo": photo,
                })
        
        # Sort by quantity descending
        items_local.sort(key=lambda x: x["qty"], reverse=True)
        return items_local

    # Build items for the selected warehouse
    items = _build_items(warehouse, show_all=False)
    
    # ????????? ???????? ????????
    total_qty = sum(item["qty"] for item in items)
    total_sum = sum(item["sum"] for item in items)
    
    # Build matrix for client-side filtering
    matrix = {
        "counts_total": dict(counts_total),
        "by_wh": {k: dict(v) for k, v in by_wh.items()},
        "revenue_total": dict(revenue_total),
        "by_wh_sum": {k: dict(v) for k, v in by_wh_sum.items()},
        "nm_by_product": nm_by_product,
        "barcode_by_product": barcode_by_product,
        "supplier_article_by_product": supplier_article_by_product,
        "stocks_by_warehouse": stocks_by_warehouse,
        "stocks_metadata": stocks_metadata,
    }

    return render_template(
        "report_orders.html",
        error=None,
        items=items,
        date_from_fmt=date_from_fmt,
        date_to_fmt=date_to_fmt,
        warehouse=warehouse,
        warehouses=warehouses,
        date_from_val=(request.args.get("date_from") or ""),
        date_to_val=(request.args.get("date_to") or ""),
        total_qty=total_qty,
        total_sum=total_sum,
        matrix=matrix,
    )
@app.route("/api/report/orders", methods=["GET"]) 
@login_required
def api_report_orders():
    cached = load_last_results()
    token = effective_wb_api_token(current_user)
    
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    warehouse = (request.args.get("warehouse") or "").strip() or None
    hide_zero_stock = request.args.get("hide_zero_stock") == "on"
    only_not_sold = request.args.get("only_not_sold") == "on"
    # ???????? force_refresh ?????????, ??? ????? ????????? ?????? ??????, ????????? ???
    force_refresh = request.args.get("force_refresh") == "1" or request.args.get("force_refresh") == "true"
    
    # ?? ????????? ??????? ??? ???????? ?????? - ??? ???????? ??????? ????? ???????
    # ??????? ??????????? ???????? ?? ???????? /stocks ??? ? ????
    # ?????????? ???????????? ??? ????????
    
    try:
        if req_from and req_to and token:
            # Use the same stable path as /orders page (period cache + missing days fetch).
            # This avoids empty report on transient WB/API errors when pressing "Загрузить".
            try:
                orders, _meta = get_orders_with_period_cache(token, req_from, req_to)
            except Exception:
                if (
                    cached
                    and current_user.is_authenticated
                    and cached.get("_user_id") == current_user.id
                    and cached.get("date_from") == req_from
                    and cached.get("date_to") == req_to
                ):
                    orders = cached.get("orders", [])
                else:
                    orders = []
            date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
            date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
        else:
            orders = (cached or {}).get("orders", [])
            date_from_fmt = (cached or {}).get("date_from_fmt") or (cached or {}).get("date_from")
            date_to_fmt = (cached or {}).get("date_to_fmt") or (cached or {}).get("date_to")
        # Build matrix for local filtering on frontend
        counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product = _normalize_and_group_orders(orders)
        warehouses = sorted({_order_row_warehouse_label(r) for r in orders}) if orders else []
        show_all = request.args.get("show_all_products") == "on"

        # build photo map from products cache
        nm_to_photo: Dict[Any, Any] = {}
        all_catalog_items: List[Dict[str, Any]] = []
        try:
            prod_cached = load_products_cache() or {}
            all_catalog_items = list(prod_cached.get("items") or [])
            for it in all_catalog_items:
                nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
                photo = it.get("photo") or it.get("img")
                if nmv is not None and nmv not in nm_to_photo:
                    nm_to_photo[nmv] = photo
        except Exception:
            nm_to_photo = {}
            all_catalog_items = []

        # ? ?????? "???????? ???? ?????" ???????????, ??? ??????? ????????? ?????????,
        # ???? ???? ???????????? ?? ???????? ???????? /products.
        if show_all and token and not all_catalog_items:
            try:
                raw_cards = fetch_all_cards(token, page_limit=100)
                all_catalog_items = normalize_cards_response({"cards": raw_cards})
                save_products_cache({"items": all_catalog_items})
                for it in all_catalog_items:
                    nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
                    photo = it.get("photo") or it.get("img")
                    if nmv is not None and nmv not in nm_to_photo:
                        nm_to_photo[nmv] = photo
            except Exception:
                # ?? ????????? ???????????? ??????, ???? ?????????? ???????? ?? ???????.
                pass
        # Load stocks data for current user
        stocks_by_warehouse = {}
        stocks_metadata = {}
        try:
            stocks_cached = load_stocks_cache()
            if stocks_cached and stocks_cached.get("_user_id"):
                items = stocks_cached.get("items", [])
                for stock_item in items:
                    barcode = stock_item.get("barcode")
                    stock_warehouse = stock_item.get("warehouse", "")
                    qty = int(stock_item.get("qty", 0) or 0)
                    vendor_code = stock_item.get("vendor_code", "")
                    nm_id = stock_item.get("nm_id")
                    
                    if barcode:
                        if barcode not in stocks_by_warehouse:
                            stocks_by_warehouse[barcode] = {}
                        if barcode not in stocks_metadata:
                            stocks_metadata[barcode] = {
                                "vendor_code": vendor_code,
                                "nm_id": nm_id,
                                "barcode": barcode
                            }
                        
                        if stock_warehouse:
                            stocks_by_warehouse[barcode][stock_warehouse] = stocks_by_warehouse[barcode].get(stock_warehouse, 0) + qty
        except Exception as e:
            stocks_by_warehouse = {}
            stocks_metadata = {}
        def _build_items(target_wh: str | None, show_all: bool = False) -> List[Dict[str, Any]]:
            items_local: List[Dict[str, Any]] = []
            all_products = set(counts_total.keys())
            if show_all:
                # ????????? ?????? ??????? ??????? (?? ???? /products), ???? ??? ?????? ? ????????.
                prod_by_nm: Dict[Any, str] = {}
                prod_by_barcode: Dict[str, str] = {}
                for prod_key in list(all_products):
                    nmv = nm_by_product.get(prod_key)
                    bcv = barcode_by_product.get(prod_key)
                    if nmv is not None:
                        prod_by_nm[nmv] = prod_key
                    if bcv is not None:
                        prod_by_barcode[str(bcv)] = prod_key

                for it in all_catalog_items:
                    nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
                    bcv = it.get("barcode")
                    vendor_code = (it.get("vendor_code") or it.get("supplier_article") or it.get("article") or "").strip()

                    prod_key = None
                    if nmv is not None and nmv in prod_by_nm:
                        prod_key = prod_by_nm[nmv]
                    elif bcv is not None and str(bcv) in prod_by_barcode:
                        prod_key = prod_by_barcode[str(bcv)]
                    else:
                        if nmv is not None:
                            prod_key = f"NM_{nmv}"
                        elif bcv:
                            prod_key = f"BARCODE_{bcv}"
                        elif vendor_code:
                            prod_key = vendor_code

                    if not prod_key:
                        continue

                    all_products.add(prod_key)
                    if nmv is not None:
                        nm_by_product[prod_key] = nmv
                    if bcv:
                        barcode_by_product[prod_key] = bcv
                    if vendor_code:
                        supplier_article_by_product[prod_key] = vendor_code

                for barcode, metadata in stocks_metadata.items():
                    found_in_orders = False
                    for prod, prod_barcode in barcode_by_product.items():
                        if prod_barcode == barcode:
                            all_products.add(prod)
                            found_in_orders = True
                            break
                    if not found_in_orders:
                        virtual_prod = metadata["vendor_code"] or f"????? ? ???????? {barcode}"
                        barcode_by_product[virtual_prod] = barcode
                        if metadata["nm_id"]:
                            nm_by_product[virtual_prod] = metadata["nm_id"]
                        if metadata["vendor_code"]:
                            supplier_article_by_product[virtual_prod] = metadata["vendor_code"]
                        all_products.add(virtual_prod)
            for prod in all_products:
                qty = (by_wh.get(prod, {}).get(target_wh, 0) if target_wh else counts_total.get(prod, 0))
                if qty > 0 or show_all:
                    s = (by_wh_sum.get(prod, {}).get(target_wh, 0.0) if target_wh else revenue_total.get(prod, 0.0))
                    barcode = barcode_by_product.get(prod)
                    stock_qty = 0
                    if barcode and barcode in stocks_by_warehouse:
                        if target_wh:
                            stock_qty = stocks_by_warehouse[barcode].get(target_wh, 0)
                        else:
                            stock_qty = sum(stocks_by_warehouse[barcode].values())
                    nm_id = nm_by_product.get(prod)
                    photo = nm_to_photo.get(nm_id) if nm_id else None
                    
                    # ?????????? ???????? ?????? ??? ???????????: ?????????? ?????????? ??????? ????????, ???? ????
                    display_name = supplier_article_by_product.get(prod)
                    if not display_name:
                        # ???? ???????? ???????? ???, ??????? ??????? ?? ????? ???????????
                        if prod.startswith("BARCODE_"):
                            display_name = f"????? ? ???????? {prod.replace('BARCODE_', '')}"
                        elif prod.startswith("NM_"):
                            display_name = f"????? ? ????????? WB {prod.replace('NM_', '')}"
                        else:
                            display_name = prod
                    
                    item = {
                        "product": display_name,
                        "qty": qty,
                        "sum": round(s, 2),
                        "warehouse": target_wh or "??? ??????",
                        "stock_qty": stock_qty,
                        "nm_id": nm_id,
                        "barcode": barcode,
                        "supplier_article": supplier_article_by_product.get(prod),
                        "photo": photo,
                    }

                    # ?????????????? ??????? ??????
                    if hide_zero_stock and int(item.get("stock_qty") or 0) == 0:
                        continue
                    if only_not_sold and int(item.get("qty") or 0) > 0:
                        continue

                    items_local.append(item)
            items_local.sort(key=lambda x: x["qty"], reverse=True)
            return items_local
        items = _build_items(warehouse, show_all)
        
        # ????????? ???????? ????????
        total_qty = sum(item["qty"] for item in items)
        total_sum = sum(item["sum"] for item in items)
        
        return jsonify({
            "items": items,
            "date_from_fmt": date_from_fmt,
            "date_to_fmt": date_to_fmt,
            "warehouses": sorted(warehouses),
            "total_qty": total_qty,
            "total_sum": total_sum,
            "matrix": {
                "counts_total": dict(counts_total),
                "by_wh": {k: dict(v) for k, v in by_wh.items()},
                "revenue_total": dict(revenue_total),
                "by_wh_sum": {k: dict(v) for k, v in by_wh_sum.items()},
                "nm_by_product": nm_by_product,
                "barcode_by_product": barcode_by_product,
                "supplier_article_by_product": supplier_article_by_product,
                "stocks_by_warehouse": stocks_by_warehouse,
                "stocks_metadata": stocks_metadata,
            }
        }), 200
    except Exception as exc:
        return jsonify({"items": [], "error": str(exc), "debug": {"force_refresh": force_refresh}}), 200

# ??????? ?????? ?? ??????? ? Excel (??? ???????? /report/orders)
@app.route("/api/report/orders/export", methods=["GET"]) 
@login_required
def api_report_orders_export():
    """??????? ?????? ?? ??????? ? Excel-????.
    ????????? ????????? ????????? ?? ???????? ?????? ?? ???????.
    """
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter

    cached = load_last_results()
    token = effective_wb_api_token(current_user)
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    warehouse = (request.args.get("warehouse") or "").strip() or None

    try:
        if req_from and req_to and token:
            if (
                cached
                and current_user.is_authenticated
                and cached.get("_user_id") == current_user.id
                and cached.get("date_from") == req_from
                and cached.get("date_to") == req_to
            ):
                orders = cached.get("orders", [])
            else:
                raw_orders = fetch_orders_range(token, req_from, req_to)
                orders = to_rows(raw_orders, req_from, req_to)
        else:
            orders = (cached or {}).get("orders", [])

        # ????????? ??? ? API/???????? ??????: ????????? ??????????
        counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product = _normalize_and_group_orders(orders)

        # ?????????? ????? ??? ????????
        def _build_items(target_wh: str | None) -> List[Dict[str, Any]]:
            items_local: List[Dict[str, Any]] = []
            for prod, total in counts_total.items():
                qty = (by_wh[prod].get(target_wh, 0) if target_wh else total)
                if qty > 0:
                    s = (by_wh_sum.get(prod, {}).get(target_wh, 0.0) if target_wh else revenue_total.get(prod, 0.0))
                    
                    # ?????????? ???????? ?????? ??? ???????????: ?????????? ?????????? ??????? ????????, ???? ????
                    display_name = supplier_article_by_product.get(prod)
                    if not display_name:
                        # ???? ???????? ???????? ???, ??????? ??????? ?? ????? ???????????
                        if prod.startswith("BARCODE_"):
                            display_name = f"????? ? ???????? {prod.replace('BARCODE_', '')}"
                        elif prod.startswith("NM_"):
                            display_name = f"????? ? ????????? WB {prod.replace('NM_', '')}"
                        else:
                            display_name = prod
                    
                    items_local.append({
                        "product": display_name,
                        "qty": qty,
                        "nm_id": nm_by_product.get(prod),
                        "barcode": barcode_by_product.get(prod),
                        "supplier_article": supplier_article_by_product.get(prod),
                        "sum": round(float(s or 0.0), 2),
                    })
            items_local.sort(key=lambda x: x["qty"], reverse=True)
            return items_local

        items = _build_items(warehouse) if orders else []

        # ????????? Excel
        wb = Workbook()
        ws = wb.active
        ws.title = "????? ?? ???????"

        headers = ["??????? WB", "??????", "?????", "???-??", "?????"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")

        for row, item in enumerate(items, 2):
            ws.cell(row=row, column=1, value=item.get("nm_id") or "")
            ws.cell(row=row, column=2, value=item.get("barcode") or "")
            ws.cell(row=row, column=3, value=item.get("product"))
            ws.cell(row=row, column=4, value=item.get("qty"))
            ws.cell(row=row, column=5, value=item.get("sum"))

        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except Exception:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width

        # ??? ?????
        now = datetime.now()
        warehouse_name = warehouse if warehouse else "???"
        filename = f"????? ?? ??????? ?? ??????? ({warehouse_name})_{now.strftime('%d.%m.%Y_%H_%M')}.xlsx"

        out = BytesIO()
        wb.save(out)
        out.seek(0)

        import urllib.parse
        encoded_filename = urllib.parse.quote(filename.encode('utf-8'))
        return out.getvalue(), 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'Content-Disposition': f"attachment; filename*=UTF-8''{encoded_filename}"
        }
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500
@app.route("/api/report/sales", methods=["GET"]) 
@login_required
def api_report_sales():
    cached = load_last_results()
    token = effective_wb_api_token(current_user)
    
    # ????????? ??????? ???? ????? (???? ??? ???????)
    if token and current_user.is_authenticated:
        _start_stocks_update_bg(current_user.id, token)
    
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    warehouse = (request.args.get("warehouse") or "").strip() or None
    try:
        if req_from and req_to and token:
            if (
                cached
                and current_user.is_authenticated
                and cached.get("_user_id") == current_user.id
                and cached.get("date_from") == req_from
                and cached.get("date_to") == req_to
            ):
                orders = cached.get("orders", [])
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
            else:
                raw_orders = fetch_orders_range(token, req_from, req_to)
                orders = to_rows(raw_orders, req_from, req_to)
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
        else:
            orders = (cached or {}).get("orders", [])
            date_from_fmt = (cached or {}).get("date_from_fmt") or (cached or {}).get("date_from")
            date_to_fmt = (cached or {}).get("date_to_fmt") or (cached or {}).get("date_to")
        # Build matrix for local filtering on frontend
        counts_total: Dict[str, int] = defaultdict(int)
        by_wh: Dict[str, Dict[str, int]] = defaultdict(lambda: defaultdict(int))
        revenue_total: Dict[str, float] = defaultdict(float)
        by_wh_sum: Dict[str, Dict[str, float]] = defaultdict(lambda: defaultdict(float))
        nm_by_product: Dict[str, Any] = {}
        barcode_by_product: Dict[str, Any] = {}
        supplier_article_by_product: Dict[str, Any] = {}
        warehouses = set()
        for r in orders:
            # ?????????? ?????????? ?????? ? ??????
            if r.get("is_cancelled", False):
                continue
            # ?????????? ??????? ???????????? (??? ??????? ?? _normalize_and_group_orders ????)
            pass
        # build photo map
        nm_to_photo: Dict[Any, Any] = {}
        try:
            prod_cached = load_products_cache() or {}
            for it in (prod_cached.get("items") or []):
                nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
                photo = it.get("photo") or it.get("img")
                if nmv is not None and nmv not in nm_to_photo:
                    nm_to_photo[nmv] = photo
        except Exception:
            nm_to_photo = {}

        # Load stocks data for current user - ????????? ??????? ?? ???????
        stocks_by_warehouse = {}
        stocks_metadata = {}  # ?????????????? ?????????? ? ??????? ?? ????????
        try:
            stocks_cached = load_stocks_cache()
            if stocks_cached and stocks_cached.get("_user_id"):
                for stock_item in stocks_cached.get("items", []):
                    barcode = stock_item.get("barcode")
                    stock_warehouse = stock_item.get("warehouse", "")
                    qty = int(stock_item.get("qty", 0) or 0)
                    vendor_code = stock_item.get("vendor_code", "")
                    nm_id = stock_item.get("nm_id")
                    
                    if barcode:
                        if barcode not in stocks_by_warehouse:
                            stocks_by_warehouse[barcode] = {}
                        if barcode not in stocks_metadata:
                            stocks_metadata[barcode] = {
                                "vendor_code": vendor_code,
                                "nm_id": nm_id,
                                "barcode": barcode
                            }
                        
                        if stock_warehouse:
                            stocks_by_warehouse[barcode][stock_warehouse] = stocks_by_warehouse[barcode].get(stock_warehouse, 0) + qty
        except Exception:
            stocks_by_warehouse = {}
            stocks_metadata = {}

        def build_items_for_wh(target_wh: str | None, show_all: bool = False) -> List[Dict[str, Any]]:
            items_local: List[Dict[str, Any]] = []
            
            # Get all products that have sales or stocks
            all_products = set(counts_total.keys())
            if show_all:
                # Add ALL products from stocks metadata (including those with zero stock)
                for barcode, metadata in stocks_metadata.items():
                    # Find product by barcode
                    found_in_sales = False
                    for prod, prod_barcode in barcode_by_product.items():
                        if prod_barcode == barcode:
                            all_products.add(prod)
                            found_in_sales = True
                            break
                    
                    # If barcode not found in sales, create a virtual product entry
                    if not found_in_sales:
                        # Use vendor_code from stocks metadata
                        virtual_prod = metadata["vendor_code"] or f"????? ? ???????? {barcode}"
                        # Add to mappings
                        barcode_by_product[virtual_prod] = barcode
                        if metadata["nm_id"]:
                            nm_by_product[virtual_prod] = metadata["nm_id"]
                        if metadata["vendor_code"]:
                            supplier_article_by_product[virtual_prod] = metadata["vendor_code"]
                        all_products.add(virtual_prod)
            
            for prod in all_products:
                qty = (by_wh.get(prod, {}).get(target_wh, 0) if target_wh else counts_total.get(prod, 0))
                
                # Include items with sales OR (if show_all) items with stocks
                if qty > 0 or (show_all and prod in barcode_by_product):
                    s = (by_wh_sum.get(prod, {}).get(target_wh, 0.0) if target_wh else revenue_total.get(prod, 0.0))
                    # Calculate stock quantity for the target warehouse
                    barcode = barcode_by_product.get(prod)
                    stock_qty = 0
                    if barcode and barcode in stocks_by_warehouse:
                        if target_wh:
                            # If specific warehouse selected, sum only for that warehouse
                            for wh_name, wh_qty in stocks_by_warehouse[barcode].items():
                                if (wh_name == target_wh or 
                                    (target_wh in wh_name) or 
                                    (wh_name in target_wh)):
                                    stock_qty += wh_qty
                        else:
                            # If no warehouse selected, sum all warehouses
                            stock_qty = sum(stocks_by_warehouse[barcode].values())
                    
                    # Only include if has sales or (if show_all) has any stock data
                    if qty > 0 or (show_all and prod in barcode_by_product):
                        items_local.append({
                            "product": prod,
                            "qty": qty,
                            "nm_id": nm_by_product.get(prod),
                            "barcode": barcode,
                            "supplier_article": supplier_article_by_product.get(prod),
                            "sum": round(float(s or 0.0), 2),
                            "photo": nm_to_photo.get(nm_by_product.get(prod)),
                            "stock_qty": stock_qty
                        })
            
            # Sort by quantity (descending), then by stock quantity (descending)
            items_local.sort(key=lambda x: (x["qty"], x["stock_qty"]), reverse=True)
            return items_local
        show_all = request.args.get("show_all_products") == "on"
        items = build_items_for_wh(warehouse, show_all)
        total_qty = sum(int(it.get("qty") or 0) for it in items)
        matrix = [{
            "product": p,
            "nm_id": nm_by_product.get(p),
            "barcode": barcode_by_product.get(p),
            "supplier_article": supplier_article_by_product.get(p),
            "total": counts_total[p],
            "by_wh": by_wh[p],
            "total_sum": round(float(revenue_total.get(p, 0.0)), 2),
            "by_wh_sum": by_wh_sum[p],
            "photo": nm_to_photo.get(nm_by_product.get(p)),
            "by_wh_stock": stocks_by_warehouse.get(barcode_by_product.get(p), {})
        } for p in counts_total.keys()]
        return jsonify({
            "items": items,
            "total_qty": total_qty,
            "date_from_fmt": date_from_fmt,
            "date_to_fmt": date_to_fmt,
            "warehouses": sorted(list(warehouses)),
            "matrix": matrix,
        }), 200
    except Exception as exc:
        return jsonify({"items": [], "error": str(exc)}), 200
@app.route("/api/report/sales/export", methods=["GET"])
@login_required
def api_report_sales_export():
    """??????? ?????? ?? ???????? ? Excel ??????"""
    from io import BytesIO
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    from openpyxl.utils import get_column_letter
    
    cached = load_last_results()
    token = effective_wb_api_token(current_user)
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    warehouse = (request.args.get("warehouse") or "").strip() or None
    
    try:
        if req_from and req_to and token:
            if (
                cached
                and current_user.is_authenticated
                and cached.get("_user_id") == current_user.id
                and cached.get("date_from") == req_from
                and cached.get("date_to") == req_to
            ):
                orders = cached.get("orders", [])
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
            else:
                raw_orders = fetch_orders_range(token, req_from, req_to)
                orders = to_rows(raw_orders, req_from, req_to)
                date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
                date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
        else:
            orders = (cached or {}).get("orders", [])
            date_from_fmt = (cached or {}).get("date_from_fmt") or (cached or {}).get("date_from")
            date_to_fmt = (cached or {}).get("date_to_fmt") or (cached or {}).get("date_to")
        
        # Build matrix for filtering
        counts_total, by_wh, revenue_total, by_wh_sum, nm_by_product, barcode_by_product, supplier_article_by_product = _normalize_and_group_orders(orders)
        warehouses = sorted({_order_row_warehouse_label(r) for r in orders}) if orders else []
        
        # Build items for export
        def _build_items(target_wh: str | None) -> List[Dict[str, Any]]:
            items_local: List[Dict[str, Any]] = []
            for prod, total in counts_total.items():
                qty = (by_wh[prod].get(target_wh, 0) if target_wh else total)
                if qty > 0:
                    s = (by_wh_sum.get(prod, {}).get(target_wh, 0.0) if target_wh else revenue_total.get(prod, 0.0))
                    items_local.append({
                        "product": prod,
                        "qty": qty,
                        "nm_id": nm_by_product.get(prod),
                        "barcode": barcode_by_product.get(prod),
                        "supplier_article": supplier_article_by_product.get(prod),
                        "sum": round(float(s or 0.0), 2),
                    })
            items_local.sort(key=lambda x: x["qty"], reverse=True)
            return items_local
        
        items = _build_items(warehouse) if orders else []
        
        # Create Excel workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "????? ?? ???????"
        
        # Headers
        headers = ["??????? WB", "??????", "?????", "???-??", "?????"]
        for col, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col, value=header)
            cell.font = Font(bold=True)
            cell.alignment = Alignment(horizontal='center')
            cell.fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        
        # Data rows
        for row, item in enumerate(items, 2):
            ws.cell(row=row, column=1, value=item["nm_id"] or "")
            ws.cell(row=row, column=2, value=item["barcode"] or "")
            ws.cell(row=row, column=3, value=item["product"])
            ws.cell(row=row, column=4, value=item["qty"])
            ws.cell(row=row, column=5, value=item["sum"])
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = get_column_letter(column[0].column)
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Generate filename
        now = datetime.now()
        warehouse_name = warehouse if warehouse else "???"
        filename = f"????? ?? ??????? ?? ??????? ({warehouse_name})_{now.strftime('%d.%m.%Y_%H_%M')}.xlsx"
        
        # Save to BytesIO
        output = BytesIO()
        wb.save(output)
        output.seek(0)
        
        # Encode filename for HTTP headers
        import urllib.parse
        encoded_filename = urllib.parse.quote(filename.encode('utf-8'))
        
        return output.getvalue(), 200, {
            'Content-Type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            'Content-Disposition': f'attachment; filename*=UTF-8\'\'{encoded_filename}'
        }
        
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/report/finance", methods=["GET"]) 
@login_required
def report_finance_page():
    # initial render without data - ?????? ?????????? ?????? ????????
    if not request.args.get("date_from") and not request.args.get("date_to"):
        return render_template(
            "finance_report.html",
            error=None,
            rows=[],
            date_from_fmt="",
            date_to_fmt="",
            date_from_val="",
            date_to_val="",
            finance_metrics={},
        ), 200
    token = effective_wb_api_token(current_user)
    if not token:
        return render_template(
            "finance_report.html",
            error="????????? API ????? (??????????)",
            items=[],
            date_from_fmt="",
            date_to_fmt="",
            date_from_val=(request.args.get("date_from") or ""),
            date_to_val=(request.args.get("date_to") or ""),
        ), 200
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    try:
        raw = fetch_finance_report(token, req_from, req_to)
        date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
        date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
    except Exception as exc:
        return render_template(
            "finance_report.html",
            error=str(exc),
            rows=[],
            date_from_fmt="",
            date_to_fmt="",
            date_from_val=req_from,
            date_to_val=req_to,
        ), 200
    return render_template(
        "finance_report.html",
        error=None,
        rows=raw,
        date_from_fmt=date_from_fmt,
        date_to_fmt=date_to_fmt,
        date_from_val=req_from,
        date_to_val=req_to,
    ), 200
@app.route("/api/report/finance/progress", methods=["GET"])
@login_required
def api_report_finance_progress():
    """Endpoint ??? ????????? ????????? ???????? ??????????? ??????"""
    if not current_user.is_authenticated:
        return jsonify({"current": 0, "total": 0, "period": "", "ready": False}), 200
    progress = _get_finance_progress(current_user.id)
    # ?????????, ?????? ?? ??????????
    user_id = current_user.id
    results_ready = user_id in FINANCE_RESULTS and not FINANCE_LOADING.get(user_id, False)
    progress["ready"] = results_ready
    if results_ready:
        progress["has_results"] = True
    return jsonify(progress), 200

@app.route("/api/report/finance/result", methods=["GET"])
@login_required
def api_report_finance_result():
    """Endpoint ??? ????????? ??????????? ??????????? ?????? ????? ??????????? ????????"""
    if not current_user.is_authenticated:
        return jsonify({"error": "Not authenticated"}), 401
    user_id = current_user.id
    if user_id in FINANCE_RESULTS:
        result = FINANCE_RESULTS[user_id]
        # ??????? ?????????? ????? ?????????
        del FINANCE_RESULTS[user_id]
        return jsonify(result), 200
    return jsonify({"error": "Results not ready"}), 404

@app.route("/api/report/finance", methods=["GET"]) 
@login_required
def api_report_finance():
    token = effective_wb_api_token(current_user)
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    async_mode = request.args.get("async", "0") == "1"  # ???????? ??? ??????????? ????????
    
    if not (token and req_from and req_to):
        return jsonify({"items": [], "error": None}), 200
    
    user_id = current_user.id if current_user.is_authenticated else None
    
    # ???? ????????? ??????????? ???????? ??? ?????? ??????? (????? 60 ????), ?????????? ??????? ??????
    from datetime import datetime as dt
    try:
        date_from_obj = dt.strptime(req_from, "%Y-%m-%d")
        date_to_obj = dt.strptime(req_to, "%Y-%m-%d")
        days_diff = (date_to_obj - date_from_obj).days
        use_async = async_mode or days_diff > 60
    except Exception:
        use_async = async_mode
    
    if use_async and user_id:
        # ??????????? ???????? ????? ??????? ??????
        # ?????????, ?? ???? ?? ??? ????????
        if FINANCE_LOADING.get(user_id, False):
            return jsonify({"loading": True, "message": "???????? ??? ???????????"}), 200
        
        # ??????? ?????????? ?????????? ? ????????
        if user_id in FINANCE_RESULTS:
            del FINANCE_RESULTS[user_id]
        _clear_finance_progress(user_id)
        FINANCE_LOADING[user_id] = True
        
        # ????????? ???????? ? ??????? ??????
        import threading
        def load_finance_background():
            try:
                # ??????? callback ??? ?????????? ?????????
                def progress_callback(current, total, period):
                    _set_finance_progress(user_id, current, total, period)
                
                # ????????? ??????
                raw = fetch_finance_report(token, req_from, req_to, progress_callback=progress_callback)
                
                # ???????????? ?????? (???????? ?????? ?? ????????? ??????)
                result = _process_finance_data(raw, req_from, req_to, user_id)
                
                # ????????? ??????????
                FINANCE_RESULTS[user_id] = result
                
                # ??????? ???????? ? ???? ????????
                _clear_finance_progress(user_id)
                FINANCE_LOADING[user_id] = False
                
                logging.info(f"??????? ???????? ??????????? ?????? ????????? ??? ???????????? {user_id}")
            except Exception as e:
                logging.error(f"?????? ??????? ???????? ??????????? ??????: {e}")
                FINANCE_LOADING[user_id] = False
                if user_id in FINANCE_RESULTS:
                    del FINANCE_RESULTS[user_id]
                _clear_finance_progress(user_id)
        
        thread = threading.Thread(target=load_finance_background)
        thread.daemon = True
        thread.start()
        
        # ????? ?????????? ????? ? ?????? ????????
        return jsonify({"loading": True, "message": "???????? ??????, ??????????? /api/report/finance/progress ??? ???????????? ?????????"}), 200
    
    # ?????????? ???????? (??? ????????? ????????)
    # ??????? ?????????? ????????
    if current_user.is_authenticated:
        _clear_finance_progress(current_user.id)
    
    try:
        # ??????? callback ??? ?????????? ?????????
        def progress_callback(current, total, period):
            if current_user.is_authenticated:
                _set_finance_progress(current_user.id, current, total, period)
        
        # Always fetch fresh report for the period (?? ???????? ?????? ??????)
        # ?????? ? ?????????? ?? ????????? ?? 7 ????
        raw = fetch_finance_report(token, req_from, req_to, progress_callback=progress_callback)
        
        # ??????? ???????? ????? ??????????
        if current_user.is_authenticated:
            _clear_finance_progress(current_user.id)
        
        # ?????????? ??????? ????????? ??????
        user_id_for_tax = current_user.id if current_user.is_authenticated else None
        result = _process_finance_data(raw, req_from, req_to, user_id_for_tax)
        
        return jsonify(result), 200
    except Exception as exc:
        return jsonify({"items": [], "error": str(exc)}), 200
        # WB ?????????? (?? retail_amount ? ????????? ?? ?????????? ??????)
        wbr_plus = 0.0
        wbr_minus = 0.0
        total_logistics = 0.0
        total_storage = 0.0
        total_acceptance = 0.0
        total_for_pay = 0.0
        total_buyouts = 0.0
        total_returns = 0.0
        total_acquiring = 0.0
        total_commission_wb = 0.0
        total_other_deductions = 0.0
        total_penalties = 0.0
        total_additional_payment = 0.0
        # ??????? ???????? (???????????): sum(ppvz_for_pay) where supplier_oper_name == "?????? ??????? ????????"
        total_paid_delivery = 0.0
        # ??????????? ????? ????????? ??????? ?? ppvz_for_pay
        x1 = x2 = x3 = x4 = x5 = x6 = x7 = x8 = 0.0
        # ???????? ?????????? K1..K9
        k1 = k2 = k3 = k4 = k5 = k6 = k7 = k8 = k9 = 0.0
        # ??????????? ?????? ?????????? U1..U14
        u1 = u2 = u3 = u4 = u5 = u6 = u7 = u8 = u9 = u10 = u11 = u12 = u13 = u14 = 0.0
        # ?????? ? ???????? ??????? ?? ??????? "supplier_oper_name"
        # ??????: one of ["???????","?????? ?????????","?????????? ???????","????????? ??????"] -> sum(retail_price)
        buyout_oper_values_lower = {"???????", "?????? ?????????", "?????????? ???????", "????????? ??????"}
        for r in raw:
            try:
                total_qty += int(r.get("quantity") or 0)
            except Exception:
                pass
            try:
                total_sum += float(r.get("retail_amount") or 0.0)
            except Exception:
                pass
            try:
                total_logistics += float(r.get("delivery_rub") or 0.0)
            except Exception:
                pass
            try:
                total_storage += float(r.get("storage_fee") or 0.0)
            except Exception:
                pass
            try:
                total_acceptance += float(r.get("acceptance") or 0.0)
            except Exception:
                pass
            try:
                total_for_pay += float(r.get("ppvz_for_pay") or 0.0)
            except Exception:
                pass
            # ??????: ????????? ????????? ???? ??? ?????? ????????? ??????
            try:
                oper = (r.get("supplier_oper_name") or "").strip()
                if oper and oper.lower() in buyout_oper_values_lower:
                    total_buyouts += float(r.get("retail_price") or 0.0)
            except Exception:
                pass
            # WB ??????????: retail_amount ? ????????????? ?? ?????????? ??????
            try:
                oper_lc = (r.get("supplier_oper_name") or "").strip().lower()
                amt = float(r.get("retail_amount") or 0.0)
                if oper_lc in {"???????","?????? ?????????","?????????? ???????","????????? ??????"}:
                    wbr_plus += amt
                elif oper_lc in {"???????","?????? ??????","?????????? ???????"}:
                    wbr_minus += amt
            except Exception:
                pass
            # ????????: supplier_oper_name == "???????"; ????????? retail_price
            try:
                oper = (r.get("supplier_oper_name") or "").strip()
                if oper == "???????":
                    total_returns += float(r.get("retail_price") or 0.0)
            except Exception:
                pass
            # ????????? ?? ???????: E1 - E2 + E3
            # E1: doc_type_name == "???????" AND acquiring_percent > 0 -> sum acquiring_fee
            # E2: doc_type_name == "???????" AND acquiring_percent > 0 -> sum acquiring_fee
            # E3: supplier_oper_name == "????????????? ??????????" -> sum ppvz_for_pay
            try:
                dt_name = (r.get("doc_type_name") or "").strip()
                acq_pct = float(r.get("acquiring_percent") or 0.0)
                afee = float(r.get("acquiring_fee") or 0.0)
                if dt_name == "???????" and acq_pct > 0:
                    total_acquiring += afee
                elif dt_name == "???????" and acq_pct > 0:
                    total_acquiring -= afee
                # ????????????? ?????????? ?? ?????????? ? "?????????", 
                # ? ??????????? ???????? ? ??????? "? ????????????"
            except Exception:
                pass
            try:
                total_commission_wb += float(r.get("ppvz_vw") or 0.0)
            except Exception:
                pass
            try:
                total_other_deductions += float(r.get("deduction") or 0.0)
                total_other_deductions += float(r.get("additional_payment") or 0.0)
            except Exception:
                pass
            try:
                total_penalties += float(r.get("penalty") or 0.0)
            except Exception:
                pass
            try:
                total_additional_payment += float(r.get("additional_payment") or 0.0)
            except Exception:
                pass
            # ??????????? ?????: ??????? X1..X8 ?? supplier_oper_name + doc_type_name, ???????? ppvz_for_pay
            try:
                oper_l = (r.get("supplier_oper_name") or "").strip().lower()
                doc_l = (r.get("doc_type_name") or "").strip().lower()
                pay_val = float(r.get("ppvz_for_pay") or 0.0)
                # ??????? ???????? ? ????????? ??? ???????????, ????????? ???????
                if oper_l == "?????? ??????? ????????":
                    total_paid_delivery += pay_val
                if oper_l == "??????????? ?????" and doc_l == "???????":
                    x1 += pay_val
                if oper_l == "?????? ?????" and doc_l == "???????":
                    x2 += pay_val
                if oper_l == "??????????? ?????" and doc_l == "???????":
                    x3 += pay_val
                if oper_l == "?????? ?????" and doc_l == "???????":
                    x4 += pay_val
                if oper_l == "????????? ??????????? ?????" and doc_l == "???????":
                    x5 += pay_val
                if oper_l == "????????? ??????????? ?????" and doc_l == "???????":
                    x6 += pay_val
                if oper_l == "???????????? ??????????? ??? ????????" and doc_l == "???????":
                    x7 += pay_val
                if oper_l == "???????????? ??????????? ??? ????????" and doc_l == "???????":
                    x8 += pay_val
                # ????????: K1..K9 ? ????? ?? ?????? ppvz_for_pay
                if oper_l == "???????":
                    k1 += pay_val
                if oper_l == "?????? ?????????":
                    k2 += pay_val
                if oper_l == "?????????? ???????":
                    k3 += pay_val
                if oper_l == "????????? ??????" and doc_l == "???????":
                    k4 += pay_val
                if oper_l == "???????":
                    k5 += pay_val
                if oper_l == "?????? ??????":
                    k6 += pay_val
                if oper_l == "????????? ??????" and doc_l == "???????":
                    k7 += pay_val
                if oper_l == "?????????? ???????":
                    k8 += pay_val
                if oper_l == "????????????? ??????????":
                    k9 += pay_val
                # ??????????? ??????: U1..U14 ?? ????????, ????????? ppvz_for_pay
                if oper_l == "?????? ??????????? ??????" and doc_l == "???????":
                    u1 += pay_val
                if oper_l == "??????????? ??????????? ??????" and doc_l == "???????":
                    u2 += pay_val
                if oper_l == "?????? ??????????? ??????" and doc_l == "???????":
                    u3 += pay_val
                if oper_l == "??????????? ??????????? ??????" and doc_l == "???????":
                    u4 += pay_val
                if oper_l == "????????? ?????? ?? ????? ??? ????????" and doc_l == "???????":
                    u5 += pay_val
                if oper_l == "????????? ?????? ?? ????? ??? ????????" and doc_l == "???????":
                    u6 += pay_val
                if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                    u7 += pay_val
                if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                    u8 += pay_val
                if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                    u9 += pay_val
                if oper_l == "??????????? ???????????? ??????" and doc_l == "???????":
                    u10 += pay_val
                if oper_l == "??????????? ??????" and doc_l == "???????":
                    u11 += pay_val
                if oper_l == "??????????? ??????" and doc_l == "???????":
                    u12 += pay_val
                if oper_l == "??????????? ??????" and doc_l == "???????":
                    u13 += pay_val
                if oper_l == "??????????? ??????" and doc_l == "???????":
                    u14 += pay_val
            except Exception:
                pass
        date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y")
        date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y")
        revenue_calc = total_buyouts - total_returns
        defect_comp = x1 + x2 - x3 - x4 + x5 - x6 + x7 - x8
        total_wb_realized = wbr_plus - wbr_minus
        # ???????? = ??????? - (K1+K2+K3+K4 - (K5+K6+K7+K8)) - ?????????
        # (K9 - ????????????? ?????????? ??????????? ???????? ? ??????? "? ????????????")
        commission_total = (
            revenue_calc
            - (k1 + k2 + k3 + k4)
            + (k5 + k6 + k7 + k8)
            - total_acquiring
        )
        damage_comp = u1 + u2 - u3 - u4 + u5 - u6 + u7 + u8 - u9 - u10 + u11 - u12 + u13 - u14
        
        # ????????? ? ??????????? WB = ???????? + ????????? + ????????? + ???????? + ?????? ????????? + ??????? - ??????????? ????? - ??????????? ?????? - ??????? ???????? + ?????? + ???????
        total_deductions = (
            commission_total + 
            total_acquiring + 
            total_logistics + 
            total_storage + 
            total_other_deductions + 
            total_acceptance - 
            defect_comp - 
            damage_comp - 
            total_paid_delivery + 
            total_penalties + 
            total_additional_payment
        )
        
        # ? ???????????? = ??????? - ????????? ? ??????????? WB + E3
        # E3: supplier_oper_name == "????????????? ??????????" -> sum ppvz_for_pay
        e3_correction = 0
        for r in raw:
            try:
                oper_name = (r.get("supplier_oper_name") or "").strip()
                if oper_name == "????????????? ??????????":
                    e3_correction += float(r.get("ppvz_for_pay") or 0.0)
            except Exception:
                pass
        
        total_for_transfer = revenue_calc - total_deductions + e3_correction
        
        # ?????? ?????? ?? ????? "WB ??????????"
        tax_amount = 0.0
        tax_rate = None
        if current_user.tax_rate is not None:
            tax_rate = float(current_user.tax_rate)
            tax_amount = (total_wb_realized * tax_rate) / 100.0
        
        # ??????????? ????????? - ?? ????????? ?????? ??????????? ??????
        
        return jsonify({
            "rows": raw,
            "total_qty": int(total_qty),
            "total_sum": round(total_sum, 2),
            "total_logistics": round(total_logistics, 2),
            "total_storage": round(total_storage, 2),
            "total_acceptance": round(total_acceptance, 2),
            "total_for_pay": round(total_for_pay, 2),
            "total_buyouts": round(total_buyouts, 2),
            "total_returns": round(total_returns, 2),
            "revenue": round(revenue_calc, 2),
            "total_wb_realized": round(total_wb_realized, 2),
            "total_commission": round(commission_total, 2),
            "total_acquiring": round(total_acquiring, 2),
            "total_commission_wb": round(total_commission_wb, 2),
            "total_other_deductions": round(total_other_deductions, 2),
            "total_penalties": round(total_penalties, 2),
            "total_defect_compensation": round(defect_comp, 2),
            "total_damage_compensation": round(damage_comp, 2),
            "total_paid_delivery": round(total_paid_delivery, 2),
            "total_additional_payment": round(total_additional_payment, 2),
            "total_deductions": round(total_deductions, 2),
            "total_for_transfer": round(total_for_transfer, 2),
            "tax_amount": round(tax_amount, 2),
            "tax_rate": tax_rate,
            "date_from_fmt": date_from_fmt,
            "date_to_fmt": date_to_fmt,
        }), 200
    except Exception as exc:
        return jsonify({"items": [], "error": str(exc)}), 200


@app.route("/report/finance/export", methods=["GET"]) 
@login_required
def export_finance_xls():
    token = effective_wb_api_token(current_user)
    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    if not (token and req_from and req_to):
        return ("????????? ???? ? ?????", 400)
    try:
        # Always fetch fresh for export
        rows = fetch_finance_report(token, req_from, req_to)
        # Build XLS (not XLSX) to match requirement "XLS"
        try:
            import xlwt  # type: ignore
        except Exception:
            return ("?? ??????? ??????????? ??????????? xlwt (??? .xls)", 500)
        wb = xlwt.Workbook()
        ws = wb.add_sheet("finance")
        header_style = xlwt.easyxf("font: bold on; align: horiz center")
        num_style = xlwt.easyxf("align: horiz right")
        cols = [
            'realizationreport_id','date_from','date_to','create_dt','currency_name','suppliercontract_code','rrd_id','gi_id','dlv_prc','fix_tariff_date_from','fix_tariff_date_to','subject_name','nm_id','brand_name','sa_name','ts_name','barcode','doc_type_name','quantity','retail_price','retail_amount','sale_percent','commission_percent','office_name','supplier_oper_name','order_dt','sale_dt','rr_dt','shk_id','retail_price_withdisc_rub','delivery_amount','return_amount','delivery_rub','gi_box_type_name','product_discount_for_report','supplier_promo','ppvz_spp_prc','ppvz_kvw_prc_base','ppvz_kvw_prc','sup_rating_prc_up','is_kgvp_v2','ppvz_sales_commission','ppvz_for_pay','ppvz_reward','acquiring_fee','acquiring_percent','payment_processing','acquiring_bank','ppvz_vw','ppvz_vw_nds','ppvz_office_name','ppvz_office_id','ppvz_supplier_id','ppvz_supplier_name','ppvz_inn','declaration_number','bonus_type_name','sticker_id','site_country','srv_dbs','penalty','additional_payment','rebill_logistic_cost','rebill_logistic_org','storage_fee','deduction','acceptance','assembly_id','kiz','srid','report_type','is_legal_entity','trbx_id','installment_cofinancing_amount','wibes_wb_discount_percent','cashback_amount','cashback_discount'
        ]
        headers_ru = [
            "????? ??????","???? ?????? ???????","???? ????? ???????","???? ????????????","??????","???????","????? ??????","????? ????????","????. ????. ??????","?????? ????????","????? ????????","???????","??????? WB","?????","??????? ????????","??????","??????","??? ?????????","??????????","???? ?????????","??????????? (??)","??????, %","???, %","?????","??????????? ??????","???? ??????","???? ???????","???? ????????","????????","????????? ? ??. ??????","???-?? ????????","???-?? ?????????","????????, ???","??? ???????","????. ???????. ??????, %","????????, %","???, %","??????? ??? ??? ???, %","???????? ??? ??? ???, %","???????? ??? (???????), %","???????? ??? (?????), %","?????????????? ? ??????","? ???????????? ????????","?????????? ???","?????????","?????????, %","??? ??????? ??????????","????-???????","?????????????? ??","??? ??","???? ????????","ID ?????","ID ????????","???????","??? ????????","? ??????????","??? ?????????/??????","ID ???????","?????? ???????","??????? ????????","??????","????????????? ??","?????????? ?????????","??????????? ?????????","????????","?????????","??????? ???????","ID ??????????","??? ??????????","SRID","??? ??????","B2B","ID ?????? ???????","????????????????","?????? Wibes, %","????? (????????)","??????????? ??????"
        ]
        for ci, title in enumerate(headers_ru, start=1):
            ws.write(0, ci-1, title, header_style)
        row_idx = 1
        for r in rows:
            for ci, key in enumerate(cols, start=1):
                val = r.get(key)
                if isinstance(val, (dict, list)):
                    try:
                        import json as _json
                        val = _json.dumps(val, ensure_ascii=False)
                    except Exception:
                        val = str(val)
                ws.write(row_idx, ci-1, val if val is not None else "")
            row_idx += 1
        out = io.BytesIO()
        wb.save(out)
        out.seek(0)
        filename = f"finance_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xls"
        return send_file(out, mimetype="application/vnd.ms-excel", as_attachment=True, download_name=filename)
    except requests.HTTPError as http_err:
        return (f"?????? API: {http_err.response.status_code}", 502)
    except Exception as exc:
        return (f"??????: {exc}", 500)

@app.route("/fbs", methods=["GET", "POST"]) 
@login_required
def fbs_page():
    error = None
    token = effective_wb_api_token(current_user)
    rows: List[Dict[str, Any]] = []
    # Load cached tasks to show immediately
    cached_tasks = load_fbs_tasks_cache() or {}
    cached_rows = cached_tasks.get("rows") or []
    rows = cached_rows

    if request.method == "POST":
        pass  # ?????? ???? ?????? ????????; ?????? ????????? ????? JS-?????? ? ?????

    # If enrichment impossible due to empty products cache
    products_hint = None
    prod_cached_now = load_products_cache()
    if not prod_cached_now or not ((prod_cached_now or {}).get("items")):
        products_hint = "??? ??????????? ???? ?????? ? ??????? ???????? ?????? ?? ???????? ??????"

    # ?? ????????? ?????? ????????: ??????? ??????? ???????? AJAX-??
    return render_template("fbs.html", error=error, rows=rows, products_hint=products_hint, current_orders=[])
@app.route("/fbs/export", methods=["POST"]) 
@login_required
def fbs_export():
    token = effective_wb_api_token(current_user)
    if not token:
        return ("????????? API ?????", 400)
    try:
        # 1) ????? ?????? ?? ???? ??????? (??, ??? ??????? ????? ?? ???????? /fbs).
        # ??? ????????? ????????, ????? WB /new ??? ??????, ?? ??? ?? ???????? ??? ???????? ???????.
        cached = load_fbs_tasks_cache() or {}
        rows = (cached.get("rows") or []) if isinstance(cached, dict) else []

        # 2) ???? ???? ??? ? ??????? ???????? ?????? ?????? (fallback).
        if not rows:
            raw = fetch_fbs_new_orders(token)
            raw_sorted = sorted(raw, key=_extract_created_at)
            rows = to_fbs_rows(raw_sorted)

            # Enrich from products cache (??? ????????/????)
            prod_cached = load_products_cache()
            items = (prod_cached or {}).get("items") or []
            by_article: Dict[str, Dict[str, Any]] = {}
            by_nm: Dict[int, Dict[str, Any]] = {}
            for it in items:
                art = (it.get("supplier_article") or it.get("vendorCode") or "").strip()
                if art:
                    by_article.setdefault(art, it)
                nmv = it.get("nm_id") or it.get("nmID")
                try:
                    if nmv:
                        by_nm[int(nmv)] = it
                except Exception:
                    pass
            for r in rows:
                art = (r.get("???????????? ??????") or "").strip()
                hit = by_article.get(art)
                if not hit and r.get("nm_id"):
                    try:
                        hit = by_nm.get(int(r["nm_id"]))
                    except Exception:
                        hit = None
                if hit:
                    if hit.get("barcode"):
                        r["barcode"] = hit.get("barcode")
                    elif isinstance(hit.get("barcodes"), list) and hit.get("barcodes"):
                        r["barcode"] = str(hit.get("barcodes")[0])
                    else:
                        sizes = hit.get("sizes") or []
                        if isinstance(sizes, list):
                            for s in sizes:
                                bar_list = s.get("skus") or s.get("barcodes")
                                if isinstance(bar_list, list) and bar_list:
                                    r["barcode"] = str(bar_list[0])
                                    break
            try:
                now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
                save_fbs_tasks_cache({"rows": rows, "updated_at": now_str})
            except Exception:
                pass

        # 3) ????????? ??????: nm_id -> barcode -> name (?????????? ????)
        agg: Dict[tuple[str, str, str], int] = {}  # (name, nm_id, barcode) -> qty
        for r in rows:
            name = (r.get("???????????? ??????") or "").strip()
            nm_id = str(r.get("nm_id") or "").strip()
            barcode = str(r.get("barcode") or "").strip()
            key = (name, nm_id, barcode)
            agg[key] = agg.get(key, 0) + 1

        # Build XLS (not XLSX)
        try:
            import xlwt  # type: ignore
        except Exception:
            return ("?? ??????? ??????????? ??????????? xlwt (??? .xls)", 500)
        wb = xlwt.Workbook()
        ws = wb.add_sheet("FBS")
        header_style = xlwt.easyxf("font: bold on; align: horiz center")
        num_style = xlwt.easyxf("align: horiz right")
        ws.write(0, 0, "????????????", header_style)
        ws.write(0, 1, "??????? WB (nmId)", header_style)
        ws.write(0, 2, "??????", header_style)
        ws.write(0, 3, "??????????", header_style)
        row_idx = 1
        for (name, nm_id, barcode), qty in sorted(agg.items(), key=lambda x: (-x[1], x[0][0][0])):
            ws.write(row_idx, 0, name)
            ws.write(row_idx, 1, nm_id)
            ws.write(row_idx, 2, barcode)
            ws.write(row_idx, 3, qty, num_style)
            row_idx += 1
        out = io.BytesIO()
        wb.save(out)
        out.seek(0)
        filename = f"fbs_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xls"
        return send_file(out, mimetype="application/vnd.ms-excel", as_attachment=True, download_name=filename)
    except requests.HTTPError as http_err:
        return (f"?????? API: {http_err.response.status_code}", 502)
    except Exception as exc:
        return (f"??????: {exc}", 500)


# --- DBS active orders cache (to track in-progress tasks) ---
try:
    BASE_DIR
except NameError:
    import os as _os
    BASE_DIR = _os.path.dirname(_os.path.abspath(__file__))

DBS_CACHE_DIR = os.path.join(BASE_DIR, "cache")
DBS_ACTIVE_IDS_PATH = os.path.join(DBS_CACHE_DIR, "dbs_active_ids.json")
DBS_KNOWN_ORDERS_PATH = os.path.join(DBS_CACHE_DIR, "dbs_known_orders.json")

def _ensure_dbs_cache_dir() -> None:
    try:
        os.makedirs(DBS_CACHE_DIR, exist_ok=True)
    except Exception:
        pass

def load_dbs_active_ids() -> Dict[str, Any]:
    _ensure_dbs_cache_dir()
    try:
        with open(DBS_ACTIVE_IDS_PATH, "r", encoding="utf-8") as f:
            import json as _json
            return _json.load(f)
    except Exception:
        return {"ids": [], "updated_at": None}

def save_dbs_active_ids(data: Dict[str, Any]) -> None:
    _ensure_dbs_cache_dir()
    try:
        with open(DBS_ACTIVE_IDS_PATH, "w", encoding="utf-8") as f:
            import json as _json
            _json.dump(data, f, ensure_ascii=False)
    except Exception:
        pass

def add_dbs_active_ids(ids: list[int]) -> None:
    if not ids:
        return
    cache = load_dbs_active_ids() or {"ids": [], "updated_at": None}
    cur_ids = set(int(x) for x in (cache.get("ids") or []))
    for i in ids:
        try:
            cur_ids.add(int(i))
        except Exception:
            continue
    cache["ids"] = sorted(cur_ids)
    cache["updated_at"] = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    save_dbs_active_ids(cache)

def load_dbs_known_orders() -> Dict[str, Any]:
    _ensure_dbs_cache_dir()
    try:
        with open(DBS_KNOWN_ORDERS_PATH, "r", encoding="utf-8") as f:
            import json as _json
            return _json.load(f)
    except Exception:
        return {"orders": {}, "updated_at": None}

def save_dbs_known_orders(data: Dict[str, Any]) -> None:
    _ensure_dbs_cache_dir()
    try:
        with open(DBS_KNOWN_ORDERS_PATH, "w", encoding="utf-8") as f:
            import json as _json
            _json.dump(data, f, ensure_ascii=False)
    except Exception:
        pass

def add_dbs_known_orders(orders: list[dict[str, Any]]) -> None:
    if not orders:
        return
    cache = load_dbs_known_orders() or {"orders": {}, "updated_at": None}
    known: Dict[str, Any] = cache.get("orders") or {}
    for it in orders:
        oid = it.get("id") or it.get("orderId") or it.get("ID")
        if oid is None:
            continue
        try:
            key = str(int(oid))
        except Exception:
            continue
        known[key] = {"item": it, "seen_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
    cache["orders"] = known
    cache["updated_at"] = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    save_dbs_known_orders(cache)


@app.route("/dbs", methods=["GET"]) 
@login_required
def dbs_page():
    """DBS page: initial render; data loaded via JS."""
    error = None
    products_hint = None
    prod_cached_now = load_products_cache()
    if not prod_cached_now or not ((prod_cached_now or {}).get("items")):
        products_hint = "??? ??????????? ???? ?????? ? ??????? ???????? ?????? ?? ???????? ??????"
    return render_template("dbs.html", error=error, products_hint=products_hint)


@app.route("/api/dbs/orders/new", methods=["GET"]) 
@login_required
def api_dbs_orders_new():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"items": [], "updated_at": None}), 200
    try:
        raw = fetch_dbs_new_orders(token)
        try:
            raw_sorted = sorted(raw, key=_extract_created_at)
        except Exception:
            raw_sorted = raw
        rows = to_dbs_rows(raw_sorted)
        # Cache known orders and IDs for tracking in-progress statuses later
        try:
            add_dbs_known_orders(raw_sorted)
            ids_to_add: list[int] = []
            for it in raw_sorted:
                oid = it.get("id") or it.get("orderId") or it.get("ID")
                try:
                    if oid is not None:
                        ids_to_add.append(int(oid))
                except Exception:
                    continue
            add_dbs_active_ids(ids_to_add)
        except Exception:
            pass
        prod_cached = load_products_cache() or {}
        items = (prod_cached.get("items") or [])
        by_nm: Dict[int, Dict[str, Any]] = {}
        for it in items:
            nmv = it.get("nm_id") or it.get("nmID")
            try:
                if nmv:
                    by_nm[int(nmv)] = it
            except Exception:
                pass
        for r in rows:
            nm = r.get("nm_id")
            try:
                nm_i = int(nm) if nm is not None else None
            except Exception:
                nm_i = None
            hit = by_nm.get(nm_i) if nm_i is not None else None
            if hit:
                r["photo"] = hit.get("photo")
                if hit.get("barcode"):
                    r["barcode"] = hit.get("barcode")
                elif isinstance(hit.get("barcodes"), list) and hit.get("barcodes"):
                    r["barcode"] = str(hit.get("barcodes")[0])
                else:
                    sizes = hit.get("sizes") or []
                    if isinstance(sizes, list):
                        for s in sizes:
                            bl = s.get("skus") or s.get("barcodes")
                            if isinstance(bl, list) and bl:
                                r["barcode"] = str(bl[0])
                                break
        now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        return jsonify({"items": rows, "updated_at": now_str}), 200
    except Exception as exc:
        return jsonify({"items": [], "error": str(exc)}), 200


@app.route("/api/dbs/orders/<order_id>/deliver", methods=["PATCH"]) 
@login_required
def api_dbs_order_deliver(order_id: str):
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "No token"}), 401
    headers_list = [
        {"Authorization": f"{token}"},
        {"Authorization": f"Bearer {token}"},
    ]
    url = f"https://marketplace-api.wildberries.ru/api/v3/dbs/orders/{order_id}/deliver"
    last_err = None
    for hdrs in headers_list:
        try:
            resp = requests.patch(url, headers=hdrs, timeout=30)
            if resp.status_code in [200, 204]:
                try:
                    add_dbs_active_ids([int(order_id)])
                except Exception:
                    pass
                return jsonify({"success": True}), 200
            else:
                last_err = f"HTTP {resp.status_code}: {resp.text}"
                continue
        except Exception as e:
            last_err = str(e)
            continue
    return jsonify({"error": last_err or "Unknown error"}), 500


@app.route("/api/dbs/orders", methods=["GET"]) 
@login_required
def api_dbs_orders_list():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"items": [], "next": None}), 200
    try:
        limit = request.args.get("limit", default="1000")
        try:
            limit_i = max(1, min(1000, int(limit)))
        except Exception:
            limit_i = 1000
        next_val = request.args.get("next")
        df_q = request.args.get("dateFrom")
        dt_q = request.args.get("dateTo")
        if df_q and dt_q:
            try:
                date_from_ts = int(df_q)
                date_to_ts = int(dt_q)
            except Exception:
                date_from_ts = None
                date_to_ts = None
        else:
            now = datetime.now(MOSCOW_TZ)
            date_to_ts = int(now.timestamp())
            date_from_ts = int((now - timedelta(days=180)).timestamp())
        
        # Strategy: collect completed orders AND in-progress orders (confirm/deliver status)
        # 1. Get completed orders from /api/v3/dbs/orders
        raw = fetch_dbs_orders(
            token,
            limit=limit_i,
            next_cursor=next_val,
            date_from_ts=date_from_ts,
            date_to_ts=date_to_ts,
        )
        orders = []
        next_cursor = None
        if isinstance(raw, dict):
            arr_top = raw.get("orders")
            if isinstance(arr_top, list):
                orders = arr_top
            elif isinstance(arr_top, dict):
                inner_items = arr_top.get("items") or arr_top.get("data") or []
                if isinstance(inner_items, list):
                    orders = inner_items
                    next_cursor = arr_top.get("next") or next_cursor
            if not orders:
                data_val = raw.get("data")
                if isinstance(data_val, list):
                    orders = data_val
                elif isinstance(data_val, dict):
                    if isinstance(data_val.get("orders"), list):
                        orders = data_val.get("orders") or []
                    elif isinstance(data_val.get("items"), list):
                        orders = data_val.get("items") or []
                    next_cursor = data_val.get("next") if next_cursor is None else next_cursor
            if next_cursor is None:
                next_cursor = raw.get("next")
        
        print(f"DBS ORDERS: fetched {len(orders)} orders from /api/v3/dbs/orders")
        
        # Check statuses of completed orders - some might still be in progress
        completed_ids: list[int] = []
        for it in orders:
            oid = it.get("id") or it.get("orderId") or it.get("ID")
            try:
                if oid is not None:
                    completed_ids.append(int(oid))
            except Exception:
                continue
        
        # Filter out orders that are actually completed (receive/cancel)
        # Keep only those that are truly completed
        # Also collect in-progress orders from this list
        truly_completed_orders: list[dict[str, Any]] = []
        in_progress_from_completed: list[dict[str, Any]] = []
        if completed_ids:
            try:
                st = fetch_dbs_statuses(token, completed_ids[:1000])
                status_arr = st.get("orders") if isinstance(st, dict) else []
                status_map: dict[int, dict[str, Any]] = {}
                if isinstance(status_arr, list):
                    for x in status_arr:
                        try:
                            status_map[int(x.get("id") or x.get("orderId") or 0)] = x
                        except Exception:
                            continue
                
                for it in orders:
                    oid = it.get("id") or it.get("orderId") or it.get("ID")
                    try:
                        oid_i = int(oid) if oid is not None else None
                    except Exception:
                        oid_i = None
                    if oid_i is None:
                        continue
                    
                    sx = status_map.get(oid_i) or {}
                    supplier_status = (
                        sx.get("supplierStatus")
                        or sx.get("status")
                        or ""
                    ).lower()
                    
                    # If status is confirm/deliver, add to in-progress
                    if supplier_status in ("confirm", "deliver"):
                        it_copy = dict(it)
                        it_copy["status"] = supplier_status
                        it_copy["supplierStatus"] = supplier_status
                        if sx.get("wbStatus"):
                            it_copy["wbStatus"] = sx.get("wbStatus")
                        in_progress_from_completed.append(it_copy)
                    else:
                        # Only include truly completed orders (receive, cancel, reject)
                        truly_completed_orders.append(it)
            except Exception:
                # If status check fails, assume all are completed
                truly_completed_orders = orders
        else:
            truly_completed_orders = orders
        
        orders = truly_completed_orders
        
        # 2. Get new orders and check their statuses, filter those in confirm/deliver
        # Also check statuses of recent orders to find in-progress ones
        in_progress_orders: list[dict[str, Any]] = []
        all_order_ids_to_check: set[int] = set()
        recent_orders: list[dict[str, Any]] = []
        new_raw: list[dict[str, Any]] | None = None
        
        # Collect IDs from new orders
        try:
            new_raw = fetch_dbs_new_orders(token)
            print(f"DBS ORDERS: fetched {len(new_raw) if new_raw else 0} new orders")
            if new_raw:
                for it in new_raw:
                    oid = it.get("id") or it.get("orderId") or it.get("ID")
                    try:
                        if oid is not None:
                            all_order_ids_to_check.add(int(oid))
                    except Exception:
                        continue
        except Exception:
            pass
        
        # Also collect IDs from recent completed orders (same period as main query) to check if they're still in progress
        # This is important because orders in deliver status might appear in completed list
        # Use the same date range as the main query (180 days)
        try:
            recent_raw = fetch_dbs_orders(
                token,
                limit=1000,
                next_cursor=0,
                date_from_ts=date_from_ts,
                date_to_ts=date_to_ts,
            )
            if isinstance(recent_raw, dict):
                arr = recent_raw.get("orders")
                if isinstance(arr, list):
                    recent_orders = arr
                elif isinstance(recent_raw.get("data"), dict) and isinstance(recent_raw.get("data", {}).get("orders"), list):
                    recent_orders = recent_raw.get("data", {}).get("orders") or []
                elif isinstance(recent_raw.get("data"), list):
                    recent_orders = recent_raw.get("data") or []
            
            print(f"DBS ORDERS: fetched {len(recent_orders)} recent orders (for status check)")
            
            for it in recent_orders:
                oid = it.get("id") or it.get("orderId") or it.get("ID")
                try:
                    if oid is not None:
                        all_order_ids_to_check.add(int(oid))
                except Exception:
                    continue
        except Exception:
            pass
        
        # Also check recent orders from /api/v3/dbs/orders (last 7 days) to find in-progress ones
        # These might be in confirm/deliver status but not yet completed
        try:
            recent_7d_from = int((datetime.now(MOSCOW_TZ) - timedelta(days=7)).timestamp())
            recent_7d_to = int(datetime.now(MOSCOW_TZ).timestamp())
            recent_7d_raw = fetch_dbs_orders(
                token,
                limit=1000,
                next_cursor=0,
                date_from_ts=recent_7d_from,
                date_to_ts=recent_7d_to,
            )
            recent_7d_orders = []
            if isinstance(recent_7d_raw, dict):
                arr = recent_7d_raw.get("orders")
                if isinstance(arr, list):
                    recent_7d_orders = arr
                elif isinstance(recent_7d_raw.get("data"), dict) and isinstance(recent_7d_raw.get("data", {}).get("orders"), list):
                    recent_7d_orders = recent_7d_raw.get("data", {}).get("orders") or []
                elif isinstance(recent_7d_raw.get("data"), list):
                    recent_7d_orders = recent_7d_raw.get("data") or []
            
            print(f"DBS ORDERS: fetched {len(recent_7d_orders)} orders from last 7 days")
            
            # Add IDs to check
            for it in recent_7d_orders:
                oid = it.get("id") or it.get("orderId") or it.get("ID")
                try:
                    if oid is not None:
                        oid_i = int(oid)
                        if oid_i not in all_order_ids_to_check:
                            all_order_ids_to_check.add(oid_i)
                except Exception:
                    continue
        except Exception:
            pass
        
        # Include cached active IDs before deciding whether to check statuses
        known_cache = None
        try:
            active_cache = load_dbs_active_ids() or {}
            active_ids = active_cache.get("ids") or []
            for aid in active_ids:
                try:
                    ai = int(aid)
                    if ai not in all_order_ids_to_check:
                        all_order_ids_to_check.add(ai)
                except Exception:
                    continue
        except Exception:
            pass

        print(f"DBS ORDERS: total IDs to check: {len(all_order_ids_to_check)}")
        
        # Check statuses of all collected orders
        if all_order_ids_to_check:
            order_ids_list = list(all_order_ids_to_check)
            all_orders_map: dict[int, dict[str, Any]] = {}
            
            # Build map of all orders (from new + recent + recent_7d)
            if new_raw:
                for it in new_raw:
                    oid = it.get("id") or it.get("orderId") or it.get("ID")
                    try:
                        if oid is not None:
                            all_orders_map[int(oid)] = it
                    except Exception:
                        continue
            
            # Also add recent orders to map (they might be in progress)
            for it in recent_orders:
                oid = it.get("id") or it.get("orderId") or it.get("ID")
                try:
                    if oid is not None:
                        oid_i = int(oid)
                        if oid_i not in all_orders_map:
                            all_orders_map[oid_i] = it
                except Exception:
                    continue
            
            # Add recent 7d orders to map
            for it in recent_7d_orders:
                oid = it.get("id") or it.get("orderId") or it.get("ID")
                try:
                    if oid is not None:
                        oid_i = int(oid)
                        if oid_i not in all_orders_map:
                            all_orders_map[oid_i] = it
                except Exception:
                    continue

            # Add known cached orders to map to enrich in-progress items
            try:
                if known_cache is None:
                    known_cache = load_dbs_known_orders() or {}
                known_map = (known_cache.get("orders") or {})
                for k, v in known_map.items():
                    try:
                        ki = int(k)
                    except Exception:
                        continue
                    if ki not in all_orders_map and isinstance(v, dict):
                        item = v.get("item") if isinstance(v.get("item"), dict) else v
                        if isinstance(item, dict):
                            all_orders_map[ki] = item
            except Exception:
                pass

            # Also enrich map with items for active cached IDs
            try:
                active_cache = load_dbs_active_ids() or {}
                active_ids = active_cache.get("ids") or []
                for aid in active_ids:
                    try:
                        ai = int(aid)
                        if ai not in all_orders_map and (known_cache or {}):
                            item = ((known_cache or {}).get("orders") or {}).get(str(ai))
                            if isinstance(item, dict):
                                it = item.get("item") if isinstance(item.get("item"), dict) else item
                                if isinstance(it, dict):
                                    all_orders_map[ai] = it
                    except Exception:
                        continue
                order_ids_list = list(all_order_ids_to_check)
            except Exception:
                pass
            
            # Check statuses in batches
            for batch_start in range(0, len(order_ids_list), 1000):
                batch = order_ids_list[batch_start:batch_start + 1000]
                try:
                    st = fetch_dbs_statuses(token, batch)
                    status_arr = st.get("orders") if isinstance(st, dict) else []
                    status_map: dict[int, dict[str, Any]] = {}
                    if isinstance(status_arr, list):
                        for x in status_arr:
                            try:
                                status_map[int(x.get("id") or x.get("orderId") or 0)] = x
                            except Exception:
                                continue
                    
                    # Find orders in progress (confirm or deliver status)
                    for oid_i in batch:
                        sx = status_map.get(oid_i) or {}
                        supplier_status = (
                            sx.get("supplierStatus")
                            or sx.get("status")
                            or ""
                        ).lower()
                        
                        # Include orders in confirm or deliver status (in progress)
                        if supplier_status in ("confirm", "deliver"):
                            order_data = all_orders_map.get(oid_i)
                            if order_data:
                                # Add status info to order data
                                order_data_copy = dict(order_data)
                                order_data_copy["status"] = supplier_status
                                order_data_copy["supplierStatus"] = supplier_status
                                if sx.get("wbStatus"):
                                    order_data_copy["wbStatus"] = sx.get("wbStatus")
                                in_progress_orders.append(order_data_copy)
                            else:
                                # If order not in our map, create minimal record
                                in_progress_orders.append({
                                    "id": oid_i,
                                    "orderId": oid_i,
                                    "ID": oid_i,
                                    "status": supplier_status,
                                    "supplierStatus": supplier_status,
                                    "wbStatus": sx.get("wbStatus"),
                                })
                                print(f"DBS ORDERS: found in-progress order {oid_i} with status {supplier_status} (no data)")
                    # Debug logging
                    if batch_start == 0:
                        in_progress_count = len([oid for oid in batch if (status_map.get(oid) or {}).get("supplierStatus", "").lower() in ("confirm", "deliver")])
                        print(f"DBS ORDERS: checked {len(batch)} orders, found {in_progress_count} in progress")
                        # Show sample statuses
                        if len(batch) > 0:
                            sample_oid = batch[0]
                            sample_status = status_map.get(sample_oid) or {}
                            print(f"DBS ORDERS: sample order {sample_oid} status: {sample_status.get('supplierStatus')} / {sample_status.get('wbStatus')}")
                except Exception:
                    continue
        
        # Add in-progress orders found from "completed" list
        in_progress_orders.extend(in_progress_from_completed)
        print(f"DBS ORDERS: in_progress_from_completed: {len(in_progress_from_completed)}, total in_progress: {len(in_progress_orders)}")
        
        # Start with completed orders from main query
        all_orders = orders
        
        # Fallback strategy: if no data returned for 180d window, retry with 60d then 30d
        if not all_orders and not next_val:
            for days in (60, 30):
                try:
                    alt_from = int((datetime.now(MOSCOW_TZ) - timedelta(days=days)).timestamp())
                    alt_to = int(datetime.now(MOSCOW_TZ).timestamp())
                    alt_raw = fetch_dbs_orders(
                        token,
                        limit=limit_i,
                        next_cursor=0,
                        date_from_ts=alt_from,
                        date_to_ts=alt_to,
                    )
                    alt_orders = []
                    if isinstance(alt_raw, dict):
                        alt_arr = alt_raw.get("orders")
                        if isinstance(alt_arr, list):
                            alt_orders = alt_arr
                        elif isinstance(alt_raw.get("data"), dict) and isinstance(alt_raw.get("data", {}).get("orders"), list):
                            alt_orders = alt_raw.get("data", {}).get("orders") or []
                        elif isinstance(alt_raw.get("data"), list):
                            alt_orders = alt_raw.get("data") or []
                    if alt_orders:
                        all_orders.extend(alt_orders)
                        break
                except Exception:
                    continue
        
        # If ???? ?????? ? ??? ?????????, ?????????? ?????? ?? 30 ???? ?? 6 ??????? ?????
        if (not next_val) and (not all_orders or len(all_orders) < 20):
            try:
                combined: list[dict[str, Any]] = []
                seen: set[int] = set()
                now_ts = int(datetime.now(MOSCOW_TZ).timestamp())
                for offset in range(0, 180, 30):
                    wnd_to = now_ts - offset * 24 * 3600
                    wnd_from = now_ts - (offset + 30) * 24 * 3600
                    r = fetch_dbs_orders(
                        token,
                        limit=1000,
                        next_cursor=0,
                        date_from_ts=wnd_from,
                        date_to_ts=wnd_to,
                    )
                    arr: list[dict[str, Any]] = []
                    if isinstance(r, dict):
                        a = r.get("orders")
                        if isinstance(a, list):
                            arr = a
                        elif isinstance(r.get("data"), dict) and isinstance(r.get("data", {}).get("orders"), list):
                            arr = r.get("data", {}).get("orders") or []
                        elif isinstance(r.get("data"), list):
                            arr = r.get("data") or []
                    for it in arr:
                        oid = it.get("id") or it.get("orderId") or it.get("ID")
                        try:
                            oi = int(oid)
                        except Exception:
                            oi = None
                        if oi is not None and oi in seen:
                            continue
                        if oi is not None:
                            seen.add(oi)
                        combined.append(it)
                if combined:
                    # Check statuses of backfilled orders to find in-progress ones
                    backfill_ids: list[int] = []
                    for it in combined:
                        oid = it.get("id") or it.get("orderId") or it.get("ID")
                        try:
                            if oid is not None:
                                backfill_ids.append(int(oid))
                        except Exception:
                            continue
                    
                    if backfill_ids:
                        print(f"DBS ORDERS: checking {len(backfill_ids)} backfilled orders for in-progress status")
                        try:
                            st = fetch_dbs_statuses(token, backfill_ids[:1000])
                            status_arr = st.get("orders") if isinstance(st, dict) else []
                            status_map: dict[int, dict[str, Any]] = {}
                            if isinstance(status_arr, list):
                                for x in status_arr:
                                    try:
                                        status_map[int(x.get("id") or x.get("orderId") or 0)] = x
                                    except Exception:
                                        continue
                            
                            # Separate completed and in-progress orders
                            truly_completed_backfill: list[dict[str, Any]] = []
                            in_progress_backfill: list[dict[str, Any]] = []
                            
                            # Log all statuses for debugging
                            print(f"DBS ORDERS: statuses from API: {[(x.get('id'), x.get('supplierStatus'), x.get('wbStatus')) for x in status_arr[:10]]}")
                            
                            for it in combined:
                                oid = it.get("id") or it.get("orderId") or it.get("ID")
                                try:
                                    oid_i = int(oid) if oid is not None else None
                                except Exception:
                                    oid_i = None
                                if oid_i is None:
                                    truly_completed_backfill.append(it)
                                    continue
                                
                                sx = status_map.get(oid_i) or {}
                                supplier_status = (
                                    sx.get("supplierStatus")
                                    or sx.get("status")
                                    or ""
                                ).lower()
                                
                                if supplier_status in ("confirm", "deliver"):
                                    it_copy = dict(it)
                                    it_copy["status"] = supplier_status
                                    it_copy["supplierStatus"] = supplier_status
                                    if sx.get("wbStatus"):
                                        it_copy["wbStatus"] = sx.get("wbStatus")
                                    in_progress_backfill.append(it_copy)
                                    print(f"DBS ORDERS: found in-progress order {oid_i} from backfill with status {supplier_status}")
                                else:
                                    truly_completed_backfill.append(it)
                            
                            # Add completed backfill orders to all_orders
                            all_orders.extend(truly_completed_backfill)
                            # Add in-progress backfill orders to in_progress_orders
                            in_progress_orders.extend(in_progress_backfill)
                            print(f"DBS ORDERS: backfill - completed: {len(truly_completed_backfill)}, in_progress: {len(in_progress_backfill)}")
                        except Exception as e:
                            print(f"DBS ORDERS: error checking backfill statuses: {e}, adding all as completed")
                            # If status check failed, add all as completed
                            all_orders.extend(combined)
                    else:
                        # No IDs to check, add all as completed
                        all_orders.extend(combined)
            except Exception:
                pass
        
        # Combine completed and in-progress orders
        # Remove duplicates by order ID (in-progress take priority if both exist)
        all_orders_dict: dict[int, dict[str, Any]] = {}
        
        # First add completed orders
        for it in all_orders:
            oid = it.get("id") or it.get("orderId") or it.get("ID")
            try:
                if oid is not None:
                    all_orders_dict[int(oid)] = it
            except Exception:
                continue
        
        # Then add in-progress orders (they override completed if same ID)
        for it in in_progress_orders:
            oid = it.get("id") or it.get("orderId") or it.get("ID")
            try:
                if oid is not None:
                    all_orders_dict[int(oid)] = it
            except Exception:
                continue
        
        all_orders = list(all_orders_dict.values())
        
        # Debug logging
        print(f"DBS ORDERS: final - completed={len(all_orders) - len(in_progress_orders)}, in_progress={len(in_progress_orders)}, total={len(all_orders)}")
        
        try:
            all_orders.sort(key=_extract_created_at, reverse=True)
        except Exception:
            pass
        
        rows = to_dbs_rows(all_orders)
        prod_cached = load_products_cache() or {}
        items = (prod_cached.get("items") or [])
        by_nm: Dict[int, Dict[str, Any]] = {}
        for it in items:
            nmv = it.get("nm_id") or it.get("nmID")
            try:
                if nmv:
                    by_nm[int(nmv)] = it
            except Exception:
                pass
        for r in rows:
            nm = r.get("nm_id")
            try:
                nm_i = int(nm) if nm is not None else None
            except Exception:
                nm_i = None
            hit = by_nm.get(nm_i) if nm_i is not None else None
            if hit:
                r["photo"] = hit.get("photo")
                if hit.get("barcode"):
                    r["barcode"] = hit.get("barcode")
                elif isinstance(hit.get("barcodes"), list) and hit.get("barcodes"):
                    r["barcode"] = str(hit.get("barcodes")[0])
                else:
                    sizes = hit.get("sizes") or []
                    if isinstance(sizes, list):
                        for s in sizes:
                            bl = s.get("skus") or s.get("barcodes")
                            if isinstance(bl, list) and bl:
                                r["barcode"] = str(bl[0])
                                break
        
        # Save statuses of in-progress orders before merge (to preserve them)
        in_progress_status_map: dict[int, dict[str, Any]] = {}
        for r in rows:
            oid = r.get("orderId")
            status_val = (r.get("status") or "").lower()
            if oid and status_val in ("confirm", "deliver"):
                try:
                    in_progress_status_map[int(oid)] = {
                        "status": r.get("status"),
                        "statusName": r.get("statusName"),
                        "supplierStatus": r.get("status"),
                    }
                except Exception:
                    pass
        
        # Merge statuses (supplier/wb) via status endpoint for reliability
        # BUT preserve in-progress statuses
        try:
            ids: list[int] = []
            for it in all_orders:
                oid = it.get("id") or it.get("orderId") or it.get("ID")
                try:
                    if oid is not None:
                        ids.append(int(oid))
                except Exception:
                    continue
            if ids:
                st = fetch_dbs_statuses(token, ids[:1000])
                arr = st.get("orders") if isinstance(st, dict) else []
                m: dict[int, dict[str, Any]] = {}
                if isinstance(arr, list):
                    for x in arr:
                        try:
                            m[int(x.get("id") or x.get("orderId") or 0)] = x
                        except Exception:
                            continue
                for r in rows:
                    try:
                        oid = int(r.get("orderId") or 0)
                        # Skip merge if order is in progress (preserve its status)
                        if oid in in_progress_status_map:
                            continue
                        sx = m.get(oid) or {}
                        status_val = sx.get("status") or sx.get("supplierStatus") or sx.get("wbStatus") or r.get("status")
                        status_name_val = (
                            sx.get("statusName")
                            or sx.get("supplierStatusName")
                            or sx.get("wbStatusName")
                            or status_val
                        )
                        if status_name_val:
                            r["statusName"] = status_name_val
                        if status_val:
                            r["status"] = status_val
                    except Exception:
                        continue
        except Exception:
            pass
        
        return jsonify({"items": rows, "next": next_cursor}), 200
    except Exception as exc:
        return jsonify({"items": [], "next": None, "error": str(exc)}), 200


@app.route("/api/fbs/tasks", methods=["GET"]) 
@login_required
def api_fbs_tasks():
    token = effective_wb_api_token(current_user)
    refresh = request.args.get("refresh") in ("1", "true", "True")
    if not token and refresh:
        return jsonify({"items": [], "updated_at": None})
    try:
        if not refresh:
            cached = load_fbs_tasks_cache() or {}
            if cached.get("rows"):
                return jsonify({"items": cached.get("rows"), "updated_at": cached.get("updated_at")})
        # Fetch fresh
        print(f"=== FETCHING FRESH FBS TASKS ===")
        raw = fetch_fbs_new_orders(token)
        print(f"Raw tasks count: {len(raw)}")
        raw_sorted = sorted(raw, key=_extract_created_at)
        rows = to_fbs_rows(raw_sorted)
        print(f"Processed rows count: {len(rows)}")
        if rows:
            print(f"First row: {rows[0]}")
        # Enrich from products cache
        prod_cached = load_products_cache() or {}
        items = (prod_cached.get("items") or [])
        by_article: Dict[str, Dict[str, Any]] = {}
        by_nm: Dict[int, Dict[str, Any]] = {}
        for it in items:
            art = (it.get("supplier_article") or it.get("vendorCode") or "").strip()
            if art:
                by_article.setdefault(art, it)
            nm = it.get("nm_id") or it.get("nmID")
            if nm:
                try:
                    by_nm[int(nm)] = it
                except Exception:
                    pass
        for r in rows:
            art = (r.get("???????????? ??????") or "").strip()
            hit = by_article.get(art)
            if not hit and r.get("nm_id"):
                try:
                    hit = by_nm.get(int(r["nm_id"]))
                except Exception:
                    hit = None
            if hit:
                if hit.get("barcode"):
                    r["barcode"] = hit.get("barcode")
                elif isinstance(hit.get("barcodes"), list) and hit.get("barcodes"):
                    r["barcode"] = str(hit.get("barcodes")[0])
                else:
                    sizes = hit.get("sizes") or []
                    if isinstance(sizes, list):
                        for s in sizes:
                            bar_list = s.get("skus") or s.get("barcodes")
                            if isinstance(bar_list, list) and bar_list:
                                r["barcode"] = str(bar_list[0])
                                break
                r["photo"] = hit.get("photo")
        now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        save_fbs_tasks_cache({"rows": rows, "updated_at": now_str})
        return jsonify({"items": rows, "updated_at": now_str})
    except Exception as exc:
        return jsonify({"items": [], "error": str(exc)}), 200
@app.route("/api/fbs/orders/load-more", methods=["POST"]) 
@login_required
def api_fbs_orders_load_more():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"items": [], "next": None}), 200
    try:
        cursor = session.get("fbs_next_cursor")
        if not cursor:
            return jsonify({"items": [], "next": None}), 200
        page = fetch_fbs_orders(token, limit=30, next_cursor=cursor)
        items, next_cursor = _normalize_fbs_orders_page(page)
        try:
            items.sort(key=_extract_created_at, reverse=True)
        except Exception:
            pass
        # Merge statuses
        need_ids: List[int] = []
        for it in items:
            oid = it.get("id") or it.get("orderId") or it.get("ID")
            try:
                if oid is not None:
                    need_ids.append(int(oid))
            except Exception:
                pass
        if need_ids:
            st = fetch_fbs_statuses(token, need_ids[:1000])
            arr = st.get("orders") or st.get("data") or st
            m: Dict[int, Any] = {}
            if isinstance(arr, list):
                for x in arr:
                    try:
                        m[int(x.get("id") or x.get("orderId") or 0)] = x
                    except Exception:
                        continue
            for it in items:
                try:
                    oid = int(it.get("id") or it.get("orderId") or it.get("ID") or 0)
                    stx = m.get(oid) or {}
                    it["statusName"] = stx.get("statusName") or stx.get("status") or it.get("statusName") or it.get("status")
                    it["status"] = stx.get("status") or it.get("status")
                except Exception:
                    pass
        session["fbs_next_cursor"] = next_cursor
        # ?????? ?????? ??? ?????????? ????: ???? ?? ??????? ????????? ???????, last_next ??? ???????? ? ??????,
        # ?? ?????? ????? ????? ????????? next, ????? ???????? ??????. ?????????? ??? ?? ??????.
        return jsonify({"items": items, "next": session.get("fbs_next_cursor")}), 200
    except Exception as exc:
        return jsonify({"items": [], "next": None, "error": str(exc)}), 200


@app.route("/api/fbs/orders", methods=["GET"]) 
@login_required
def api_fbs_orders_first():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"items": [], "next": None}), 200
    try:
        items, last_next = get_orders_with_status(token, need_count=30, start_next=None)
        session["fbs_next_cursor"] = last_next
        return jsonify({"items": items, "next": last_next}), 200
    except Exception as exc:
        return jsonify({"items": [], "next": None, "error": str(exc)}), 200


@app.route("/api/fbs/orders/with-status", methods=["GET"]) 
@login_required
def api_fbs_orders_with_status():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"items": [], "next": None}), 200
    try:
        limit = request.args.get("limit", default="30")
        try:
            limit_i = max(1, min(200, int(limit)))
        except Exception:
            limit_i = 30
        next_val = request.args.get("next")
        items, last_next = get_orders_with_status(token, need_count=limit_i, start_next=next_val)
        session["fbs_next_cursor"] = last_next
        return jsonify({"items": items, "next": last_next}), 200
    except Exception as exc:
        return jsonify({"items": [], "next": None, "error": str(exc)}), 200


def _collect_fbs_orders_for_supplies(token: str, max_pages: int = 5, limit: int = 200) -> List[Dict[str, Any]]:
    collected: List[Dict[str, Any]] = []
    seen: set[int] = set()
    cursor: Any = 0
    pages = 0
    while pages < max_pages:
        page = fetch_fbs_orders(token, limit=limit, next_cursor=cursor)
        items, next_cursor = _normalize_fbs_orders_page(page)
        if not items:
            break
        for it in items:
            oid = it.get("id") or it.get("orderId") or it.get("ID")
            try:
                if oid is not None:
                    oid_i = int(oid)
                    if oid_i in seen:
                        continue
                    seen.add(oid_i)
            except Exception:
                pass
            collected.append(it)
        if not next_cursor or next_cursor == cursor:
            break
        cursor = next_cursor
        pages += 1
    return collected
def _aggregate_fbs_supplies(items: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    groups: Dict[str, Dict[str, Any]] = {}
    for it in items:
        sid = it.get("supplyId") or it.get("supplyID") or it.get("supply_id")
        if not sid:
            continue
        d_raw = it.get("createdAt") or it.get("dateCreated") or it.get("date")
        dt = parse_wb_datetime(str(d_raw)) if d_raw else None
        dt_msk = to_moscow(dt) if dt else None
        status_name = (it.get("statusName") or it.get("status") or "").strip()
        g = groups.get(str(sid))
        if not g:
            g = {"supplyId": str(sid), "count": 0, "last_dt": dt_msk, "status_counts": {}}
            groups[str(sid)] = g
        g["count"] = int(g.get("count", 0)) + 1
        # last date in the supply
        if dt_msk and (g.get("last_dt") is None or dt_msk > g.get("last_dt")):
            g["last_dt"] = dt_msk
        # collect status counts
        if status_name:
            sc = g.get("status_counts") or {}
            sc[status_name] = int(sc.get(status_name, 0)) + 1
            g["status_counts"] = sc

    result: List[Dict[str, Any]] = []
    for g in groups.values():
        dt_val = g.get("last_dt")
        date_str = dt_val.strftime("%d.%m.%Y %H:%M") if dt_val else ""
        ts = int(dt_val.timestamp()) if dt_val else 0
        # determine supply status: priority rules, then most frequent
        status_counts: Dict[str, int] = g.get("status_counts") or {}
        status_final = ""
        if status_counts:
            # Priority: ????????? ???????? -> ???????? ??????? -> otherwise most frequent
            lowered = {k.lower(): k for k in status_counts.keys()}
            # Look for '??????'
            for lk, orig in lowered.items():
                if "??????" in lk:
                    status_final = orig
                    break
            if not status_final:
                for lk, orig in lowered.items():
                    if "??????" in lk or "???????" in lk:
                        status_final = orig
                        break
            if not status_final:
                status_final = max(status_counts.items(), key=lambda x: x[1])[0]
        result.append({
            "supplyId": g["supplyId"],
            "date": date_str,
            "ts": ts,
            "count": g["count"],
            "status": status_final,
        })
    result.sort(key=lambda x: x.get("ts", 0), reverse=True)
    return result
# DEPRECATED: ???? ???? ????????? ? blueprints/fbs_supplies.py
# @app.route("/api/fbs/supplies", methods=["GET"]) 
# @login_required
# def api_fbs_supplies():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"items": [], "lastUpdated": None}), 200

    # Support two modes: cached read (default) and refresh=1 to re-fetch
    refresh_flag = request.args.get("refresh") in ("1", "true", "True")
    limit_param = request.args.get("limit", default="5")
    offset_param = request.args.get("offset", default="0")
    try:
        limit_i = max(1, min(1000, int(limit_param)))
    except Exception:
        limit_i = 5
    try:
        offset_i = int(offset_param)
    except Exception:
        offset_i = 0

    # ?????? ????????? ?? API (??? ???? ??????), ?? ? ?????????? ?????? ? fallback ?? ???
    try:
        # Load ALL supplies at once (fast - 262ms as per user's test)
        headers_list = [
            {"Authorization": f"{token}"},
            {"Authorization": f"Bearer {token}"},
        ]
        all_supplies_raw = []
        
        for hdrs in headers_list:
            try:
                # ?????????? ??????? 10 ?????? ??? ????????? ??????? (???????)
                resp = get_with_retry(FBS_SUPPLIES_LIST_URL, hdrs, params={"limit": 1000, "next": 0}, timeout_s=10)
                data = resp.json()
                
                # Handle both list and dict response formats
                if isinstance(data, list):
                    all_supplies_raw = data
                elif isinstance(data, dict):
                    all_supplies_raw = data.get("supplies", []) or data.get("data", []) or []
                else:
                    continue
                break
            except (requests.Timeout, requests.RequestException) as e:
                # ??? ???????? ??? ?????? ?????? ???? ??????? ????????? ?????????
                # ???? ??? ????????? ?????????, ?????????? ???
                if hdrs == headers_list[-1]:
                    cached = load_fbs_supplies_cache() or {}
                    cached_supplies = cached.get("all_supplies_raw", [])
                    if cached_supplies:
                        all_supplies_raw = cached_supplies
                        break
                continue
            except Exception as e:
                # ???? ??? ????????? ????????? ? ????????? ??????, ??????? ???
                if hdrs == headers_list[-1]:
                    cached = load_fbs_supplies_cache() or {}
                    cached_supplies = cached.get("all_supplies_raw", [])
                    if cached_supplies:
                        all_supplies_raw = cached_supplies
                        break
                continue
        
        if not all_supplies_raw:
            # ???? API ?? ?????? ??????, ??????? ???????????? ??? ?????? ?????????? ????? ?? ???????
            cached = load_fbs_supplies_cache() or {}
            cached_supplies = cached.get("all_supplies_raw", [])
            if cached_supplies:
                all_supplies_raw = cached_supplies
            else:
                # ?????? ???? ???? ???, ???????? ??????? ?? ??????? (?? ??? ????? ???? ????????)
                try:
                    orders = _collect_fbs_orders_for_supplies(token, max_pages=5, limit=100)  # ????????? ????????? ??? ????????
                    all_supplies_raw = _aggregate_fbs_supplies(orders)
                except Exception:
                    # ???? ? ??? ?? ???????, ?????????? ?????? ??????
                    all_supplies_raw = []
        
        # Sort all supplies by creation date (newest first)
        all_supplies_raw.sort(key=lambda x: x.get("createdAt", ""), reverse=True)
        
        # Get only the supplies we need to process (based on offset and limit)
        supplies_to_process = all_supplies_raw[offset_i:offset_i + limit_i]
        
        # Process only the supplies we need
        processed_supplies = []
        for s in supplies_to_process:
            if not isinstance(s, dict):
                continue
            supply_id = s.get("id") or s.get("supplyId") or s.get("supply_id")
            if not supply_id:
                continue

            # ?????????? ?(???? API ?? ?????? ???? ? 0, ????? ??? ????? ??????? ?????? ??? ?????????)
            count = s.get("orderCount") or s.get("ordersCount") or s.get("count") or 0

            # ?????????? ? ????? ? ???????
            created_at = s.get("createdAt") or s.get("dateCreated") or s.get("date")
            closed_at = s.get("closedAt") or s.get("doneAt")
            raw_status = str(s.get("status") or "").upper()
            done_flag = bool(s.get("done")) or raw_status in ("DONE", "CLOSED", "COMPLETED", "FINISHED", "SHIPPED")

            # Status label for UI
            if done_flag:
                status_label = "?????????"
                try:
                    _sdt = parse_wb_datetime(str(closed_at)) if closed_at else None
                    _sdt_msk = to_moscow(_sdt) if _sdt else None
                    status_dt = _sdt_msk.strftime("%d.%m.%Y %H:%M") if _sdt_msk else (str(closed_at) if closed_at else "")
                except Exception:
                    status_dt = str(closed_at) if closed_at else ""
            else:
                status_label = "?? ?????????"
                status_dt = ""
                
            # Date column should use createdAt
            try:
                date_dt = parse_wb_datetime(str(created_at)) if created_at else None
                date_msk = to_moscow(date_dt) if date_dt else None
                date_str = date_msk.strftime("%d.%m.%Y %H:%M") if date_msk else ""
            except Exception:
                date_str = str(created_at) if created_at else ""
            
            processed_supplies.append({
                "supplyId": str(supply_id),
                "date": date_str,
                "count": count,
                "status": status_label,
                "statusDt": status_dt,
            })

        # Save raw supplies to cache (for future pagination)
        now_msk = datetime.now(MOSCOW_TZ)
        cache_payload = {
            "all_supplies_raw": all_supplies_raw,  # Store raw data for pagination
            "lastUpdated": now_msk.strftime("%d.%m.%Y %H:%M"),
            "ts": int(now_msk.timestamp())
        }
        save_fbs_supplies_cache(cache_payload)

        return jsonify({
            "items": processed_supplies,
            "lastUpdated": cache_payload["lastUpdated"],
            "total": len(all_supplies_raw),
            "hasMore": offset_i + limit_i < len(all_supplies_raw)
        }), 200
    except Exception as exc:
        # ??? ?????? (??????? ???????) ?????????? ???????????? ??????, ???? ??? ????
        cached = load_fbs_supplies_cache() or {}
        cached_items = cached.get("all_supplies_raw", [])
        
        if cached_items:
            # ???????????? ???????????? ??????
            cached_items.sort(key=lambda x: x.get("createdAt", ""), reverse=True)
            supplies_to_process = cached_items[offset_i:offset_i + limit_i]
            
            processed_supplies = []
            for s in supplies_to_process:
                supply_id = s.get("id")
                if supply_id:
                    # ??? ???????????? ?????? ?????????? ??????????????? ?????????? ??? 0
                    count = s.get("count", 0)
                    
                    created_at = s.get("createdAt")
                    done = s.get("done")
                    closed_at = s.get("closedAt")
                    
                    if done:
                        status_label = "?????????"
                        try:
                            _sdt = parse_wb_datetime(str(closed_at))
                            _sdt_msk = to_moscow(_sdt) if _sdt else None
                            status_dt = _sdt_msk.strftime("%d.%m.%Y %H:%M") if _sdt_msk else str(closed_at)
                        except Exception:
                            status_dt = str(closed_at)
                    else:
                        status_label = "?? ?????????"
                        status_dt = None
                    
                    date_dt = parse_wb_datetime(str(created_at)) if created_at else None
                    date_msk = to_moscow(date_dt) if date_dt else None
                    date_str = date_msk.strftime("%d.%m.%Y %H:%M") if date_msk else ""
                    
                    processed_supplies.append({
                        "supplyId": str(supply_id),
                        "date": date_str,
                        "count": count,
                        "status": status_label,
                        "statusDt": status_dt or "",
                    })
            
            return jsonify({
                "items": processed_supplies,
                "lastUpdated": cached.get("lastUpdated", "??????????"),
                "total": len(cached_items),
                "hasMore": offset_i + limit_i < len(cached_items),
                "cached": True,
                "error": str(exc)
            }), 200
        
        # ???? ???? ???, ?????????? ?????? ??????
        return jsonify({"items": [], "error": str(exc), "lastUpdated": None}), 200


# DEPRECATED: ???? ???? ????????? ? blueprints/fbs_supplies.py
# @app.route("/api/fbs/supplies/<supply_id>/orders", methods=["GET"]) 
# @login_required
# def api_fbs_supply_orders(supply_id: str):
#     token = effective_wb_api_token(current_user)
#     if not token:
#         return jsonify({"items": []}), 200
#     headers_list = [
#         {"Authorization": f"{token}"},
#         {"Authorization": f"Bearer {token}"},
#     ]
#     last_err = None
#     try:
#         for hdrs in headers_list:
#             try:
#                 url = FBS_SUPPLY_ORDERS_URL.replace("{supplyId}", str(supply_id))
#                 resp = get_with_retry(url, hdrs, params={})
#                 data = resp.json()
#                 # API ????? ??????? ?????? ??????? ???????? ??? ????????? ? ???????
#                 if isinstance(data, dict):
#                     items = data.get("orders")
#                     if items is None and isinstance(data.get("data"), dict):
#                         items = data["data"].get("orders")
#                     if items is None and isinstance(data.get("data"), list):
#                         items = data["data"]
#                 elif isinstance(data, list):
#                     items = data
#                 else:
#                     items = []
#                 if not isinstance(items, list):
#                     items = []
#                 # Minimal normalization for frontend: id, article, barcode, nmId, photo
#                 norm = []
#                 prod_cached = load_products_cache() or {}
#                 by_nm: Dict[int, Dict[str, Any]] = {}
#                 try:
#                     for it in (prod_cached.get("items") or []):
#                         nmv = it.get("nm_id") or it.get("nmID")
#                         if nmv:
#                             by_nm[int(nmv)] = it
#                 except Exception:
#                     pass
#                 for it in items:
#                     if not isinstance(it, dict):
#                         continue
#                     nm = it.get("nmId") or it.get("nmID")
#                     photo = None
#                     barcode = None
#                     # format createdAt for item
#                     created_raw = it.get("createdAt") or it.get("dateCreated") or it.get("date")
#                     try:
#                         _dt = parse_wb_datetime(str(created_raw)) if created_raw else None
#                         _dt_msk = to_moscow(_dt) if _dt else None
#                         created_str = _dt_msk.strftime("%d.%m.%Y %H:%M") if _dt_msk else (str(created_raw) if created_raw else "")
#                     except Exception:
#                         created_str = str(created_raw) if created_raw else ""
#                     if nm:
#                         try:
#                             hit = by_nm.get(int(nm))
#                         except Exception:
#                             hit = None
#                         if hit:
#                             photo = hit.get("photo")
#                             if hit.get("barcode"):
#                                 barcode = hit.get("barcode")
#                             elif isinstance(hit.get("barcodes"), list) and hit.get("barcodes"):
#                                 barcode = str(hit.get("barcodes")[0])
#                             else:
#                                 sizes = hit.get("sizes") or []
#                                 if isinstance(sizes, list):
#                                     for s in sizes:
#                                         bl = s.get("skus") or s.get("barcodes")
#                                         if isinstance(bl, list) and bl:
#                                             barcode = str(bl[0])
#                                             break
#                     norm.append({
#                         "id": it.get("id") or it.get("orderId") or it.get("ID"),
#                         "article": it.get("article") or it.get("vendorCode") or "",
#                         "barcode": (it.get("skus")[0] if isinstance(it.get("skus"), list) and it.get("skus") else None) or barcode or "",
#                         "nmId": nm,
#                         "photo": photo,
#                         "createdAt": created_str,
#                     })
#                 return jsonify({"items": norm}), 200
#             except Exception as e:
#                 last_err = e
#                 continue
#         if last_err:
#             raise last_err
#         return jsonify({"items": []}), 200
#     except Exception as exc:
#         return jsonify({"items": [], "error": str(exc)}), 200


# DEPRECATED: ???? ???? ????????? ? blueprints/fbs_supplies.py
# @app.route("/api/fbs/supplies/<supply_id>/orders/<order_id>", methods=["PATCH"])
# @login_required
# def api_fbs_add_order_to_supply(supply_id: str, order_id: str):
#     """???????? ????????? ??????? ? ????????"""
#     token = effective_wb_api_token(current_user)
#     if not token:
#         return jsonify({"error": "No token"}), 401
#     
#     headers_list = [
#         {"Authorization": f"{token}"},
#         {"Authorization": f"Bearer {token}"},
#     ]
#     
#     # URL ??? ?????????? ??????? ? ????????
#     url = f"https://marketplace-api.wildberries.ru/api/v3/supplies/{supply_id}/orders/{order_id}"
#     
#     last_err = None
#     for hdrs in headers_list:
#         try:
#             resp = requests.patch(url, headers=hdrs, timeout=30)
#             if resp.status_code in [200, 204]:  # 204 = No Content (???????)
#                 return jsonify({"success": True}), 200
#             elif resp.status_code == 409:
#                 # ??????? ??? ? ????????
#                 return jsonify({"error": "Order already in supply"}), 409
#             else:
#                 last_err = f"HTTP {resp.status_code}: {resp.text}"
#                 continue
#         except Exception as e:
#             last_err = str(e)
#             continue
#     
#     return jsonify({"error": last_err or "Unknown error"}), 500


@app.route("/api/fbs/supplies/create", methods=["POST"])
@login_required
def api_fbs_create_supply():
    """??????? ????? ????????"""
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "No token"}), 401
    
    headers_list = [
        {"Authorization": f"{token}"},
        {"Authorization": f"Bearer {token}"},
    ]
    
    # URL ??? ???????? ????????
    url = "https://marketplace-api.wildberries.ru/api/v3/supplies"
    
    last_err = None
    for hdrs in headers_list:
        try:
            # ????????? Content-Type ??? JSON
            hdrs_with_content_type = hdrs.copy()
            hdrs_with_content_type["Content-Type"] = "application/json"
            
            resp = requests.post(url, headers=hdrs_with_content_type, json={}, timeout=30)
            if resp.status_code in [200, 201]:
                data = resp.json()
                supply_id = data.get("id") or data.get("supplyId") or "??????????"
                return jsonify({"success": True, "supplyId": supply_id}), 200
            else:
                last_err = f"HTTP {resp.status_code}: {resp.text}"
                continue
        except Exception as e:
            last_err = str(e)
            continue
    
    return jsonify({"error": last_err or "Unknown error"}), 500


@app.route("/coefficients", methods=["GET", "POST"]) 
@login_required
def coefficients_page():
    error = None
    token = effective_wb_api_token(current_user)
    items: List[Dict[str, Any]] | None = None
    warehouses: List[str] = []
    date_keys: List[str] = []
    date_labels: List[str] = []
    grid: Dict[str, Dict[str, Dict[str, Any]]] = {}
    generated_at = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
    period_label = None

    if not token:
        error = "??????? ????? API ?? ???????? ?????????"
    else:
        try:
            items = fetch_acceptance_coefficients(token)
            if not isinstance(items, list):
                items = []
            warehouses, date_keys, date_labels, grid = build_acceptance_grid(items, days=14)
            if date_keys:
                try:
                    start = datetime.strptime(date_keys[0], "%Y-%m-%d").date()
                    end = datetime.strptime(date_keys[-1], "%Y-%m-%d").date()
                    period_label = f"{start.strftime('%d.%m')} ?? {end.strftime('%d.%m')}"
                except Exception:
                    period_label = None
        except requests.HTTPError as http_err:
            error = f"?????? API: {http_err.response.status_code}"
        except Exception as exc:
            error = f"??????: {exc}"

    return render_template(
        "coefficients.html",
        error=error,
        warehouses=warehouses,
        date_labels=date_labels,
        date_keys=date_keys,
        grid=grid,
        generated_at=generated_at,
        period_label=period_label,
    )
@app.route("/api/acceptance-coefficients", methods=["GET"]) 
@login_required
def api_acceptance_coefficients():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        items = fetch_acceptance_coefficients(token) or []
        if not isinstance(items, list):
            items = []
        warehouses, date_keys, date_labels, grid = build_acceptance_grid(items, days=14)
        return jsonify({
            "warehouses": warehouses,
            "date_keys": date_keys,
            "date_labels": date_labels,
            "grid": grid,
            "lastUpdated": datetime.now().strftime("%d.%m.%Y %H:%M:%S"),
        })
    except requests.HTTPError as http_err:
        status = 502
        retry_after = None
        try:
            if http_err.response is not None:
                if http_err.response.status_code:
                    status = http_err.response.status_code
                retry_after = http_err.response.headers.get("Retry-After")
        except Exception:
            status = 502
        cached = load_stocks_cache()
        return jsonify({
            "error": "http",
            "status": status,
            "retry_after": retry_after,
            "updated_at": (cached.get("updated_at") if cached and cached.get("_user_id") == current_user.id else None)
        }), status
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


def post_with_retry(url: str, headers: Dict[str, str], json_body: Dict[str, Any], max_retries: int = 3) -> requests.Response:
    last_exc: Exception | None = None
    last_resp: requests.Response | None = None
    for attempt in range(max_retries):
        try:
            resp = requests.post(url, headers=headers, json=json_body, timeout=30)
            last_resp = resp
            if resp.status_code in (429, 500, 502, 503, 504):
                retry_after = resp.headers.get("Retry-After")
                if retry_after is not None:
                    try:
                        sleep_s = float(retry_after)
                    except ValueError:
                        sleep_s = 1.0
                else:
                    sleep_s = min(15, 0.8 * (2 ** attempt) + random.uniform(0, 0.7))
                time.sleep(sleep_s)
                continue
            resp.raise_for_status()
            return resp
        except requests.RequestException as exc:
            last_exc = exc
            time.sleep(min(8, 0.5 * (2 ** attempt) + random.uniform(0, 0.5)))
            continue
    if last_exc:
        raise last_exc
    if last_resp is not None:
        raise requests.HTTPError(f"HTTP {last_resp.status_code} after {max_retries} retries", response=last_resp)
    raise RuntimeError("Request failed after retries")


def fetch_cards_list(
    token: str,
    nm_ids: List[int] | None = None,
    cursor: Dict[str, Any] | None = None,
    limit: int = 100,
    text_search: str | None = None,
    vendor_codes: List[str] | None = None,
) -> Dict[str, Any]:
    # Build request body per WB docs: settings.cursor + settings.filter
    base_cursor = {"limit": limit, "nmID": 0}
    if cursor:
        base_cursor.update(cursor)
    body: Dict[str, Any] = {
        "settings": {
            "cursor": base_cursor,
            "filter": {
                "textSearch": (text_search or ""),
                "withPhoto": -1,  # -1 ? ?? ??????????? ?? ??????? ????
            },
        }
    }
    if nm_ids:
        # ?????? ?????? ???????? ?? ?????? nmID ? ???? ???????? ??????
        body["nmID"] = nm_ids
    if vendor_codes:
        # ?????? ?????? ?? ???????? ????????
        body["settings"]["filter"]["vendorCode"] = vendor_codes
    # Try with Bearer first, then raw token (Content API ????? ????????? ??? Bearer)
    headers1 = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    try:
        resp = post_with_retry(WB_CARDS_LIST_URL, headers1, body)
        return resp.json()
    except requests.HTTPError as err:
        if err.response is not None and err.response.status_code in (401, 403):
            headers2 = {"Authorization": f"{token}", "Content-Type": "application/json"}
            resp2 = post_with_retry(WB_CARDS_LIST_URL, headers2, body)
            return resp2.json()
        raise


def fetch_all_cards(token: str, page_limit: int = 1000) -> List[Dict[str, Any]]:
    all_cards: List[Dict[str, Any]] = []
    seen_keys: set[tuple] = set()
    cursor: Dict[str, Any] = {"limit": page_limit, "nmID": 0}
    safety = 0
    while True:
        safety += 1
        if safety > 5000:
            break
        data = fetch_cards_list(token, cursor=cursor, limit=page_limit)
        payload = data.get("data") or data
        cards = payload.get("cards") or []
        if not cards:
            break
        all_cards.extend(cards)
        cur = payload.get("cursor") or {}
        key = (cur.get("updatedAt"), cur.get("nmID"), cur.get("nmIDNext"))
        if key in seen_keys:
            break
        seen_keys.add(key)
        # Prepare next cursor
        next_nm = cur.get("nmIDNext") or cur.get("nmID")
        next_cursor: Dict[str, Any] = {"limit": page_limit}
        if cur.get("updatedAt"):
            next_cursor["updatedAt"] = cur.get("updatedAt")
        if next_nm is not None:
            next_cursor["nmID"] = next_nm
        cursor = next_cursor
        # If ???????? ?????? ??????, ????????, ????????? ?????
        if len(cards) < page_limit:
            break
    return all_cards
def normalize_cards_response(data: Dict[str, Any]) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    try:
        payload = data.get("data") or data
        cards = payload.get("cards") or []
        for c in cards:
            nm_id = c.get("nmID") or c.get("nmId") or c.get("nm")
            supplier_article = c.get("supplierArticle") or c.get("vendorCode") or c.get("article")
            photo = None
            try:
                photos = c.get("mediaFiles") or c.get("photos") or []
                if isinstance(photos, list) and photos:
                    p0 = photos[0]
                    if isinstance(p0, str):
                        photo = p0
                    elif isinstance(p0, dict):
                        photo = p0.get("small") or p0.get("preview") or p0.get("url") or p0.get("big")
                if isinstance(photo, str) and photo.startswith("//"):
                    photo = "https:" + photo
                if isinstance(photo, str) and not (photo.startswith("http://") or photo.startswith("https://")):
                    photo = "https://" + photo.lstrip("/")
            except Exception:
                photo = None
            barcode = None
            size_info = None
            try:
                sizes = c.get("sizes") or []
                for s in sizes:
                    # ???????? ?????? ?????? ??? ???????? chrtID ? ?????????
                    chrt_id = s.get("chrtID")
                    skus = s.get("skus") or s.get("barcodes") or []
                    if skus and not barcode:
                        barcode = str(skus[0])
                    if chrt_id:
                        size_info = {
                            "chrtID": chrt_id,
                            "skus": [str(x) for x in (s.get("skus") or [])]
                        }
                        break
            except Exception:
                barcode = None
            # ???????? ???????? ??????
            name = c.get("name") or c.get("title") or c.get("subject") or "??? ????????"
            
            # ???????? subject_id ??? ????????
            subject_id = c.get("subjectID") or c.get("subjectId") or c.get("subject_id")
            
            # ???????? ??????? ??????
            dimensions = c.get("dimensions") or {}
            length = dimensions.get("length", 0)
            width = dimensions.get("width", 0)
            height = dimensions.get("height", 0)
            weight = dimensions.get("weightBrutto", 0)
            volume = (length * width * height) / 1000 if all([length, width, height]) else 0
            
            items.append({
                "photo": photo,
                "supplier_article": supplier_article,
                "nm_id": nm_id,
                "barcode": barcode,
                "chrt_id": size_info.get("chrtID") if size_info else None,
                "name": name,
                "subject_id": subject_id,
                "dimensions": {
                    "length": length,
                    "width": width,
                    "height": height,
                    "weight": weight,
                    "volume": round(volume, 2)
                },
                "size_info": size_info
            })
    except Exception:
        pass
    return items
def update_cards_dimensions(token: str, updates: List[Dict[str, Any]]) -> Dict[str, Any]:
    """
    ????????? ??????? ? ??? ??????? ????? API Wildberries
    
    Args:
        token: API ????? Wildberries
        updates: ?????? ?????????? ? ??????? [{"nmID": int, "vendorCode": str, "barcode": str, "chrtID": int, "dimensions": {...}}]
    
    Returns:
        ????????? ??????????
    """
    if not updates:
        return {"error": "??? ?????? ??? ??????????"}
    
    # ???????? ??? ??????? ??? fallback ??????
    try:
        cached = load_products_cache()
        products_cache = {item.get("nm_id"): item for item in cached.get("items", [])} if cached else {}
    except Exception as e:
        print(f"?????? ???????? ???? ???????: {e}")
        products_cache = {}
    
    # ?????????????? ?????? ??? API (?????? ????????? ??? ???? ????????, ?????? ?????? dimensions)
    cards_data = []
    skipped = []

    for update in updates:
        nm_id = update.get("nmID")
        vendor_from_ui = (update.get("vendorCode") or "").strip()
        dimensions = update.get("dimensions", {})

        if not nm_id:
            skipped.append({"nmID": nm_id, "reason": "no_nmID"})
            continue

        # 1) ???????? ?????? ????????: ??????? ?????? ?? nmID; ???? ?? ????? ? ?? vendorCode (UI/???);
        #    ???? ??? ??? ??? ? ??????? textSearch ?? nmID ???????. ???? ? ?????? ????? ?????? ???????? ??? merge.
        base = None
        last_err = None
        try:
            fc = fetch_cards_list(token, nm_ids=[nm_id], limit=2)
            payload = fc.get("data") or fc
            cards = payload.get("cards") or []
            exact = [c for c in cards if (c.get("nmID") == nm_id or c.get("nmId") == nm_id)]
            if exact:
                base = exact[0]
        except Exception as e_fetch:
            last_err = e_fetch
            print(f"fetch by nmID failed for {nm_id}: {e_fetch}")

        if base is None:
            # fallback by vendorCode
            vendor_fallback = vendor_from_ui
            if not vendor_fallback:
                try:
                    vendor_fallback = (products_cache.get(nm_id) or {}).get("supplier_article") or ""
                except Exception:
                    vendor_fallback = ""
            if vendor_fallback:
                try:
                    fc_v = fetch_cards_list(token, vendor_codes=[str(vendor_fallback)], limit=100)
                    pl_v = fc_v.get("data") or fc_v
                    cards_v = pl_v.get("cards") or []
                    # ??????? ?? ??????? nmID, ????? ?? ?????????? vendorCode
                    exact_nm = [c for c in cards_v if (c.get("nmID") == nm_id or c.get("nmId") == nm_id)]
                    base = exact_nm[0] if exact_nm else next((c for c in cards_v if str(c.get("vendorCode", "")).strip() == str(vendor_fallback).strip()), None)
                except Exception as e_v:
                    last_err = e_v
                    print(f"fetch by vendorCode failed for {nm_id}/{vendor_fallback}: {e_v}")

        if base is None:
            # last resort: textSearch by nmID string
            try:
                fc_t = fetch_cards_list(token, text_search=str(nm_id), limit=50)
                pl_t = fc_t.get("data") or fc_t
                cards_t = pl_t.get("cards") or []
                base = next((c for c in cards_t if (c.get("nmID") == nm_id or c.get("nmId") == nm_id)), None)
            except Exception as e_t:
                last_err = e_t
                print(f"fetch by textSearch failed for {nm_id}: {e_t}")

        if base is None:
            reason = "not_found"
            if last_err is not None:
                reason = f"fetch_failed: {last_err}"
            skipped.append({"nmID": nm_id, "reason": reason})
            continue

        # 2) ????????? ?????? ?????? ???? dimensions
        base_dimensions = (base.get("dimensions") or {}).copy()
        if "length" in dimensions:
            base_dimensions["length"] = dimensions["length"]
        if "width" in dimensions:
            base_dimensions["width"] = dimensions["width"]
        if "height" in dimensions:
            base_dimensions["height"] = dimensions["height"]
        if "weightBrutto" in dimensions:
            base_dimensions["weightBrutto"] = dimensions["weightBrutto"]
        base_dimensions["isValid"] = True

        # 3) ???????? ????????, ???????? ??? ????????? ???? ??? ?????????
        card_data = {
            "nmID": base.get("nmID") or nm_id,
            "vendorCode": base.get("vendorCode", ""),
            "dimensions": base_dimensions,
            "sizes": base.get("sizes", []),
            "characteristics": base.get("characteristics", []),
            "title": base.get("title", ""),
            "description": base.get("description", ""),
            "brand": base.get("brand", "")
        }

        cards_data.append(card_data)
        print(f"????????? ???????? ??? ?????????? (merge): nmID={nm_id}")
    
    if not cards_data:
        return {"ok": True, "skipped": skipped, "message": "??? ???????? ?????? ??? ??????????"}
    
    # ?????????? ??????? ?? ??????
    headers1 = {"Authorization": f"Bearer {token}", "Content-Type": "application/json"}
    aggregated_results = []
    aggregated_errors = []

    for card in cards_data:
        nm = card.get("nmID")
        try:
            print(f"?????????? ?????? ? {WB_CARDS_UPDATE_URL} ??? nmID {nm}")
            print(f"?????? ??? ????????: {[card]}")
            resp = post_with_retry(WB_CARDS_UPDATE_URL, headers1, [card])
            print(f"????? API: ?????? {resp.status_code}, ?????: {resp.text}")
            
            if resp.status_code == 200:
                try:
                    result = resp.json()
                    aggregated_results.append(result)
                    print(f"??????? ???????? nmID {nm}")
                except Exception as e:
                    aggregated_results.append({"nmID": nm, "raw": resp.text, "error": str(e)})
            else:
                error_info = {"nmID": nm, "status": resp.status_code, "text": resp.text}
                aggregated_errors.append(error_info)
                print(f"?????? ?????????? nmID {nm}: {resp.status_code} - {resp.text}")

        except requests.HTTPError as err:
            if err.response is not None and err.response.status_code in (401, 403):
                # ??????? ? ?????? ???????? ???????????
                headers2 = {"Authorization": f"{token}", "Content-Type": "application/json"}
                try:
                    resp2 = post_with_retry(WB_CARDS_UPDATE_URL, headers2, [card])
                    if resp2.status_code == 200:
                        try:
                            result = resp2.json()
                            aggregated_results.append(result)
                            print(f"??????? ???????? nmID {nm} (?????? ??????)")
                        except Exception as e:
                            aggregated_results.append({"nmID": nm, "raw": resp2.text, "error": str(e)})
                    else:
                        error_info = {"nmID": nm, "status": resp2.status_code, "text": resp2.text}
                        aggregated_errors.append(error_info)
                        print(f"?????? ?????????? nmID {nm} (?????? ??????): {resp2.status_code} - {resp2.text}")
                except Exception as e2:
                    error_info = {"nmID": nm, "error": f"?????? ???????????: {e2}"}
                    aggregated_errors.append(error_info)
                    print(f"?????? ??????????? ??? nmID {nm}: {e2}")
            else:
                error_info = {"nmID": nm, "error": f"HTTP ?????? {err.response.status_code}: {err.response.text}" if err.response else f"HTTP ??????: {err}"}
                aggregated_errors.append(error_info)
                print(f"HTTP ?????? ??? nmID {nm}: {error_info['error']}")
        except Exception as e:
            error_info = {"nmID": nm, "error": f"?????? ???????: {e}"}
            aggregated_errors.append(error_info)
            print(f"?????? ??????? ??? nmID {nm}: {e}")

    # ?????????? ?????????
    result = {
        "ok": True,
        "updated_count": len(aggregated_results),
        "skipped": skipped,
        "results": aggregated_results
    }
    
    if aggregated_errors:
        result["errors"] = aggregated_errors
    
    return result


def fetch_commission_data(token: str) -> Dict[int, Dict[str, Any]]:
    """???????? ?????? ? ????????? Wildberries ?? ???? ??????????"""
    try:
        print("???????? ?????? ? ?????????...")
        
        # ????????? ?????? ???????? ??????????
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        response = requests.get(
            COMMISSION_API_URL,
            headers=headers,
            timeout=30
        )
        
        print(f"?????? ?????? ????????: {response.status_code}")
        print(f"????? ????????: {response.text[:500]}...")
        
        if response.status_code == 200:
            result = response.json()
            # API ?????????? ?????? ? ??????? {"report": [...]}
            report_data = result.get("report", [])
            print(f"???????? {len(report_data)} ??????? ? ?????????")
            if report_data:
                print(f"?????? ??????: {report_data[0]}")
            
            # ??????? ??????? ??? ???????? ?????? ?? subjectID
            commission_data = {}
            for item in report_data:
                subject_id = item.get("subjectID")
                if subject_id:
                    commission_data[subject_id] = {
                        "parent_name": item.get("parentName", ""),
                        "subject_name": item.get("subjectName", ""),
                        "fbs_commission": item.get("kgvpMarketplace", 0),  # FBS
                        "cc_commission": item.get("kgvpPickup", 0),  # C&C
                        "dbs_dbw_commission": item.get("kgvpSupplier", 0),  # DBS + DBW
                        "edbs_commission": item.get("kgvpSupplierExpress", 0),  # EDBS
                        "fbw_commission": item.get("paidStorageKgvp", 0),  # FBW
                    }
            
            print(f"?????????? {len(commission_data)} ????????")
            if commission_data:
                print(f"?????? ????????: {list(commission_data.values())[0]}")
                print(f"????????? subject_id: {list(commission_data.keys())[:10]}...")  # ?????? 10 ID
            return commission_data
        else:
            print(f"?????? ????????? ????????: {response.status_code} - {response.text}")
            
    except Exception as e:
        print(f"?????? ??? ????????? ????????: {e}")
        import traceback
        traceback.print_exc()
    
    return {}


def fetch_warehouses_data(token: str) -> List[Dict[str, Any]]:
    """????????? ?????? ? ??????? ????? API Wildberries"""
    try:
        from datetime import datetime
        
        # ???????? ??????? ???? ? ??????? ????-??-??
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        print(f"???????? ?????? ? ??????? ?? ???? {current_date}")
        
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }
        
        # ????????? ???????? date ? URL
        url = f"{WAREHOUSES_API_URL}?date={current_date}"
        
        response = requests.get(url, headers=headers, timeout=30)
        print(f"?????? ?????? API ???????: {response.status_code}")
        response.raise_for_status()
        data = response.json()
        
        print(f"????? API ???????: {data}")
        
        # ???????????? ?????
        warehouses = []
        if isinstance(data, list):
            warehouses = data
        elif isinstance(data, dict):
            # ????????? ?????? ????????? ????????? ??????
            if 'response' in data and isinstance(data['response'], dict):
                # ?????????: {'response': {'data': {'warehouseList': [...]}}}
                response_data = data['response']
                if 'data' in response_data and isinstance(response_data['data'], dict):
                    if 'warehouseList' in response_data['data']:
                        warehouses = response_data['data']['warehouseList']
                elif 'data' in response_data and isinstance(response_data['data'], list):
                    warehouses = response_data['data']
            elif 'data' in data:
                if isinstance(data['data'], list):
                    warehouses = data['data']
                elif isinstance(data['data'], dict) and 'warehouseList' in data['data']:
                    warehouses = data['data']['warehouseList']
            elif 'warehouseList' in data:
                warehouses = data['warehouseList']
        
        print(f"??????? {len(warehouses)} ??????? ? ?????? API")
        
        # ?????????? ?????????? ? ????????? ??????
        if warehouses:
            print(f"?????? ????? (?????????): {warehouses[0]}")
            print(f"????? ??????? ??????: {list(warehouses[0].keys())}")
        
        # ????????? ?????? ??????? ? ??????? ???????
        warehouses_list = []
        for i, warehouse in enumerate(warehouses):
            # ??????? ?????? ????????? ???????? ?????
            warehouse_name = (warehouse.get('warehouseName') or 
                            warehouse.get('name') or 
                            warehouse.get('warehouse_name') or 
                            warehouse.get('title') or '')
            
            box_delivery_coef = (warehouse.get('boxDeliveryCoefExpr') or 
                               warehouse.get('coefficient') or 
                               warehouse.get('coef') or 0)
            
            print(f"????? {i+1}: name='{warehouse_name}', coef='{box_delivery_coef}'")
            print(f"  ??? ????: {list(warehouse.keys())}")
            
            if warehouse_name:
                try:
                    # ??????????? ??????????? ? ????? ?????
                    coef_value = int(float(box_delivery_coef)) if box_delivery_coef else 0
                    warehouses_list.append({
                        'name': warehouse_name,
                        'coefficient': coef_value
                    })
                except (ValueError, TypeError):
                    print(f"  ?????? ?????????????? ????????????: {box_delivery_coef}")
                    warehouses_list.append({
                        'name': warehouse_name,
                        'coefficient': 0
                    })
        
        print(f"????????? {len(warehouses_list)} ???????")
        return warehouses_list
        
    except Exception as e:
        print(f"?????? ???????? ???????: {e}")
        import traceback
        traceback.print_exc()
        return []


def fetch_dimensions_data(token: str, nm_ids: List[int]) -> Dict[int, Dict[str, Any]]:
    """????????? ?????? ? ???????? ??????? ????? API Wildberries"""
    if not nm_ids:
        return {}
    
    try:
        print(f"???????? ??????? ??? {len(nm_ids)} ???????")
        
        # ????????? ?? ????? ?? 100 ??????? (????? API)
        batch_size = 100
        dimensions_dict = {}
        
        for i in range(0, len(nm_ids), batch_size):
            batch_nm_ids = nm_ids[i:i + batch_size]
            
            payload = {
                "id": [str(nm_id) for nm_id in batch_nm_ids]
            }
            
            headers = {
                "Authorization": f"Bearer {token}",
                "Content-Type": "application/json"
            }
            
            response = requests.post(DIMENSIONS_API_URL, json=payload, headers=headers, timeout=30)
            print(f"?????? ?????? API ????????: {response.status_code}")
            response.raise_for_status()
            data = response.json()
            
            print(f"????? API ????????: {data}")
            print(f"????? ? ??????: {list(data.keys()) if isinstance(data, dict) else '?? ???????'}")
            
            # ???????????? ????? - ????????? ?????? ????????? ?????????
            items = []
            if 'data' in data:
                items = data['data']
            elif isinstance(data, list):
                items = data
            elif 'cards' in data:
                items = data['cards']
            
            print(f"??????? {len(items)} ????????? ??? ?????????")
            
            for item in items:
                print(f"???????????? ???????: {item}")
                nm_id = item.get('id') or item.get('nm_id')
                dimensions = item.get('dimensions', {})
                
                print(f"nm_id: {nm_id}, dimensions: {dimensions}")
                
                if nm_id and dimensions:
                    length = dimensions.get('length', 0)
                    width = dimensions.get('width', 0)
                    height = dimensions.get('height', 0)
                    
                    print(f"???????: length={length}, width={width}, height={height}")
                    
                    # ????????? ????? ?? ???????: (length * width * height) / 1000
                    volume = (length * width * height) / 1000 if all([length, width, height]) else 0
                    
                    dimensions_dict[int(nm_id)] = {
                        'length': length,
                        'width': width,
                        'height': height,
                        'volume': round(volume, 2)
                    }
                    print(f"???????? ????? {volume} ??? nm_id {nm_id}")
                else:
                    print(f"???????? ???????: nm_id={nm_id}, dimensions={dimensions}")
            
            # ????????? ????? ????? ?????????
            time.sleep(0.1)
        
        print(f"????????? {len(dimensions_dict)} ??????? ????????")
        return dimensions_dict
        
    except Exception as e:
        print(f"?????? ???????? ????????: {e}")
        import traceback
        traceback.print_exc()
        return {}
# ----------------------------
# ???????? ? API: ????????? ??????
# ----------------------------
@app.route("/product/<int:nm_id>", methods=["GET"]) 
@login_required
def product_analytics_page(nm_id: int):
    token = effective_wb_api_token(current_user)
    # ?????? ?? query, ????? ???? ?????? ??? ?????? ??????
    date_from = (request.args.get("date_from") or "").strip()
    date_to = (request.args.get("date_to") or "").strip()

    # ????????? ??????? ?????? ?????? ?? ???? ???????? ??? ????????? ????? ?????
    photo = None
    supplier_article = None
    barcode = None
    product_name = None
    try:
        prod_cached = load_products_cache() or {}
        for it in (prod_cached.get("items") or []):
            nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
            if str(nmv) == str(nm_id):
                photo = it.get("photo") or it.get("img")
                supplier_article = it.get("supplier_article") or it.get("vendorCode") or it.get("vendor_code")
                barcode = it.get("barcode")
                product_name = it.get("name") or it.get("title") or it.get("subject") or "??? ????????"
                break
    except Exception:
        photo = None
        supplier_article = None
        barcode = None
        product_name = None

    # ??????????????? ???? ??? ?????????? (????? ???? ???????)
    try:
        date_from_fmt = datetime.strptime(date_from, "%Y-%m-%d").strftime("%d.%m.%Y") if date_from else ""
        date_to_fmt = datetime.strptime(date_to, "%Y-%m-%d").strftime("%d.%m.%Y") if date_to else ""
    except Exception:
        date_from_fmt = date_from
        date_to_fmt = date_to

    return render_template(
        "product_analytics.html",
        nm_id=nm_id,
        photo=photo,
        product_name=product_name,
        supplier_article=supplier_article,
        barcode=barcode,
        date_from=date_from,
        date_to=date_to,
        date_from_fmt=date_from_fmt,
        date_to_fmt=date_to_fmt,
        token=token,
    )
@app.route("/api/product/<int:nm_id>", methods=["GET"]) 
@login_required
def api_product_analytics(nm_id: int):
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 400

    req_from = (request.args.get("date_from") or "").strip()
    req_to = (request.args.get("date_to") or "").strip()
    force_refresh = (request.args.get("force_refresh") or "").strip() in ("1", "true", "True", "on")

    if not req_from or not req_to:
        return jsonify({"error": "bad_dates"}), 400

    try:
        if force_refresh:
            raw_orders = fetch_orders_range(token, req_from, req_to)
            orders = to_rows(raw_orders, req_from, req_to)
            _update_period_cache_with_data(token, req_from, req_to, orders)
            cache_info = {"used_cache_days": 0, "fetched_days": len(list(_daterange_inclusive(req_from, req_to)))}
        else:
            orders, cache_info = get_orders_with_period_cache(token, req_from, req_to)

        # ?????????? ?? nm_id - ??????? ???????? ??? ?????? ?? ??????
        all_product_orders = []
        for r in orders:
            nmv = r.get("??????? WB") or r.get("nmId") or r.get("nmID")
            if str(nmv) == str(nm_id):
                all_product_orders.append(r)

        # ????????? ?? ???????? ? ??????????
        total_orders = len(all_product_orders)
        total_cancelled = len([r for r in all_product_orders if r.get("is_cancelled", False)])
        total_active = total_orders - total_cancelled
        
        # ??? ????????? ???????? ?????????? ?????? ???????? ??????
        filtered = [r for r in all_product_orders if not r.get("is_cancelled", False)]
        total_revenue = round(sum(float(r.get("???? ?? ??????? ????????") or 0) for r in filtered), 2)
        
        # ????????? ??????? ??????? ? ????
        try:
            from datetime import datetime
            date_from_dt = datetime.strptime(req_from, "%Y-%m-%d")
            date_to_dt = datetime.strptime(req_to, "%Y-%m-%d")
            days_count = (date_to_dt - date_from_dt).days + 1
            avg_daily_sales = round(total_active / days_count, 1) if days_count > 0 else 0
        except Exception:
            avg_daily_sales = 0

        # ????? ???????? - ?????????? ??? ?????? ??? ??????????? ??????????? ??????????
        o_counts_map, o_rev_map, o_cancelled_counts_map = aggregate_daily_counts_and_revenue(all_product_orders)
        labels = sorted(k for k in o_counts_map if k)
        series_orders = [o_counts_map.get(d, 0) for d in labels]
        series_cancelled = [o_cancelled_counts_map.get(d, 0) for d in labels]
        series_revenue = [round(o_rev_map.get(d, 0.0), 2) for d in labels]
        
        # ?????? (?? ???????? ???????)
        wh_summary = aggregate_by_warehouse_orders_only(filtered)

        # ??????? ??????? ?? ??????? ??? ????? ??????
        # ???????? ??? ????????? ????????? (SKU) ??? ??????? nm_id ?? ???? ????????
        product_barcodes: set[str] = set()
        try:
            prod_cached_full = load_products_cache() or {}
            for it in (prod_cached_full.get("items") or []):
                nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
                if str(nmv) == str(nm_id):
                    if it.get("barcode"):
                        product_barcodes.add(str(it.get("barcode")))
                    bars = it.get("barcodes") or []
                    if isinstance(bars, list):
                        for b in bars:
                            if b:
                                product_barcodes.add(str(b))
                    sizes = it.get("sizes") or []
                    if isinstance(sizes, list):
                        for s in sizes:
                            bl = s.get("skus") or s.get("barcodes")
                            if isinstance(bl, list):
                                for b in bl:
                                    if b:
                                        product_barcodes.add(str(b))
                    break
        except Exception:
            pass

        # ????????? ??? ???????? ? ?????????? ?? ??????? ?????? ??? ?????????? ????? ??????
        stock_by_warehouse: dict[str, int] = {}
        try:
            stocks_cached = load_stocks_cache() or {}
            for stock_item in (stocks_cached.get("items") or []):
                bc = str(stock_item.get("barcode") or "")
                if not bc or (product_barcodes and bc not in product_barcodes):
                    continue
                wname = stock_item.get("warehouse", "") or ""
                qty = int(stock_item.get("qty", 0) or 0)
                if wname:
                    stock_by_warehouse[wname] = stock_by_warehouse.get(wname, 0) + qty
        except Exception:
            stock_by_warehouse = {}

        # ?????????? ??????? ? ??????? ?? ???????
        warehouses_merged: list[dict[str, Any]] = []
        seen_warehouses: set[str] = set()
        for wh in (wh_summary or []):
            wname = wh.get("warehouse") or ""
            orders_count = int(wh.get("orders") or 0)
            warehouses_merged.append({
                "warehouse": wname,
                "orders": orders_count,
                "stock": int(stock_by_warehouse.get(wname, 0)),
            })
            seen_warehouses.add(wname)
        # ????????? ??????, ??? ???? ?????? ???????
        for wname, qty in stock_by_warehouse.items():
            if wname not in seen_warehouses:
                warehouses_merged.append({
                    "warehouse": wname,
                    "orders": 0,
                    "stock": int(qty),
                })
        # ??????? ?????? ??? ?????? ? ??? ????????
        warehouses_merged = [w for w in warehouses_merged if int(w.get("orders", 0)) > 0 or int(w.get("stock", 0)) > 0]
        # ?????????? ?? ???????? ????., ????? ?? ??????? ????.
        warehouses_merged.sort(key=lambda x: (int(x.get("orders", 0)), int(x.get("stock", 0))), reverse=True)

        # ???????? ???????? ? ?????? ???????? ???????? (?? ????)
        product_income_data = {}
        try:
            # ?????????? ???????????? ?????? ? ?????????
            supplies_cache = load_fbw_supplies_detailed_cache()
            if supplies_cache and supplies_cache.get("supplies_by_date"):
                supplies_by_date = supplies_cache["supplies_by_date"]
                
                # ???????? ??????? ??? ??????? ??????
                for date_str, barcodes_data in supplies_by_date.items():
                    total_income = 0
                    for barcode, qty in barcodes_data.items():
                        if barcode in product_barcodes:
                            total_income += qty
                    
                    if total_income > 0:
                        product_income_data[date_str] = total_income
                        
        except Exception:
            product_income_data = {}
        
        # ???????????? ???????? ???????? ? ?????? ???????? ? ??????
        current_stock_total = sum(stock_by_warehouse.values()) if stock_by_warehouse else 0
        series_stock = []
        running_stock = current_stock_total
        
        # ???? ?? ???? ? ???????? ???????
        for day in reversed(labels):
            daily_sales = o_counts_map.get(day, 0)
            daily_income = product_income_data.get(day, 0)
            
            # ????????? ??????? ? ??????? (?????????? ????????? ?? ??????)
            running_stock += daily_sales
            # ???????? ??????? (?????????? ????????? ?? ???????)
            running_stock -= daily_income
            
            series_stock.append(max(0, running_stock))  # ??????? ?? ????? ???? ?????????????
        
        # ????????????? ?????? ???????
        series_stock.reverse()

        # ?????????? ? ?????? - ????? ?????????? ?????? ?? ???? ???????
        photo = None
        product_name = None
        supplier_article = None
        barcode = None
        try:
            prod_cached = load_products_cache() or {}
            for it in (prod_cached.get("items") or []):
                nmv = it.get("nm_id") or it.get("nmId") or it.get("nmID")
                if str(nmv) == str(nm_id):
                    photo = it.get("photo") or it.get("img")
                    product_name = it.get("name") or it.get("title") or it.get("subject") or "??? ????????"
                    supplier_article = it.get("supplier_article") or it.get("vendorCode") or it.get("vendor_code")
                    barcode = it.get("barcode")
                    break
        except Exception:
            photo = None
            product_name = None
            supplier_article = None
            barcode = None

        date_from_fmt = datetime.strptime(req_from, "%Y-%m-%d").strftime("%d.%m.%Y") if req_from else ""
        date_to_fmt = datetime.strptime(req_to, "%Y-%m-%d").strftime("%d.%m.%Y") if req_to else ""

        # ???????? ??????? ????????? ???? ?? ?????? ???????
        price_history = []
        try:
            # ??????? ????????? ???????? ??????? ????
            prices_data = fetch_prices_data(token, [nm_id])
            print(f"DEBUG: Prices data for nm_id {nm_id}: {prices_data}")
            
            # ??????? ??????? ??? ?? ?????? ?????? ???????
            # ??????? ??????? ???? = ??????? ?? ??????? "???? ? ?????? ???? ??????" ?? ????
            print(f"DEBUG: Building daily average prices from {len(filtered)} active orders")
            print(f"DEBUG: Period: {req_from} to {req_to}")
            print(f"DEBUG: Product nm_id: {nm_id}")
            
            # ???????? ?????? ????????? ??????? ??? ???????
            for i, order in enumerate(filtered[:3]):
                print(f"DEBUG: Sample order {i+1}: {order}")
            
            daily_sum_client: Dict[str, float] = {}  # ???? ? ?????? ???? ?????? (???? ???????)
            daily_sum_seller: Dict[str, float] = {}  # ???? ?? ??????? ???????? (??? ????)
            daily_qty: Dict[str, int] = {}
            
            # ??????? ???????? ?????? ?? ???? ? ?????????
            for i, order in enumerate(filtered):
                # 1) ????
                order_date = (
                    order.get("????")
                    or order.get("???? ??????")
                    or order.get("???? ???????")
                    or order.get("sale_date")
                    or order.get("date")
                    or order.get("orderDate")
                    or order.get("lastChangeDate")
                )
                if not order_date:
                    continue
                try:
                    if isinstance(order_date, str):
                        # ????? ?????? ???? ?? 'T'
                        date_str = str(order_date).split('T')[0].split()[0]
                        parsed_date = None
                        for fmt in ("%Y-%m-%d", "%d.%m.%Y", "%Y/%m/%d"):
                            try:
                                parsed_date = datetime.strptime(date_str, fmt)
                                break
                            except Exception:
                                continue
                        if not parsed_date:
                            # ??????? ???????? ?? ?????? 10 ????????
                            try:
                                parsed_date = datetime.strptime(str(order_date)[:10], "%Y-%m-%d")
                            except Exception:
                                continue
                    else:
                        parsed_date = order_date
                    date_key = parsed_date.strftime("%d.%m.%Y")
                except Exception as date_err:
                    print(f"DEBUG: date parse error: {order_date} -> {date_err}")
                    continue

                # 2) ???? (??? ?????? ????)
                price_all_discounts = (
                    order.get("???? ? ?????? ???? ??????")
                    or order.get("priceWithAllDiscounts")
                    or order.get("retail_price")  # ?????? ???????? ???? ?? ??????? ??? retail_price
                )
                
                price_seller_discount = (
                    order.get("???? ?? ??????? ????????")
                    or order.get("priceWithSellerDiscount")
                    or order.get("seller_price")
                )

                # ???? ???? total 'retail_amount' ? ??????????, ???????? ????
                retail_amount = order.get("retail_amount") or order.get("sumWithAllDiscounts")

                # ?????????? ?????? ? ??????/???????
                qty_val_raw = order.get("quantity") or order.get("???-??") or order.get("qty") or order.get("sale_qty") or 1
                
                print(f"DEBUG: Order {i+1} - Date: {order_date}, Client price: {price_all_discounts}, Seller price: {price_seller_discount}, Qty: {qty_val_raw}")

                qty_val = 1
                try:
                    qty_val = int(qty_val_raw) if qty_val_raw is not None else 1
                except Exception:
                    qty_val = 1

                # ???? ??????? (? ?????? ???? ??????)
                price_client = None
                try:
                    if price_all_discounts is not None:
                        price_client = float(price_all_discounts)
                except Exception:
                    price_client = None

                if price_client is None and retail_amount is not None:
                    try:
                        amt = float(retail_amount)
                        if qty_val > 0:
                            price_client = amt / qty_val
                    except Exception:
                        pass

                # ???? ???????? (?? ??????? ????????)
                price_seller = None
                try:
                    if price_seller_discount is not None:
                        price_seller = float(price_seller_discount)
                except Exception:
                    price_seller = None

                if price_client is None and price_seller is None:
                    # ??? ???????? ??? ? ??????????
                    continue

                # ??????????? ????? ??? ????? ???
                if price_client is not None:
                    daily_sum_client[date_key] = daily_sum_client.get(date_key, 0.0) + price_client * qty_val
                if price_seller is not None:
                    daily_sum_seller[date_key] = daily_sum_seller.get(date_key, 0.0) + price_seller * qty_val
                
                daily_qty[date_key] = daily_qty.get(date_key, 0) + qty_val
            
            print(f"DEBUG: Daily sums client: {daily_sum_client}")
            print(f"DEBUG: Daily sums seller: {daily_sum_seller}")
            print(f"DEBUG: Daily qty: {daily_qty}")
            
            # ??????? ??????? ??? ?????? ??? ???? ? ????????? ????????? (??? ????????)
            all_dates = set(daily_sum_client.keys()) | set(daily_sum_seller.keys())
            sorted_dates = sorted(all_dates, key=lambda x: datetime.strptime(x, "%d.%m.%Y"))
            
            for date_key in sorted_dates:
                if daily_qty.get(date_key, 0) > 0:
                    price_entry = {"date": date_key}
                    
                    # ??????? ???? ???????
                    if date_key in daily_sum_client:
                        avg_price_client = round(daily_sum_client[date_key] / daily_qty[date_key], 2)
                        price_entry["price_client"] = avg_price_client
                        print(f"DEBUG: {date_key} - Client price: {avg_price_client} (qty {daily_qty[date_key]})")
                    
                    # ??????? ???? ????????
                    if date_key in daily_sum_seller:
                        avg_price_seller = round(daily_sum_seller[date_key] / daily_qty[date_key], 2)
                        price_entry["price_seller"] = avg_price_seller
                        print(f"DEBUG: {date_key} - Seller price: {avg_price_seller} (qty {daily_qty[date_key]})")
                    
                    price_history.append(price_entry)
                else:
                    print(f"DEBUG: {date_key} - No quantity data, skipping")
            
            print(f"DEBUG: Extracted price history from orders: {price_history}")
            
            # ?? ??????? ???????????????? ?????? - ?????????? ?????? ???????? ?????? ?? ???????
            print(f"DEBUG: Final price history length: {len(price_history)}")
            if len(price_history) == 0:
                print("DEBUG: No price history found from orders - this might be normal if no sales in period")
                        
        except Exception as e:
            print(f"?????? ????????? ??????? ???: {e}")
            price_history = []

        # ?? ??????? ???????? ??????? ??? - ?????????? ?????? ??? ? ????????? ?????????

        return jsonify({
            "kpi": {
                "total_orders": total_orders,
                "total_active_orders": total_active,
                "total_cancelled_orders": total_cancelled,
                "total_revenue": total_revenue,
                "total_stock": current_stock_total,
                "avg_daily_sales": avg_daily_sales,
                "updated_at": cache_info.get("updated_at") if isinstance(cache_info, dict) else "",
            },
            "series": {
                "labels": labels,
                "orders": series_orders,
                "cancelled": series_cancelled,
                "revenue": series_revenue,
                "stock": series_stock,
                "income": [product_income_data.get(day, 0) for day in labels],
            },
            "warehouses": wh_summary,
            "warehouses_sales_vs_stock": warehouses_merged,
            "product": {
                "nm_id": nm_id,
                "name": product_name,
                "supplier_article": supplier_article,
                "barcode": barcode,
                "photo": photo,
            },
            "date_from_fmt": date_from_fmt,
            "date_to_fmt": date_to_fmt,
            "cache_info": cache_info,
            "price_history": price_history,
        }), 200
    except requests.HTTPError as http_err:
        return jsonify({"error": f"api:{http_err.response.status_code}"}), 400
    except Exception as exc:
        return jsonify({"error": str(exc)}), 400
def fetch_prices_data(token: str, nm_ids: List[int]) -> Dict[int, Dict[str, Any]]:
    """???????? ?????? ? ????? ??????? ????? API Wildberries"""
    if not nm_ids:
        return {}
    
    try:
        print(f"??????? ???????? ???? ??? {len(nm_ids)} ???????")
        
        # ?????????? ?????????? API ???????? Postman
        headers = {"Authorization": f"Bearer {token}"}
        
        # GET ?????? ? ?????????? limit
        params = {"limit": 500}
        
        response = requests.get(
            DISCOUNTS_PRICES_API_URL,
            headers=headers,
            params=params,
            timeout=30
        )
        
        print(f"?????? ??????: {response.status_code}")
        
        if response.status_code == 200:
            result = response.json()
            print(f"????? API: {result}")
            
            prices_data = {}
            
            # ???????????? ????? ???????? ????????? ?? Postman
            if isinstance(result, dict) and "data" in result:
                list_goods = result["data"].get("listGoods", [])
                print(f"??????? {len(list_goods)} ??????? ? listGoods")
                
                for goods_item in list_goods:
                    nm_id = goods_item.get("nmID")
                    if nm_id:
                        # ???? ????????? ? ??????? sizes
                        sizes = goods_item.get("sizes", [])
                        print(f"????? {nm_id}: ??????? {len(sizes)} ????????")
                        
                        if sizes:
                            first_size = sizes[0]
                            price = first_size.get("price", 0)  # ??????? ????
                            discounted_price = first_size.get("discountedPrice", price)  # ???? ?? ???????
                            
                            # ???????? ??? ????????? ???? ??? ???????
                            print(f"DEBUG: ??? ???? ? first_size ??? ?????? {nm_id}: {list(first_size.keys())}")
                            
                            if price > 0:
                                club_discounted_price = first_size.get("clubDiscountedPrice", discounted_price)  # ???? ?? ??????? WB ????????
                                
                                # ????????? ?????? ???????? ??? ???????? ????? ??????? ????? ? ????? ?? ???????
                                seller_discount_amount = price - discounted_price if price > discounted_price else 0
                                
                                # ????????? ??????? ?????? ????????
                                seller_discount_percent = (seller_discount_amount / price * 100) if price > 0 else 0
                                
                                prices_data[nm_id] = {
                                    "price": price,  # ???? ?? ??????
                                    "discount_price": discounted_price,  # ???? ?? ???????
                                    "club_discount_price": club_discounted_price,  # ???? ?? ??????? WB ????????
                                    "seller_discount_amount": round(seller_discount_amount, 2),  # ?????? ???????? ? ??????
                                    "seller_discount_percent": round(seller_discount_percent, 2)  # ?????? ???????? ? ?????????
                                }
                                print(f"????? {nm_id}: ???? ?? ?????? {price}, ???? ?? ??????? {discounted_price}, ???? WB ???????? {club_discounted_price}, ?????? ???????? {seller_discount_amount} ???. ({seller_discount_percent}%)")
                        else:
                            print(f"????? {nm_id}: ??? ????????")
            
            if prices_data:
                print(f"??????? ???????? {len(prices_data)} ??? ?? API")
                return prices_data
            else:
                print("??? ?????? ? ????? ? ?????? ?? API")
        else:
            print(f"?????? {response.status_code}: {response.text}")
            
    except Exception as e:
        print(f"?????? ??? ????????? ???: {e}")
    
    print("?? ??????? ???????? ???? ?? API")
    return {}
@app.route("/products", methods=["GET"]) 
@login_required
def products_page():
    token = effective_wb_api_token(current_user)
    error = None
    products: List[Dict[str, Any]] = []
    if not token:
        error = "??????? ????? API ? ???????"
    else:
        try:
            # ?????? ?????? ?????? ?????? ? WB ??? ????????????? ?????????? ????
            raw_cards = fetch_all_cards(token, page_limit=100)
            products = normalize_cards_response({"cards": raw_cards})
            # ??????? ??? ? ???? ??? ?????? ???????, ?? ???????? /products ?????? ?????????? live-??????
            try:
                save_products_cache({"items": products})
            except Exception:
                pass
        except requests.HTTPError as http_err:
            error = f"?????? API: {http_err.response.status_code}"
        except Exception as exc:
            error = f"??????: {exc}"
    return render_template("products.html", error=error, items=products, items_count=len(products))


# -------------------------
# Stocks page
# -------------------------

def fetch_stocks_all(token: str) -> List[Dict[str, Any]]:
    """/supplier/stocks ?????? ??????? ??????? ????? ??????? ??? ?????????. ????? ?????? ?????? ?? ???? ??????."""
    headers1 = {"Authorization": f"Bearer {token}"}
    # WB ?????? ?????? 502/504 ? ??????? ????????? ???????? ? ?????????????? ?????????
    try:
        # ???? ?????? ??? ??????????? ???????, ????? ?? ??????? 429 ?? ????????
        resp = get_with_retry(STOCKS_API_URL, headers1, params={}, max_retries=1, timeout_s=30)
        return resp.json()
    except requests.HTTPError as err:
        # ???? ??????????? ? ????????? ??? Bearer
        if err.response is not None and err.response.status_code in (401, 403):
            headers2 = {"Authorization": f"{token}"}
            resp2 = get_with_retry(STOCKS_API_URL, headers2, params={}, max_retries=1, timeout_s=30)
            return resp2.json()
        # 429 ??????? ?????? ??? ???????? ? ????? ????? ??????? ??????
        raise


def fetch_stocks_paginated(token: str, start_iso: str = "1970-01-01T00:00:00") -> List[Dict[str, Any]]:
    headers = {"Authorization": f"Bearer {token}"}
    cursor = start_iso
    collected: List[Dict[str, Any]] = []
    safety = 0
    while True:
        safety += 1
        if safety > 5000:
            break
        params = {"dateFrom": cursor, "flag": 0}
        try:
            resp = get_with_retry(STOCKS_API_URL, headers, params, max_retries=3, timeout_s=30)
        except requests.HTTPError as err:
            if err.response is not None and err.response.status_code in (401, 403):
                alt_headers = {"Authorization": f"{token}"}
                resp = get_with_retry(STOCKS_API_URL, alt_headers, params, max_retries=3, timeout_s=30)
            else:
                raise
        page = resp.json()
        if not isinstance(page, list) or not page:
            break
        try:
            page.sort(key=lambda x: parse_wb_datetime(str(x.get("lastChangeDate"))) or datetime.min)
        except Exception:
            pass
        collected.extend(page)
        last_lcd = None
        try:
            last_lcd = page[-1].get("lastChangeDate")
        except Exception:
            last_lcd = None
        if not last_lcd:
            break
        cursor = str(last_lcd)
        time.sleep(0.1)
    return collected


# ?????????? ?????????? ??? ?????????????? ????????????? ???????? ? API ????????
_stocks_api_lock = threading.Lock()
_last_stocks_request_time = 0

def fetch_stocks_resilient(token: str) -> List[Dict[str, Any]]:
    global _last_stocks_request_time
    
    # ?????????? ?????????? ??? ?????????????? ????????????? ????????
    with _stocks_api_lock:
        # ?????????, ?? ?????? ?? ?? ?????? ??????? ??????? (??????? 1 ??????? ????? ?????????)
        current_time = time.time()
        if current_time - _last_stocks_request_time < 1.0:
            sleep_time = 1.0 - (current_time - _last_stocks_request_time)
            print(f"=== RATE LIMITING: ???? {sleep_time:.2f} ??? ????? ???????? ? API ???????? ===")
            time.sleep(sleep_time)
        
        _last_stocks_request_time = time.time()
        
        try:
            data = fetch_stocks_all(token)
            if isinstance(data, list) and data:
                return data
        except requests.HTTPError as e:
            # ???? 429 ? ?? ?????? ? ?????????, ?????????? 429
            try:
                if e.response is not None and e.response.status_code == 429:
                    raise
            except Exception:
                pass
            # ????? ????????? ??????????? (?????? ?????? ?????????????? ????????)
            return fetch_stocks_paginated(token)
        # Fallback to paginated flow
        return fetch_stocks_paginated(token)


def fetch_product_price_history(token: str, nm_id: int) -> List[Dict[str, Any]]:
    """???????? ??????? ????????? ???? ?????? ????? API Wildberries"""
    if not token or not nm_id:
        return []
    
    try:
        headers = {"Authorization": f"Bearer {token}"}
        url = f"{PRODUCT_HISTORY_API_URL}/{nm_id}"
        
        response = requests.get(url, headers=headers, timeout=30)
        
        if response.status_code == 200:
            data = response.json()
            # API ?????????? ?????? ???????? ? ???????? ?????????
            if isinstance(data, list):
                return data
            elif isinstance(data, dict) and 'data' in data:
                return data['data']
            else:
                return [data] if data else []
        else:
            print(f"?????? ????????? ??????? ???? ??? ?????? {nm_id}: {response.status_code}")
            return []
            
    except Exception as e:
        print(f"?????????? ??? ????????? ??????? ???? ??? ?????? {nm_id}: {e}")
        return []
def normalize_stocks(rows: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
    items: List[Dict[str, Any]] = []
    for r in rows or []:
        # On-hand stock per WB statistics API is in field 'quantity'.
        # In-transit can be represented by 'inWayToClient' and possibly 'inWayFromClient'.
        qty_val = r.get("quantity") or r.get("qty") or 0
        try:
            qty_int = int(qty_val)
        except Exception:
            try:
                qty_int = int(float(qty_val))
            except Exception:
                qty_int = 0
        # Collect in-transit (both directions if provided by API)
        in_way_to_client = r.get("inWayToClient") or 0
        in_way_from_client = r.get("inWayFromClient") or 0
        try:
            in_way_to_client = int(in_way_to_client)
        except Exception:
            try:
                in_way_to_client = int(float(in_way_to_client))
            except Exception:
                in_way_to_client = 0
        try:
            in_way_from_client = int(in_way_from_client)
        except Exception:
            try:
                in_way_from_client = int(float(in_way_from_client))
            except Exception:
                in_way_from_client = 0
        in_transit_total = max(0, in_way_to_client + in_way_from_client)
        items.append({
            "vendor_code": r.get("supplierArticle") or r.get("vendorCode") or r.get("article"),
            "barcode": r.get("barcode") or r.get("skus") or r.get("sku"),
            "nm_id": r.get("nmId") or r.get("nmID") or r.get("nm") or None,
            # Keep 'qty' as on-hand to not break existing aggregations that expect it
            "qty": qty_int,
            "in_transit": in_transit_total,
            "warehouse": r.get("warehouseName") or r.get("warehouse") or r.get("warehouse_name"),
        })
    return items
def update_stocks_if_needed(user_id: int, token: str, force_update: bool = False) -> bool:
    """
    ????????? ??????? ???? ????? (???? ??? ??????? ??? ?????????????)
    ?????????? True ???? ??????? ???? ?????????, False ???? ????????????? ???
    """
    try:
        cached = load_stocks_cache_for_user(user_id)
        should_refresh = force_update
        
        if not should_refresh and cached:
            # ?????????, ????? ????????? ??? ??????????? ???????
            updated_at = cached.get("updated_at")
            if updated_at:
                try:
                    from datetime import datetime
                    # ?????? ????? ?????????? ?? ????
                    cache_time = datetime.strptime(updated_at, "%d.%m.%Y %H:%M:%S")
                    # ???? ??????? ??????????? ????? 10 ????? ?????, ?????????? ???
                    if (datetime.now() - cache_time).total_seconds() < 600:  # 10 ?????
                        should_refresh = False
                        print(f"=== ????? ?? ???????: ?????????? ???????????? ??????? ===")
                        print(f"??? ????????: {updated_at}")
                    else:
                        should_refresh = True
                        print(f"=== ????? ?? ???????: ??? ???????, ????????? ??????? ===")
                except Exception as e:
                    print(f"?????? ???????? ??????? ????: {e}")
                    should_refresh = True
        else:
            should_refresh = True
            print(f"=== ????? ?? ???????: ??? ???? ??? ?????????????? ?????????? ===")
        
        if should_refresh:
            print(f"????????? ??????? ??? ???????????? {user_id}")
            try:
                raw_stocks = fetch_stocks_resilient(token)
                stocks = normalize_stocks(raw_stocks)
                from datetime import datetime
                now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
                save_stocks_cache_for_user(user_id, {"items": stocks, "updated_at": now_str})
                print(f"??????? ????????? ??? ?????? ?? ???????: {len(stocks)} ??????? ? {now_str}")
                return True
            except requests.HTTPError as e:
                if e.response and e.response.status_code == 429:
                    print("=== ????? ?? ???????: ?????? 429, ?????????? ??? ===")
                    if cached:
                        print(f"?????????? ???????????? ???????: {len(cached.get('items', []))} ???????")
                        return False
                    else:
                        print("??? ???? ? ?????? 429 - ?? ????? ???????? ???????")
                        return False
                else:
                    print(f"?????? ??? ?????????? ????????: {e}")
                    return False
        else:
            return False
            
    except Exception as e:
        print(f"?????? ? update_stocks_if_needed: {e}")
        return False


# Background stocks refresh guard (to avoid blocking page requests)
_stocks_bg_update_lock = threading.Lock()
_stocks_bg_update_users: set[int] = set()


def _start_stocks_update_bg(user_id: int, token: str) -> None:
    """
    ????????? update_stocks_if_needed ? ????, ?? ?? ?????????
    ???????????? ?????????? ???????? ??? ?????? ????????????.
    """
    if not token:
        return

    with _stocks_bg_update_lock:
        if user_id in _stocks_bg_update_users:
            return
        _stocks_bg_update_users.add(user_id)

    def _worker() -> None:
        try:
            update_stocks_if_needed(user_id, token, force_update=False)
        finally:
            with _stocks_bg_update_lock:
                _stocks_bg_update_users.discard(user_id)

    threading.Thread(target=_worker, daemon=True).start()


@app.route("/api/stocks/update-time", methods=["GET"])
@login_required
def api_stocks_update_time():
    """API ??? ????????? ??????? ?????????? ?????????? ????????"""
    try:
        cached = load_stocks_cache()
        if cached and cached.get("_user_id") == current_user.id:
            return jsonify({
                "updated_at": cached.get("updated_at", "??????????")
            })
        else:
            return jsonify({
                "updated_at": "??????? ?? ?????????"
            })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/stocks", methods=["GET"]) 
@login_required
def stocks_page():
    token = effective_wb_api_token(current_user)
    error = None
    items: List[Dict[str, Any]] = []
    if not token:
        error = "??????? ????? API ? ???????"
    else:
        try:
            cached = load_stocks_cache()
            if cached and cached.get("_user_id") == current_user.id:
                items = cached.get("items", [])
            else:
                raw = fetch_stocks_resilient(token)
                items = normalize_stocks(raw)
                save_stocks_cache({"items": items, "updated_at": datetime.now().strftime("%d.%m.%Y %H:%M:%S")})
        except requests.HTTPError as http_err:
            error = f"?????? API: {http_err.response.status_code}"
        except Exception as exc:
            error = f"??????: {exc}"
    # Aggregations
    total_qty_all = sum(int(it.get("qty", 0) or 0) for it in items)
    total_in_transit_all = sum(int(it.get("in_transit", 0) or 0) for it in items)
    # by product
    prod_map: Dict[tuple, Dict[str, Any]] = {}
    for it in items:
        key = (it.get("vendor_code") or "", it.get("barcode") or "")
        rec = prod_map.get(key)
        if not rec:
            rec = {"vendor_code": key[0], "barcode": key[1], "nm_id": it.get("nm_id"), "total_qty": 0, "total_in_transit": 0, "warehouses": []}
            prod_map[key] = rec
        rec["total_qty"] += int(it.get("qty", 0) or 0)
        rec["total_in_transit"] += int(it.get("in_transit", 0) or 0)
        rec["warehouses"].append({
            "warehouse": it.get("warehouse"),
            "qty": int(it.get("qty", 0) or 0),
            "in_transit": int(it.get("in_transit", 0) or 0),
        })
    for rec in prod_map.values():
        from collections import defaultdict as _dd
        qty_acc = _dd(int)
        transit_acc = _dd(int)
        for w in rec["warehouses"]:
            name = w.get("warehouse") or ""
            qty_acc[name] += int(w.get("qty", 0) or 0)
            transit_acc[name] += int(w.get("in_transit", 0) or 0)
        wh_list = [
            {"warehouse": name, "qty": qty, "in_transit": transit_acc.get(name, 0)}
            for name, qty in qty_acc.items() if qty > 0 or transit_acc.get(name, 0) > 0
        ]
        wh_list.sort(key=lambda x: (-x["qty"], x["warehouse"]))
        rec["warehouses"] = wh_list
    # Enrich with product photos from products cache
    try:
        nm_to_photo: Dict[Any, Any] = {}
        prod_cached = load_products_cache() or {}
        for it_p in (prod_cached.get("items") or []):
            nmv = it_p.get("nm_id") or it_p.get("nmId") or it_p.get("nmID")
            if nmv is None:
                continue
            photo = it_p.get("photo") or it_p.get("img")
            if photo:
                nm_to_photo[int(nmv)] = photo
        # attach photo to products_agg items
        for rec in prod_map.values():
            nm = rec.get("nm_id")
            if nm is not None:
                try:
                    rec["photo"] = nm_to_photo.get(int(nm))
                except Exception:
                    rec["photo"] = nm_to_photo.get(nm)
    except Exception:
        pass
    products_agg = sorted(prod_map.values(), key=lambda x: (-x["total_qty"], x["vendor_code"] or ""))
    # by warehouse
    wh_map: Dict[str, Dict[str, Any]] = {}
    for it in items:
        w = it.get("warehouse") or ""
        rec = wh_map.get(w)
        if not rec:
            rec = {"warehouse": w, "total_qty": 0, "total_in_transit": 0, "products": []}
            wh_map[w] = rec
        qty_i = int(it.get("qty", 0) or 0)
        in_transit_i = int(it.get("in_transit", 0) or 0)
        rec["total_qty"] += qty_i
        rec["total_in_transit"] += in_transit_i
        rec["products"].append({
            "vendor_code": it.get("vendor_code"),
            "barcode": it.get("barcode"),
            "nm_id": it.get("nm_id"),
            "qty": qty_i,
            "in_transit": in_transit_i,
        })
    for rec in wh_map.values():
        rec["products"].sort(key=lambda x: (-x["qty"], x["vendor_code"] or ""))
    # Enrich nested products with photos as well
    try:
        nm_to_photo: Dict[Any, Any] = {}
        prod_cached = load_products_cache() or {}
        for it_p in (prod_cached.get("items") or []):
            nmv = it_p.get("nm_id") or it_p.get("nmId") or it_p.get("nmID")
            if nmv is None:
                continue
            photo = it_p.get("photo") or it_p.get("img")
            if photo:
                nm_to_photo[int(nmv)] = photo
        for wh in wh_map.values():
            for p in wh.get("products", []):
                nm = p.get("nm_id")
                if nm is not None:
                    try:
                        p["photo"] = nm_to_photo.get(int(nm))
                    except Exception:
                        p["photo"] = nm_to_photo.get(nm)
    except Exception:
        pass
    warehouses_agg = sorted(wh_map.values(), key=lambda x: (-x["total_qty"], x["warehouse"] or ""))
    updated_at = None
    try:
        cached = load_stocks_cache()
        if cached and cached.get("_user_id") == current_user.id:
            updated_at = cached.get("updated_at")
    except Exception:
        updated_at = None
    return render_template(
        "stocks.html",
        error=error,
        items=items,
        items_count=len(items),
        total_qty_all=total_qty_all,
        updated_at=updated_at,
        products_agg=products_agg,
        warehouses_agg=warehouses_agg,
    )


@app.route("/api/stocks/refresh", methods=["POST"]) 
@login_required
def api_stocks_refresh():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        raw = fetch_stocks_resilient(token)
        items = normalize_stocks(raw)
        now_str = datetime.now().strftime("%d.%m.%Y %H:%M:%S")
        save_stocks_cache({"items": items, "updated_at": now_str})
        return jsonify({"ok": True, "count": len(items), "updated_at": now_str})
    except requests.HTTPError as http_err:
        return jsonify({"error": "http", "status": http_err.response.status_code}), 502
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500
@app.route("/api/stocks/data", methods=["GET"]) 
@login_required
def api_stocks_data():
    cached = load_stocks_cache()
    if not cached or not (current_user.is_authenticated and cached.get("_user_id") == current_user.id):
        return jsonify({"products": [], "warehouses": [], "total_qty_all": 0, "updated_at": None})
    items = cached.get("items", [])
    # Total
    try:
        total_qty_all = sum(int((it.get("qty") or 0)) for it in items)
    except Exception:
        total_qty_all = 0
    try:
        total_in_transit_all = sum(int((it.get("in_transit") or 0)) for it in items)
    except Exception:
        total_in_transit_all = 0
    # by product (same shape as on page)
    prod_map: Dict[tuple, Dict[str, Any]] = {}
    for it in items:
        key = (it.get("vendor_code") or "", it.get("barcode") or "")
        rec = prod_map.get(key)
        if not rec:
            rec = {"vendor_code": key[0], "barcode": key[1], "nm_id": it.get("nm_id"), "total_qty": 0, "total_in_transit": 0, "warehouses": []}
            prod_map[key] = rec
        qty_i = int(it.get("qty", 0) or 0)
        rec["total_qty"] += qty_i
        rec["total_in_transit"] += int(it.get("in_transit", 0) or 0)
        rec["warehouses"].append({
            "warehouse": it.get("warehouse"),
            "qty": qty_i,
            "in_transit": int(it.get("in_transit", 0) or 0),
        })
    from collections import defaultdict as _dd
    products_agg = []
    for rec in prod_map.values():
        qty_acc = _dd(int)
        transit_acc = _dd(int)
        for w in rec["warehouses"]:
            name = w.get("warehouse") or ""
            qty_acc[name] += int(w.get("qty", 0) or 0)
            transit_acc[name] += int(w.get("in_transit", 0) or 0)
        rec["total_in_transit"] = sum(transit_acc.values())
        rec["warehouses"] = [
            {"warehouse": name, "qty": qty, "in_transit": transit_acc.get(name, 0)}
            for name, qty in qty_acc.items() if qty > 0 or transit_acc.get(name, 0) > 0
        ]
        rec["warehouses"].sort(key=lambda x: (-x["qty"], x["warehouse"]))
        products_agg.append(rec)
    # Enrich with product photos from products cache
    try:
        nm_to_photo: Dict[Any, Any] = {}
        prod_cached = load_products_cache() or {}
        for it_p in (prod_cached.get("items") or []):
            nmv = it_p.get("nm_id") or it_p.get("nmId") or it_p.get("nmID")
            if nmv is None:
                continue
            photo = it_p.get("photo") or it_p.get("img")
            if photo:
                nm_to_photo[int(nmv)] = photo
        for rec in products_agg:
            nm = rec.get("nm_id")
            if nm is not None:
                try:
                    rec["photo"] = nm_to_photo.get(int(nm))
                except Exception:
                    rec["photo"] = nm_to_photo.get(nm)
    except Exception:
        pass
    products_agg.sort(key=lambda x: (-x["total_qty"], x["vendor_code"] or ""))

    # by warehouse
    wh_map: Dict[str, Dict[str, Any]] = {}
    for it in items:
        w = it.get("warehouse") or ""
        rec = wh_map.get(w)
        if not rec:
            rec = {"warehouse": w, "total_qty": 0, "total_in_transit": 0, "products": []}
            wh_map[w] = rec
        qty_i = int(it.get("qty", 0) or 0)
        in_transit_i = int(it.get("in_transit", 0) or 0)
        rec["total_qty"] += qty_i
        rec["total_in_transit"] += in_transit_i
        rec["products"].append({
            "vendor_code": it.get("vendor_code"),
            "nm_id": it.get("nm_id"),
            "barcode": it.get("barcode"),
            "qty": qty_i,
            "in_transit": in_transit_i,
        })
    for rec in wh_map.values():
        rec["products"].sort(key=lambda x: (-x["qty"], x["vendor_code"] or ""))
    # Enrich nested products with photos as well
    try:
        nm_to_photo: Dict[Any, Any] = {}
        prod_cached = load_products_cache() or {}
        for it_p in (prod_cached.get("items") or []):
            nmv = it_p.get("nm_id") or it_p.get("nmId") or it_p.get("nmID")
            if nmv is None:
                continue
            photo = it_p.get("photo") or it_p.get("img")
            if photo:
                nm_to_photo[int(nmv)] = photo
        for wh in wh_map.values():
            for p in wh.get("products", []):
                nm = p.get("nm_id")
                if nm is not None:
                    try:
                        p["photo"] = nm_to_photo.get(int(nm))
                    except Exception:
                        p["photo"] = nm_to_photo.get(nm)
    except Exception:
        pass
    warehouses_agg = sorted(wh_map.values(), key=lambda x: (-x["total_qty"], x["warehouse"] or ""))

    return jsonify({
        "products": products_agg,
        "warehouses": warehouses_agg,
        "total_qty_all": total_qty_all,
        "updated_at": cached.get("updated_at"),
        "total_in_transit_all": total_in_transit_all,
    })

@app.route("/stocks/export", methods=["POST"]) 
@login_required
def stocks_export():
    token = effective_wb_api_token(current_user)
    if not token:
        return redirect(url_for("stocks_page"))
    try:
        cached = load_stocks_cache()
        if cached and cached.get("_user_id") == current_user.id:
            items = cached.get("items", [])
        else:
            raw = fetch_stocks_all(token)
            items = normalize_stocks(raw)
            save_stocks_cache({"items": items})

        wb = Workbook()
        ws = wb.active
        ws.title = "stocks"
        headers = ["??????? ????????", "??????", "???????", "? ????", "?????"]
        ws.append(headers)
        for it in items:
            ws.append([
                it.get("vendor_code", ""),
                it.get("barcode", ""),
                it.get("qty", 0),
                it.get("in_transit", 0),
                it.get("warehouse", ""),
            ])
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        from datetime import datetime as _dt
        filename = f"wb_stocks_{_dt.now().strftime('%Y%m%d_%H%M')}.xlsx"
        return send_file(output, as_attachment=True, download_name=filename, mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
    except Exception:
        return redirect(url_for("stocks_page"))


@app.route("/api/products/refresh", methods=["POST"]) 
@login_required
def api_products_refresh():
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    try:
        raw_cards = fetch_all_cards(token, page_limit=100)
        items = normalize_cards_response({"cards": raw_cards})
        save_products_cache({"items": items})
        return jsonify({"ok": True, "count": len(items)})
    except requests.HTTPError as http_err:
        return jsonify({"error": "http", "status": http_err.response.status_code}), 502
    except Exception as exc:
        return jsonify({"error": str(exc)}), 500

@app.route("/api/products/update-dimensions", methods=["POST"])
@login_required
def api_products_update_dimensions():
    """API ??? ?????????? ???????? ? ???? ???????"""
    token = effective_wb_api_token(current_user)
    if not token:
        return jsonify({"error": "no_token"}), 401
    
    try:
        data = request.get_json()
        updates = data.get("updates", [])
        
        # ????????? chrtID ???? ?????? ? ??????
        for u in updates:
            if not isinstance(u, dict):
                continue
            u.setdefault("vendorCode", "")
            u.setdefault("barcode", "")
            u.setdefault("chrtID", None)
        
        if not updates:
            return jsonify({"error": "??? ?????? ??? ??????????"}), 400
        
        # ????????? ???????? ????? API Wildberries
        result = update_cards_dimensions(token, updates)
        
        # ?????????, ???? ?? ???????? ??????
        if result.get("error") and result.get("error") != "??? ?????? ??? ??????????":
            return jsonify({"error": result.get("error", "??????????? ?????? API")}), 400
        
        # ????????? ??? ???????: ??????? ?????????? ?????? ????????? ???
        try:
            cached = load_products_cache() or {}
            items = cached.get("items") or []
            if items:
                # ??????? ????????? ?????????? ????????? ? ???? ??? ??????????? ???????????
                nm_to_item = {it.get("nm_id"): it for it in items}
                for upd in updates:
                    nm_id_u = upd.get("nmID")
                    dims_u = upd.get("dimensions") or {}
                    it = nm_to_item.get(nm_id_u)
                    if not it:
                        continue
                    it_dims = (it.get("dimensions") or {}).copy()
                    if "length" in dims_u:
                        it_dims["length"] = dims_u["length"]
                    if "width" in dims_u:
                        it_dims["width"] = dims_u["width"]
                    if "height" in dims_u:
                        it_dims["height"] = dims_u["height"]
                    if "weightBrutto" in dims_u:
                        it_dims["weight"] = dims_u["weightBrutto"]
                    # ???????? ?????? (?)
                    length = float(it_dims.get("length") or 0)
                    width = float(it_dims.get("width") or 0)
                    height = float(it_dims.get("height") or 0)
                    volume = (length * width * height) / 1000 if (length and width and height) else 0
                    it_dims["volume"] = round(volume, 2)
                    it["dimensions"] = it_dims
                
                save_products_cache({"items": items})
        except Exception as cache_patch_err:
            print(f"?????? ???????????? ????? ????: {cache_patch_err}")

        # ????? ???????? ????????? ?????? ?????? ?? WB (????? ??????????? ? ?????????)
        try:
            raw_cards = fetch_all_cards(token, page_limit=100)
            items_full = normalize_cards_response({"cards": raw_cards})
            if items_full:
                save_products_cache({"items": items_full})
        except Exception as cache_err:
            # ???????? ?????? ????, ?? ?? ????????? ??????????
            print(f"?????? ?????????? ????: {cache_err}")
        
        return jsonify({
            "ok": True, 
            "updated_count": result.get("updated_count", 0),
            "skipped": result.get("skipped", []),
            "errors": result.get("errors", []),
            "result": result
        })

    except Exception as exc:
        return jsonify({"error": str(exc)}), 500


@app.route("/api/products/export-excel", methods=["POST"])
@login_required
def api_products_export_excel():
    """??????? ??????? ? Excel ??????"""
    try:
        data = request.get_json()
        if not data or 'products' not in data:
            return jsonify({"error": "??? ?????? ??? ????????"}), 400
        
        products = data['products']
        
        if not products:
            return jsonify({"error": "??? ??????? ??? ????????"}), 400
        
        print(f"??????? Excel: ???????? {len(products)} ???????")
        
        # ??????? Excel ????
        wb = Workbook()
        ws = wb.active
        ws.title = "??????"
        
        # ?????????
        headers = [
            "?", 
            "??????? ????????", 
            "??????? WB",
            "??????",
            "?????, ??", 
            "??????, ??", 
            "??????, ??", 
            "????? ?.", 
            "???, ??"
        ]
        ws.append(headers)
        
        # ????? ??? ??????????
        from openpyxl.styles import Font, Alignment, PatternFill
        header_font = Font(bold=True)
        header_fill = PatternFill(start_color="CCCCCC", end_color="CCCCCC", fill_type="solid")
        header_alignment = Alignment(horizontal="center", vertical="center")
        
        for col_num, header in enumerate(headers, 1):
            cell = ws.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # ?????? ???????
        for product in products:
            nm_val = product.get("nm_id") if "nm_id" in product else product.get("nmID")
            row = [
                product.get("index", ""),
                product.get("supplier_article", ""),
                nm_val or "",
                product.get("barcode", ""),
                product.get("length", 0),
                product.get("width", 0),
                product.get("height", 0),
                product.get("volume", 0),
                product.get("weight", 0)
            ]
            ws.append(row)
        
        # ?????????? ?????? ???????
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # ????????? ? ??????
        output = io.BytesIO()
        wb.save(output)
        output.seek(0)
        
        # ?????????? ??? ?????
        from datetime import datetime
        filename = f"wb_products_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        
        return send_file(
            output, 
            as_attachment=True, 
            download_name=filename, 
            mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        )
        
    except Exception as e:
        print(f"?????? ???????? Excel: {e}")
        return jsonify({"error": f"?????? ????????: {str(e)}"}), 500

@app.route("/api/products/save-images", methods=["POST"])
@login_required
def api_products_save_images():
    """??????? ZIP-????? ? ?????????? ???????, ?????????? ?? ????????."""
    try:
        data = request.get_json()
        if not data or not data.get("products"):
            return jsonify({"error": "??? ?????? ? ???????"}), 400
        
        products = data["products"]
        if not products:
            return jsonify({"error": "?????? ??????? ????"}), 400
        
        # ??????????? ??????????? ??????
        import zipfile
        from PIL import Image
        import tempfile
        
        # ??????? ????????? ???? ??? ZIP-??????
        temp_zip = tempfile.NamedTemporaryFile(delete=False, suffix='.zip')
        temp_zip.close()
        
        # ??????? ZIP-?????
        with zipfile.ZipFile(temp_zip.name, 'w', zipfile.ZIP_DEFLATED) as zip_file:
            for product in products:
                barcode = product.get("barcode")
                image_url = product.get("imageUrl")
                
                if not barcode or not image_url:
                    continue
                
                try:
                    # ????????? ???????????
                    response = requests.get(image_url, timeout=30)
                    response.raise_for_status()
                    
                    # ???????????? ? JPEG ???? ?????
                    image = Image.open(io.BytesIO(response.content))
                    if image.mode in ('RGBA', 'LA', 'P'):
                        image = image.convert('RGB')
                    
                    # ????????? ? ?????? ??? JPEG
                    img_buffer = io.BytesIO()
                    image.save(img_buffer, format='JPEG', quality=95)
                    img_buffer.seek(0)
                    
                    # ????????? ? ZIP ? ?????? ?? ???????
                    filename = f"{barcode}.jpeg"
                    zip_file.writestr(filename, img_buffer.getvalue())
                    
                except Exception as e:
                    print(f"?????? ????????? ??????????? ??? ??????? {barcode}: {e}")
                    continue
        
        # ?????? ????????? ZIP-????
        with open(temp_zip.name, 'rb') as f:
            zip_data = f.read()
        
        # ??????? ????????? ????
        os.unlink(temp_zip.name)
        
        # ?????????? ZIP-????
        return send_file(
            io.BytesIO(zip_data),
            as_attachment=True,
            download_name=f"product_images_{time.strftime('%Y%m%d_%H%M')}.zip",
            mimetype="application/zip"
        )
        
    except Exception as e:
        print(f"?????? ???????? ?????? ? ??????????: {e}")
        return jsonify({"error": f"?????? ???????? ??????: {str(e)}"}), 500

@app.route('/favicon.ico')
def favicon():
    return send_from_directory(app.root_path, 'favicon.ico', mimetype='image/vnd.microsoft.icon')

@app.route('/logo.png')
def logo():
    return send_from_directory(os.path.join(app.root_path, 'templates'), 'logo.png', mimetype='image/png')


# -------------------------
# Auth routes
# -------------------------

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        username = request.form.get("username", "").strip()
        password = request.form.get("password", "")
        remember_flag = request.form.get("remember") in ("on", "true", "1")
        user = User.query.filter_by(username=username).first()
        # Authentication and account state checks
        if not user or user.password != password:
            return render_template("login.html", error="???????? ????? ??? ??????")
        if not user.is_active:
            return render_template("login.html", error="???? ??????? ?????? ?????????????, ?????????? ? ????????????")
        # Check validity dates
        today = datetime.now(MOSCOW_TZ).date()
        if user.valid_from and today < user.valid_from:
            return render_template("login.html", error="??????? ?????? ??? ?? ???????")
        if user.valid_to and today > user.valid_to:
            return render_template("login.html", error="???? ???????? ????? ???????? ?????")
        # ????????? ???????? "????????? ????":
        # - ???????? ?????????? ?????? ??? ????????? cookie ??????????
        # - ?????????? remember ??? Flask-Login, ????? ??????? ???????? ????? ?? 30 ????
        session.permanent = bool(remember_flag)
        login_user(user, remember=bool(remember_flag))
        return redirect(request.args.get("next") or url_for("index"))
    return render_template("login.html")


@app.route("/logout", methods=["GET"]) 
def logout():
    logout_user()
    return redirect(url_for("login"))


def admin_required():
    if not current_user.is_authenticated or not current_user.is_admin:
        return False
    return True


@app.route("/admin/users", methods=["GET"]) 
@login_required
def admin_users():
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        flash("? ??? ??? ???? ??? ??????? ? ???? ????????", "error")
        return redirect(url_for("index"))
    
    users = User.query.order_by(User.id.asc()).all()
    wb_expiry = {u.id: wb_api_key_expiry_summary(getattr(u, "wb_token", None)) for u in users}
    return render_template("admin_users.html", users=users, wb_expiry=wb_expiry, message=None)
@app.route("/admin/users/create", methods=["POST"]) 
@login_required
def admin_users_create():
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    username = request.form.get("username", "").strip()
    password = request.form.get("password", "")
    display_name = request.form.get("display_name", "").strip() or None
    is_admin = bool(request.form.get("is_admin"))
    vf = request.form.get("valid_from") or None
    vt = request.form.get("valid_to") or None
    if not username or not password:
        flash("??????? ????? ? ??????")
        return redirect(url_for("admin_users"))
    if User.query.filter_by(username=username).first():
        flash("????? ????? ??? ??????????")
        return redirect(url_for("admin_users"))
    try:
        from datetime import date
        vf_d = None
        vt_d = None
        if vf:
            try:
                vf_d = datetime.strptime(vf, "%Y-%m-%d").date()
            except Exception:
                vf_d = None
        if vt:
            try:
                vt_d = datetime.strptime(vt, "%Y-%m-%d").date()
            except Exception:
                vt_d = None
        u = User(username=username, password=password, display_name=display_name, is_admin=is_admin, is_active=True, valid_from=vf_d, valid_to=vt_d)
        db.session.add(u)
        db.session.commit()
        flash("???????????? ??????")
    except Exception as exc:
        try:
            db.session.rollback()
        except Exception:
            pass
        app.logger.exception("admin_users_create failed")
        flash(f"?????? ???????? ????????????: {exc}")
    return redirect(url_for("admin_users"))
@app.route("/admin/users/<int:user_id>/block", methods=["POST"]) 
@login_required
def admin_users_block(user_id: int):
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    u = db.session.get(User, user_id)
    if u:
        try:
            u.is_active = False
            db.session.commit()
            flash("???????????? ????????????")
        except Exception:
            db.session.rollback()
            flash("??????")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/unblock", methods=["POST"]) 
@login_required
def admin_users_unblock(user_id: int):
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    u = db.session.get(User, user_id)
    if u:
        try:
            u.is_active = True
            db.session.commit()
            flash("???????????? ?????????????")
        except Exception:
            db.session.rollback()
            flash("??????")
    return redirect(url_for("admin_users"))
@app.route("/admin/users/<int:user_id>/reset", methods=["POST"]) 
@login_required
def admin_users_reset(user_id: int):
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    new_pass = request.form.get("password", "")
    if not new_pass:
        flash("??????? ????? ??????")
        return redirect(url_for("admin_users"))
    u = db.session.get(User, user_id)
    if u:
        try:
            u.password = new_pass
            db.session.commit()
            flash("?????? ????????")
        except Exception:
            db.session.rollback()
            flash("??????")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/delete", methods=["POST"]) 
@login_required
def admin_users_delete(user_id: int):
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    if current_user.id == user_id:
        flash("?????? ??????? ???? ??????? ??????")
        return redirect(url_for("admin_users"))
    try:
        if delete_user_with_related(user_id):
            try:
                cache_path = _cache_path_for_user_id(user_id)
                if os.path.isfile(cache_path):
                    os.remove(cache_path)
            except Exception:
                pass
            flash("???????????? ??????")
        else:
            flash("???????????? ?? ??????")
    except Exception as exc:
        flash(f"?????? ????????: {exc}")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/validity", methods=["POST"]) 
@login_required
def admin_users_validity(user_id: int):
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    vf = request.form.get("valid_from") or None
    vt = request.form.get("valid_to") or None
    u = db.session.get(User, user_id)
    if u:
        try:
            vf_d = None
            vt_d = None
            if vf:
                try:
                    vf_d = datetime.strptime(vf, "%Y-%m-%d").date()
                except Exception:
                    vf_d = None
            if vt:
                try:
                    vt_d = datetime.strptime(vt, "%Y-%m-%d").date()
                except Exception:
                    vt_d = None
            u.valid_from = vf_d
            u.valid_to = vt_d
            db.session.commit()
            flash("???? ???????? ????????")
        except Exception as exc:
            try:
                db.session.rollback()
            except Exception:
                pass
            app.logger.exception("admin_users_validity failed")
            flash(f"?????? ?????????? ?????: {exc}")
    return redirect(url_for("admin_users"))


@app.route("/admin/users/<int:user_id>/display_name", methods=["POST"]) 
@login_required
def admin_users_display_name(user_id: int):
    # ????????? ????? ??????????????
    if not current_user.is_authenticated or not current_user.is_admin:
        return jsonify({"success": False, "error": "? ??? ??? ???? ??? ?????????? ????? ????????"}), 403
    display_name = request.form.get("display_name", "").strip() or None
    u = db.session.get(User, user_id)
    if u:
        try:
            u.display_name = display_name
            db.session.commit()
            flash("??? ???????????? ?????????")
        except Exception as exc:
            try:
                db.session.rollback()
            except Exception:
                pass
            app.logger.exception("admin_users_display_name failed")
            flash(f"?????? ?????????? ????? ????????????: {exc}")
    return redirect(url_for("admin_users"))


# -------------------------
# Template context: subscription banner
# -------------------------

def _banner_visible_today(cookie_name: str, today) -> bool:
    """False ???? ???????????? ????? ?????? cookie ?? ??????? ??? ????? (??? ? ????????)."""
    hide_until = request.cookies.get(cookie_name)
    if not hide_until:
        return True
    try:
        hide_date = datetime.strptime(hide_until, "%Y-%m-%d").date()
        return hide_date < today
    except Exception:
        return True


@app.context_processor
def inject_subscription_banner():
    banner = {"show": False}
    api_token_banner = {"show": False}
    try:
        if current_user.is_authenticated:
            today = datetime.now(MOSCOW_TZ).date()
            if current_user.valid_to:
                days_left = (current_user.valid_to - today).days
                if 0 <= days_left <= 5 and _banner_visible_today("hide_sub_banner_until", today):
                    banner = {
                        "show": True,
                        "days_left": days_left,
                        "end_date": current_user.valid_to.strftime("%d.%m.%Y"),
                    }
            b = wb_api_key_expiry_banner(getattr(current_user, "wb_token", None))
            if b.get("show") and _banner_visible_today("hide_wb_token_banner_until", today):
                api_token_banner = {
                    "show": True,
                    "days_left": b.get("days_left"),
                    "end_date": b.get("end_date") or "",
                }
    except Exception:
        pass
    return {
        "subscription_banner": banner,
        "api_token_banner": api_token_banner,
        "app_version": _read_version(),
    }


@app.route("/changelog")
def changelog_page():
    md = _read_changelog_md()
    return render_template("changelog.html", app_version=_read_version(), md=md)


@app.route("/changelog/edit", methods=["GET", "POST"]) 
@login_required
def changelog_edit():
    if not current_user.is_admin:
        return redirect(url_for("changelog_page"))
    error = None
    message = None
    current_version = _read_version()
    md_content = _read_changelog_md()
    if request.method == "POST":
        try:
            new_version = (request.form.get("version") or "").strip()
            new_md = request.form.get("md_content")
            if new_md is not None:
                _write_changelog_md(new_md)
                md_content = new_md
            if new_version:
                _write_version(new_version)
                current_version = new_version
            message = "?????????"
        except Exception as exc:
            error = f"??????: {exc}"
    return render_template(
        "changelog_edit.html",
        app_version=current_version,
        md_content=md_content,
        default_date=datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y"),
        error=error,
        message=message,
    )


_DB_INIT_DONE = False

@app.before_request
def _init_db_once_per_process():
    global _DB_INIT_DONE
    if _DB_INIT_DONE:
        # ?? ?????? ?????????? ????? ???????/?????? ? ??? ?????????? ????????:
        # ????????? ??????? ??????? ?????????????? ? ??????? ??? ?????????????.
        try:
            from sqlalchemy import inspect as _inspect
            if not _inspect(db.engine).has_table("margin_calculations"):
                db.create_all()
        except Exception:
            pass
        return
    try:
        db.create_all()
        _ensure_schema_users_validity_columns()
        if not User.query.filter_by(username="admin").first():
            db.session.add(User(username="admin", password="admin", is_admin=True, is_active=True))
            db.session.commit()
        
        # Start notification monitoring
        start_notification_monitoring()
        
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
    finally:
        _DB_INIT_DONE = True


# -------------------------
# Tools: Labels for boxes
# -------------------------

@app.route("/tools/labels", methods=["GET"]) 
@login_required
def tools_labels_page():
    token = effective_wb_api_token(current_user)
    warehouses = []
    if token:
        try:
            headers = {"Authorization": f"Bearer {token}"}
            resp = get_with_retry(SUPPLIES_WAREHOUSES_URL, headers, params={})
            warehouses = resp.json() or []
        except Exception:
            try:
                headers2 = {"Authorization": f"{token}"}
                resp = get_with_retry(SUPPLIES_WAREHOUSES_URL, headers2, params={})
                warehouses = resp.json() or []
            except Exception:
                warehouses = []
    return render_template(
        "tools_labels.html",
        warehouses=warehouses,
    )
@app.route("/tools/labels/download", methods=["POST"]) 
@login_required
def tools_labels_download():
    # Inputs
    warehouse_name = (request.form.get("warehouse") or "").strip()
    boxes = int(request.form.get("boxes") or 0)
    if boxes <= 0:
        return jsonify({"error": "bad_boxes"}), 400

    shipper_name = (request.form.get("shipper_name") or current_user.shipper_name or "").strip()
    contact_person = (request.form.get("contact_person") or getattr(current_user, 'contact_person', None) or "").strip()
    phone = (request.form.get("phone") or current_user.phone or "").strip()
    email = (request.form.get("email") or current_user.email or "").strip()
    address = (request.form.get("shipper_address") or current_user.shipper_address or "").strip()

    # Build DOCX
    doc = Document()
    section = doc.sections[0]
    section.left_margin = Cm(1.2)
    section.right_margin = Cm(1.2)
    section.top_margin = Cm(0.7)
    section.bottom_margin = Cm(0.7)

    # Flow labels as table rows; prevent row splitting so ?? ???????? ???????? ?????? ????? ????????
    labels_table = doc.add_table(rows=0, cols=1)
    labels_table.autofit = True

    for n in range(1, boxes + 1):
        row = labels_table.add_row()
        # ?????? ??????? ?????? ??????? ????? ??????????
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        cantSplit = OxmlElement('w:cantSplit')
        trPr.append(cantSplit)

        cell = row.cells[0]

        # Title line: increase font (approx 28pt)
        p1 = cell.add_paragraph()
        p1.paragraph_format.space_before = Pt(0)
        p1.paragraph_format.space_after = Pt(0)
        p1.paragraph_format.line_spacing = 1.0
        r1 = p1.add_run(f"????????? ?? WB {warehouse_name}")
        r1.bold = True
        r1.font.size = Pt(28)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Counter line: increase font (approx 32pt)
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(2)
        p2.paragraph_format.line_spacing = 1.0
        r2 = p2.add_run(f"{n} ?? {boxes} ???????")
        r2.bold = True
        r2.font.size = Pt(32)
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Supplier line in one sentence
        supplier_line = ", ".join(filter(None, [shipper_name, f"?????????? ????: {contact_person}" if contact_person else None, phone, email, address]))
        p3 = cell.add_paragraph(supplier_line)
        p3.paragraph_format.space_before = Pt(0)
        p3.paragraph_format.space_after = Pt(2)
        p3.paragraph_format.line_spacing = 1.0
        p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p3.runs:
            run.font.size = Pt(11)

        # Horizontal line separator in the cell
        p_sep = cell.add_paragraph()
        p_sep.paragraph_format.space_before = Pt(2)
        p_sep.paragraph_format.space_after = Pt(2)
        p_sep.paragraph_format.line_spacing = 1.0
        p_sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = p_sep._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '8')
        bottom.set(qn('w:space'), '0')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    out = io.BytesIO()
    doc.save(out)
    out.seek(0)
    fname = f"labels_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return send_file(out, as_attachment=True, download_name=fname, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


# Notification API endpoints
@app.route("/api/notifications/count", methods=["GET"])
@login_required
def api_notifications_count():
    """Get count of unread notifications"""
    count = get_unread_notifications_count(current_user.id)
    return jsonify({"count": count})


@app.route("/api/notifications", methods=["GET"])
@login_required
def api_notifications():
    """Get user notifications"""
    limit = int(request.args.get('limit', 20))
    notifications = get_user_notifications(current_user.id, limit)
    
    result = []
    for notif in notifications:
        data = None
        if notif.data:
            try:
                data = json.loads(notif.data)
            except:
                pass
                
        # ??????????? ????? ???????? ??????????? ? ?????????? ???????
        # ????? ??? ????????? ? ?????????? ???????, ?? ????? ???? naive
        if notif.created_at.tzinfo is None:
            # ???? ????? naive, ??????? ??? ??????????
            moscow_time = notif.created_at.replace(tzinfo=MOSCOW_TZ)
        else:
            # ???? ????? ??? ? timezone, ???????????? ? ??????????
            moscow_time = notif.created_at.astimezone(MOSCOW_TZ)
        
        formatted_time = moscow_time.strftime('%d.%m.%Y %H:%M')
        
        
        
        result.append({
            'id': notif.id,
            'title': notif.title,
            'message': notif.message,
            'type': notif.notification_type,
            'is_read': notif.is_read,
            'created_at': formatted_time,
            'data': data
        })
    
    return jsonify({"notifications": result})
@app.route("/api/notifications/<int:notification_id>/read", methods=["POST"])
@login_required
def api_mark_notification_read(notification_id: int):
    """Mark a notification as read"""
    success = mark_notification_as_read(notification_id, current_user.id)
    if success:
        return jsonify({"success": True})
    else:
        return jsonify({"error": "Notification not found"}), 404


@app.route("/api/notifications/read-all", methods=["POST"])
@login_required
def api_mark_all_notifications_read():
    """Mark all notifications as read"""
    count = mark_all_notifications_as_read(current_user.id)
    return jsonify({"success": True, "count": count})


@app.route("/api/notifications/<int:notification_id>/delete", methods=["DELETE"])
@login_required
def api_delete_notification(notification_id: int):
    """Delete a notification"""
    notification = Notification.query.filter_by(id=notification_id, user_id=current_user.id).first()
    if notification:
        db.session.delete(notification)
        db.session.commit()
        return jsonify({"success": True})
    else:
        return jsonify({"error": "Notification not found"}), 404



@app.route("/api/notifications/delete-all", methods=["DELETE"])
@login_required
def api_delete_all_notifications():
    """Delete all notifications for current user"""
    count = Notification.query.filter_by(user_id=current_user.id).delete()
    db.session.commit()
    return jsonify({"success": True, "count": count})


@app.route("/api/notifications/test", methods=["POST"])
@login_required
def test_notification():
    """Create a test notification for debugging"""
    try:
        create_notification(
            user_id=current_user.id,
            title="???????? ???????????",
            message="??? ???????? ??????????? ??? ???????? ???????",
            notification_type="test",
            created_at=datetime.now(MOSCOW_TZ)
        )
        return jsonify({"success": True, "message": "Test notification created"})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/tools/prices", methods=["GET"])
@login_required
def tools_prices_page():
    """???????? ?????????? ??????"""
    token = effective_wb_api_token(current_user)
    error = None
    products = []
    prices_data = {}
    
    if not token:
        error = "??????? ????? API ? ???????"
    else:
        try:
            # ????????? ?????? ?? ???? ??? ? API
            cached = load_products_cache()
            if cached and cached.get("_user_id") == current_user.id:
                products = cached.get("items", [])
            else:
                # ????????? ??? ???????? ???????
                raw_cards = fetch_all_cards(token, page_limit=100)
                products = normalize_cards_response({"cards": raw_cards})
                save_products_cache({"items": products, "_user_id": current_user.id})
            
            # ???????? ???? ??????? ??? ???????
            if products:
                nm_ids = []
                for p in products:
                    nm_id = p.get("nm_id")
                    if nm_id:
                        try:
                            # ??????????? ? int, ???? ??? ??????
                            nm_ids.append(int(nm_id))
                        except (ValueError, TypeError):
                            print(f"?? ??????? ????????????? nm_id ? int: {nm_id}")
                            continue
                
                print(f"??????? {len(nm_ids)} ???????? nm_id ??? ??????? ???")
                if nm_ids:
                    # ???????? ???? ??? ???? ???????
                    prices_data = fetch_prices_data(token, nm_ids)
            
            # ???????? ?????? ? ?????????
            commission_data = {}
            try:
                commission_data = fetch_commission_data(token)
                print(f"????????? {len(commission_data)} ????????")
            except Exception as e:
                print(f"?????? ??? ???????? ????????: {e}")
                commission_data = {}
            
            # ???????? ?????? ? ???????? ??????? ?? ????????
            dimensions_data = {}
            try:
                for product in products:
                    nm_id = product.get('nm_id')
                    dimensions = product.get('dimensions', {})
                    if nm_id and dimensions:
                        dimensions_data[nm_id] = dimensions
                print(f"????????? {len(dimensions_data)} ??????? ???????? ?? ????????")
            except Exception as e:
                print(f"?????? ??? ????????? ????????: {e}")
                dimensions_data = {}
            
            # ???????? ?????? ? ???????
            warehouses_data = []
            try:
                warehouses_data = fetch_warehouses_data(token)
                print(f"????????? {len(warehouses_data)} ???????")
                if warehouses_data:
                    print(f"?????? ?????: {warehouses_data[0]}")
                else:
                    print("?????? ??????? ????!")
            except Exception as e:
                print(f"?????? ??? ???????? ???????: {e}")
                import traceback
                traceback.print_exc()
                warehouses_data = []
            
            # ???????? ?????? ?? ???????? FBW
            stocks_data = {}
            try:
                stocks_cached = load_stocks_cache()
                if stocks_cached and stocks_cached.get("_user_id") == current_user.id:
                    items = stocks_cached.get("items", [])
                    # ?????????? ??????? ?? ?????????
                    for stock_item in items:
                        barcode = stock_item.get("barcode")
                        qty = int(stock_item.get("qty", 0) or 0)
                        if barcode:
                            if barcode not in stocks_data:
                                stocks_data[barcode] = 0
                            stocks_data[barcode] += qty
                    print(f"????????? ???????? ??? {len(stocks_data)} ???????")
                else:
                    print("??? ???????? ?? ?????? ??? ???????")
            except Exception as e:
                print(f"?????? ??? ???????? ????????: {e}")
                stocks_data = {}

            # ???????? ?????? ?? ???????? FBS (????? ?? ???? FBS-???????)
            fbs_stocks_data: dict[str, int] = {}
            try:
                # ???????? ?????? ???????? (SKU) ?? ???? ???????
                prod_cached = load_products_cache() or {}
                products_all = prod_cached.get("products") or prod_cached.get("items") or []
                skus: list[str] = []
                for p in products_all:
                    if isinstance(p.get("barcodes"), list):
                        skus.extend([str(x) for x in p.get("barcodes") if x])
                    elif p.get("barcode"):
                        skus.append(str(p.get("barcode")))
                # ?????????? SKU
                skus = list({s for s in skus if s})
                if skus:
                    wlist = fetch_fbs_warehouses(token)
                    for w in wlist or []:
                        wid = w.get("id") or w.get("warehouseId") or w.get("warehouseID")
                        if not wid:
                            continue
                        try:
                            stocks = fetch_fbs_stocks_by_warehouse(token, int(wid), skus)
                        except Exception:
                            stocks = []
                        for s in stocks or []:
                            bc = str(s.get("sku") or s.get("barcode") or "").strip()
                            amount = int(s.get("amount") or 0)
                            if not bc:
                                continue
                            fbs_stocks_data[bc] = fbs_stocks_data.get(bc, 0) + amount
                print(f"????????? FBS ???????? ??? {len(fbs_stocks_data)} ???????")
            except Exception as e:
                print(f"?????? ??? ???????? FBS ????????: {e}")
                fbs_stocks_data = {}
            
            # ????????? ????? ????????????
            margin_settings = load_user_margin_settings(current_user.id)

            # ?????????? ?????????? ? ???????
            if products:
                print(f"????????? subject_id ? ??????? ? ???????:")
                found_commissions = 0
                found_dimensions = 0
                for i, product in enumerate(products[:5]):  # ?????? 5 ???????
                    subject_id = product.get('subject_id')
                    nm_id = product.get('nm_id')
                    dimensions = product.get('dimensions', {})
                    print(f"????? {i+1}: subject_id = {subject_id}, nm_id = {nm_id}")
                    print(f"  -> ???????: {dimensions}")
                    if dimensions and dimensions.get('volume', 0) > 0:
                        found_dimensions += 1
                    if subject_id and commission_data:
                        if subject_id in commission_data:
                            print(f"  -> ??????? ????????: {commission_data[subject_id]}")
                            found_commissions += 1
                        else:
                            print(f"  -> ???????? ?? ??????? ??? subject_id {subject_id}")
                print(f"??????? ???????? ??? {found_commissions} ?? {min(5, len(products))} ???????")
                print(f"??????? ???????? ??? {found_dimensions} ?? {min(5, len(products))} ???????")
            
            # ???????? ??????????? ?????????? ???? ?? ???? ??????
            purchase_prices = {}
            try:
                saved_prices = PurchasePrice.query.filter_by(user_id=current_user.id).all()
                print(f"??????? {len(saved_prices)} ??????? ?????????? ??? ? ?? ??? ???????????? {current_user.id}")
                for price_record in saved_prices:
                    if price_record.barcode:
                        barcode_str = str(price_record.barcode).strip()
                        if barcode_str.endswith('.0'):
                            try:
                                int(float(barcode_str))
                                barcode_str = barcode_str[:-2]
                            except (ValueError, TypeError):
                                pass
                        purchase_prices[barcode_str] = float(price_record.price)
                print(f"????????? {len(purchase_prices)} ??????????? ?????????? ???")
                if purchase_prices:
                    print(f"??????? ??????????? ???????? ?? ??: {list(purchase_prices.keys())[:5]}")
                if products:
                    product_barcodes = [str(p.get('barcode', '')).strip() for p in products[:5] if p.get('barcode')]
                    print(f"??????? ???????? ?? ???????: {product_barcodes}")
            except Exception as e:
                print(f"?????? ??? ???????? ?????????? ???: {e}")
                purchase_prices = {}
                    
        except requests.HTTPError as http_err:
            error = f"?????? API: {http_err.response.status_code}"
        except Exception as exc:
            error = f"??????: {exc}"
    
    # ????? ?????????? ?????????? ??? ?? ?????? ??????? ????????
    prices_last_updated = datetime.now(MOSCOW_TZ).strftime("%d.%m.%Y %H:%M")

    return render_template(
        "tools_prices.html",
        products=products,
        prices_data=prices_data,
        commission_data=commission_data,
        dimensions_data=dimensions_data if 'dimensions_data' in locals() else {},
        warehouses_data=warehouses_data if 'warehouses_data' in locals() else [],
        stocks_data=stocks_data if 'stocks_data' in locals() else {},
            fbs_stocks_data=fbs_stocks_data if 'fbs_stocks_data' in locals() else {},
        purchase_prices=purchase_prices,
        margin_settings=margin_settings if 'margin_settings' in locals() else load_user_margin_settings(current_user.id),
        prices_last_updated=prices_last_updated,
        error=error,
        token=token
    )
@app.route("/api/prices/upload", methods=["POST"])
@login_required
def api_prices_upload():
    """???????? ?????????? ??? ?? Excel ?????"""
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "???? ?? ??????"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "???? ?? ??????"}), 400
        
        if not file.filename.lower().endswith(('.xlsx', '.xls')):
            return jsonify({"success": False, "error": "?????????????? ?????? Excel ????? (.xlsx, .xls)"}), 400
        
        # ?????? Excel ???? ? ??????????? ?? ???????
        prices = {}
        updated_count = 0
        
        if file.filename.lower().endswith('.xlsx'):
            # ????? ?????? Excel (.xlsx)
            workbook = load_workbook(file, data_only=True)
            worksheet = workbook.active
            
            # ?????? ?????? ?? ?????? ???? ??????? (?????? ? ????)
            for row in worksheet.iter_rows(min_row=2, max_col=2, values_only=True):
                if len(row) >= 2 and row[0] and row[1]:
                    barcode = str(row[0]).strip()
                    if barcode.endswith('.0'):
                        try:
                            int(float(barcode))
                            barcode = barcode[:-2]
                        except (ValueError, TypeError):
                            pass
                    try:
                        price = float(row[1])
                        if price > 0:
                            prices[barcode] = price
                            updated_count += 1
                    except (ValueError, TypeError):
                        continue
        else:
            # ?????? ?????? Excel (.xls)
            file.seek(0)  # ?????????? ??????? ?????
            workbook = xlrd.open_workbook(file_contents=file.read())
            worksheet = workbook.sheet_by_index(0)
            
            # ?????? ?????? ?? ?????? ???? ??????? (?????? ? ????)
            for row_idx in range(1, worksheet.nrows):  # ?????????? ?????????
                if worksheet.ncols >= 2:
                    barcode_cell = worksheet.cell_value(row_idx, 0)
                    price_cell = worksheet.cell_value(row_idx, 1)
                    
                    if barcode_cell and price_cell:
                        barcode = str(barcode_cell).strip()
                        if barcode.endswith('.0'):
                            try:
                                int(float(barcode))
                                barcode = barcode[:-2]
                            except (ValueError, TypeError):
                                pass
                        try:
                            price = float(price_cell)
                            if price > 0:
                                prices[barcode] = price
                                updated_count += 1
                        except (ValueError, TypeError):
                            continue
        
        return jsonify({
            "success": True,
            "prices": prices,
            "updated_count": updated_count
        })
        
    except Exception as e:
        return jsonify({"success": False, "error": f"?????? ????????? ?????: {str(e)}"}), 500
@app.route("/api/prices/save", methods=["POST"])
@login_required
def api_prices_save():
    """?????????? ?????????? ???"""
    try:
        data = request.get_json() or {}
        prices = data.get('prices', {})
        
        if not prices:
            return jsonify({"success": False, "error": "??? ?????? ??? ??????????"}), 400
        
        saved_count = 0
        
        # ????????? ?????? ???? ? ???? ??????
        for barcode, price in prices.items():
            try:
                barcode_str = str(barcode).strip()
                if barcode_str.endswith('.0'):
                    try:
                        int(float(barcode_str))
                        barcode_str = barcode_str[:-2]
                    except (ValueError, TypeError):
                        pass

                existing_price = PurchasePrice.query.filter_by(
                    user_id=current_user.id, 
                    barcode=barcode_str
                ).first()
                
                if existing_price:
                    existing_price.price = float(price)
                    existing_price.updated_at = datetime.now(MOSCOW_TZ)
                else:
                    new_price = PurchasePrice(
                        user_id=current_user.id,
                        barcode=barcode_str,
                        price=float(price)
                    )
                    db.session.add(new_price)
                
                saved_count += 1
                
            except Exception as e:
                print(f"?????? ??? ?????????? ???? ??? ??????? {barcode}: {e}")
                continue
        
        # ????????? ????????? ? ???? ??????
        db.session.commit()
        
        return jsonify({
            "success": True,
            "saved_count": saved_count,
            "message": f"????????? {saved_count} ???"
        })
        
    except Exception as e:
        db.session.rollback()
        return jsonify({"success": False, "error": f"?????? ??????????: {str(e)}"}), 500


@app.route("/api/prices/export-excel", methods=["POST"])
@login_required
def api_prices_export_excel():
    """??????? ?????? ?????????? ?????? ? Excel ?????? XLS"""
    try:
        data = request.get_json()
        if not data or 'table_data' not in data:
            return jsonify({"error": "??? ?????? ??? ????????"}), 400
        
        table_data = data['table_data']
        visible_columns = data.get('visible_columns', [])
        margin_settings = data.get('margin_settings', {})
        
        print(f"??????? Excel: ???????? {len(table_data)} ????? ???????")
        print(f"??????? Excel: ??????? ???????: {visible_columns}")
        
        if not table_data:
            return jsonify({"error": "??? ?????? ??? ????????"}), 400
        
        # ??????? Excel ???? ? ??????? XLS (Excel 97-2003)
        workbook = xlwt.Workbook(encoding='utf-8')
        worksheet = workbook.add_sheet('?????????? ??????')
        
        # ?????
        header_style = xlwt.easyxf('font: bold on; align: horiz center;')
        number_style = xlwt.easyxf('align: horiz right;')
        text_style = xlwt.easyxf('align: horiz left;')
        
        # ?????????? ????????? ?? ?????? ??????? ???????
        column_mapping = {
            'index': '?',
            'photo': '????',
            'name': '????????????',
            'nm_id': '??????? WB',
            'barcode': '??????',
            'purchase': '?????????? ????',
            'price_before': '???? ?? ??????',
            'seller_discount': '?????? ????????',
            'price_discount': '???? ?? ??????',
            'price_wallet': '???? ?? ??????? (WB ????)',
            'category': '?????????',
            'volume': '????? ?.',
            'stocks': '???????',
            'commission_pct': '???????? WB, %',
            'commission_rub': '???????? WB ???.',
            'tax_rub': '????? ???.',
            'logistics_rub': '????????? ???.',
            'storage_rub': '???????? ???.',
            'receiving_rub': '??????? ???.',
            'acquiring_rub': '????????? ???.',
            'total_expenses': '????? ????????: ???.',
            'expenses_pct': '???????? ? %',
            'price_to_receive': '???? ? ????????? ???.',
            'profit_net': '??????? ?????? ???.',
            'profit_pct': '??????? %'
        }
        
        # ???? ?? ??????? ??????? ???????, ?????????? ???
        if not visible_columns:
            visible_columns = list(column_mapping.keys())
        
        # ????????? ????????? ?????? ??? ??????? ???????
        headers = [column_mapping.get(col, col) for col in visible_columns if col in column_mapping]
        
        for col, header in enumerate(headers):
            worksheet.write(0, col, header, header_style)
        
        # ?????? ?? ???????
        for row, row_data in enumerate(table_data, 1):
            # ?????????? ?????? ??????? ???????
            col_index = 0
            for col_key in visible_columns:
                if col_key in row_data:
                    value = row_data[col_key]
                    if col_key in ['commission_pct', 'expenses_pct', 'profit_pct', 'seller_discount']:
                        worksheet.write(row, col_index, value, number_style)
                    elif col_key in ['commission_rub', 'tax_rub', 'logistics_rub', 'storage_rub', 'receiving_rub', 'acquiring_rub', 'total_expenses', 'price_to_receive', 'profit_net', 'purchase', 'price_before', 'price_discount', 'price_wallet']:
                        worksheet.write(row, col_index, value, number_style)
                    elif col_key in ['nm_id', 'barcode']:
                        worksheet.write(row, col_index, str(value), text_style)
                    else:
                        worksheet.write(row, col_index, str(value), text_style)
                    col_index += 1
        
        # ?????????? ?????? ??????? ?? ?????? ??????? ???????
        column_widths_map = {
            'index': 1000,
            'photo': 1000,
            'name': 8000,
            'nm_id': 2000,
            'barcode': 2000,
            'purchase': 2000,
            'price_before': 2000,
            'seller_discount': 2000,
            'price_discount': 2000,
            'price_wallet': 2000,
            'category': 3000,
            'volume': 1500,
            'commission_pct': 1500,
            'commission_rub': 2000,
            'tax_rub': 2000,
            'logistics_rub': 2000,
            'storage_rub': 2000,
            'receiving_rub': 2000,
            'acquiring_rub': 2000,
            'total_expenses': 2000,
            'expenses_pct': 1500,
            'price_to_receive': 2000,
            'profit_net': 2000,
            'profit_pct': 1500
        }
        
        for col, col_key in enumerate(visible_columns):
            if col_key in column_widths_map:
                worksheet.col(col).width = column_widths_map[col_key]
        
        # ??????? ???? ? ??????
        output = BytesIO()
        workbook.save(output)
        output.seek(0)
        
        # ?????????? ??? ?????
        now = datetime.now()
        day = now.strftime("%d.%m.%Y")
        time = now.strftime("%H_%M")
        filename = f"??????????_??????_{day}_{time}.xls"
        
        return send_file(
            output,
            as_attachment=True,
            download_name=filename,
            mimetype='application/vnd.ms-excel'
        )
        
    except Exception as e:
        print(f"?????? ???????? ? Excel: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"error": f"?????? ????????: {str(e)}"}), 500



@app.route("/tools/prices/template", methods=["GET"])
@login_required
def download_prices_template():
    """??????? ?????? Excel ??? ???????? ?????????? ??? (XLS)."""
    try:
        workbook = xlwt.Workbook(encoding='utf-8')
        worksheet = workbook.add_sheet('?????? ?????????? ???')

        header_style = xlwt.easyxf('font: bold on; align: horiz center;')
        text_style = xlwt.easyxf('align: horiz left;')

        # ?????????
        worksheet.write(0, 0, '??????', header_style)
        worksheet.write(0, 1, '?????????? ????', header_style)

        # ?????? ??????
        worksheet.write(1, 0, '2001234567890', text_style)
        worksheet.write(1, 1, '123.45', text_style)

        worksheet.col(0).width = 5000
        worksheet.col(1).width = 5000

        output = BytesIO()
        workbook.save(output)
        output.seek(0)

        return send_file(
            output,
            as_attachment=True,
            download_name='??????_??????????_???.xls',
            mimetype='application/vnd.ms-excel'
        )
    except Exception as e:
        return jsonify({"error": f"?? ??????? ??????? ??????: {str(e)}"}), 500


@app.route("/api/tools/prices/margin-settings", methods=["GET", "POST"])
@login_required
def api_margin_settings():
    """?????????/?????????? ???????? ????? ???????? ????????????."""
    try:
        if request.method == "GET":
            settings = load_user_margin_settings(current_user.id)
            return jsonify({"success": True, "settings": settings})
        else:
            data = request.get_json() or {}
            saved = save_user_margin_settings(current_user.id, data)
            return jsonify({"success": True, "settings": saved})
    except Exception as e:
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/tools/prices/margins/save", methods=["POST"])
@login_required
def api_tools_prices_margins_save():
    """?????????? ?????????????? (profit_pct/profit_net) ?? /tools/prices ? ??."""
    try:
        payload = request.get_json() or {}
        items = payload.get("items")
        margins = payload.get("margins")

        records: list[dict[str, Any]] = []
        if isinstance(items, list):
            records = items
        elif isinstance(margins, dict):
            for nm_id, v in margins.items():
                if not isinstance(v, dict):
                    continue
                records.append(
                    {
                        "nm_id": nm_id,
                        "barcode": v.get("barcode"),
                        "profit_pct": v.get("profit_pct"),
                        "profit_net": v.get("profit_net"),
                    }
                )
        else:
            return jsonify({"success": False, "error": "??? ??????"}), 400

        from models import MarginCalculation

        saved_count = 0
        for rec in records:
            try:
                nm_id_raw = rec.get("nm_id")
                if nm_id_raw is None or nm_id_raw == "":
                    continue
                nm_id = int(float(nm_id_raw))

                profit_pct = rec.get("profit_pct")
                profit_net = rec.get("profit_net")
                barcode = rec.get("barcode")

                profit_pct_f = float(profit_pct) if profit_pct is not None and profit_pct != "" else None
                profit_net_f = float(profit_net) if profit_net is not None and profit_net != "" else None
                barcode_str = str(barcode).strip() if barcode else None

                existing = MarginCalculation.query.filter_by(user_id=current_user.id, nm_id=nm_id).first()
                if existing:
                    existing.barcode = barcode_str
                    existing.profit_pct = profit_pct_f
                    existing.profit_net = profit_net_f
                else:
                    db.session.add(
                        MarginCalculation(
                            user_id=current_user.id,
                            nm_id=nm_id,
                            barcode=barcode_str,
                            profit_pct=profit_pct_f,
                            profit_net=profit_net_f,
                        )
                    )
                saved_count += 1
            except Exception:
                continue

        db.session.commit()
        return jsonify({"success": True, "saved_count": saved_count})
    except Exception as e:
        try:
            db.session.rollback()
        except Exception:
            pass
        return jsonify({"success": False, "error": str(e)}), 500


@app.route("/api/tools/prices/warehouses", methods=["GET"])
@login_required
def api_tools_prices_warehouses():
    """?????????? ?????? ??????? ? ?????????????? ??? ??????????? ??????."""
    try:
        token = effective_wb_api_token(current_user)
        if not token:
            return jsonify({"success": False, "error": "?? ?????? WB ?????"}), 400
        warehouses = fetch_warehouses_data(token)
        return jsonify({"success": True, "warehouses": warehouses})
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)}), 500


if __name__ == "__main__":
    # ??????? ??????? ? ???? ??????
    with app.app_context():
        db.create_all()
    app.run(host="0.0.0.0", port=5000, debug=True)